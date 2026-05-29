
import base64
import csv
import hashlib
import hmac
import io
import json
import logging
import math
import os
import re
import difflib
import smtplib
import tempfile
import time
import zipfile
import uuid
import xml.etree.ElementTree as _StdET
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from email.message import EmailMessage
from html import escape
from pathlib import Path
from textwrap import dedent
from urllib.parse import quote
from typing import Any, Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

import pandas as pd
import requests
import streamlit as st
import streamlit.components.v1 as components
from openai import OpenAI

LOGGER = logging.getLogger(__name__)

try:
    import defusedxml.ElementTree as ET
except ImportError:
    ET = _StdET
    LOGGER.warning("defusedxml is not installed. Falling back to xml.etree.ElementTree.")

try:
    from managed_ai_router import ai_json_items, select_ai_route
except Exception as exc:
    LOGGER.warning("managed_ai_router import failed: %s", exc)
    ai_json_items = None
    select_ai_route = None

# ErrorSweep backend-only translation router:
# Built-in route = self-hosted commercial-safe MT router
# IndicTrans2 for Indian languages + OPUS-MT for supported global pairs
try:
    from translator_router import (
        translate_batch as builtin_translate_batch,
        current_builtin_engine_label,
        builtin_engine_status,
        smoke_test_builtin_engines,
    )
except Exception as exc:
    LOGGER.warning("translator_router import failed: %s", exc)
    builtin_translate_batch = None
    current_builtin_engine_label = None
    builtin_engine_status = None
    smoke_test_builtin_engines = None

# Speech-to-text helper for subtitle/transcription editor.
# v32 policy: auto transcription only uses the user's own API key; no-key users get blank manual rows.
try:
    from speech_transcription import transcribe_media_file_to_rows, transcribe_media_to_rows, speech_engine_label
except Exception as exc:
    LOGGER.warning("speech_transcription import failed: %s", exc)
    transcribe_media_file_to_rows = None
    transcribe_media_to_rows = None
    speech_engine_label = None



from openpyxl import load_workbook, Workbook, Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from docx import Document

# External editor job storage for v41.
# Stores Pro translated rows so a separate browser tab can open the CAT editor by job_id.
try:
    from editor_job_store import save_editor_job, load_editor_job, update_editor_job
except Exception as exc:
    LOGGER.warning("editor_job_store import failed: %s", exc)
    save_editor_job = None
    load_editor_job = None
    update_editor_job = None


# v42 production persistence. Supabase is used when configured; local JSON fallback keeps the MVP working.
try:
    from production_persistence import (
        save_persistent_editor_job,
        load_persistent_editor_job,
        update_persistent_editor_job,
        log_persistent_usage_event,
        fetch_persistent_usage_events,
        fetch_persistent_editor_jobs,
        save_saas_record,
        fetch_saas_records,
        persistence_health,
    )
except Exception as exc:
    LOGGER.warning("production_persistence import failed: %s", exc)
    save_persistent_editor_job = None
    load_persistent_editor_job = None
    update_persistent_editor_job = None
    log_persistent_usage_event = None
    fetch_persistent_usage_events = None
    fetch_persistent_editor_jobs = None
    save_saas_record = None
    fetch_saas_records = None
    persistence_health = None

try:
    from qa_engine_global_v15 import deterministic_checks_v2
except Exception as exc:
    LOGGER.warning("qa_engine_global_v15 import failed: %s", exc)
    deterministic_checks_v2 = None

try:
    from cloud_object_storage import (
        build_object_key,
        object_storage_status,
        put_file as put_object_file,
        signed_url_for_key,
    )
except Exception as exc:
    LOGGER.warning("cloud_object_storage import failed: %s", exc)
    build_object_key = None
    object_storage_status = None
    put_object_file = None
    signed_url_for_key = None

try:
    from async_worker_queue import async_backend_status, enqueue_async_task
except Exception as exc:
    LOGGER.warning("async_worker_queue import failed: %s", exc)
    async_backend_status = None
    enqueue_async_task = None

try:
    from billing_webhooks import normalize_billing_webhook, verify_billing_webhook_signature
except Exception as exc:
    LOGGER.warning("billing_webhooks import failed: %s", exc)
    normalize_billing_webhook = None
    verify_billing_webhook_signature = None



# ==========================================================
# ErrorSweep Platform
# Security hardening + production persistence + usage tracking + external CAT/media editor launcher
# Built-in MT: self-hosted IndicTrans2, MADLAD-400 when enabled, and OPUS-MT fallback.
# Editor jobs and usage persist to Supabase when configured, with local JSON fallback
# ==========================================================

APP_VERSION = "v46 Security + QA Workflow Hardening"
DEFAULT_MODEL = "gpt-4o-mini"
SESSION_TTL_SECONDS = 60 * 60 * 24 * 7
DEFAULT_SESSION_SECRET = "errorsweep-dev-session-secret-change-me"
PASSWORD_HASH_ITERATIONS = 260_000
SESSION_HISTORY_LIMIT = 500
RULE_ZIP_MAX_FILES = int(os.getenv("ERRORSWEEP_RULE_ZIP_MAX_FILES", "250"))
RULE_ZIP_MAX_BYTES = int(os.getenv("ERRORSWEEP_RULE_ZIP_MAX_BYTES", str(25 * 1024 * 1024)))
MEDIA_PREVIEW_TTL_SECONDS = int(os.getenv("ERRORSWEEP_MEDIA_PREVIEW_TTL_SECONDS", str(60 * 60 * 24 * 2)))
SESSION_COLLECTION_LIMITS = {
    "ai_usage_events": 500,
    "audit_logs": 500,
    "jobs": 500,
    "owner_recent_editor_jobs": 100,
    "payments": 500,
    "projects": 500,
    "files": 1000,
    "tm": 5000,
    "glossary": 5000,
    "dnt": 5000,
    "rule_instructions": 1000,
    "users": 1000,
    "workspaces": 1000,
    "notifications": 1000,
    "task_queue": 1000,
    "subscriptions": 500,
    "checkout_sessions": 500,
    "billing_events": 500,
    "auth_tokens": 500,
}
LANGUAGE_CATALOG = [
    "English",
    "French",
    "Spanish",
    "German",
    "Italian",
    "Portuguese",
    "Hindi",
    "Bengali",
    "Tamil",
    "Telugu",
    "Kannada",
    "Malayalam",
    "Marathi",
    "Gujarati",
    "Punjabi",
    "Urdu",
    "Arabic",
    "Persian",
    "Hebrew",
    "Russian",
    "Ukrainian",
    "Polish",
    "Turkish",
    "Greek",
    "Dutch",
    "Norwegian",
    "Swedish",
    "Danish",
    "Finnish",
    "Afrikaans",
    "Swahili",
    "Hausa",
    "Sinhala",
    "Zulu",
    "Amharic",
    "Yoruba",
    "Chinese",
    "Japanese",
    "Korean",
    "Thai",
    "Indonesian",
    "Malay",
    "Tagalog",
    "Burmese",
    "Khmer",
    "Lao",
    "Mongolian",
    "Vietnamese",
]
PLAN_CATALOG = [
    {
        "name": "Trial",
        "monthly": 0,
        "annual": 0,
        "currency": "INR",
        "trial_days": 14,
        "seats": 2,
        "segments": 500,
        "characters": 100_000,
        "label": "Free trial with mandate",
        "description": "Validate QA, Pro review, scorecards, and workspace setup. Card or UPI mandate required; cancel anytime before trial ends.",
    },
    {
        "name": "Pro",
        "monthly": 3999,
        "annual": 39990,
        "currency": "INR",
        "seats": 5,
        "segments": 10_000,
        "characters": 2_000_000,
        "label": "Growing teams",
        "description": "Production localization QA, Pro translation routing, and reviewer workflows.",
    },
    {
        "name": "Agency",
        "monthly": 11999,
        "annual": 119990,
        "currency": "INR",
        "seats": 20,
        "segments": 50_000,
        "characters": 10_000_000,
        "label": "Multi-client delivery",
        "description": "Higher-volume project, QA, subtitle, scorecard, and team management workflows.",
    },
    {
        "name": "Enterprise",
        "monthly": 0,
        "annual": 0,
        "currency": "INR",
        "seats": 100,
        "segments": 250_000,
        "characters": 50_000_000,
        "label": "Custom plan",
        "description": "Custom usage, SSO, security review, dedicated deployment, and guided onboarding.",
    },
]
EMAIL_DISPATCH_BATCH_LIMIT = int(os.getenv("ERRORSWEEP_EMAIL_DISPATCH_BATCH_LIMIT", "25"))
AUTH_TOKEN_TTL_SECONDS = int(os.getenv("ERRORSWEEP_AUTH_TOKEN_TTL_SECONDS", str(60 * 60 * 24)))
COMPLIANCE_ACK_LABEL = "I accept the [Terms of Service](?public=terms), [Privacy Policy](?public=privacy), and NDA/confidentiality obligations for this workspace."
SENSITIVE_TEXT_RE = re.compile(
    r"("
    r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}"
    r"|\b(?:\+?\d[\d .()-]{7,}\d)\b"
    r"|\b(?:api[_-]?key|secret|token|password|passwd|pwd)\b\s*[:=]"
    r"|\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b"
    r")",
    re.I,
)
LOCALIZATION_PROTECTED_RE = re.compile(
    r"(\{\{[^{}]+\}\}|\{[^{}]+\}|<[^>]+>|%[0-9$.\-+]*[sdif]|https?://\S+|www\.\S+|[\w.+-]+@[\w.-]+\.[A-Za-z]{2,})"
)
LOCALIZATION_VISUAL_PREFIX_RE = re.compile(
    r"^(\s*(?:(?:[•∙◦▪▫●○\-–—*]+)|(?:[\U0001F300-\U0001FAFF\u2600-\u27BF]\ufe0f?))+\s*)"
)
LOCALIZATION_EMOJI_RE = re.compile(r"[\U0001F300-\U0001FAFF\u2600-\u27BF]\ufe0f?")
TIME_DISPLAY_COLUMNS = {"created_at", "updated_at", "created", "updated", "time", "date", "login_at"}

st.set_page_config(
    page_title="ErrorSweep",
    page_icon="🧹",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# ==========================================================
# Visual system
# ==========================================================

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&family=Space+Mono:wght@400;700&display=swap');

:root {
  --es-bg: #070914;
  --es-panel: rgba(18, 21, 38, .86);
  --es-card: rgba(18, 21, 38, .78);
  --es-border: rgba(84, 105, 180, .35);
  --es-border-soft: rgba(84, 105, 180, .20);
  --es-text: #f7fbff;
  --es-muted: #a8b0d6;
  --es-green: #00d985;
  --es-cyan: #34bdf6;
  --es-purple: #8b5cf6;
  --es-red: #ff4b33;
  --es-amber: #f59e0b;
}

html, body, [class*="css"] {
  font-family: Inter, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
}

.stApp {
  background:
    radial-gradient(circle at 8% 12%, rgba(0, 217, 133, .18), transparent 26%),
    radial-gradient(circle at 88% 8%, rgba(52, 189, 246, .12), transparent 28%),
    radial-gradient(circle at 55% 96%, rgba(139, 92, 246, .10), transparent 34%),
    linear-gradient(180deg, #070914 0%, #060711 100%);
  color: var(--es-text);
}

#MainMenu, footer, header,
[data-testid="stToolbar"], [data-testid="stDecoration"], [data-testid="stStatusWidget"],
[data-testid="stDeployButton"], .stAppDeployButton {
  visibility: hidden !important;
  display: none !important;
}

.block-container {
  padding: 1.05rem 2.2rem 2rem !important;
  max-width: min(1760px, calc(100vw - 56px)) !important;
}

[data-testid="stAppViewContainer"] .main .block-container,
[data-testid="stMainBlockContainer"] {
  max-width: min(1760px, calc(100vw - 56px)) !important;
}

.es-shell {
  min-height: calc(100vh - 80px);
}

.es-rail {
  background:
    linear-gradient(180deg, rgba(19, 25, 48, .94), rgba(8, 12, 25, .92)),
    radial-gradient(circle at 20% 0%, rgba(0,217,133,.16), transparent 38%);
  border: 1px solid rgba(97, 119, 198, .32);
  border-radius: 16px;
  padding: 14px;
  position: sticky;
  top: 18px;
  box-shadow: 0 26px 70px rgba(0,0,0,.34), inset 0 1px 0 rgba(255,255,255,.05);
  backdrop-filter: blur(18px);
}

.es-logo {
  display: flex;
  align-items: center;
  gap: 10px;
  font-weight: 900;
  font-size: 18px;
  margin-bottom: 4px;
  letter-spacing: -0.02em;
}

.es-logo-badge {
  width: 34px;
  height: 34px;
  border-radius: 10px;
  display: grid;
  place-items: center;
  background:
    linear-gradient(135deg, rgba(0,217,133,.98), rgba(52,189,246,.92));
  color: #05131c;
  font-weight: 950;
  font-size: 14px;
  letter-spacing: -.08em;
  position: relative;
  box-shadow: 0 10px 30px rgba(52, 189, 246, .22);
}

.es-small {
  color: var(--es-muted);
  font-size: 12px;
  line-height: 1.4;
}

.es-nav-label {
  font-family: "Space Mono", monospace;
  font-size: 10px;
  letter-spacing: .12em;
  text-transform: uppercase;
  color: #8ea1dc;
  margin: 16px 0 8px;
}

.es-hero {
  background:
    linear-gradient(135deg, rgba(0, 217, 133, .14), rgba(52, 189, 246, .08) 42%, rgba(245, 158, 11, .08) 68%, rgba(139, 92, 246, .18)),
    rgba(17, 20, 38, .88);
  border: 1px solid rgba(52, 189, 246, .26);
  border-radius: 14px;
  padding: 28px 30px;
  margin-bottom: 20px;
  box-shadow: 0 28px 84px rgba(0,0,0,.30), inset 0 1px 0 rgba(255,255,255,.06);
  position: relative;
  overflow: hidden;
}

.es-hero::after {
  content: "";
  position: absolute;
  inset: 0;
  pointer-events: none;
  background-image:
    linear-gradient(rgba(255,255,255,.035) 1px, transparent 1px),
    linear-gradient(90deg, rgba(255,255,255,.035) 1px, transparent 1px);
  background-size: 34px 34px;
  mask-image: linear-gradient(90deg, transparent, black 32%, black 78%, transparent);
  opacity: .65;
}

.es-kicker {
  display: inline-flex;
  align-items: center;
  gap: 8px;
  padding: 7px 12px;
  border-radius: 999px;
  border: 1px solid rgba(0,217,133,.28);
  background: rgba(0,217,133,.10);
  color: #33f2aa;
  font-family: "Space Mono", monospace;
  font-size: 11px;
  font-weight: 700;
  text-transform: uppercase;
}

.es-title {
  font-size: clamp(34px, 3.4vw, 60px);
  line-height: 1.02;
  font-weight: 900;
  letter-spacing: -.04em;
  margin: 16px 0 10px;
  color: #f8fbff;
}

.es-title span {
  background: linear-gradient(90deg, #dfffee, #7dd3fc, #c4b5fd);
  -webkit-background-clip: text;
  -webkit-text-fill-color: transparent;
}

.es-subtitle {
  color: #c2c9e9;
  font-size: 16px;
  max-width: 980px;
}

.es-card {
  background: var(--es-card);
  border: 1px solid var(--es-border-soft);
  border-radius: 12px;
  padding: 20px;
  box-shadow: 0 18px 54px rgba(0,0,0,.20), inset 0 1px 0 rgba(255,255,255,.04);
}

.es-card h3, .es-card h4 {
  margin-top: 0;
}

.es-bento {
  display: grid;
  grid-template-columns: minmax(360px, 1.45fr) repeat(3, minmax(180px, .75fr));
  gap: 14px;
  margin: 16px 0 22px;
}

.es-bento-card {
  background:
    linear-gradient(180deg, rgba(23, 29, 54, .82), rgba(12, 16, 31, .86));
  border: 1px solid rgba(88,113,190,.28);
  border-radius: 12px;
  padding: 18px;
  min-height: 132px;
  box-shadow: 0 18px 46px rgba(0,0,0,.20), inset 0 1px 0 rgba(255,255,255,.045);
  position: relative;
  overflow: hidden;
}

.es-bento-card.wide {
  grid-column: span 1;
}

.es-bento-card::after {
  content: "";
  position: absolute;
  inset: auto 0 0;
  height: 3px;
  background: linear-gradient(90deg, #00d985, #34bdf6, #f59e0b);
  opacity: .65;
}

.es-spark {
  width: 100%;
  height: 34px;
  margin-top: 8px;
}

.es-spark polyline {
  fill: none;
  stroke: #34bdf6;
  stroke-width: 3;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.es-activity-drawer {
  border: 1px solid rgba(84,105,180,.24);
  border-radius: 8px;
  background: rgba(9,12,25,.76);
  padding: 12px;
  max-height: 260px;
  overflow-y: auto;
}

.es-activity-item {
  display: grid;
  grid-template-columns: 76px 1fr;
  gap: 10px;
  border-bottom: 1px solid rgba(255,255,255,.07);
  padding: 8px 0;
}

.es-command-strip {
  border: 1px solid rgba(52,189,246,.26);
  background: rgba(13,16,31,.70);
  border-radius: 8px;
  padding: 10px;
  margin: 10px 0 14px;
}

.es-landing-shell {
  min-height: 100vh;
}

.es-landing-nav {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 18px;
  padding: 14px 4px 18px;
}

.es-landing-brand {
  display: inline-flex;
  align-items: center;
  gap: 10px;
  font-weight: 900;
  font-size: 20px;
}

.es-landing-links {
  display: flex;
  align-items: center;
  gap: 18px;
  color: #c2c9e9;
  font-size: 13px;
  font-weight: 800;
}

.es-landing-hero {
  min-height: 76vh;
  display: grid;
  grid-template-columns: minmax(320px, .86fr) minmax(420px, 1.14fr);
  gap: 34px;
  align-items: center;
  padding: 22px 0 36px;
}

.es-landing-kicker {
  display: inline-flex;
  align-items: center;
  gap: 8px;
  color: #75f7c4;
  font-family: "Space Mono", monospace;
  text-transform: uppercase;
  font-size: 11px;
  font-weight: 800;
  letter-spacing: .08em;
}

.es-landing-dot {
  width: 9px;
  height: 9px;
  border-radius: 999px;
  background: #00d985;
  box-shadow: 0 0 0 7px rgba(0,217,133,.12);
}

.es-landing-title {
  color: #f8fbff;
  font-size: clamp(44px, 7vw, 84px);
  line-height: .96;
  font-weight: 950;
  letter-spacing: -.04em;
  margin: 18px 0 16px;
}

.es-landing-title span {
  color: #7dd3fc;
}

.es-landing-copy {
  color: #c2c9e9;
  font-size: 18px;
  line-height: 1.6;
  max-width: 680px;
}

.es-landing-actions {
  display: flex;
  gap: 12px;
  flex-wrap: wrap;
  margin-top: 24px;
}

.es-landing-action {
  border-radius: 8px;
  border: 1px solid rgba(52,189,246,.28);
  background: rgba(18,21,38,.76);
  color: #f8fbff;
  padding: 11px 14px;
  font-weight: 900;
  font-size: 13px;
}

.es-landing-action.primary {
  background: linear-gradient(90deg, #00bf75, #2094f3);
  border-color: rgba(0,217,133,.34);
}

.es-product-scene {
  position: relative;
  min-height: 520px;
  border: 1px solid rgba(84,105,180,.24);
  border-radius: 8px;
  overflow: hidden;
  background:
    linear-gradient(180deg, rgba(11,14,27,.96), rgba(6,8,18,.96)),
    repeating-linear-gradient(90deg, rgba(255,255,255,.045) 0 1px, transparent 1px 82px),
    repeating-linear-gradient(0deg, rgba(255,255,255,.035) 0 1px, transparent 1px 58px);
  box-shadow: 0 34px 90px rgba(0,0,0,.34);
}

.es-scene-top {
  display: flex;
  justify-content: space-between;
  gap: 14px;
  padding: 13px 15px;
  border-bottom: 1px solid rgba(255,255,255,.08);
  background: rgba(18,21,38,.88);
}

.es-scene-pill {
  border: 1px solid rgba(84,105,180,.26);
  border-radius: 999px;
  padding: 5px 9px;
  color: #c2c9e9;
  font-size: 11px;
  font-weight: 800;
}

.es-scene-grid {
  display: grid;
  grid-template-columns: 1fr 1fr 160px;
  gap: 1px;
  padding: 16px;
}

.es-scene-head,
.es-scene-cell {
  background: rgba(255,255,255,.045);
  padding: 10px;
  min-height: 40px;
  color: #dce8ff;
  font-size: 12px;
}

.es-scene-head {
  color: #9aa7da;
  font-family: "Space Mono", monospace;
  text-transform: uppercase;
  font-size: 10px;
}

.es-scene-cell.mark {
  border-left: 3px solid #00d985;
}

.es-scene-cell.warn {
  border-left: 3px solid #f59e0b;
}

.es-scene-panel {
  position: absolute;
  right: 18px;
  bottom: 18px;
  width: 310px;
  border: 1px solid rgba(52,189,246,.28);
  border-radius: 8px;
  background: rgba(13,16,31,.94);
  padding: 14px;
  box-shadow: 0 24px 70px rgba(0,0,0,.38);
}

.es-landing-section {
  padding: 26px 0;
}

.es-landing-grid {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 14px;
}

.es-landing-card {
  border: 1px solid rgba(84,105,180,.24);
  border-radius: 8px;
  background: rgba(18,21,38,.72);
  padding: 17px;
  min-height: 142px;
}

.es-price-band {
  display: grid;
  grid-template-columns: 1fr 1fr 1fr;
  gap: 14px;
  margin-top: 12px;
}

.es-lp {
  min-height: 100vh;
  overflow: hidden;
  color: #fff;
  background:
    radial-gradient(circle at 20% 10%, rgba(17,245,181,.18), transparent 30%),
    radial-gradient(circle at 85% 5%, rgba(156,92,255,.28), transparent 28%),
    radial-gradient(circle at 50% 95%, rgba(74,168,255,.14), transparent 34%);
}

.es-lp-inner {
  max-width: 1280px;
  margin: 0 auto;
  padding: 0 20px;
}

.es-lp-nav {
  position: sticky;
  top: 0;
  z-index: 12;
  margin: 0 -20px;
  border-bottom: 1px solid rgba(255,255,255,.10);
  background: rgba(5,7,19,.70);
  backdrop-filter: blur(18px);
}

.es-lp-nav-inner {
  max-width: 1280px;
  margin: 0 auto;
  padding: 16px 20px;
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 20px;
}

.es-lp-brand {
  display: inline-flex;
  align-items: center;
  gap: 12px;
  color: #fff;
}

.es-lp-logo {
  width: 40px;
  height: 40px;
  border-radius: 18px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  color: #050713;
  font-weight: 950;
  background: linear-gradient(135deg, #11f5b5, #4aa8ff, #9c5cff);
  box-shadow: 0 0 80px rgba(17,245,181,.16);
}

.es-lp-brand-name {
  font-size: 18px;
  font-weight: 950;
  line-height: 1.05;
}

.es-lp-brand-sub {
  margin-top: 1px;
  color: rgba(255,255,255,.45);
  font-size: 11px;
  font-weight: 800;
  letter-spacing: .24em;
  text-transform: uppercase;
}

.es-lp-links {
  display: flex;
  align-items: center;
  gap: 30px;
  color: rgba(255,255,255,.70);
  font-size: 14px;
  font-weight: 800;
}

.es-lp-link {
  color: rgba(255,255,255,.70) !important;
  text-decoration: none !important;
}

.es-lp-actions {
  display: flex;
  align-items: center;
  gap: 12px;
}

.es-lp-btn {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  border-radius: 999px;
  padding: 12px 20px;
  min-height: 44px;
  color: #fff !important;
  text-decoration: none !important;
  font-size: 14px;
  font-weight: 950;
  border: 1px solid rgba(255,255,255,.15);
  background: rgba(255,255,255,.05);
}

.es-lp-btn.primary {
  color: #050713 !important;
  border-color: transparent;
  background: linear-gradient(90deg, #11f5b5, #4aa8ff);
  box-shadow: 0 0 80px rgba(17,245,181,.16);
}

.es-lp-hero {
  position: relative;
  text-align: center;
  padding: 82px 0 56px;
}

.es-lp-grid-bg {
  position: absolute;
  inset: 20px 0 0;
  opacity: .60;
  pointer-events: none;
  background-image:
    linear-gradient(rgba(255,255,255,.04) 1px, transparent 1px),
    linear-gradient(90deg, rgba(255,255,255,.04) 1px, transparent 1px);
  background-size: 34px 34px;
  mask-image: radial-gradient(circle at center, black, transparent 72%);
}

.es-lp-hero-content {
  position: relative;
  max-width: 1120px;
  margin: 0 auto;
}

.es-lp-badge {
  display: inline-flex;
  align-items: center;
  gap: 8px;
  border: 1px solid rgba(17,245,181,.25);
  border-radius: 999px;
  background: rgba(17,245,181,.10);
  color: #11f5b5;
  padding: 8px 16px;
  font-size: 14px;
  font-weight: 950;
}

.es-lp-dot {
  width: 8px;
  height: 8px;
  border-radius: 999px;
  background: #11f5b5;
  box-shadow: 0 0 18px rgba(17,245,181,.95);
}

.es-lp-title {
  max-width: 1120px;
  margin: 24px auto 0;
  color: #fff;
  font-size: clamp(48px, 6.4vw, 84px);
  line-height: 1.02;
  font-weight: 950;
  letter-spacing: -.04em;
}

.es-lp-copy {
  max-width: 760px;
  margin: 28px auto 0;
  color: rgba(255,255,255,.70);
  font-size: 20px;
  line-height: 1.6;
}

.es-lp-hero-actions {
  display: flex;
  justify-content: center;
  gap: 16px;
  flex-wrap: wrap;
  margin-top: 36px;
}

.es-lp-product {
  padding: 0 0 80px;
}

.es-lp-glass {
  position: relative;
  overflow: hidden;
  border-radius: 32px;
  padding: 12px;
  background: linear-gradient(180deg, rgba(255,255,255,.09), rgba(255,255,255,.045));
  border: 1px solid rgba(255,255,255,.12);
  box-shadow: 0 0 80px rgba(156,92,255,.24), 0 24px 90px rgba(0,0,0,.35);
  backdrop-filter: blur(18px);
}

.es-lp-glow-a,
.es-lp-glow-b {
  position: absolute;
  width: 288px;
  height: 288px;
  border-radius: 999px;
  filter: blur(48px);
}

.es-lp-glow-a {
  right: -96px;
  top: -96px;
  background: rgba(156,92,255,.30);
}

.es-lp-glow-b {
  left: -96px;
  bottom: -96px;
  background: rgba(17,245,181,.20);
}

.es-lp-window {
  position: relative;
  border-radius: 24px;
  border: 1px solid rgba(255,255,255,.10);
  background: #090b19;
  padding: 24px;
}

.es-lp-window-top {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 16px;
  border-bottom: 1px solid rgba(255,255,255,.10);
  padding-bottom: 16px;
  margin-bottom: 20px;
}

.es-lp-traffic {
  display: flex;
  gap: 8px;
}

.es-lp-traffic span {
  width: 12px;
  height: 12px;
  border-radius: 999px;
  display: inline-block;
}

.es-lp-traffic span:nth-child(1) { background: #f87171; }
.es-lp-traffic span:nth-child(2) { background: #facc15; }
.es-lp-traffic span:nth-child(3) { background: #4ade80; }

.es-lp-window-label {
  border: 1px solid rgba(255,255,255,.10);
  border-radius: 999px;
  background: rgba(255,255,255,.05);
  color: rgba(255,255,255,.55);
  padding: 7px 16px;
  font-size: 12px;
  font-weight: 900;
}

.es-lp-workflow {
  display: grid;
  grid-template-columns: 1.15fr .85fr;
  gap: 20px;
}

.es-lp-panel {
  border-radius: 24px;
  border: 1px solid rgba(255,255,255,.10);
  background: rgba(255,255,255,.035);
  padding: 20px;
}

.es-lp-panel-head {
  display: flex;
  align-items: flex-start;
  justify-content: space-between;
  gap: 16px;
  margin-bottom: 16px;
}

.es-lp-eyebrow {
  color: #11f5b5;
  font-size: 13px;
  font-weight: 950;
  letter-spacing: .18em;
  text-transform: uppercase;
}

.es-lp-panel h3 {
  margin: 4px 0 0;
  color: #fff;
  font-size: 26px;
  font-weight: 950;
}

.es-lp-status {
  white-space: nowrap;
  border-radius: 999px;
  background: rgba(17,245,181,.10);
  color: #11f5b5;
  padding: 6px 12px;
  font-size: 12px;
  font-weight: 950;
}

.es-lp-task {
  border-radius: 16px;
  border: 1px solid rgba(255,255,255,.10);
  background: rgba(255,255,255,.05);
  padding: 16px;
  margin-top: 12px;
}

.es-lp-task-row,
.es-lp-pipeline-row {
  display: flex;
  justify-content: space-between;
  gap: 16px;
  color: rgba(255,255,255,.80);
  font-size: 14px;
  font-weight: 800;
}

.es-lp-green { color: #11f5b5; }
.es-lp-sky { color: #4aa8ff; }
.es-lp-violet { color: #ddd6fe; }
.es-lp-yellow { color: #fde047; }

.es-lp-progress {
  height: 8px;
  border-radius: 999px;
  background: rgba(255,255,255,.10);
  overflow: hidden;
  margin-top: 12px;
}

.es-lp-progress span {
  display: block;
  height: 100%;
  border-radius: inherit;
  background: linear-gradient(90deg, #11f5b5, #4aa8ff);
}

.es-lp-findings {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 12px;
  margin-top: 12px;
  text-align: center;
  color: rgba(255,255,255,.70);
  font-size: 12px;
  font-weight: 900;
}

.es-lp-finding {
  border-radius: 12px;
  padding: 12px;
}

.es-lp-finding.red { background: rgba(239,68,68,.10); color: #fecaca; }
.es-lp-finding.yellow { background: rgba(234,179,8,.10); color: #fef08a; }
.es-lp-finding.blue { background: rgba(59,130,246,.10); color: #bfdbfe; }

.es-lp-score {
  margin-top: 20px;
  display: flex;
  align-items: flex-end;
  gap: 16px;
}

.es-lp-score-number {
  color: #fff;
  font-size: 72px;
  line-height: .9;
  font-weight: 950;
}

.es-lp-muted {
  color: rgba(255,255,255,.55);
  font-size: 14px;
  font-weight: 800;
}

.es-lp-mini-grid {
  display: grid;
  grid-template-columns: repeat(2, minmax(0, 1fr));
  gap: 12px;
  margin-top: 22px;
}

.es-lp-mini {
  border-radius: 16px;
  background: rgba(255,255,255,.05);
  padding: 16px;
  color: rgba(255,255,255,.70);
  font-size: 13px;
  font-weight: 850;
}

.es-lp-pipeline {
  margin-top: 20px;
}

.es-lp-pipeline .es-lp-eyebrow {
  color: rgba(255,255,255,.40);
}

.es-lp-pipeline-row {
  margin-top: 12px;
}

.es-lp-social {
  border-top: 1px solid rgba(255,255,255,.10);
  border-bottom: 1px solid rgba(255,255,255,.10);
  background: rgba(255,255,255,.025);
  padding: 56px 0;
  text-align: center;
}

.es-lp-social-title {
  color: rgba(255,255,255,.75);
  font-size: 18px;
  font-weight: 850;
}

.es-lp-logo-grid {
  display: grid;
  grid-template-columns: repeat(6, minmax(0, 1fr));
  gap: 16px;
  margin-top: 32px;
}

.es-lp-logo-card,
.es-lp-card,
.es-lp-stat {
  border: 1px solid rgba(255,255,255,.13);
  background: rgba(255,255,255,.06);
}

.es-lp-logo-card {
  border-radius: 16px;
  padding: 20px;
  color: rgba(255,255,255,.55);
  font-size: 14px;
  font-weight: 950;
}

.es-lp-awards {
  display: flex;
  justify-content: center;
  gap: 12px;
  flex-wrap: wrap;
  margin-top: 28px;
}

.es-lp-award {
  border-radius: 999px;
  padding: 8px 16px;
  font-size: 12px;
  font-weight: 950;
  border: 1px solid rgba(255,255,255,.18);
  background: rgba(255,255,255,.07);
}

.es-lp-section {
  padding: 80px 0;
}

.es-lp-stats,
.es-lp-card-grid {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 20px;
}

.es-lp-stat {
  border-radius: 24px;
  padding: 32px;
  text-align: center;
  box-shadow: 0 24px 90px rgba(0,0,0,.25);
}

.es-lp-stat-number {
  font-size: 52px;
  line-height: 1;
  font-weight: 950;
}

.es-lp-stat p {
  margin: 12px 0 0;
  color: rgba(255,255,255,.70);
  font-weight: 850;
}

.es-lp-section-title {
  max-width: 760px;
}

.es-lp-section-title .es-lp-eyebrow {
  margin-bottom: 16px;
}

.es-lp-section-title h2,
.es-lp-feature h2,
.es-lp-cta h2 {
  margin: 0;
  color: #fff;
  font-size: clamp(38px, 4vw, 56px);
  line-height: 1.06;
  font-weight: 950;
}

.es-lp-card-grid {
  margin-top: 40px;
}

.es-lp-card {
  border-radius: 32px;
  padding: 28px;
  transition: transform .18s ease, border-color .18s ease;
  box-shadow: 0 24px 90px rgba(0,0,0,.25);
}

.es-lp-card:hover {
  transform: translateY(-4px);
  border-color: rgba(17,245,181,.40);
}

.es-lp-icon {
  color: #fff;
  font-size: 36px;
  font-weight: 950;
}

.es-lp-card h3 {
  margin: 20px 0 12px;
  color: #fff;
  font-size: 24px;
  font-weight: 950;
}

.es-lp-card p,
.es-lp-feature p,
.es-lp-cta p {
  color: rgba(255,255,255,.65);
  line-height: 1.65;
}

.es-lp-card-link {
  display: inline-flex;
  margin-top: 18px;
  color: #11f5b5;
  font-weight: 950;
}

.es-lp-feature-grid {
  display: grid;
  grid-template-columns: repeat(2, minmax(0, 1fr));
  gap: 40px;
  align-items: center;
  margin-top: 64px;
}

.es-lp-feature-grid:first-child {
  margin-top: 0;
}

.es-lp-feature p {
  font-size: 18px;
}

.es-lp-list {
  margin: 28px 0 0;
  padding: 0;
  list-style: none;
  color: rgba(255,255,255,.70);
}

.es-lp-list li {
  display: flex;
  gap: 12px;
  margin-top: 12px;
}

.es-lp-mock {
  border-radius: 32px;
  padding: 20px;
  background: linear-gradient(180deg, rgba(255,255,255,.09), rgba(255,255,255,.045));
  border: 1px solid rgba(255,255,255,.12);
  box-shadow: 0 24px 90px rgba(0,0,0,.35);
}

.es-lp-mock-inner {
  border-radius: 24px;
  background: #090b19;
  padding: 20px;
}

.es-lp-mock-row,
.es-lp-issue {
  border-radius: 16px;
  background: rgba(255,255,255,.05);
  padding: 16px;
  margin-top: 12px;
  color: rgba(255,255,255,.72);
}

.es-lp-mock-row:first-child {
  margin-top: 0;
}

.es-lp-mock-row b {
  color: #fff;
}

.es-lp-three {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 12px;
}

.es-lp-count {
  border-radius: 16px;
  padding: 16px;
  text-align: center;
  background: rgba(255,255,255,.05);
}

.es-lp-count strong {
  display: block;
  font-size: 26px;
  color: #fff;
}

.es-lp-issue {
  border: 1px solid rgba(255,255,255,.10);
}

.es-lp-code {
  border-radius: 4px;
  background: rgba(255,255,255,.10);
  padding: 1px 5px;
}

.es-lp-cta {
  border-radius: 32px;
  border: 1px solid rgba(255,255,255,.10);
  background: linear-gradient(135deg, rgba(17,245,181,.15), rgba(74,168,255,.10), rgba(156,92,255,.20));
  padding: 48px;
  text-align: center;
  box-shadow: 0 0 80px rgba(156,92,255,.24);
}

.es-lp-cta p {
  max-width: 680px;
  margin: 20px auto 0;
  font-size: 18px;
}

.es-lp-footer {
  border-top: 1px solid rgba(255,255,255,.10);
  padding: 40px 0;
  color: rgba(255,255,255,.45);
  font-size: 14px;
}

.es-lp-footer-row {
  display: flex;
  justify-content: space-between;
  gap: 20px;
  flex-wrap: wrap;
}

.es-lp-footer-links {
  display: flex;
  gap: 20px;
}

.es-auth-shell {
  max-width: 920px;
  margin: 40px auto 24px;
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 18px;
  padding: 18px;
  border: 1px solid rgba(255,255,255,.12);
  border-radius: 24px;
  background: rgba(18,21,38,.72);
  box-shadow: 0 24px 90px rgba(0,0,0,.25);
}

.es-auth-links {
  display: flex;
  gap: 12px;
  flex-wrap: wrap;
}

.es-grid-4 {
  display: grid;
  grid-template-columns: repeat(4, minmax(0, 1fr));
  gap: 14px;
  margin: 16px 0 22px;
}

.es-grid-3 {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 14px;
  margin: 16px 0 22px;
}

.es-metric-label {
  font-family: "Space Mono", monospace;
  color: #9aa7da;
  font-size: 11px;
  letter-spacing: .10em;
  text-transform: uppercase;
}

.es-metric-value {
  font-size: 30px;
  font-weight: 900;
  color: #fff;
  margin: 7px 0 4px;
}

.es-chip {
  display:inline-block;
  padding: 5px 10px;
  border-radius: 999px;
  border: 1px solid rgba(52,189,246,.22);
  background: rgba(52,189,246,.08);
  color: #bfeaff;
  font-size: 12px;
  font-weight: 700;
}

.es-chip.green {
  border-color: rgba(0,217,133,.24);
  background: rgba(0,217,133,.10);
  color: #77ffc9;
}

.es-chip.amber {
  border-color: rgba(245,158,11,.30);
  background: rgba(245,158,11,.12);
  color: #ffd18a;
}

.es-chip.red {
  border-color: rgba(255,75,51,.30);
  background: rgba(255,75,51,.12);
  color: #ffb0a5;
}

.es-row-card {
  border: 1px solid var(--es-border-soft);
  border-radius: 16px;
  padding: 12px;
  background: rgba(255,255,255,.03);
  margin: 8px 0;
}

.es-highlight {
  color: #f8fbff;
  line-height: 1.6;
}

.es-mark-ph {
  color: #7dd3fc;
  background: rgba(52,189,246,.12);
  border: 1px solid rgba(52,189,246,.22);
  border-radius: 4px;
  padding: 1px 4px;
  font-family: "Space Mono", monospace;
}

.es-mark-dnt {
  color: #ffd18a;
  background: rgba(245,158,11,.12);
  border: 1px solid rgba(245,158,11,.28);
  border-radius: 4px;
  padding: 1px 4px;
}

.es-diff-del {
  color: #ffb0a5;
  background: rgba(255,75,51,.14);
  text-decoration: line-through;
  border-radius: 4px;
  padding: 1px 4px;
}

.es-diff-add {
  color: #9fffd3;
  background: rgba(0,217,133,.14);
  border-radius: 4px;
  padding: 1px 4px;
}

.es-timeline {
  position: relative;
  height: 68px;
  border: 1px solid rgba(84,105,180,.28);
  border-radius: 18px;
  background: linear-gradient(90deg, rgba(0,217,133,.08), rgba(52,189,246,.06), rgba(139,92,246,.08));
  overflow: hidden;
  margin-top: 12px;
}

.es-timebar {
  position: absolute;
  top: 18px;
  height: 26px;
  border-radius: 999px;
  background: linear-gradient(90deg, var(--es-green), var(--es-cyan));
  border: 1px solid rgba(255,255,255,.22);
  box-shadow: 0 12px 24px rgba(0,217,133,.16);
}

.es-timebar.current {
  background: linear-gradient(90deg, #f59e0b, #ff4b33);
}

.es-waveform {
  height: 82px;
  border: 1px solid rgba(84,105,180,.28);
  border-radius: 8px;
  background: rgba(8,10,19,.86);
  display: flex;
  align-items: center;
  gap: 3px;
  padding: 9px 10px;
  overflow: hidden;
}

.es-wavebar {
  flex: 1;
  min-width: 3px;
  border-radius: 999px;
  background: linear-gradient(180deg, #34bdf6, #00d985);
  opacity: .82;
}

.es-radar-wrap {
  display: grid;
  grid-template-columns: minmax(180px, 260px) 1fr;
  gap: 16px;
  align-items: center;
}

.es-radar {
  width: 100%;
  aspect-ratio: 1;
}

.es-heatmap {
  display: grid;
  grid-template-columns: repeat(12, minmax(0, 1fr));
  gap: 5px;
}

.es-heat-cell {
  min-height: 28px;
  border-radius: 6px;
  border: 1px solid rgba(255,255,255,.08);
}

.stButton > button, .stDownloadButton > button {
  width: 100%;
  border-radius: 10px !important;
  border: 1px solid rgba(0,217,133,.24) !important;
  background: linear-gradient(90deg, #00bf75, #2094f3) !important;
  color: white !important;
  font-weight: 800 !important;
  min-height: 42px;
  box-shadow: 0 12px 30px rgba(0,0,0,.20);
}

.stButton > button:hover, .stDownloadButton > button:hover {
  transform: translateY(-1px);
  box-shadow: 0 18px 38px rgba(52,189,246,.24);
}


.es-nav-link {
  display: flex;
  align-items: center;
  justify-content: space-between;
  text-decoration: none !important;
  color: #dce8ff !important;
  border: 1px solid rgba(84,105,180,.20);
  background: rgba(9,13,27,.54);
  border-radius: 10px;
  padding: 0.62rem 0.78rem;
  margin: 0.35rem 0;
  font-weight: 800;
  font-size: 13px;
  box-shadow: inset 0 1px 0 rgba(255,255,255,.035);
  transition: border-color .15s ease, background .15s ease, transform .15s ease;
}
.es-nav-link::after {
  content: ">";
  color: rgba(168,176,214,.45);
  font-family: "Space Mono", monospace;
  font-size: 11px;
}
.es-nav-link:hover {
  background: rgba(52,189,246,.12);
  border-color: rgba(52,189,246,.45);
  color: #ffffff !important;
  transform: translateX(2px);
}
.es-nav-link.active {
  background:
    linear-gradient(90deg, rgba(0,191,117,.95), rgba(32,148,243,.95));
  color: #ffffff !important;
  border-color: rgba(0,217,133,.45);
  box-shadow: 0 14px 30px rgba(32,148,243,.22), inset 0 1px 0 rgba(255,255,255,.18);
}
.es-nav-link.active::after {
  color: #fff;
}

.es-topnav {
  position: sticky;
  top: 10px;
  z-index: 500;
  width: 100%;
  margin: 0 0 22px;
  border: 1px solid rgba(84,105,180,.34);
  border-radius: 12px;
  background:
    linear-gradient(135deg, rgba(16,22,43,.96), rgba(11,15,31,.94)),
    radial-gradient(circle at 12% 0%, rgba(0,217,133,.18), transparent 38%),
    radial-gradient(circle at 90% 12%, rgba(52,189,246,.16), transparent 36%);
  color: #f7fbff;
  box-shadow: 0 24px 68px rgba(0,0,0,.34), inset 0 1px 0 rgba(255,255,255,.06);
  backdrop-filter: blur(16px);
  overflow: visible;
}

.es-topnav-row {
  position: relative;
  z-index: 4;
  min-height: 76px;
  display: flex;
  align-items: stretch;
  gap: 12px;
  padding: 0 18px;
}

.es-topnav-brand {
  display: flex;
  align-items: center;
  gap: 10px;
  min-width: 250px;
  color: #f7fbff;
  text-decoration: none !important;
}

.es-topnav-mark {
  width: 44px;
  height: 44px;
  border-radius: 11px;
  display: grid;
  place-items: center;
  background: linear-gradient(135deg, rgba(0,217,133,.98), rgba(52,189,246,.92));
  color: #04131c;
  font-weight: 950;
  letter-spacing: -.06em;
  box-shadow: 0 16px 38px rgba(52,189,246,.22);
}

.es-topnav-name {
  font-size: 25px;
  line-height: 1;
  font-weight: 950;
  letter-spacing: -.04em;
  color: #4aa8ff;
}

.es-topnav-name span {
  color: #34bdf6;
  font-weight: 850;
}

.es-topnav-sub {
  margin-top: 3px;
  color: #a8b0d6;
  font-size: 12px;
  font-weight: 800;
}

.es-topnav-links {
  flex: 1;
  display: flex;
  align-items: stretch;
  justify-content: center;
  flex-wrap: wrap;
  gap: 0;
}

.es-topnav-link {
  min-height: 76px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  position: relative;
  padding: 0 14px;
  color: #dce8ff !important;
  text-decoration: none !important;
  font-size: 13px;
  font-weight: 950;
  text-transform: uppercase;
  white-space: nowrap;
}

.es-topnav-link:hover {
  color: #75f7c4 !important;
  background: rgba(52,189,246,.08);
}

.es-topnav-link.active {
  color: #75f7c4 !important;
}

.es-topnav-link.active::after {
  content: "";
  position: absolute;
  left: 14px;
  right: 14px;
  bottom: 14px;
  height: 4px;
  border-radius: 999px;
  background: linear-gradient(90deg, var(--es-green), var(--es-cyan));
  box-shadow: 0 0 22px rgba(52,189,246,.48);
}

.es-topnav-tools {
  display: flex;
  align-items: stretch;
  min-height: 76px;
  border-left: 1px solid rgba(84,105,180,.26);
}

.es-topnav-tool {
  min-width: 58px;
  padding: 0 12px;
  display: grid;
  place-items: center;
  border-right: 1px solid rgba(84,105,180,.24);
  color: #dce8ff !important;
  font-weight: 900;
  text-decoration: none !important;
  position: relative;
}

.es-topnav-badge {
  position: absolute;
  top: 11px;
  right: 12px;
  min-width: 16px;
  height: 16px;
  border-radius: 999px;
  display: grid;
  place-items: center;
  background: linear-gradient(135deg, var(--es-green), var(--es-cyan));
  color: #05131c;
  font-size: 10px;
  font-weight: 900;
}

.es-topnav-user-wrap {
  position: relative;
  display: flex;
  align-items: stretch;
}

.es-topnav-user {
  display: flex;
  align-items: center;
  gap: 10px;
  padding: 0 14px;
  min-width: 220px;
  color: #f7fbff;
  cursor: pointer;
  outline: none;
}

.es-topnav-user-wrap:hover .es-account-menu,
.es-topnav-user-wrap:focus-within .es-account-menu {
  opacity: 1;
  transform: translateY(0);
  pointer-events: auto;
}

.es-topnav-avatar {
  width: 42px;
  height: 42px;
  border-radius: 50%;
  display: grid;
  place-items: center;
  background: rgba(255,255,255,.08);
  color: #75f7c4;
  border: 1px solid rgba(52,189,246,.28);
  font-weight: 950;
}

.es-account-caret {
  color: #8ea1dc;
  font-weight: 900;
}

.es-account-menu {
  position: absolute;
  top: calc(100% - 2px);
  right: 10px;
  min-width: 230px;
  border: 1px solid rgba(84,105,180,.34);
  border-radius: 10px;
  background: linear-gradient(180deg, rgba(18,24,48,.98), rgba(9,13,27,.98));
  box-shadow: 0 24px 70px rgba(0,0,0,.42), inset 0 1px 0 rgba(255,255,255,.06);
  padding: 8px;
  opacity: 0;
  transform: translateY(-6px);
  pointer-events: none;
  transition: opacity .15s ease, transform .15s ease;
  z-index: 1200;
}

.es-account-menu a {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 12px;
  color: #dce8ff !important;
  text-decoration: none !important;
  font-size: 13px;
  font-weight: 850;
  padding: 10px 11px;
  border-radius: 8px;
}

.es-account-menu a span {
  color: #75f7c4;
  font-size: 11px;
}

.es-account-menu a:hover {
  background: rgba(52,189,246,.12);
  color: #75f7c4 !important;
}

.es-account-menu a.logout {
  color: #ffd18a !important;
  border-top: 1px solid rgba(84,105,180,.22);
  margin-top: 4px;
  padding-top: 11px;
}

.es-owner-strip {
  position: relative;
  z-index: 1;
  display: flex;
  align-items: center;
  gap: 8px;
  flex-wrap: wrap;
  padding: 9px 18px 11px;
  border-top: 1px solid rgba(84,105,180,.24);
  background: rgba(7,9,20,.38);
}

.es-owner-strip span {
  font-family: "Space Mono", monospace;
  font-size: 10px;
  letter-spacing: .10em;
  text-transform: uppercase;
  color: #8ea1dc;
  margin-right: 6px;
}

.es-owner-link {
  color: #c2c9e9 !important;
  text-decoration: none !important;
  font-weight: 900;
  font-size: 12px;
  padding: 7px 10px;
  border-radius: 8px;
}

.es-owner-link:hover,
.es-owner-link.active {
  color: #75f7c4 !important;
  background: rgba(52,189,246,.10);
}

.es-dashboard-grid {
  display: grid;
  grid-template-columns: minmax(0, 1.55fr) minmax(340px, .85fr);
  gap: 16px;
  align-items: start;
}

.es-personal-hero {
  position: relative;
  overflow: hidden;
  border: 1px solid rgba(88,113,190,.30);
  border-radius: 16px;
  padding: 28px;
  margin-bottom: 16px;
  background:
    radial-gradient(circle at 18% 18%, rgba(0,217,133,.30), transparent 30%),
    radial-gradient(circle at 78% 8%, rgba(52,189,246,.24), transparent 28%),
    radial-gradient(circle at 88% 86%, rgba(245,158,11,.14), transparent 28%),
    linear-gradient(135deg, rgba(18,24,48,.94), rgba(17,11,39,.92));
  box-shadow: 0 30px 90px rgba(0,0,0,.32), inset 0 1px 0 rgba(255,255,255,.06);
}

.es-personal-hero::before {
  content:"";
  position:absolute;
  inset:-40%;
  background:
    conic-gradient(from 90deg, rgba(0,217,133,.14), rgba(52,189,246,.16), rgba(139,92,246,.14), rgba(245,158,11,.10), rgba(0,217,133,.14));
  filter: blur(42px);
  opacity:.62;
  animation: esMeshDrift 18s ease-in-out infinite alternate;
}

@keyframes esMeshDrift {
  from { transform: translate3d(-2%, -1%, 0) rotate(0deg) scale(1); }
  to { transform: translate3d(3%, 2%, 0) rotate(9deg) scale(1.05); }
}

.es-personal-hero > * {
  position: relative;
  z-index: 1;
}

.es-hero-row {
  display:grid;
  grid-template-columns: minmax(0, 1fr) auto;
  gap: 24px;
  align-items:center;
}

.es-welcome-title {
  font-size: clamp(36px, 4.2vw, 64px);
  line-height: .98;
  letter-spacing: -.045em;
  margin: 10px 0;
  font-weight: 950;
}

.es-hero-summary {
  max-width: 760px;
  color:#cbd5f5;
  font-size: 16px;
}

.es-hero-orb {
  width: 168px;
  height: 168px;
  border-radius: 999px;
  display:grid;
  place-items:center;
  color:#f8fbff;
  font-family:"Space Mono", monospace;
  font-size: 12px;
  text-transform:uppercase;
  background:
    radial-gradient(circle at 35% 30%, rgba(255,255,255,.28), transparent 18%),
    conic-gradient(from 210deg, #00d985, #34bdf6, #8b5cf6, #f59e0b, #00d985);
  box-shadow: 0 0 70px rgba(52,189,246,.28);
}

.es-fab-row {
  display:flex;
  flex-wrap:wrap;
  gap: 10px;
  margin-top: 20px;
}

.es-fab-action {
  display:inline-flex;
  align-items:center;
  justify-content:center;
  gap:8px;
  min-height:44px;
  padding: 0 16px;
  border-radius: 999px;
  text-decoration:none !important;
  color:#06121f !important;
  font-weight:900;
  background: linear-gradient(90deg, #00d985, #34bdf6);
  box-shadow: 0 16px 40px rgba(52,189,246,.22);
  transition: transform .18s ease, box-shadow .18s ease;
}

.es-fab-action.secondary {
  color:#f8fbff !important;
  background: rgba(255,255,255,.075);
  border: 1px solid rgba(255,255,255,.12);
}

.es-fab-action:hover {
  transform: translateY(-2px) scale(1.015);
  box-shadow: 0 24px 58px rgba(52,189,246,.30);
}

.es-area-chart {
  width: 100%;
  height: 76px;
  margin-top: 10px;
}

.es-area-chart path.line {
  fill: none;
  stroke: #55d7ff;
  stroke-width: 2.4;
  stroke-linecap: round;
  stroke-linejoin: round;
  stroke-dasharray: 220;
  stroke-dashoffset: 220;
  animation: esDrawLine 1.25s ease forwards;
}

.es-area-chart path.fill {
  opacity: .56;
  animation: esFadeIn .9s ease forwards;
}

@keyframes esDrawLine { to { stroke-dashoffset: 0; } }
@keyframes esFadeIn { from { opacity: 0; } to { opacity: .56; } }

.es-radial-wrap {
  display:flex;
  align-items:center;
  gap: 14px;
}

.es-radial {
  width: 118px;
  height: 118px;
}

.es-radial .track {
  fill:none;
  stroke:rgba(255,255,255,.10);
  stroke-width:10;
}

.es-radial .meter {
  fill:none;
  stroke:url(#esRadialGradient);
  stroke-width:10;
  stroke-linecap:round;
  transform: rotate(-90deg);
  transform-origin: 50% 50%;
  stroke-dasharray: 283;
  stroke-dashoffset: var(--dash);
  animation: esRadial 1.2s ease forwards;
}

@keyframes esRadial { from { stroke-dashoffset: 283; } }

.es-attention-list {
  display:grid;
  gap: 10px;
}

.es-attention-item {
  display:grid;
  grid-template-columns: 12px 1fr auto;
  gap: 12px;
  align-items:center;
  border:1px solid rgba(255,255,255,.08);
  background:rgba(255,255,255,.04);
  border-radius:12px;
  padding:12px;
}

.es-attention-dot {
  width:10px;
  height:10px;
  border-radius:999px;
  background:#34bdf6;
  box-shadow:0 0 18px rgba(52,189,246,.8);
}

.es-attention-item.critical .es-attention-dot {
  background:#ff4b33;
  box-shadow:0 0 20px rgba(255,75,51,.85);
  animation: esPulseDot 1.6s ease-in-out infinite;
}

.es-attention-item.major .es-attention-dot {
  background:#f59e0b;
  box-shadow:0 0 20px rgba(245,158,11,.78);
  animation: esPulseDot 2s ease-in-out infinite;
}

@keyframes esPulseDot {
  0%, 100% { transform: scale(1); opacity: .78; }
  50% { transform: scale(1.45); opacity: 1; }
}

.es-empty-state {
  border:1px dashed rgba(88,113,190,.34);
  border-radius:14px;
  padding:22px;
  min-height:220px;
  display:grid;
  place-items:center;
  text-align:center;
  background:
    radial-gradient(circle at 50% 18%, rgba(52,189,246,.16), transparent 28%),
    rgba(255,255,255,.025);
}

.es-empty-icon {
  width:72px;
  height:72px;
  margin:0 auto 12px;
  border-radius:20px;
  transform: rotate(-8deg);
  background: linear-gradient(135deg, rgba(0,217,133,.95), rgba(52,189,246,.9));
  box-shadow: 18px 18px 0 rgba(139,92,246,.20), 0 24px 60px rgba(52,189,246,.22);
}

.es-global-search {
  border: 1px solid rgba(52,189,246,.28);
  background: rgba(8,12,25,.72);
  border-radius: 12px;
  padding: 10px;
  margin: 8px 0 12px;
}

.es-search-result {
  display: grid;
  grid-template-columns: 54px 1fr;
  gap: 10px;
  align-items: center;
  border: 1px solid rgba(255,255,255,.08);
  background: rgba(255,255,255,.035);
  border-radius: 10px;
  padding: 9px;
  margin-top: 8px;
}

.es-kanban {
  display: grid;
  grid-template-columns: repeat(5, minmax(180px, 1fr));
  gap: 12px;
  margin: 14px 0 20px;
}

.es-kanban-col {
  min-height: 260px;
  border: 1px solid rgba(88,113,190,.24);
  border-radius: 12px;
  background: rgba(8,12,25,.60);
  padding: 12px;
}

.es-kanban-head {
  display:flex;
  align-items:center;
  justify-content:space-between;
  margin-bottom: 10px;
}

.es-kanban-card {
  border: 1px solid rgba(255,255,255,.09);
  background: linear-gradient(180deg, rgba(23,29,54,.88), rgba(13,17,32,.90));
  border-radius: 10px;
  padding: 11px;
  margin-bottom: 9px;
  box-shadow: 0 12px 32px rgba(0,0,0,.20);
}

.es-flyout {
  border: 1px solid rgba(52,189,246,.22);
  background: linear-gradient(180deg, rgba(18,24,48,.92), rgba(8,12,25,.92));
  border-radius: 12px;
  padding: 16px;
  box-shadow: 0 22px 64px rgba(0,0,0,.24);
}

.es-dropzone {
  border: 1.5px dashed rgba(52,189,246,.52);
  background:
    radial-gradient(circle at 18% 8%, rgba(0,217,133,.13), transparent 26%),
    rgba(8,12,25,.58);
  border-radius: 14px;
  padding: 18px;
  margin: 8px 0 12px;
  box-shadow: inset 0 1px 0 rgba(255,255,255,.045);
}

.es-stepper {
  display: grid;
  grid-template-columns: repeat(3, 1fr);
  gap: 10px;
  margin: 10px 0 18px;
}

.es-step {
  border: 1px solid rgba(255,255,255,.09);
  border-radius: 12px;
  background: rgba(255,255,255,.035);
  padding: 12px;
}

.es-step.active {
  border-color: rgba(0,217,133,.42);
  background: rgba(0,217,133,.08);
}

.es-tag-cloud {
  display:flex;
  flex-wrap:wrap;
  gap:8px;
  margin: 10px 0;
}

.es-tag {
  display:inline-flex;
  align-items:center;
  gap:7px;
  border:1px solid rgba(52,189,246,.24);
  background:rgba(52,189,246,.08);
  border-radius:999px;
  padding:7px 10px;
  color:#dff7ff;
  font-weight:800;
  font-size:12px;
}

.es-avatar-grid {
  display:grid;
  grid-template-columns: repeat(auto-fit, minmax(210px, 1fr));
  gap: 12px;
  margin-bottom: 20px;
}

.es-avatar-card {
  display:grid;
  grid-template-columns: 44px 1fr;
  gap: 12px;
  align-items:center;
  border:1px solid rgba(88,113,190,.24);
  background:rgba(255,255,255,.035);
  border-radius:12px;
  padding:12px;
}

.es-avatar {
  width:44px;
  height:44px;
  border-radius:14px;
  display:grid;
  place-items:center;
  color:#06121f;
  font-weight:950;
  background:linear-gradient(135deg, #00d985, #34bdf6);
}

.es-role-badge {
  display:inline-block;
  margin-top:5px;
  border:1px solid rgba(0,217,133,.24);
  border-radius:999px;
  color:#77f4bf;
  background:rgba(0,217,133,.08);
  padding:4px 8px;
  font-size:10px;
  font-family:"Space Mono", monospace;
  text-transform:uppercase;
}

.es-topology {
  display:grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 12px;
  margin: 14px 0;
}

.es-topology-node {
  position:relative;
  border:1px solid rgba(88,113,190,.28);
  background:rgba(255,255,255,.035);
  border-radius:12px;
  padding:14px;
}

.es-topology-node::before {
  content:"";
  width:10px;
  height:10px;
  border-radius:999px;
  display:inline-block;
  margin-right:8px;
  background:#00d985;
  box-shadow:0 0 18px rgba(0,217,133,.8);
}

.es-topology-node.warn::before {
  background:#f59e0b;
  box-shadow:0 0 18px rgba(245,158,11,.75);
}

.es-dashboard-panel {
  border: 1px solid rgba(88,113,190,.26);
  border-radius: 12px;
  background:
    linear-gradient(180deg, rgba(20,25,47,.84), rgba(10,13,27,.86));
  box-shadow: 0 22px 62px rgba(0,0,0,.24), inset 0 1px 0 rgba(255,255,255,.04);
  padding: 18px;
}

.es-dashboard-title {
  display:flex;
  align-items:center;
  justify-content:space-between;
  gap:14px;
  margin-bottom: 14px;
}

.es-dashboard-title h3 {
  margin:0;
  font-size: 19px;
  letter-spacing: -.02em;
}

.es-status-pill {
  display:inline-flex;
  align-items:center;
  gap:7px;
  border:1px solid rgba(0,217,133,.24);
  background:rgba(0,217,133,.09);
  color:#6ff5bf;
  border-radius:999px;
  padding:6px 10px;
  font-family:"Space Mono", monospace;
  font-size:10px;
  font-weight:800;
  text-transform:uppercase;
}

.es-status-pill::before {
  content:"";
  width:7px;
  height:7px;
  border-radius:999px;
  background:#00d985;
  box-shadow:0 0 14px rgba(0,217,133,.85);
}

.es-flow {
  display:grid;
  grid-template-columns: repeat(4, minmax(0,1fr));
  gap: 10px;
}

.es-flow-step {
  border:1px solid rgba(255,255,255,.08);
  background:rgba(255,255,255,.035);
  border-radius:10px;
  padding:13px;
  min-height:108px;
}

.es-flow-step b {
  display:block;
  margin:6px 0 6px;
}

.es-code-chip {
  color:#86efac;
  font-family:"Space Mono", monospace;
  font-size:10px;
  text-transform:uppercase;
}

.es-mini-table {
  display:grid;
  gap: 8px;
}

.es-mini-row {
  display:grid;
  grid-template-columns: 1fr auto;
  gap: 14px;
  align-items:center;
  border:1px solid rgba(255,255,255,.08);
  background:rgba(255,255,255,.035);
  border-radius:10px;
  padding:10px 12px;
}

.es-orbit {
  position:relative;
  min-height: 300px;
  overflow:hidden;
}

.es-orbit::before {
  content:"";
  position:absolute;
  width:300px;
  height:300px;
  border:1px solid rgba(52,189,246,.16);
  border-radius:999px;
  right:-70px;
  top:-40px;
  box-shadow: inset 0 0 54px rgba(52,189,246,.08);
}

.es-node {
  position:relative;
  z-index:1;
  border:1px solid rgba(255,255,255,.09);
  background:rgba(8,12,25,.72);
  border-radius:10px;
  padding:11px 12px;
  margin-bottom:10px;
}

.es-node strong {
  display:block;
  color:#f7fbff;
}

.es-node span {
  color:#a8b0d6;
  font-size:12px;
}

[data-testid="stDataFrame"] {
  border: 1px solid rgba(88,113,190,.24) !important;
  border-radius: 12px !important;
  box-shadow: 0 18px 46px rgba(0,0,0,.18);
}
.es-video-compact-note {
  color: #a8b0d6;
  font-size: 12px;
  margin: 6px 0 14px;
}
[data-testid="stVideo"] video {
  max-height: 220px !important;
  object-fit: contain !important;
  background: #000 !important;
  border-radius: 16px !important;
}
[data-testid="stVideo"] {
  max-width: 520px !important;
  margin-left: auto !important;
  margin-right: auto !important;
}

textarea, input, select {
  border-radius: 12px !important;
}

/* Streamlit's dataframe/data_editor uses an overlay input for the active cell.
   In long pages it can remain visible while scrolling, which looks like a
   selected Excel cell floating across the screen. Clip grid paint and hide
   stale overlay inputs when they are no longer being edited. */
div[data-testid="stDataFrame"],
div[data-testid="stDataEditor"] {
  contain: paint !important;
  isolation: isolate !important;
  overflow: hidden !important;
  border-radius: 10px;
}

div[data-testid="stDataFrame"] canvas,
div[data-testid="stDataEditor"] canvas {
  user-select: none !important;
}

div[data-testid="stDataFrame"] textarea:not(:focus),
div[data-testid="stDataFrame"] input:not(:focus),
div[data-testid="stDataEditor"] textarea:not(:focus),
div[data-testid="stDataEditor"] input:not(:focus),
textarea.gdg-input:not(:focus),
input.gdg-input:not(:focus),
textarea[class*="gdg"]:not(:focus),
input[class*="gdg"]:not(:focus),
textarea[class*="dvn"]:not(:focus),
input[class*="dvn"]:not(:focus) {
  opacity: 0 !important;
  pointer-events: none !important;
}

body:has(div[data-testid="stDataFrame"]:hover) textarea[class*="gdg"]:focus,
body:has(div[data-testid="stDataEditor"]:hover) textarea[class*="gdg"]:focus,
body:has(div[data-testid="stDataEditor"]:hover) input[class*="gdg"]:focus {
  opacity: 1 !important;
  pointer-events: auto !important;
}

[data-testid="stFileUploader"] {
  background: rgba(18,21,38,.72);
  border: 1px solid rgba(84,105,180,.24);
  border-radius: 18px;
  padding: 10px 14px;
}

[data-testid="stExpander"] {
  background: rgba(18,21,38,.70) !important;
  border: 1px solid rgba(84,105,180,.24) !important;
  border-radius: 18px !important;
}

@media (max-width: 1100px) {
  .es-grid-4, .es-grid-3 { grid-template-columns: 1fr; }
  .es-bento { grid-template-columns: 1fr; }
  .es-bento-card.wide { grid-column: span 1; }
  .es-radar-wrap { grid-template-columns: 1fr; }
  .es-landing-hero { grid-template-columns: 1fr; min-height: auto; }
  .es-product-scene { min-height: 440px; }
  .es-landing-grid, .es-price-band { grid-template-columns: 1fr; }
  .es-landing-links { display: none; }
  .es-topnav-row { align-items: flex-start; flex-direction: column; padding: 12px; }
  .es-topnav-brand { min-width: 0; }
  .es-topnav-links { justify-content: flex-start; }
  .es-topnav-link { min-height: 42px; padding: 0 10px; border-radius: 6px; }
  .es-topnav-link.active::after { display: none; }
  .es-topnav-tools { min-height: 52px; align-self: stretch; border-left: 0; border-top: 1px solid rgba(84,105,180,.24); }
  .es-topnav-user { min-width: 0; margin-left: auto; }
  .es-account-menu { right: 0; }
}
</style>
""",
    unsafe_allow_html=True,
)

components.html(
    """
    <script>
    const parentDoc = window.parent && window.parent.document;
    if (parentDoc && !parentDoc.__errorsweepGridScrollBlur) {
      parentDoc.__errorsweepGridScrollBlur = true;
      const isGridOverlay = (el) => {
        if (!el) return false;
        const tag = (el.tagName || "").toLowerCase();
        if (tag !== "textarea" && tag !== "input") return false;
        const cls = String(el.className || "");
        const aria = String(el.getAttribute("aria-label") || "");
        return cls.includes("gdg") || cls.includes("dvn") || aria.includes("cell") ||
          !!el.closest('[data-testid="stDataFrame"],[data-testid="stDataEditor"]');
      };
      parentDoc.addEventListener("scroll", () => {
        const el = parentDoc.activeElement;
        if (isGridOverlay(el)) el.blur();
      }, true);
      parentDoc.addEventListener("wheel", () => {
        const el = parentDoc.activeElement;
        if (isGridOverlay(el)) el.blur();
      }, true);
    }
    </script>
    """,
    height=0,
)


# ==========================================================
# Session and config helpers
# ==========================================================

def secret(name: str, default: str = "") -> str:
    value = os.environ.get(name)
    if value:
        return value
    try:
        value = st.secrets.get(name)
        if value:
            return value
    except Exception as exc:
        LOGGER.debug("Unable to read secret %s: %s", name, exc)
    return default


def is_production_mode() -> bool:
    mode = secret("ERRORSWEEP_ENV", secret("ENVIRONMENT", secret("APP_ENV", ""))).strip().lower()
    return mode in {"prod", "production"}


def session_secret() -> str:
    value = secret("ERRORSWEEP_SESSION_SECRET", DEFAULT_SESSION_SECRET)
    if is_production_mode() and value == DEFAULT_SESSION_SECRET:
        raise RuntimeError("ERRORSWEEP_SESSION_SECRET must be configured in production.")
    return value


def b64url(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).decode("utf-8").rstrip("=")


def hash_password(password: str, salt: Optional[str] = None, iterations: int = PASSWORD_HASH_ITERATIONS) -> str:
    salt = salt or b64url(os.urandom(16))
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), iterations)
    return f"pbkdf2_sha256${iterations}${salt}${b64url(digest)}"


def verify_password(password: str, stored_hash: str) -> bool:
    parts = str(stored_hash or "").split("$")
    if len(parts) != 4 or parts[0] != "pbkdf2_sha256":
        return False
    try:
        iterations = int(parts[1])
        digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), parts[2].encode("utf-8"), iterations)
        return hmac.compare_digest(b64url(digest), parts[3])
    except Exception as exc:
        LOGGER.warning("Password hash verification failed: %s", exc)
        return False


def verify_login_password(password: str, hash_secret_name: str, legacy_secret_name: str) -> bool:
    stored_hash = secret(hash_secret_name, "")
    if stored_hash:
        return verify_password(password, stored_hash)

    legacy_password = secret(legacy_secret_name, "")
    if not legacy_password:
        return False
    if is_production_mode():
        LOGGER.error("%s is plaintext and is not accepted in production. Configure %s.", legacy_secret_name, hash_secret_name)
        return False
    LOGGER.warning("%s is plaintext and should be replaced with %s.", legacy_secret_name, hash_secret_name)
    return hmac.compare_digest(password, legacy_password)


def password_configured(hash_secret_name: str, legacy_secret_name: str) -> bool:
    return bool(secret(hash_secret_name, "") or secret(legacy_secret_name, ""))


def trim_session_list(key: str, limit: int = SESSION_HISTORY_LIMIT) -> None:
    items = st.session_state.get(key)
    if isinstance(items, list) and len(items) > limit:
        del items[limit:]


def trim_session_collections() -> None:
    for key, limit in SESSION_COLLECTION_LIMITS.items():
        trim_session_list(key, limit)


def persist_saas_record(collection: str, record: Dict[str, Any]) -> Dict[str, Any]:
    if save_saas_record is None:
        return record
    try:
        return save_saas_record(collection, record, user=current_user() or {})
    except Exception as exc:
        LOGGER.warning("Unable to persist SaaS record %s: %s", collection, exc)
        return record


def load_saas_records(collection: str, workspace: str = "", include_all_workspaces: bool = False, limit: int = 500) -> List[Dict[str, Any]]:
    if fetch_saas_records is None:
        return []
    try:
        return fetch_saas_records(collection, workspace=workspace, include_all_workspaces=include_all_workspaces, limit=limit)
    except Exception as exc:
        LOGGER.warning("Unable to load SaaS records %s: %s", collection, exc)
        return []


def public_auth_link(route: str, token: str) -> str:
    query = f"?public={quote(route)}&token={quote(token)}"
    base_url = secret("ERRORSWEEP_PUBLIC_BASE_URL", "").strip().rstrip("/")
    if base_url:
        return f"{base_url}/{query}"
    return query


def auth_token_hash(token: str) -> str:
    return hmac.new(session_secret().encode("utf-8"), safe_text(token).encode("utf-8"), hashlib.sha256).hexdigest()


def auth_token_expired(record: Dict[str, Any]) -> bool:
    raw = safe_text(record.get("expires_at"))
    if not raw:
        return False
    try:
        expires = datetime.fromisoformat(raw.replace("Z", "+00:00"))
        if expires.tzinfo is None:
            expires = expires.replace(tzinfo=timezone.utc)
        return datetime.now(timezone.utc) > expires
    except Exception as exc:
        LOGGER.warning("Unable to parse auth token expiry %s: %s", raw, exc)
        return True


def upsert_session_record(key: str, record: Dict[str, Any]) -> None:
    rows = st.session_state.setdefault(key, [])
    record_id = safe_text(record.get("id"))
    for idx, item in enumerate(rows):
        if record_id and safe_text(item.get("id")) == record_id:
            rows[idx] = record
            break
    else:
        rows.insert(0, record)
    trim_session_list(key, SESSION_COLLECTION_LIMITS.get(key, SESSION_HISTORY_LIMIT))


def create_auth_token(email: str, token_type: str, workspace: str = "", metadata: Optional[Dict[str, Any]] = None) -> Tuple[str, Dict[str, Any]]:
    raw_token = b64url(os.urandom(32))
    expires_at = datetime.fromtimestamp(time.time() + AUTH_TOKEN_TTL_SECONDS, timezone.utc).isoformat()
    record = persist_saas_record("auth_tokens", {
        "workspace": workspace or "Demo Workspace",
        "user_email": email,
        "email": email,
        "token_hash": auth_token_hash(raw_token),
        "token_type": token_type,
        "status": "active",
        "expires_at": expires_at,
        "used_at": None,
        "metadata_json": metadata or {},
        "created_at": now_stamp(),
        "updated_at": now_stamp(),
    })
    upsert_session_record("auth_tokens", record)
    return raw_token, record


def auth_token_records() -> List[Dict[str, Any]]:
    records = list(st.session_state.get("auth_tokens", []))
    for record in load_saas_records("auth_tokens", include_all_workspaces=True, limit=SESSION_COLLECTION_LIMITS.get("auth_tokens", 500)):
        if not any(safe_text(item.get("id")) == safe_text(record.get("id")) for item in records):
            records.append(record)
    return records


def find_auth_token(raw_token: str, token_type: str) -> Optional[Dict[str, Any]]:
    if not safe_text(raw_token):
        return None
    digest = auth_token_hash(raw_token)
    for record in auth_token_records():
        if not hmac.compare_digest(safe_text(record.get("token_hash")), digest):
            continue
        if safe_text(record.get("token_type")) != token_type:
            continue
        if safe_text(record.get("status")).lower() != "active":
            return None
        if auth_token_expired(record):
            record["status"] = "expired"
            record["updated_at"] = now_stamp()
            persisted = persist_saas_record("auth_tokens", record)
            upsert_session_record("auth_tokens", persisted)
            return None
        return record
    return None


def consume_auth_token(record: Dict[str, Any]) -> Dict[str, Any]:
    updated = dict(record)
    updated["status"] = "used"
    updated["used_at"] = now_stamp()
    updated["updated_at"] = now_stamp()
    persisted = persist_saas_record("auth_tokens", updated)
    upsert_session_record("auth_tokens", persisted)
    return persisted


def update_stored_user(email: str, changes: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    email_key = safe_text(email).lower()
    users = list(st.session_state.get("users", []))
    for record in load_saas_records("users", include_all_workspaces=True, limit=SESSION_COLLECTION_LIMITS.get("users", 1000)):
        if not any(safe_text(item.get("id")) == safe_text(record.get("id")) for item in users):
            users.append(record)
    matched = None
    for item in users:
        if safe_text(item.get("email")).lower() == email_key:
            matched = dict(item)
            break
    if matched is None:
        return None
    matched.update(changes)
    matched["updated_at"] = now_stamp()
    persisted = persist_saas_record("users", matched)
    upsert_session_record("users", persisted)
    return persisted


def queue_verification_email(email: str, workspace: str, name: str = "") -> str:
    token, _ = create_auth_token(email, "email_verification", workspace, metadata={"name": name})
    link = public_auth_link("verify", token)
    queue_email_notification(
        email,
        "Verify your ErrorSweep email",
        f"Verify your ErrorSweep workspace email for '{workspace}' using this link: {link}",
        "auth.email_verification",
        metadata={"workspace": workspace, "verify_url": link},
        workspace=workspace,
    )
    return link


def queue_password_reset_email(email: str, workspace: str = "") -> str:
    token, _ = create_auth_token(email, "password_reset", workspace or "Demo Workspace")
    link = public_auth_link("reset", token)
    queue_email_notification(
        email,
        "Reset your ErrorSweep password",
        f"Reset your ErrorSweep password using this link: {link}",
        "auth.password_reset",
        metadata={"reset_url": link},
        workspace=workspace or "Demo Workspace",
    )
    return link


def hydrate_saas_state_for_user() -> None:
    user = current_user()
    if not user or st.session_state.get("_saas_state_hydrated"):
        return
    workspace = user.get("workspace", "")
    include_all = is_owner()
    for key in [
        "users",
        "workspaces",
        "projects",
        "jobs",
        "payments",
        "audit_logs",
        "notifications",
        "task_queue",
        "files",
        "subscriptions",
        "checkout_sessions",
        "billing_events",
        "auth_tokens",
    ]:
        records = load_saas_records(key, workspace=workspace, include_all_workspaces=include_all, limit=SESSION_COLLECTION_LIMITS.get(key, 500))
        if records:
            st.session_state[key] = records
    if fetch_persistent_usage_events is not None:
        try:
            usage_records = fetch_persistent_usage_events(SESSION_COLLECTION_LIMITS.get("ai_usage_events", 500))
            if usage_records:
                st.session_state["ai_usage_events"] = usage_records
        except Exception as exc:
            LOGGER.warning("Unable to hydrate usage events: %s", exc)
    st.session_state["_saas_state_hydrated"] = True


def b64url_decode(data: str) -> bytes:
    pad = "=" * (-len(data) % 4)
    return base64.urlsafe_b64decode(data + pad)


def sign_payload(payload: Dict[str, Any]) -> str:
    raw = json.dumps(payload, separators=(",", ":"), sort_keys=True).encode("utf-8")
    body = b64url(raw)
    sig = hmac.new(session_secret().encode("utf-8"), body.encode("utf-8"), hashlib.sha256).hexdigest()
    return f"{body}.{sig}"


def verify_payload(token: str) -> Optional[Dict[str, Any]]:
    try:
        body, sig = token.split(".", 1)
        expected = hmac.new(session_secret().encode("utf-8"), body.encode("utf-8"), hashlib.sha256).hexdigest()
        if not hmac.compare_digest(expected, sig):
            return None
        data = json.loads(b64url_decode(body))
        if int(data.get("exp", 0)) < int(time.time()):
            return None
        return data
    except RuntimeError:
        raise
    except Exception as exc:
        LOGGER.debug("Session token verification failed: %s", exc)
        return None


def query_get(name: str) -> str:
    try:
        val = st.query_params.get(name, "")
        if isinstance(val, list):
            return val[0] if val else ""
        return val or ""
    except Exception as exc:
        LOGGER.debug("Unable to read query param %s: %s", name, exc)
        return ""


def query_set(name: str, value: str) -> None:
    try:
        st.query_params[name] = value
    except Exception as exc:
        LOGGER.debug("Unable to set query param %s: %s", name, exc)


def query_clear(name: str) -> None:
    try:
        if name in st.query_params:
            del st.query_params[name]
    except Exception as exc:
        LOGGER.debug("Unable to clear query param %s: %s", name, exc)


def login_user(email: str, role: str, account_type: str, workspace: str = "Demo Workspace") -> None:
    user = {
        "email": email,
        "role": role,
        "account_type": account_type,
        "workspace": workspace,
        "login_at": datetime.now(timezone.utc).isoformat(),
    }
    st.session_state["user"] = user
    payload = {**user, "exp": int(time.time()) + SESSION_TTL_SECONDS}
    query_clear("public")
    query_set("es_session", sign_payload(payload))


def restore_session_from_query() -> None:
    if st.session_state.get("user"):
        return
    token = query_get("es_session")
    if not token:
        return
    data = verify_payload(token)
    if data:
        st.session_state["user"] = {
            "email": data.get("email", ""),
            "role": data.get("role", "User"),
            "account_type": data.get("account_type", "user"),
            "workspace": data.get("workspace", "Demo Workspace"),
            "login_at": data.get("login_at", ""),
        }


def logout() -> None:
    st.session_state.pop("user", None)
    st.session_state.pop("_saas_state_hydrated", None)
    query_clear("es_session")
    query_clear("es_page")
    query_clear("es_logout")
    st.rerun()


restore_session_from_query()


# ==========================================================
# Data initialization
# ==========================================================

def init_state() -> None:
    defaults = {
        "page": "Dashboard",
        "projects": [],
        "jobs": [],
        "tm": [],
        "glossary": [
            {"source": "Docflow", "target": "Docflow", "notes": "Product name / DNT"},
            {"source": "FitJourney", "target": "FitJourney", "notes": "Product name / DNT"},
        ],
        "dnt": ["Docflow", "FitJourney", "{{email}}", "{{password}}", "{{user_name}}"],
        "rule_instructions": [
            {"text": "Preserve placeholders, product names, emojis, numbers, and client locked terms unless explicitly instructed otherwise.", "source": "Default workspace policy"}
        ],
        "review_segments": [],
        "subtitle_segments": [],
        "payments": [
            {"date": "2026-05-01", "workspace": "Demo Workspace", "user": "demo@errorsweep.local", "plan": "Trial", "amount": 0, "currency": "INR", "status": "Demo"}
        ],
        "workspaces": [
            {"workspace": "Demo Workspace", "owner": "demo@errorsweep.local", "plan": "Trial", "status": "Active", "users": 3, "jobs": 0}
        ],
        "users": [
            {"email": "owner@errorsweep.local", "workspace": "Platform", "role": "Platform Owner", "plan": "Owner", "status": "Active"},
            {"email": "demo@errorsweep.local", "workspace": "Demo Workspace", "role": "Workspace Owner", "plan": "Trial", "status": "Active"},
            {"email": "reviewer@errorsweep.local", "workspace": "Demo Workspace", "role": "Reviewer", "plan": "Trial", "status": "Active"},
        ],
        "audit_logs": [],
        "ai_usage_events": [],
        "files": [],
        "notifications": [],
        "task_queue": [],
        "subscriptions": [],
        "checkout_sessions": [],
        "billing_events": [],
        "auth_tokens": [],
        "selected_review_index": 0,
        "selected_subtitle_index": 0,
        "subtitle_editor_active": False,
        "subtitle_workflow": "Transcription",
        "subtitle_video_bytes": None,
        "subtitle_video_metadata": {},
        "subtitle_video_name": "",
        "subtitle_video_type": "video/mp4",
        "show_timing_grid": False,
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val


init_state()
trim_session_collections()


# ==========================================================
# Permissions and navigation
# ==========================================================

OWNER_PAGES = [
    "Owner Console",
    "Payments Received",
    "User Access Matrix",
    "All Workspaces",
    "Platform Settings",
    "Platform Audit Logs",
]

WORKSPACE_PAGES = [
    "Dashboard",
    "Projects",
    "Jobs",
    "ErrorSweep QA",
    "ErrorSweep Pro",
    "Subtitle / Transcription Editor",
    "Scorecards",
    "Memory & Rules",
    "Team & Roles",
    "Billing",
    "Account",
    "Admin",
]

# Hidden route pages. They are not shown as navigation buttons,
# but they let editors open as dedicated professional workspaces.
HIDDEN_EDITOR_PAGES = [
    "Human Review Workspace",
    "Subtitle Workspace",
    "Transcription Workspace",
]

ROLE_PAGE_ACCESS = {
    "Platform Owner": OWNER_PAGES + WORKSPACE_PAGES,
    "Workspace Owner": WORKSPACE_PAGES,
    "Workspace Admin": ["Dashboard", "Projects", "Jobs", "ErrorSweep QA", "ErrorSweep Pro", "Subtitle / Transcription Editor", "Scorecards", "Memory & Rules", "Team & Roles", "Account", "Admin"],
    "Project Manager": ["Dashboard", "Projects", "Jobs", "ErrorSweep QA", "ErrorSweep Pro", "Subtitle / Transcription Editor", "Scorecards", "Memory & Rules", "Account"],
    "Translator": ["Dashboard", "Jobs", "Subtitle / Transcription Editor", "Account"],
    "Reviewer": ["Dashboard", "Jobs", "ErrorSweep QA", "Subtitle / Transcription Editor", "Scorecards", "Memory & Rules", "Account"],
    "Client Viewer": ["Dashboard", "Jobs", "Account"],
    "Billing Admin": ["Dashboard", "Billing", "Account"],
    "User": ["Dashboard", "Projects", "Jobs", "ErrorSweep QA", "ErrorSweep Pro", "Subtitle / Transcription Editor", "Scorecards", "Memory & Rules", "Account"],
}


def current_user() -> Optional[Dict[str, Any]]:
    return st.session_state.get("user")


def current_role() -> str:
    user = current_user() or {}
    return user.get("role", "User")


def allowed_pages() -> List[str]:
    pages = list(ROLE_PAGE_ACCESS.get(current_role(), ROLE_PAGE_ACCESS["User"]))
    # Add hidden editor pages without showing them in the left navigation.
    for page in HIDDEN_EDITOR_PAGES:
        if page not in pages:
            pages.append(page)
    return pages


def is_owner() -> bool:
    return current_role() == "Platform Owner"


def page_link(page: str) -> str:
    token = query_get("es_session")
    page_param = quote(page)
    if token:
        return f"?es_session={token}&es_page={page_param}"
    return f"?es_page={page_param}"


def public_page_link(page: str) -> str:
    return f"?public={quote(page)}"


def open_page(page: str) -> None:
    """Open an internal ErrorSweep route as a dedicated page in the same session."""
    st.session_state.page = page
    query_set("es_page", page)
    if page == "Human Review Workspace":
        session_id = st.session_state.get("active_review_session_id")
        if session_id:
            query_set("review_id", str(session_id))
    st.rerun()


def nav_button(page: str, key_prefix: str = "nav") -> None:
    active = st.session_state.get("page") == page
    cls = "es-nav-link active" if active else "es-nav-link"
    label = page
    if st.session_state.get("nav_compact"):
        label = "".join(part[:1] for part in re.split(r"[\s/&]+", page) if part)[:4].upper() or page[:3].upper()
    st.markdown(
        f'<a class="{cls}" href="{page_link(page)}" target="_self" title="{escape(page)}">{escape(label)}</a>',
        unsafe_allow_html=True,
    )


def render_navigation() -> None:
    user = current_user() or {}
    pages = allowed_pages()
    label_map = {
        "Dashboard": "Dashboard",
        "Projects": "Projects",
        "Jobs": "Jobs",
        "ErrorSweep QA": "QA Tasks",
        "ErrorSweep Pro": "Pro",
        "Subtitle / Transcription Editor": "Subtitles",
        "Scorecards": "Scorecards",
        "Memory & Rules": "Rules",
        "Team & Roles": "Team",
        "Billing": "Billing",
        "Account": "Account",
        "Admin": "Admin",
    }
    workspace_links = []
    for page in WORKSPACE_PAGES:
        if page not in pages:
            continue
        active = " active" if st.session_state.get("page") == page else ""
        workspace_links.append(
            f'<a class="es-topnav-link{active}" href="{page_link(page)}" target="_self">{escape(label_map.get(page, page))}</a>'
        )
    owner_links = []
    if is_owner():
        for page in OWNER_PAGES:
            active = " active" if st.session_state.get("page") == page else ""
            owner_links.append(
                f'<a class="es-owner-link{active}" href="{page_link(page)}" target="_self">{escape(page)}</a>'
            )
    open_count = sum(
        1
        for job in st.session_state.get("jobs", [])
        if safe_text(job.get("status")).lower() in {"draft", "needs human review", "needs_review", "running"}
    )
    notification_count = len(st.session_state.get("notifications", []))
    user_email = safe_text(user.get("email", "user@errorsweep.local"))
    user_name = user_email.split("@", 1)[0].replace("_", " ").replace(".", " ").title() or "User"
    role = current_role()
    settings_page = "Platform Settings" if is_owner() else ("Admin" if "Admin" in pages else "Account")
    billing_item = (
        f'<a href="{page_link("Billing")}" target="_self">Billing <span>Plan</span></a>'
        if "Billing" in pages
        else ""
    )
    topnav = f"""
    <nav class="es-topnav">
      <div class="es-topnav-row">
        <a class="es-topnav-brand" href="{page_link('Dashboard')}" target="_self">
          <div class="es-topnav-mark">ES</div>
          <div>
            <div class="es-topnav-name">error<span>sweep</span></div>
            <div class="es-topnav-sub">by Nawin Corp</div>
          </div>
        </a>
        <div class="es-topnav-links">
          {''.join(workspace_links)}
        </div>
        <div class="es-topnav-tools">
          <a class="es-topnav-tool" href="{page_link('Jobs')}" target="_self" title="Jobs">JOBS<span class="es-topnav-badge">{open_count}</span></a>
          <a class="es-topnav-tool" href="{page_link('Account')}" target="_self" title="Notifications">NOTES<span class="es-topnav-badge">{notification_count}</span></a>
          <div class="es-topnav-tool" title="Language">EN</div>
          <div class="es-topnav-user-wrap">
            <div class="es-topnav-user" tabindex="0" title="Account menu">
              <div>
                <div style="font-weight:900;white-space:nowrap;">{escape(user_name)}</div>
                <div style="font-size:11px;color:#8ea1dc;font-weight:800;">{escape(role)}</div>
              </div>
              <div class="es-topnav-avatar">{escape(monogram(user_name))}</div>
              <span class="es-account-caret">v</span>
            </div>
            <div class="es-account-menu">
              <a href="{page_link('Account')}" target="_self">Profile <span>Account</span></a>
              <a href="{page_link(settings_page)}" target="_self">Settings <span>{escape(settings_page)}</span></a>
              {billing_item}
              <a href="{page_link('Jobs')}" target="_self">Jobs <span>{open_count}</span></a>
              <a class="logout" href="?es_logout=1" target="_self">Logout <span>Exit</span></a>
            </div>
          </div>
        </div>
      </div>
      {f'<div class="es-owner-strip"><span>Owner only</span>{"".join(owner_links)}</div>' if owner_links else ''}
    </nav>
    """
    st.markdown(topnav, unsafe_allow_html=True)


# ==========================================================
# General helpers
# ==========================================================

def now_stamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def local_timezone() -> timezone:
    configured = secret("ERRORSWEEP_DISPLAY_TIMEZONE", secret("TZ", "")).strip()
    if configured:
        try:
            return ZoneInfo(configured)
        except Exception as exc:
            LOGGER.warning("Invalid display timezone %s: %s", configured, exc)
    try:
        browser_timezone = getattr(st.context, "timezone", None)
        if browser_timezone:
            return ZoneInfo(str(browser_timezone))
    except Exception as exc:
        LOGGER.debug("Browser timezone unavailable: %s", exc)
    return datetime.now().astimezone().tzinfo or timezone.utc


def format_local_time(value: Any) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, (int, float)) and value > 1_000_000_000:
        dt = datetime.fromtimestamp(float(value), tz=timezone.utc)
    else:
        raw = str(value).strip()
        if not raw:
            return ""
        if re.fullmatch(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}", raw):
            return raw
        try:
            dt = datetime.fromisoformat(raw.replace("Z", "+00:00"))
        except ValueError:
            return raw
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=local_timezone())
    local_dt = dt.astimezone(local_timezone())
    zone = local_dt.tzname() or ""
    if " " in zone:
        zone = "".join(part[0] for part in zone.split() if part[:1]).upper()
    return f"{local_dt.strftime('%Y-%m-%d %I:%M %p')} {zone}".strip()


def display_records(records: List[Dict[str, Any]], time_columns: Optional[set] = None) -> List[Dict[str, Any]]:
    cols = time_columns or TIME_DISPLAY_COLUMNS
    output = []
    for record in records or []:
        item = dict(record)
        for key in list(item.keys()):
            if key in cols or key.endswith("_at"):
                item[key] = format_local_time(item.get(key))
        output.append(item)
    return output


def display_dataframe(records: List[Dict[str, Any]], **kwargs) -> None:
    st.dataframe(pd.DataFrame(display_records(records)), use_container_width=True, hide_index=True, **kwargs)


def add_audit(action: str, details: str = "") -> None:
    record = {
        "time": now_stamp(),
        "actor": (current_user() or {}).get("email", "unknown"),
        "workspace": (current_user() or {}).get("workspace", "Demo Workspace"),
        "action": action,
        "details": details,
    }
    record = persist_saas_record("audit_logs", record)
    st.session_state.audit_logs.insert(0, record)
    trim_session_list("audit_logs")


def email_provider_label() -> str:
    return secret("ERRORSWEEP_EMAIL_PROVIDER", "").strip().lower() or "not_configured"


def email_from_address() -> str:
    return (
        secret("ERRORSWEEP_EMAIL_FROM", "")
        or secret("SENDGRID_FROM_EMAIL", "")
        or secret("RESEND_FROM_EMAIL", "")
        or "no-reply@errorsweep.local"
    ).strip()


def queue_email_notification(
    recipient: str,
    subject: str,
    body: str,
    event_type: str,
    metadata: Optional[Dict[str, Any]] = None,
    workspace: str = "",
) -> Dict[str, Any]:
    """Create a workspace-scoped notification record for later email dispatch."""
    recipient = safe_text(recipient).strip()
    if not recipient:
        recipient = safe_text((current_user() or {}).get("email", ""))
    provider = email_provider_label()
    record = {
        "id": uuid.uuid4().hex,
        "created": now_stamp(),
        "updated_at": now_stamp(),
        "workspace": workspace or safe_text((current_user() or {}).get("workspace", "Demo Workspace")),
        "recipient": recipient,
        "subject": safe_text(subject),
        "body": safe_text(body),
        "event_type": safe_text(event_type),
        "provider": provider,
        "status": "provider_pending" if provider in {"", "manual", "not_configured"} else "queued",
        "error": "",
        "sent_at": "",
        "metadata_json": metadata or {},
    }
    record = persist_saas_record("notifications", record)
    st.session_state.setdefault("notifications", [])
    st.session_state.notifications.insert(0, record)
    trim_session_list("notifications")
    return record


def dispatch_email_notification(record: Dict[str, Any]) -> Dict[str, Any]:
    """Send one queued notification via Resend, SendGrid, or SMTP when configured."""
    provider = email_provider_label()
    if provider in {"", "manual", "not_configured"}:
        record["status"] = "provider_pending"
        record["error"] = "Email provider is not configured."
        record["updated_at"] = now_stamp()
        return record

    recipient = safe_text(record.get("recipient", "")).strip()
    subject = safe_text(record.get("subject", "")).strip()
    body = safe_text(record.get("body", "")).strip()
    sender = email_from_address()
    try:
        if not recipient or "@" not in recipient:
            raise ValueError("Notification recipient is missing or invalid.")
        if provider == "resend":
            api_key = secret("RESEND_API_KEY", secret("ERRORSWEEP_RESEND_API_KEY", "")).strip()
            if not api_key:
                raise RuntimeError("RESEND_API_KEY is not configured.")
            response = requests.post(
                "https://api.resend.com/emails",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={"from": sender, "to": [recipient], "subject": subject, "text": body},
                timeout=20,
            )
            if response.status_code >= 300:
                raise RuntimeError(f"Resend returned {response.status_code}: {response.text[:300]}")
        elif provider == "sendgrid":
            api_key = secret("SENDGRID_API_KEY", secret("ERRORSWEEP_SENDGRID_API_KEY", "")).strip()
            if not api_key:
                raise RuntimeError("SENDGRID_API_KEY is not configured.")
            response = requests.post(
                "https://api.sendgrid.com/v3/mail/send",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={
                    "personalizations": [{"to": [{"email": recipient}]}],
                    "from": {"email": sender},
                    "subject": subject,
                    "content": [{"type": "text/plain", "value": body}],
                },
                timeout=20,
            )
            if response.status_code >= 300:
                raise RuntimeError(f"SendGrid returned {response.status_code}: {response.text[:300]}")
        elif provider == "smtp":
            host = secret("SMTP_HOST", secret("ERRORSWEEP_SMTP_HOST", "")).strip()
            port = int(secret("SMTP_PORT", secret("ERRORSWEEP_SMTP_PORT", "587")) or "587")
            username = secret("SMTP_USER", secret("ERRORSWEEP_SMTP_USER", "")).strip()
            password = secret("SMTP_PASSWORD", secret("ERRORSWEEP_SMTP_PASSWORD", "")).strip()
            if not host:
                raise RuntimeError("SMTP_HOST is not configured.")
            message = EmailMessage()
            message["From"] = sender
            message["To"] = recipient
            message["Subject"] = subject
            message.set_content(body)
            with smtplib.SMTP(host, port, timeout=20) as client:
                if secret("SMTP_TLS", secret("ERRORSWEEP_SMTP_TLS", "true")).strip().lower() not in {"0", "false", "no", "off"}:
                    client.starttls()
                if username:
                    client.login(username, password)
                client.send_message(message)
        else:
            raise RuntimeError(f"Unsupported email provider: {provider}. Use resend, sendgrid, smtp, or manual.")
        record["status"] = "sent"
        record["sent_at"] = now_stamp()
        record["error"] = ""
    except Exception as exc:
        record["status"] = "failed"
        record["error"] = safe_text(exc)[:700]
        LOGGER.warning("Email notification dispatch failed: %s", exc)
    record["provider"] = provider
    record["updated_at"] = now_stamp()
    return record


def dispatch_pending_notifications(limit: int = EMAIL_DISPATCH_BATCH_LIMIT) -> Tuple[int, int]:
    sent = 0
    failed = 0
    pending_statuses = {"queued", "failed"}
    for record in st.session_state.get("notifications", []):
        if sent + failed >= max(1, int(limit or 1)):
            break
        if safe_text(record.get("status")).lower() not in pending_statuses:
            continue
        updated = dispatch_email_notification(record)
        if updated.get("status") == "sent":
            sent += 1
        else:
            failed += 1
    return sent, failed


def create_task_record(
    task_type: str,
    label: str,
    total_units: int = 0,
    metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Create a durable lifecycle record for a long-running workflow."""
    user = current_user() or {}
    now = now_stamp()
    record = {
        "id": uuid.uuid4().hex,
        "workspace": user.get("workspace", "Demo Workspace"),
        "user_email": user.get("email", ""),
        "task_type": safe_text(task_type),
        "label": safe_text(label),
        "status": "queued",
        "progress": 0,
        "total_units": int(total_units or 0),
        "processed_units": 0,
        "result_ref": "",
        "error": "",
        "metadata_json": metadata or {},
        "started_at": "",
        "finished_at": "",
        "created_at": now,
        "updated_at": now,
    }
    record = persist_saas_record("task_queue", record)
    st.session_state.setdefault("task_queue", [])
    st.session_state.task_queue.insert(0, record)
    trim_session_list("task_queue")
    return record


def queue_external_workflow_if_configured(
    task: Dict[str, Any],
    workflow: str,
    primary_file: Any,
    rules_file: Optional[Any] = None,
    parameters: Optional[Dict[str, Any]] = None,
) -> bool:
    """Queue a heavy workflow on an external worker when one is configured.

    Local development keeps running inline. In production, the worker receives
    object-storage manifests for the uploaded source/rules files so Streamlit
    does not need to hold the request open while work runs.
    """
    if async_backend_status is None or enqueue_async_task is None:
        return False
    status = async_backend_status()
    if status.get("mode") != "external":
        return False

    task_id = safe_text(task.get("id")) or uuid.uuid4().hex
    input_manifests = save_job_attachment_files(task_id, [primary_file] if primary_file is not None else [], purpose=f"{workflow}_input")
    rules_manifests = save_job_attachment_files(task_id, [rules_file] if rules_file is not None else [], purpose=f"{workflow}_rules")
    payload = {
        "workflow": workflow,
        "input_files": input_manifests,
        "rules_files": rules_manifests,
        "parameters": parameters or {},
        "object_storage": object_storage_status() if object_storage_status is not None else {"provider": "local"},
        "created_at": now_stamp(),
    }
    try:
        queued = enqueue_async_task(task, payload)
    except Exception as exc:
        LOGGER.warning("External async enqueue failed for task %s: %s", task_id, exc)
        update_task_record(task_id, status="failed", progress=100, error=f"External queue failed: {exc}")
        st.warning("External worker queue is configured, but enqueue failed. The task was not run inline to avoid duplicate production processing.")
        return True

    if queued.get("queued"):
        update_task_record(
            task_id,
            status="queued",
            progress=2,
            processed_units=0,
            result_ref=safe_text(queued.get("external_id", "")),
            metadata_json={
                **(task.get("metadata_json") or {}),
                "async_provider": queued.get("provider", ""),
                "external_id": queued.get("external_id", ""),
                "input_files": input_manifests,
                "rules_files": rules_manifests,
                "queued_message": queued.get("message", ""),
            },
        )
        add_audit("External task queued", f"{workflow}: {queued.get('provider')} {queued.get('external_id')}")
        st.success(f"Task queued on external worker ({queued.get('provider')}). You can track it from Jobs.")
        if input_manifests:
            st.caption(f"Stored input file: {safe_text(input_manifests[0].get('storage_key', ''))[:180]}")
        return True

    return False


def update_task_record(task_id: str, **changes: Any) -> Dict[str, Any]:
    """Update a task lifecycle record in session state and persistence."""
    task_id = safe_text(task_id)
    task_rows = st.session_state.setdefault("task_queue", [])
    task = next((item for item in task_rows if safe_text(item.get("id")) == task_id), None)
    if task is None:
        task = {"id": task_id, "created_at": now_stamp()}
        task_rows.insert(0, task)
    task.update(changes)
    task["updated_at"] = now_stamp()
    status = safe_text(task.get("status")).lower()
    if status == "running" and not task.get("started_at"):
        task["started_at"] = now_stamp()
    if status in {"completed", "failed", "cancelled", "needs_review"} and not task.get("finished_at"):
        task["finished_at"] = now_stamp()
    try:
        task["progress"] = max(0, min(100, int(task.get("progress", 0) or 0)))
    except Exception:
        task["progress"] = 0
    persisted = persist_saas_record("task_queue", dict(task))
    task.update(persisted)
    return task


def task_status_summary(tasks: List[Dict[str, Any]]) -> Dict[str, int]:
    summary = {"queued": 0, "running": 0, "completed": 0, "failed": 0, "needs_review": 0}
    for task in tasks or []:
        status = safe_text(task.get("status")).lower() or "queued"
        summary[status] = summary.get(status, 0) + 1
    return summary


def render_task_queue_panel(limit: int = 12) -> None:
    tasks = st.session_state.get("task_queue", [])
    st.markdown("### Async task queue")
    if not tasks:
        st.info("No task lifecycle records yet. QA and Pro runs will appear here with progress, status, and retry visibility.")
        return
    summary = task_status_summary(tasks)
    metrics([
        ("Queued", summary.get("queued", 0), "waiting"),
        ("Running", summary.get("running", 0), "active"),
        ("Completed", summary.get("completed", 0), "finished"),
        ("Failed", summary.get("failed", 0), "needs retry"),
    ])
    latest = tasks[:limit]
    for idx, task in enumerate(latest):
        label = safe_text(task.get("label") or task.get("task_type") or "Task")
        status = safe_text(task.get("status") or "queued")
        progress = max(0, min(100, int(task.get("progress", 0) or 0)))
        with st.container(border=True):
            c1, c2, c3 = st.columns([0.46, 0.24, 0.30])
            c1.markdown(f"**{label}**")
            c1.caption(f"{safe_text(task.get('task_type', 'workflow'))} · {format_local_time(task.get('updated_at', task.get('created_at', '')))}")
            c2.markdown(f"`{status}`")
            c2.caption(f"{int(task.get('processed_units') or 0)}/{int(task.get('total_units') or 0)} units")
            c3.progress(progress / 100, text=f"{progress}%")
            if safe_text(task.get("error")):
                st.caption(f"Error: {safe_text(task.get('error'))[:240]}")
            if status.lower() in {"failed", "cancelled"}:
                if st.button("Queue retry request", key=f"retry_task_{idx}_{task.get('id')}", use_container_width=True):
                    update_task_record(
                        safe_text(task.get("id")),
                        status="queued",
                        progress=0,
                        processed_units=0,
                        error="Retry requested from Jobs page. Re-run the source workflow to execute.",
                    )
                    add_audit("Task retry requested", label)
                    st.rerun()
    with st.expander("Raw task records", expanded=False):
        rows = []
        for task in tasks[:100]:
            row = dict(task)
            if isinstance(row.get("metadata_json"), (dict, list)):
                row["metadata_json"] = json.dumps(row["metadata_json"], ensure_ascii=False)
            rows.append(row)
        st.dataframe(pd.DataFrame(display_records(rows)), use_container_width=True, hide_index=True)


def hero(kicker: str, title: str, subtitle: str) -> None:
    st.markdown(
        f"""
        <section class="es-hero">
          <div class="es-kicker">{escape(kicker)}</div>
          <div class="es-title">{title}</div>
          <div class="es-subtitle">{escape(subtitle)}</div>
        </section>
        """,
        unsafe_allow_html=True,
    )


def metrics(cards: List[Tuple[str, Any, str]]) -> None:
    """Render metric cards without raw HTML."""
    if not cards:
        return
    cols = st.columns(min(4, len(cards)))
    for idx, (label, value, note) in enumerate(cards):
        with cols[idx % len(cols)]:
            with st.container(border=True):
                st.caption(str(label).upper())
                st.markdown(f"### {escape(str(value))}")
                if note:
                    st.caption(str(note))


def sparkline_svg(values: List[int]) -> str:
    values = [int(v or 0) for v in values] or [0]
    if len(values) == 1:
        values = [0, values[0]]
    max_v = max(values) or 1
    min_v = min(values)
    spread = max(max_v - min_v, 1)
    points = []
    for idx, val in enumerate(values):
        x = 4 + idx * (92 / max(1, len(values) - 1))
        y = 30 - ((val - min_v) / spread) * 24
        points.append(f"{x:.1f},{y:.1f}")
    return f'<svg class="es-spark" viewBox="0 0 100 34" preserveAspectRatio="none"><polyline points="{" ".join(points)}" /></svg>'


def area_chart_svg(values: List[int], chart_id: str = "chart") -> str:
    values = [int(v or 0) for v in values] or [0]
    if len(values) == 1:
        values = [0, values[0]]
    max_v = max(values) or 1
    min_v = min(values)
    spread = max(max_v - min_v, 1)
    points = []
    for idx, val in enumerate(values):
        x = 4 + idx * (92 / max(1, len(values) - 1))
        y = 62 - ((val - min_v) / spread) * 48
        points.append((x, y))
    line_points = " ".join(f"{x:.1f},{y:.1f}" for x, y in points)
    area_points = f"4,72 {line_points} 96,72"
    gradient_id = f"esArea{re.sub(r'[^a-zA-Z0-9]+', '', chart_id) or 'Chart'}"
    return f"""
    <svg class="es-area-chart" viewBox="0 0 100 76" preserveAspectRatio="none">
      <defs>
        <linearGradient id="{gradient_id}" x1="0" x2="0" y1="0" y2="1">
          <stop offset="0%" stop-color="#34bdf6" stop-opacity=".52"/>
          <stop offset="100%" stop-color="#00d985" stop-opacity="0"/>
        </linearGradient>
      </defs>
      <path class="fill" d="M {area_points} Z" fill="url(#{gradient_id})"/>
      <path class="line" d="M {line_points}"/>
    </svg>
    """


def radial_progress_svg(score: int, label: str = "TQI") -> str:
    score = max(0, min(100, int(score or 0)))
    dash = 283 - (score / 100) * 283
    return f"""
    <div class="es-radial-wrap">
      <svg class="es-radial" viewBox="0 0 120 120" style="--dash:{dash:.1f};">
        <defs>
          <linearGradient id="esRadialGradient" x1="0" x2="1" y1="0" y2="1">
            <stop offset="0%" stop-color="#00d985"/>
            <stop offset="55%" stop-color="#34bdf6"/>
            <stop offset="100%" stop-color="#8b5cf6"/>
          </linearGradient>
        </defs>
        <circle class="track" cx="60" cy="60" r="45"/>
        <circle class="meter" cx="60" cy="60" r="45"/>
        <text x="60" y="57" text-anchor="middle" fill="#f7fbff" font-size="24" font-weight="900">{score}</text>
        <text x="60" y="76" text-anchor="middle" fill="#a8b0d6" font-size="10" font-family="Space Mono">{escape(label)}</text>
      </svg>
      <div>
        <div class="es-metric-label">Aggregate Quality</div>
        <div class="es-small">Derived from current queue health, review load, and completed work.</div>
      </div>
    </div>
    """


def first_name_from_user(user: Dict[str, Any]) -> str:
    email = safe_text(user.get("email", ""))
    if not email:
        return "there"
    name = email.split("@", 1)[0].replace("_", " ").replace(".", " ").strip()
    return name.split()[0].title() if name else "there"


def dashboard_attention_items() -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    for job in st.session_state.get("jobs", []):
        status = safe_text(job.get("status", ""))
        severity = "critical" if status.lower() in {"blocked", "critical", "failed"} else "major" if "review" in status.lower() or "needs" in status.lower() else ""
        if severity:
            items.append({
                "severity": severity,
                "title": f"{safe_text(job.get('type', 'Job'))}: {status}",
                "meta": f"{safe_text(job.get('language', '')) or safe_text(job.get('workspace', 'Workspace'))} · {safe_text(job.get('segments', '')) or 'open'} segments",
            })
    if st.session_state.get("review_segments"):
        pending = sum(1 for r in st.session_state.review_segments if r.get("status") not in ("Approved", "Rejected"))
        if pending:
            items.insert(0, {"severity": "major", "title": "Human Review queue waiting", "meta": f"{pending} segment(s) need a decision"})
    return items[:6]


def render_stepper(steps: List[str], active_idx: int = 0) -> None:
    cards = []
    for idx, label in enumerate(steps):
        cls = "es-step active" if idx == active_idx else "es-step"
        cards.append(f'<div class="{cls}"><span class="es-code-chip">Step {idx + 1}</span><br><b>{escape(label)}</b></div>')
    st.html(f'<div class="es-stepper">{"".join(cards)}</div>')


def render_upload_dropzone(title: str, subtitle: str, file_types: str) -> None:
    st.html(
        f"""
        <div class="es-dropzone">
          <div class="es-code-chip">{escape(file_types)}</div>
          <h3 style="margin:6px 0 4px;">{escape(title)}</h3>
          <div class="es-small">{escape(subtitle)}</div>
        </div>
        """
    )


def monogram(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9 ]+", " ", safe_text(value)).strip()
    parts = cleaned.split()
    if not parts and "@" in value:
        parts = value.split("@", 1)[0].replace(".", " ").replace("_", " ").split()
    if not parts:
        return "ES"
    return "".join(part[:1].upper() for part in parts[:2])


def render_tag_cloud(items: List[str], empty: str = "No saved tags yet.") -> None:
    tags = [safe_text(item) for item in items if safe_text(item)]
    if not tags:
        st.info(empty)
        return
    st.html('<div class="es-tag-cloud">' + "".join(f'<span class="es-tag">{escape(tag)}</span>' for tag in tags[:120]) + '</div>')


def render_topology_map(health: Optional[Dict[str, Any]] = None) -> None:
    health = health or {}
    supabase_ready = bool(health.get("supabase_configured"))
    jobs_ready = safe_text(health.get("editor_jobs_table", "")).lower() == "ok"
    usage_ready = safe_text(health.get("usage_events_table", "")).lower() == "ok"
    nodes = [
        ("Streamlit App", "Ready", ""),
        ("Supabase Persistence", "Connected" if supabase_ready and jobs_ready and usage_ready else "Fallback mode", "" if supabase_ready else "warn"),
        ("Self-hosted MT", "Configured" if current_builtin_engine_label else "Available when worker is running", "" if current_builtin_engine_label else "warn"),
    ]
    html = "".join(
        f'<div class="es-topology-node {cls}"><b>{escape(name)}</b><br><span class="es-small">{escape(status)}</span></div>'
        for name, status, cls in nodes
    )
    st.html(f'<div class="es-topology">{html}</div>')


def job_pipeline_status(job: Dict[str, Any]) -> str:
    status = safe_text(job.get("status", "")).lower()
    job_type = safe_text(job.get("type", "")).lower()
    if status in {"draft", ""}:
        return "Draft"
    if "translation" in job_type or "translat" in status:
        return "Translating"
    if "review" in status or "review" in job_type or "needs" in status:
        return "Human Review"
    if "qa" in job_type or "qa" in status:
        return "QA"
    if "complete" in status or "deliver" in status:
        return "Delivered"
    return "QA"


def render_jobs_kanban(jobs: List[Dict[str, Any]]) -> None:
    columns = ["Draft", "Translating", "Human Review", "QA", "Delivered"]
    grouped = {col: [] for col in columns}
    for job in jobs:
        grouped.setdefault(job_pipeline_status(job), []).append(job)
    html_cols = []
    for col in columns:
        cards = []
        for job in grouped.get(col, [])[:8]:
            cards.append(
                f'<div class="es-kanban-card"><span class="es-code-chip">{escape(safe_text(job.get("type", "Job")))}</span>'
                f'<b>{escape(safe_text(job.get("language") or job.get("workspace") or "Workspace"))}</b>'
                f'<div class="es-small">{escape(safe_text(job.get("status", "")))} · {escape(format_local_time(job.get("created", job.get("created_at", ""))))}</div></div>'
            )
        if not cards:
            cards.append('<div class="es-small">No jobs in this lane.</div>')
        html_cols.append(
            f'<section class="es-kanban-col"><div class="es-kanban-head"><b>{escape(col)}</b><span class="es-code-chip">{len(grouped.get(col, []))}</span></div>{"".join(cards)}</section>'
        )
    st.html(f'<div class="es-kanban">{"".join(html_cols)}</div>')


def render_command_palette() -> None:
    pages = allowed_pages()
    if not pages:
        return
    st.html('<div class="es-command-strip"><span class="es-small">COMMAND PALETTE / GLOBAL SEARCH</span></div>')
    query = st.text_input("Search pages, projects, jobs, TM, or settings", key="global_omnibar_query", placeholder="Type to search...", label_visibility="collapsed")
    selected = st.selectbox("Jump to workspace", pages, index=pages.index(st.session_state.page) if st.session_state.page in pages else 0, key="command_palette_page", label_visibility="collapsed")
    if selected != st.session_state.page:
        open_page(selected)
    if query:
        q = query.lower().strip()
        results: List[Tuple[str, str, str]] = []
        for page in pages:
            if q in page.lower():
                results.append(("PAGE", page, page_link(page)))
        for project in st.session_state.get("projects", []):
            label = safe_text(project.get("project") or project.get("workspace") or project.get("client"))
            if q in label.lower():
                results.append(("PROJECT", label, page_link("Projects")))
        for job in st.session_state.get("jobs", []):
            label = f"{safe_text(job.get('type', 'Job'))} {safe_text(job.get('language', ''))} {safe_text(job.get('status', ''))}"
            if q in label.lower():
                results.append(("JOB", label, page_link("Jobs")))
        for tm in st.session_state.get("tm", [])[:200]:
            label = f"{safe_text(tm.get('source', ''))} -> {safe_text(tm.get('target', ''))}"
            if q in label.lower():
                results.append(("TM", label[:90], page_link("Memory & Rules")))
        result_html = "".join(
            f'<a class="es-search-result" href="{href}" target="_self"><span class="es-code-chip">{escape(kind)}</span><span>{escape(label)}</span></a>'
            for kind, label, href in results[:6]
        ) or '<div class="es-small">No matches yet.</div>'
        st.html(f'<div class="es-global-search">{result_html}</div>')


def highlight_localization_text(text: str) -> str:
    output = escape(safe_text(text))
    for term in sorted([safe_text(t) for t in st.session_state.get("dnt", []) if safe_text(t)], key=len, reverse=True):
        output = re.sub(re.escape(escape(term)), lambda m: f'<span class="es-mark-dnt">{m.group(0)}</span>', output, flags=re.I)
    output = re.sub(r"(\{\{[^{}]+\}\}|\{[^{}]+\}|&lt;[^&]+&gt;|%[0-9$.\-+]*[sdif]|https?://\S+|www\.\S+)", r'<span class="es-mark-ph">\1</span>', output)
    return f'<div class="es-highlight">{output}</div>'


def inline_diff_html(before: str, after: str) -> str:
    before_words = safe_text(before).split()
    after_words = safe_text(after).split()
    diff = difflib.ndiff(before_words, after_words)
    parts = []
    for item in diff:
        marker = item[:2]
        word = escape(item[2:])
        if marker == "- ":
            parts.append(f'<span class="es-diff-del">{word}</span>')
        elif marker == "+ ":
            parts.append(f'<span class="es-diff-add">{word}</span>')
        elif marker == "  ":
            parts.append(word)
    return f'<div class="es-highlight">{" ".join(parts)}</div>'


def render_waveform_preview(rows: List[Dict[str, Any]], current_idx: int = 0) -> None:
    bars = []
    for idx in range(40):
        source_idx = min(len(rows) - 1, int(idx / 40 * max(len(rows), 1))) if rows else 0
        row = rows[source_idx] if rows else {}
        text_len = len(safe_text(row.get("target", "")) or safe_text(row.get("source", "")))
        height = 18 + ((text_len + idx * 7) % 52)
        opacity = "1" if source_idx == current_idx else ".46"
        bars.append(f'<span class="es-wavebar" style="height:{height}px;opacity:{opacity};"></span>')
    st.markdown(f'<div class="es-waveform">{"".join(bars)}</div>', unsafe_allow_html=True)


def render_segment_timeline(rows: List[Dict[str, Any]], current_idx: int = 0) -> None:
    if not rows:
        return
    max_end = max(float(r.get("end", 0) or 0) for r in rows) or 1.0
    bars = []
    for idx, row in enumerate(rows[:80]):
        start = max(0.0, float(row.get("start", 0) or 0))
        end = max(start + 0.1, float(row.get("end", start + 0.1) or start + 0.1))
        left = min(98.0, (start / max_end) * 100)
        width = max(1.5, min(100 - left, ((end - start) / max_end) * 100))
        cls = "es-timebar current" if idx == current_idx else "es-timebar"
        bars.append(f'<span class="{cls}" style="left:{left:.2f}%;width:{width:.2f}%;" title="Segment {idx + 1}"></span>')
    st.markdown(f'<div class="es-timeline">{"".join(bars)}</div>', unsafe_allow_html=True)


def render_lqa_visuals(records: List[Dict[str, Any]], summary: Dict[str, Any]) -> None:
    category_counts = summary.get("category_counts", {})
    axes = ERROR_CATEGORIES
    totals = [sum(int(category_counts.get(cat, {}).get(sev, 0) or 0) for sev in ERROR_SEVERITIES) for cat in axes]
    max_total = max(totals) or 1
    center = 50
    points = []
    axis_lines = []
    labels = []
    for idx, (cat, total) in enumerate(zip(axes, totals)):
        angle = -math.pi / 2 + idx * (2 * math.pi / len(axes))
        radius = 12 + (total / max_total) * 34
        x = center + math.cos(angle) * radius
        y = center + math.sin(angle) * radius
        points.append(f"{x:.1f},{y:.1f}")
        ax = center + math.cos(angle) * 44
        ay = center + math.sin(angle) * 44
        lx = center + math.cos(angle) * 49
        ly = center + math.sin(angle) * 49
        axis_lines.append(f'<line x1="50" y1="50" x2="{ax:.1f}" y2="{ay:.1f}" stroke="rgba(168,176,214,.35)" stroke-width="1" />')
        labels.append(f'<text x="{lx:.1f}" y="{ly:.1f}" fill="#c2c9e9" font-size="5" text-anchor="middle">{escape(cat.split()[0])}</text>')
    radar = f"""
    <svg class="es-radar" viewBox="0 0 100 100">
      <polygon points="50,8 90,36 76,86 24,86 10,36" fill="none" stroke="rgba(84,105,180,.32)" stroke-width="1" />
      {''.join(axis_lines)}
      <polygon points="{' '.join(points)}" fill="rgba(52,189,246,.24)" stroke="#34bdf6" stroke-width="2" />
      {''.join(labels)}
    </svg>
    """
    heat_cells = []
    changed_flags = [r.get("Changed") == "Yes" for r in records[:120]]
    for changed in changed_flags or [False]:
        color = "rgba(255,75,51,.72)" if changed else "rgba(0,217,133,.18)"
        heat_cells.append(f'<span class="es-heat-cell" style="background:{color};"></span>')
    st.markdown(
        f"""
        <div class="es-bento-card">
          <div class="es-radar-wrap">
            <div>{radar}</div>
            <div>
              <div class="es-metric-label">LQA Radar</div>
              <div class="es-small">Error distribution across Accuracy, Readability, Style, Grammar, and Country Standards.</div>
              <div class="es-heatmap" style="margin-top:14px;">{''.join(heat_cells)}</div>
              <div class="es-small" style="margin-top:8px;">Heatmap: red cells are changed/error rows.</div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def safe_text(x: Any) -> str:
    if x is None:
        return ""
    return str(x).replace("\u00A0", " ").strip()


def plan_record(name: str) -> Dict[str, Any]:
    requested = safe_text(name).lower()
    return next((plan for plan in PLAN_CATALOG if plan["name"].lower() == requested), PLAN_CATALOG[0])


def format_money(amount: Any, currency: str = "INR", decimals: bool = False) -> str:
    try:
        value = float(amount or 0)
    except Exception:
        value = 0.0
    precision = 2 if decimals and value != int(value) else 0
    return f"{safe_text(currency or 'INR')} {value:,.{precision}f}"


def configured_trial_days() -> int:
    try:
        return max(1, int(secret("ERRORSWEEP_TRIAL_DAYS", safe_text(plan_record("Trial").get("trial_days", 14)))))
    except Exception:
        return 14


def plan_env_slug(plan_name: str) -> str:
    return re.sub(r"[^A-Z0-9]+", "_", safe_text(plan_name).upper()).strip("_")


def payment_link_for_plan(plan_name: str, billing_cycle: str = "monthly") -> str:
    """Resolve a plan-specific hosted payment link from secrets/env."""
    slug = plan_env_slug(plan_name)
    cycle = "ANNUAL" if safe_text(billing_cycle).lower().startswith("annual") else "MONTHLY"
    for key in (
        f"ERRORSWEEP_PAYMENT_LINK_{slug}_{cycle}",
        f"ERRORSWEEP_PAYMENT_LINK_{slug}",
        "ERRORSWEEP_PAYMENT_LINK",
    ):
        value = secret(key, "").strip()
        if value:
            return value
    return ""


def monthly_mandate_link_for_plan(plan_name: str, billing_cycle: str = "monthly") -> str:
    """Resolve a recurring card/UPI monthly mandate link for checkout."""
    slug = plan_env_slug(plan_name or "Pro")
    cycle = "ANNUAL" if safe_text(billing_cycle).lower().startswith("annual") else "MONTHLY"
    for key in (
        f"ERRORSWEEP_MONTHLY_MANDATE_LINK_{slug}",
        f"ERRORSWEEP_CARD_UPI_MANDATE_LINK_{slug}",
        f"ERRORSWEEP_TRIAL_MANDATE_LINK_{slug}_{cycle}",
        f"ERRORSWEEP_TRIAL_MANDATE_LINK_{slug}",
        "ERRORSWEEP_MONTHLY_MANDATE_LINK",
        "ERRORSWEEP_CARD_UPI_MANDATE_LINK",
        "ERRORSWEEP_TRIAL_MANDATE_LINK",
        "ERRORSWEEP_UPI_MANDATE_LINK",
    ):
        value = secret(key, "").strip()
        if value:
            return value
    return payment_link_for_plan(plan_name, billing_cycle)


def trial_mandate_link_for_plan(post_trial_plan: str, billing_cycle: str = "monthly") -> str:
    """Compatibility wrapper for the free-trial post-trial monthly mandate."""
    return monthly_mandate_link_for_plan(post_trial_plan, billing_cycle)


def sanitize_payment_link(url: str) -> str:
    value = safe_text(url)
    if not value:
        return ""
    if re.match(r"^https?://", value, flags=re.I):
        return value
    return ""


def render_pricing_graphic(active_plan: str = "Trial", billing_cycle: str = "monthly") -> None:
    cards = []
    accents = [
        ("#00d985", "#34bdf6", "TRIAL"),
        ("#34bdf6", "#8b5cf6", "PRO"),
        ("#f59e0b", "#34bdf6", "TEAM"),
        ("#8b5cf6", "#00d985", "SSO"),
    ]
    for idx, plan in enumerate(PLAN_CATALOG):
        accent_a, accent_b, code = accents[idx % len(accents)]
        is_active = safe_text(plan.get("name")).lower() == safe_text(active_plan).lower()
        cycle = "annual" if safe_text(billing_cycle).lower().startswith("annual") else "monthly"
        if plan["name"] == "Enterprise":
            price = "Custom"
        elif plan["name"] == "Trial":
            price = f"INR 0 / {configured_trial_days()} days"
        else:
            monthly_equivalent = plan.get("monthly", 0)
            price = f"{format_money(monthly_equivalent, plan.get('currency', 'INR'))}/mo"
        mandate_note = (
            f"<span><b>Card/UPI</b> monthly mandate</span><span><b>Cancel</b> before day {configured_trial_days()}</span>"
            if plan["name"] == "Trial" else
            "<span><b>Card/UPI</b> monthly mandate</span><span><b>Recurring</b> monthly deduction</span>"
        )
        payment_link = monthly_mandate_link_for_plan("Pro" if plan["name"] == "Trial" else plan["name"], "monthly")
        cta = "Mandate link ready" if payment_link else "Mandate link not configured"
        link_html = (
            f'<a class="es-pricing-link" href="{escape(payment_link)}" target="_blank" rel="noopener">Open mandate link</a>'
            if payment_link else
            '<span class="es-pricing-link muted">Add link below</span>'
        )
        active_badge = '<span class="es-pricing-active">Current</span>' if is_active else ""
        cards.append(
            f"""
            <article class="es-pricing-card {'active' if is_active else ''}" style="--accent-a:{accent_a}; --accent-b:{accent_b};">
              <div class="es-pricing-orb">
                <svg viewBox="0 0 120 120" role="img" aria-label="{escape(plan['name'])} plan graphic">
                  <defs>
                    <linearGradient id="planGrad{idx}" x1="0%" y1="0%" x2="100%" y2="100%">
                      <stop offset="0%" stop-color="{accent_a}" />
                      <stop offset="100%" stop-color="{accent_b}" />
                    </linearGradient>
                  </defs>
                  <circle cx="60" cy="60" r="52" fill="none" stroke="rgba(255,255,255,.09)" stroke-width="10" />
                  <circle cx="60" cy="60" r="52" fill="none" stroke="url(#planGrad{idx})" stroke-width="10" stroke-linecap="round" stroke-dasharray="{62 + idx * 28} 330" transform="rotate(-90 60 60)" />
                  <circle cx="60" cy="60" r="34" fill="rgba(8,12,25,.74)" stroke="rgba(255,255,255,.10)" />
                  <text x="60" y="65" text-anchor="middle" font-size="16" font-weight="900" fill="#f8fbff">{code}</text>
                </svg>
              </div>
              <div class="es-pricing-top">
                <span>{escape(safe_text(plan['label']).upper())}</span>
                {active_badge}
              </div>
              <h3>{escape(plan['name'])}</h3>
              <div class="es-pricing-price">{escape(price)}</div>
              <p>{escape(plan['description'])}</p>
              <div class="es-pricing-stats">
                <span><b>{int(plan['seats']):,}</b> seats</span>
                <span><b>{int(plan['segments']):,}</b> segments</span>
                <span><b>{int(plan['characters']):,}</b> chars</span>
                {mandate_note}
              </div>
              <div class="es-pricing-footer">
                <span>{cta}</span>
                {link_html}
              </div>
            </article>
            """
        )

    st.html(
        f"""
        <style>
          .es-pricing-scene {{
            display:grid;
            grid-template-columns: repeat(4, minmax(0, 1fr));
            gap:16px;
            margin: 16px 0 22px;
          }}
          .es-pricing-card {{
            position:relative;
            overflow:hidden;
            min-height: 430px;
            border:1px solid rgba(91,113,190,.32);
            border-radius:16px;
            padding:22px 20px 18px;
            background:
              radial-gradient(circle at 80% 8%, color-mix(in srgb, var(--accent-a) 18%, transparent), transparent 34%),
              linear-gradient(160deg, rgba(18,24,48,.92), rgba(8,10,22,.94) 58%, rgba(16,8,30,.88));
            box-shadow: 0 28px 80px rgba(0,0,0,.34), inset 0 1px 0 rgba(255,255,255,.055);
          }}
          .es-pricing-card::before {{
            content:"";
            position:absolute;
            inset:0;
            background-image:
              linear-gradient(rgba(255,255,255,.035) 1px, transparent 1px),
              linear-gradient(90deg, rgba(255,255,255,.035) 1px, transparent 1px);
            background-size: 28px 28px;
            mask-image: linear-gradient(180deg, black, transparent 78%);
            pointer-events:none;
          }}
          .es-pricing-card.active {{
            border-color: color-mix(in srgb, var(--accent-a) 72%, white 8%);
            box-shadow: 0 32px 94px color-mix(in srgb, var(--accent-a) 18%, transparent), inset 0 1px 0 rgba(255,255,255,.08);
          }}
          .es-pricing-orb {{
            width: 132px;
            height: 132px;
            margin-bottom: 18px;
            filter: drop-shadow(0 18px 26px color-mix(in srgb, var(--accent-a) 22%, transparent));
          }}
          .es-pricing-top,
          .es-pricing-footer,
          .es-pricing-stats {{
            position:relative;
            z-index:1;
          }}
          .es-pricing-top {{
            display:flex;
            align-items:center;
            justify-content:space-between;
            gap:10px;
            color:#a8b0d6;
            font-family:"Space Mono", monospace;
            font-size:11px;
            font-weight:900;
            letter-spacing:.06em;
          }}
          .es-pricing-active {{
            color:#07101c;
            background: linear-gradient(90deg, var(--accent-a), var(--accent-b));
            border-radius:999px;
            padding:4px 8px;
            letter-spacing:0;
          }}
          .es-pricing-card h3 {{
            position:relative;
            z-index:1;
            color:#f8fbff;
            font-size: 34px;
            line-height:1;
            margin: 18px 0 10px;
          }}
          .es-pricing-price {{
            position:relative;
            z-index:1;
            color:#ffffff;
            font-size: 24px;
            font-weight: 900;
            margin-bottom: 14px;
          }}
          .es-pricing-card p {{
            position:relative;
            z-index:1;
            color:#b8c0df;
            min-height: 76px;
            line-height: 1.55;
            margin: 0 0 16px;
          }}
          .es-pricing-stats {{
            display:grid;
            gap:8px;
            margin-top: 16px;
          }}
          .es-pricing-stats span {{
            display:flex;
            justify-content:space-between;
            border:1px solid rgba(255,255,255,.08);
            background:rgba(255,255,255,.04);
            border-radius:10px;
            padding:9px 10px;
            color:#dce6ff;
          }}
          .es-pricing-footer {{
            display:flex;
            align-items:center;
            justify-content:space-between;
            gap:12px;
            margin-top: 18px;
            color:#9da8d6;
            font-size:12px;
          }}
          .es-pricing-link {{
            color:#07101c !important;
            background: linear-gradient(90deg, var(--accent-a), var(--accent-b));
            border-radius:8px;
            padding:8px 10px;
            font-weight:900;
            text-decoration:none !important;
            white-space:nowrap;
          }}
          .es-pricing-link.muted {{
            color:#c7d2fe !important;
            background:rgba(255,255,255,.07);
          }}
          @media (max-width: 1250px) {{
            .es-pricing-scene {{ grid-template-columns: repeat(2, minmax(0, 1fr)); }}
          }}
          @media (max-width: 760px) {{
            .es-pricing-scene {{ grid-template-columns: 1fr; }}
          }}
        </style>
        <div class="es-pricing-scene">
          {''.join(cards)}
        </div>
        """
    )


def billing_provider_label() -> str:
    provider = secret("ERRORSWEEP_BILLING_PROVIDER", "").strip().lower()
    if provider:
        return provider
    if secret("STRIPE_SECRET_KEY", ""):
        return "stripe"
    if secret("RAZORPAY_KEY_ID", "") and secret("RAZORPAY_KEY_SECRET", ""):
        return "razorpay"
    return "manual"


def billing_provider_ready(provider: Optional[str] = None) -> bool:
    provider = safe_text(provider or billing_provider_label()).lower()
    if provider == "stripe":
        return bool(secret("STRIPE_SECRET_KEY", ""))
    if provider == "razorpay":
        return bool(secret("RAZORPAY_KEY_ID", "") and secret("RAZORPAY_KEY_SECRET", ""))
    return False


def workspace_subscription(workspace: str = "") -> Dict[str, Any]:
    workspace = safe_text(workspace or (current_user() or {}).get("workspace") or "Demo Workspace")
    subscriptions = [
        item for item in st.session_state.get("subscriptions", [])
        if safe_text(item.get("workspace")) == workspace
    ]
    active = next((item for item in subscriptions if safe_text(item.get("status")).lower() == "active"), None)
    if active:
        return active
    if subscriptions:
        return subscriptions[0]
    plan = plan_record("Trial")
    return {
        "workspace": workspace,
        "plan": plan["name"],
        "status": "Active",
        "billing_cycle": "monthly",
        "currency": plan["currency"],
        "base_amount": plan["monthly"],
        "included_segments": plan["segments"],
        "included_characters": plan["characters"],
        "included_seats": plan["seats"],
        "provider": "trial",
    }


def _usage_int(value: Any) -> int:
    try:
        return max(0, int(float(value or 0)))
    except Exception:
        return 0


def workspace_usage_estimate(rows: List[Dict[str, Any]]) -> Dict[str, int]:
    segments = len(rows or [])
    characters = 0
    for row in rows or []:
        characters += len(safe_text(row.get("source", "")))
        characters += len(safe_text(row.get("target", "")))
    return {"segments": segments, "characters": characters}


def workspace_usage_totals(workspace: str = "") -> Dict[str, int]:
    workspace = safe_text(workspace or (current_user() or {}).get("workspace") or "Demo Workspace")
    totals = {"segments": 0, "characters": 0}
    for item in st.session_state.get("ai_usage_events", []):
        item_workspace = safe_text(item.get("workspace") or workspace)
        if item_workspace != workspace:
            continue
        metadata = item.get("metadata") or item.get("metadata_json") or {}
        if isinstance(metadata, str):
            try:
                metadata = json.loads(metadata)
            except Exception:
                metadata = {}
        billable = item.get("billable", metadata.get("billable") if isinstance(metadata, dict) else None)
        if billable is False or safe_text(billable).lower() in {"false", "0", "no"}:
            continue
        totals["segments"] += _usage_int(item.get("segments") or item.get("row_count"))
        totals["characters"] += _usage_int(item.get("characters"))
    return totals


def workspace_usage_allowance(workspace: str = "") -> Dict[str, Any]:
    subscription = workspace_subscription(workspace)
    plan = plan_record(subscription.get("plan", "Trial"))
    return {
        "subscription": subscription,
        "plan": plan,
        "segments": _usage_int(subscription.get("included_segments") or plan.get("segments")),
        "characters": _usage_int(subscription.get("included_characters") or plan.get("characters")),
        "seats": _usage_int(subscription.get("included_seats") or plan.get("seats")),
    }


def workspace_user_count(workspace: str = "", include_pending: bool = True) -> int:
    workspace = safe_text(workspace or (current_user() or {}).get("workspace") or "Demo Workspace")
    if safe_text(workspace).lower() == "platform":
        return 0
    active_statuses = {"active", "invited"} if include_pending else {"active"}
    emails = set()
    for item in st.session_state.get("users", []):
        if safe_text(item.get("workspace")) != workspace:
            continue
        status = safe_text(item.get("status") or "Active").lower()
        if status not in active_statuses:
            continue
        email = safe_text(item.get("email")).lower()
        if email:
            emails.add(email)
    return len(emails)


def workspace_seat_state(workspace: str = "") -> Dict[str, Any]:
    workspace = safe_text(workspace or (current_user() or {}).get("workspace") or "Demo Workspace")
    allowance = workspace_usage_allowance(workspace)
    used = workspace_user_count(workspace)
    return {
        "workspace": workspace,
        "subscription": allowance["subscription"],
        "plan": allowance["plan"],
        "used": used,
        "limit": allowance["seats"],
        "available": max(0, allowance["seats"] - used),
    }


def check_workspace_seat_allowance(workspace: str, email: str, status: str = "Active") -> Tuple[bool, str, Dict[str, Any]]:
    workspace = safe_text(workspace or (current_user() or {}).get("workspace") or "Demo Workspace")
    email_key = safe_text(email).lower()
    state = workspace_seat_state(workspace)
    duplicate = next(
        (
            item for item in st.session_state.get("users", [])
            if safe_text(item.get("workspace")) == workspace
            and safe_text(item.get("email")).lower() == email_key
            and safe_text(item.get("status") or "Active").lower() not in {"deleted", "removed"}
        ),
        None,
    )
    if duplicate:
        return False, f"{email} is already listed in the {workspace} workspace.", state

    subscription_status = safe_text(state["subscription"].get("status") or "Active").lower()
    if subscription_status in {"cancelled", "canceled", "expired", "past_due", "unpaid", "inactive"}:
        return False, f"Cannot add users because the {workspace} subscription is {state['subscription'].get('status', subscription_status)}. Renew or activate the plan from Billing.", state

    counted_status = safe_text(status or "Active").lower() in {"active", "invited"}
    if safe_text(workspace).lower() != "platform" and counted_status and state["used"] >= max(1, state["limit"]):
        return False, f"{workspace} has reached the {state['plan']['name']} seat limit ({state['used']:,}/{state['limit']:,}). Upgrade the plan or suspend an inactive user before adding another active seat.", state
    return True, "", state


def check_workspace_usage_allowance(rows: List[Dict[str, Any]], purpose: str, workspace: str = "") -> Tuple[bool, str, Dict[str, Any]]:
    workspace = safe_text(workspace or (current_user() or {}).get("workspace") or "Demo Workspace")
    usage = workspace_usage_totals(workspace)
    allowance = workspace_usage_allowance(workspace)
    estimate = workspace_usage_estimate(rows)
    subscription = allowance["subscription"]
    plan_name = safe_text(subscription.get("plan") or allowance["plan"].get("name") or "Trial")
    status = safe_text(subscription.get("status") or "Active").lower()
    details = {
        "workspace": workspace,
        "plan": plan_name,
        "status": status,
        "used_segments": usage["segments"],
        "used_characters": usage["characters"],
        "requested_segments": estimate["segments"],
        "requested_characters": estimate["characters"],
        "projected_segments": usage["segments"] + estimate["segments"],
        "projected_characters": usage["characters"] + estimate["characters"],
        "segment_limit": allowance["segments"],
        "character_limit": allowance["characters"],
    }
    if status in {"cancelled", "canceled", "expired", "past_due", "unpaid", "inactive"}:
        return False, f"{purpose} is blocked because the {workspace} subscription is {subscription.get('status', status)}. Activate or renew the plan from Billing.", details

    over_segments = details["projected_segments"] > max(1, allowance["segments"])
    over_characters = details["projected_characters"] > max(1, allowance["characters"])
    if over_segments or over_characters:
        limits = (
            f"{details['projected_segments']:,}/{allowance['segments']:,} segments, "
            f"{details['projected_characters']:,}/{allowance['characters']:,} characters"
        )
        return False, f"{purpose} would exceed the {plan_name} plan allowance ({limits}). Upgrade the plan or reduce the upload size before running.", details

    segment_ratio = details["projected_segments"] / max(1, allowance["segments"])
    character_ratio = details["projected_characters"] / max(1, allowance["characters"])
    if max(segment_ratio, character_ratio) >= 0.85:
        return True, f"{purpose} can run, but this workspace will be above 85% of the {plan_name} allowance after this job.", details
    return True, "", details


def record_billable_workflow_usage(purpose: str, rows: List[Dict[str, Any]], provider: str = "errorsweep", model: str = "workflow") -> None:
    estimate = workspace_usage_estimate(rows)
    log_ai_usage_event(
        {
            "provider": provider,
            "model": model,
            "managed": True,
            "requests": 1,
            "success": True,
            "characters": estimate["characters"],
            "billable": True,
            "usage_kind": "workflow",
        },
        purpose,
        estimate["segments"],
    )


def create_checkout_intent(
    plan_name: str,
    billing_cycle: str,
    payment_link: str = "",
    post_trial_plan: str = "",
) -> Dict[str, Any]:
    user = current_user() or {}
    workspace = safe_text(user.get("workspace") or "Demo Workspace")
    plan = plan_record(plan_name)
    is_trial = safe_text(plan["name"]).lower() == "trial"
    cycle = "monthly"
    post_plan = plan_record(post_trial_plan or "Pro") if is_trial else None
    amount = 0.0 if is_trial else float(plan["monthly"])
    mandate_amount = float((post_plan or plan)["monthly"])
    trial_days = configured_trial_days() if is_trial else 0
    trial_ends_at = (datetime.now(timezone.utc) + timedelta(days=trial_days)).isoformat() if is_trial else ""
    provider = billing_provider_label()
    checkout_base = secret("ERRORSWEEP_CHECKOUT_BASE_URL", "").strip().rstrip("/")
    configured_payment_link = sanitize_payment_link(payment_link) or sanitize_payment_link(
        monthly_mandate_link_for_plan(post_plan["name"], cycle) if is_trial and post_plan else monthly_mandate_link_for_plan(plan["name"], cycle)
    )
    checkout_url = configured_payment_link or (f"{checkout_base}?plan={quote(plan['name'])}&cycle={quote(cycle)}&workspace={quote(workspace)}" if checkout_base else "")
    if is_trial:
        status = "trial_mandate_pending" if checkout_url else "trial_mandate_link_missing"
    else:
        status = "monthly_mandate_pending" if checkout_url else "monthly_mandate_link_missing"
    intent = persist_saas_record("checkout_sessions", {
        "id": uuid.uuid4().hex,
        "workspace": workspace,
        "user_email": user.get("email", ""),
        "plan": plan["name"],
        "billing_cycle": cycle,
        "currency": (post_plan or plan)["currency"],
        "amount": amount,
        "provider": provider,
        "status": status,
        "checkout_url": checkout_url,
        "provider_session_id": "",
        "metadata_json": {
            "included_segments": plan["segments"],
            "included_characters": plan["characters"],
            "included_seats": plan["seats"],
            "description": plan["description"],
            "trial_days": trial_days,
            "trial_ends_at": trial_ends_at,
            "cancel_anytime_until": trial_ends_at,
            "mandate_required": True,
            "mandate_type": "card_or_upi_monthly",
            "post_trial_plan": post_plan["name"] if post_plan else "",
            "post_trial_billing_cycle": cycle if is_trial else "",
            "monthly_mandate_amount": mandate_amount,
            "post_trial_monthly_amount": mandate_amount if is_trial else 0,
            "post_trial_currency": (post_plan or plan)["currency"],
        },
        "created_at": now_stamp(),
        "updated_at": now_stamp(),
    })
    st.session_state.setdefault("checkout_sessions", [])
    st.session_state.checkout_sessions.insert(0, intent)
    trim_session_list("checkout_sessions")
    queue_email_notification(
        user.get("email", ""),
        "ErrorSweep trial mandate created" if is_trial else "ErrorSweep monthly mandate created",
        (
            f"Trial setup started for {workspace}. A card or UPI mandate is required for {post_plan['name']} "
            f"at {format_money(mandate_amount, post_plan['currency'])}/month after the {trial_days}-day trial. "
            f"You can cancel anytime before {trial_ends_at}."
            if is_trial and post_plan else
            (
                f"Monthly card/UPI mandate created for {workspace}: {plan['name']} custom monthly amount."
                if plan["name"] == "Enterprise" else
                f"Monthly card/UPI mandate created for {workspace}: {plan['name']} at {format_money(amount, plan['currency'])}/month."
            )
        ),
        "billing.checkout_intent",
        metadata={
            "plan": plan["name"],
            "billing_cycle": cycle,
            "amount": amount,
            "provider": provider,
            "status": status,
            "post_trial_plan": post_plan["name"] if post_plan else "",
            "trial_days": trial_days,
        },
        workspace=workspace,
    )
    audit_detail = f"{workspace}: trial -> {post_plan['name']} monthly mandate via {provider}" if is_trial and post_plan else f"{workspace}: {plan['name']} monthly mandate via {provider}"
    add_audit("Checkout intent created", audit_detail)
    return intent


def activate_subscription(plan_name: str, billing_cycle: str, provider: str = "") -> Dict[str, Any]:
    user = current_user() or {}
    workspace = safe_text(user.get("workspace") or "Demo Workspace")
    plan = plan_record(plan_name)
    cycle = "monthly"
    subscription = persist_saas_record("subscriptions", {
        "workspace": workspace,
        "user_email": user.get("email", ""),
        "plan": plan["name"],
        "status": "Active",
        "billing_cycle": cycle,
        "currency": plan["currency"],
        "base_amount": float(plan["annual" if cycle == "annual" else "monthly"]),
        "included_segments": plan["segments"],
        "included_characters": plan["characters"],
        "included_seats": plan["seats"],
        "provider": safe_text(provider or billing_provider_label()),
        "provider_customer_id": "",
        "provider_subscription_id": "",
        "current_period_start": datetime.now(timezone.utc).isoformat(),
        "current_period_end": None,
        "cancel_at_period_end": False,
        "cancelled_at": None,
        "cancellation_reason": "",
        "metadata_json": {},
    })
    st.session_state.setdefault("subscriptions", [])
    st.session_state.subscriptions = [
        item for item in st.session_state.subscriptions
        if safe_text(item.get("workspace")) != workspace
    ]
    st.session_state.subscriptions.insert(0, subscription)
    trim_session_list("subscriptions")
    for workspace_record in st.session_state.get("workspaces", []):
        if safe_text(workspace_record.get("workspace")) == workspace:
            workspace_record["plan"] = plan["name"]
            break
    add_audit("Subscription activated", f"{workspace}: {plan['name']} {cycle}")
    queue_email_notification(
        user.get("email", ""),
        "ErrorSweep subscription updated",
        f"Your workspace '{workspace}' is now on the {plan['name']} plan.",
        "billing.subscription_updated",
        metadata={"plan": plan["name"], "billing_cycle": cycle},
        workspace=workspace,
    )
    return subscription


def _replace_session_record(collection_key: str, updated: Dict[str, Any]) -> None:
    rows = st.session_state.setdefault(collection_key, [])
    updated_id = safe_text(updated.get("id"))
    for idx, item in enumerate(rows):
        if updated_id and safe_text(item.get("id")) == updated_id:
            rows[idx] = updated
            return
    rows.insert(0, updated)
    trim_session_list(collection_key, SESSION_COLLECTION_LIMITS.get(collection_key, SESSION_HISTORY_LIMIT))


def cancel_current_billing(reason: str = "") -> Tuple[str, Dict[str, Any]]:
    """Cancel the current subscription or pending trial mandate for the workspace."""
    user = current_user() or {}
    workspace = safe_text(user.get("workspace") or "Demo Workspace")
    cancelled_at = datetime.now(timezone.utc).isoformat()
    reason = safe_text(reason) or "Cancelled by user"
    subscriptions = [
        item for item in st.session_state.get("subscriptions", [])
        if safe_text(item.get("workspace")) == workspace
    ]
    cancellable_subscription = next(
        (
            item for item in subscriptions
            if safe_text(item.get("status")).lower() not in {"cancelled", "canceled", "expired"}
        ),
        None,
    )
    if cancellable_subscription:
        updated = dict(cancellable_subscription)
        metadata = updated.get("metadata_json") if isinstance(updated.get("metadata_json"), dict) else {}
        updated.update({
            "status": "Cancelled",
            "cancel_at_period_end": True,
            "cancelled_at": cancelled_at,
            "current_period_end": updated.get("current_period_end") or cancelled_at,
            "cancellation_reason": reason,
            "metadata_json": {
                **metadata,
                "cancelled_by": user.get("email", ""),
                "cancelled_at": cancelled_at,
                "cancellation_source": "billing_page",
            },
        })
        persisted = persist_saas_record("subscriptions", updated)
        _replace_session_record("subscriptions", persisted)
        queue_email_notification(
            user.get("email", ""),
            "ErrorSweep subscription cancelled",
            f"Your ErrorSweep {safe_text(updated.get('plan'))} subscription for '{workspace}' was cancelled. Reason: {reason}.",
            "billing.subscription_cancelled",
            metadata={"workspace": workspace, "plan": updated.get("plan"), "cancelled_at": cancelled_at, "reason": reason},
            workspace=workspace,
        )
        add_audit("Subscription cancelled", f"{workspace}: {updated.get('plan')} - {reason}")
        return "subscription", persisted

    checkout_rows = [
        item for item in st.session_state.get("checkout_sessions", [])
        if safe_text(item.get("workspace")) == workspace
        and safe_text(item.get("status")).lower() not in {"cancelled", "canceled", "expired", "paid", "completed"}
    ]
    if checkout_rows:
        updated_checkout = dict(checkout_rows[0])
        metadata = updated_checkout.get("metadata_json") if isinstance(updated_checkout.get("metadata_json"), dict) else {}
        updated_checkout["status"] = "cancelled"
        updated_checkout["metadata_json"] = {
            **metadata,
            "cancelled_by": user.get("email", ""),
            "cancelled_at": cancelled_at,
            "cancellation_reason": reason,
        }
        persisted = persist_saas_record("checkout_sessions", updated_checkout)
        _replace_session_record("checkout_sessions", persisted)
        queue_email_notification(
            user.get("email", ""),
            "ErrorSweep trial mandate cancelled",
            f"Your pending ErrorSweep trial/payment mandate for '{workspace}' was cancelled before activation. Reason: {reason}.",
            "billing.trial_mandate_cancelled",
            metadata={"workspace": workspace, "plan": updated_checkout.get("plan"), "cancelled_at": cancelled_at, "reason": reason},
            workspace=workspace,
        )
        add_audit("Trial/payment mandate cancelled", f"{workspace}: {updated_checkout.get('plan')} - {reason}")
        return "checkout", persisted

    return "", {}


def billing_webhook_secret_for_provider(provider: str) -> str:
    provider_key = safe_text(provider).upper()
    candidates = [
        f"ERRORSWEEP_{provider_key}_WEBHOOK_SECRET",
        f"{provider_key}_WEBHOOK_SECRET",
        "ERRORSWEEP_BILLING_WEBHOOK_SECRET",
    ]
    for key in candidates:
        value = secret(key, "").strip()
        if value:
            return value
    return ""


def _billing_event_workspace(normalized: Dict[str, Any]) -> str:
    return safe_text(normalized.get("workspace")) or safe_text((current_user() or {}).get("workspace")) or "Demo Workspace"


def _find_checkout_for_billing_event(normalized: Dict[str, Any], workspace: str) -> Optional[Dict[str, Any]]:
    checkout_id = safe_text(normalized.get("checkout_id"))
    provider_subscription_id = safe_text(normalized.get("provider_subscription_id"))
    provider_payment_id = safe_text(normalized.get("provider_payment_id"))
    plan_name = safe_text(normalized.get("plan"))
    terminal = {"cancelled", "canceled", "expired", "paid", "completed", "failed"}
    candidates = [
        item for item in st.session_state.get("checkout_sessions", [])
        if safe_text(item.get("workspace")) == workspace
    ]
    for item in candidates:
        metadata = item.get("metadata_json") if isinstance(item.get("metadata_json"), dict) else {}
        if checkout_id and checkout_id in {safe_text(item.get("id")), safe_text(item.get("provider_session_id")), safe_text(metadata.get("checkout_id"))}:
            return item
        if provider_subscription_id and provider_subscription_id == safe_text(metadata.get("provider_subscription_id")):
            return item
        if provider_payment_id and provider_payment_id == safe_text(metadata.get("provider_payment_id")):
            return item
    active_candidates = [
        item for item in candidates
        if safe_text(item.get("status")).lower() not in terminal
    ]
    if plan_name:
        for item in active_candidates:
            metadata = item.get("metadata_json") if isinstance(item.get("metadata_json"), dict) else {}
            if safe_text(item.get("plan")) == plan_name or safe_text(metadata.get("post_trial_plan")) == plan_name:
                return item
    return active_candidates[0] if active_candidates else None


def _activate_subscription_from_billing_event(normalized: Dict[str, Any], checkout: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    user = current_user() or {}
    workspace = _billing_event_workspace(normalized)
    checkout_metadata = checkout.get("metadata_json") if checkout and isinstance(checkout.get("metadata_json"), dict) else {}
    checkout_plan = safe_text(checkout.get("plan")) if checkout else ""
    plan_name = safe_text(normalized.get("plan")) or checkout_plan or "Pro"
    if checkout_plan == "Trial":
        plan_name = "Trial"
    plan = plan_record(plan_name)
    now = datetime.now(timezone.utc)
    is_trial = plan["name"] == "Trial"
    trial_days = int(checkout_metadata.get("trial_days") or configured_trial_days()) if is_trial else 0
    period_end = (now + timedelta(days=trial_days if is_trial else 30)).isoformat()
    amount = 0.0 if is_trial else float(normalized.get("amount") or plan.get("monthly") or 0)
    metadata = {
        **checkout_metadata,
        "activated_from_billing_event": normalized.get("event_id", ""),
        "mandate_type": "card_or_upi_monthly",
        "provider_payment_id": normalized.get("provider_payment_id", ""),
        "provider_order_id": normalized.get("provider_order_id", ""),
        "monthly_mandate_amount": normalized.get("amount") or checkout_metadata.get("monthly_mandate_amount") or plan.get("monthly", 0),
    }
    subscription = persist_saas_record("subscriptions", {
        "workspace": workspace,
        "user_email": safe_text(normalized.get("user_email")) or user.get("email", ""),
        "plan": plan["name"],
        "status": "Active",
        "billing_cycle": "monthly",
        "currency": safe_text(normalized.get("currency")) or plan["currency"],
        "base_amount": amount,
        "included_segments": plan["segments"],
        "included_characters": plan["characters"],
        "included_seats": plan["seats"],
        "provider": safe_text(normalized.get("provider")) or billing_provider_label(),
        "provider_customer_id": safe_text(normalized.get("provider_customer_id")),
        "provider_subscription_id": safe_text(normalized.get("provider_subscription_id")),
        "current_period_start": now.isoformat(),
        "current_period_end": period_end,
        "cancel_at_period_end": False,
        "cancelled_at": None,
        "cancellation_reason": "",
        "metadata_json": metadata,
    })
    st.session_state.setdefault("subscriptions", [])
    st.session_state.subscriptions = [
        item for item in st.session_state.subscriptions
        if safe_text(item.get("workspace")) != workspace
    ]
    st.session_state.subscriptions.insert(0, subscription)
    trim_session_list("subscriptions")
    for workspace_record in st.session_state.get("workspaces", []):
        if safe_text(workspace_record.get("workspace")) == workspace:
            workspace_record["plan"] = plan["name"]
            break
    queue_email_notification(
        safe_text(normalized.get("user_email")) or user.get("email", ""),
        "ErrorSweep mandate activated",
        f"Your ErrorSweep {plan['name']} monthly mandate for '{workspace}' is active.",
        "billing.mandate_activated",
        metadata={"plan": plan["name"], "workspace": workspace, "event_id": normalized.get("event_id")},
        workspace=workspace,
    )
    add_audit("Billing mandate activated", f"{workspace}: {plan['name']} via {normalized.get('provider')}")
    return subscription


def _record_payment_from_billing_event(normalized: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    amount = float(normalized.get("amount") or 0)
    if amount <= 0:
        return None
    workspace = _billing_event_workspace(normalized)
    payment = persist_saas_record("payments", {
        "date": now_stamp(),
        "workspace": workspace,
        "user": safe_text(normalized.get("user_email")),
        "plan": safe_text(normalized.get("plan")),
        "amount": amount,
        "currency": safe_text(normalized.get("currency")) or "INR",
        "status": "Recorded from webhook",
    })
    st.session_state.setdefault("payments", [])
    st.session_state.payments.insert(0, payment)
    trim_session_list("payments")
    return payment


def apply_billing_webhook_event(normalized: Dict[str, Any]) -> List[str]:
    workspace = _billing_event_workspace(normalized)
    status = safe_text(normalized.get("status")).lower()
    messages: List[str] = []
    checkout = _find_checkout_for_billing_event(normalized, workspace)
    success_statuses = {"paid", "active", "authorized", "authenticated", "captured", "complete", "succeeded"}
    failed_statuses = {"failed", "failure", "declined"}
    cancelled_statuses = {"cancelled", "canceled"}

    if checkout:
        metadata = checkout.get("metadata_json") if isinstance(checkout.get("metadata_json"), dict) else {}
        updated_checkout = dict(checkout)
        if status in success_statuses:
            updated_checkout["status"] = "mandate_active"
        elif status in failed_statuses:
            updated_checkout["status"] = "failed"
        elif status in cancelled_statuses:
            updated_checkout["status"] = "cancelled"
        else:
            updated_checkout["status"] = status or updated_checkout.get("status", "received")
        updated_checkout["provider_session_id"] = safe_text(normalized.get("provider_payment_id") or normalized.get("provider_subscription_id") or updated_checkout.get("provider_session_id"))
        updated_checkout["metadata_json"] = {
            **metadata,
            "last_billing_event_id": normalized.get("event_id", ""),
            "provider_payment_id": normalized.get("provider_payment_id", ""),
            "provider_subscription_id": normalized.get("provider_subscription_id", ""),
            "provider_order_id": normalized.get("provider_order_id", ""),
        }
        persisted_checkout = persist_saas_record("checkout_sessions", updated_checkout)
        _replace_session_record("checkout_sessions", persisted_checkout)
        messages.append(f"Checkout updated to {persisted_checkout.get('status')}.")

    if status in success_statuses:
        subscription = _activate_subscription_from_billing_event(normalized, checkout)
        messages.append(f"Subscription activated: {subscription.get('plan')}.")
        payment_plan = safe_text(normalized.get("plan")) or (safe_text(checkout.get("plan")) if checkout else "")
        payment_context = {**normalized, "plan": payment_plan}
        payment = _record_payment_from_billing_event(payment_context)
        if payment:
            messages.append(f"Payment recorded: {format_money(payment.get('amount'), payment.get('currency', 'INR'))}.")
    elif status in cancelled_statuses:
        cancelled_kind, _ = cancel_current_billing(f"Provider event {normalized.get('event_id')} marked billing cancelled.")
        messages.append(f"Cancellation applied: {cancelled_kind or 'no active record'}.")
    elif status in failed_statuses:
        add_audit("Billing webhook failed", f"{workspace}: {normalized.get('event_type')} {normalized.get('event_id')}")
        messages.append("Billing event recorded as failed.")
    else:
        messages.append("Billing event stored; no lifecycle change was applied.")
    return messages


def record_billing_webhook_event(
    provider: str,
    raw_payload: str,
    signature_header: str = "",
    webhook_secret: str = "",
    apply_updates: bool = True,
) -> Tuple[Dict[str, Any], List[str]]:
    if normalize_billing_webhook is None:
        raise RuntimeError("Billing webhook helpers are unavailable.")
    normalized = normalize_billing_webhook(provider, raw_payload)
    resolved_provider = safe_text(normalized.get("provider") or provider or "manual").lower()
    secret_value = safe_text(webhook_secret) or billing_webhook_secret_for_provider(resolved_provider)
    signature_status = "not_checked"
    if signature_header:
        if verify_billing_webhook_signature is None:
            signature_status = "unavailable"
        elif not secret_value:
            signature_status = "secret_missing"
        else:
            signature_status = "verified" if verify_billing_webhook_signature(resolved_provider, raw_payload, signature_header, secret_value) else "invalid"

    workspace = _billing_event_workspace(normalized)
    event_record = persist_saas_record("billing_events", {
        "workspace": workspace,
        "user_email": safe_text(normalized.get("user_email")) or (current_user() or {}).get("email", ""),
        "provider": resolved_provider,
        "event_id": safe_text(normalized.get("event_id")),
        "event_type": safe_text(normalized.get("event_type")),
        "status": safe_text(normalized.get("status")),
        "plan": safe_text(normalized.get("plan")),
        "amount": float(normalized.get("amount") or 0),
        "currency": safe_text(normalized.get("currency") or "INR"),
        "provider_payment_id": safe_text(normalized.get("provider_payment_id")),
        "provider_subscription_id": safe_text(normalized.get("provider_subscription_id")),
        "provider_order_id": safe_text(normalized.get("provider_order_id")),
        "provider_customer_id": safe_text(normalized.get("provider_customer_id")),
        "checkout_id": safe_text(normalized.get("checkout_id")),
        "signature_status": signature_status,
        "applied": False,
        "raw_sha256": safe_text(normalized.get("raw_sha256")),
        "metadata_json": {
            "provider_metadata": normalized.get("metadata") or {},
            "normalized": normalized,
        },
    })
    _replace_session_record("billing_events", event_record)

    messages: List[str] = []
    if apply_updates and signature_status not in {"invalid", "secret_missing"}:
        messages = apply_billing_webhook_event(normalized)
        event_record["applied"] = True
        event_record["metadata_json"] = {
            **(event_record.get("metadata_json") if isinstance(event_record.get("metadata_json"), dict) else {}),
            "applied_messages": messages,
        }
        event_record = persist_saas_record("billing_events", event_record)
        _replace_session_record("billing_events", event_record)
    elif signature_status == "invalid":
        messages = ["Signature verification failed. Event was stored but not applied."]
    elif signature_status == "secret_missing":
        messages = ["Signature header was provided, but no webhook secret was available. Event was stored but not applied."]
    else:
        messages = ["Event stored without applying billing updates."]
    return event_record, messages


def launch_readiness_rows(health: Optional[Dict[str, Any]] = None) -> List[Dict[str, str]]:
    health = health or {}
    supabase_ready = bool(health.get("supabase_configured")) and all(
        safe_text(value) == "ok" for value in (health.get("saas_tables") or {}).values()
    )
    email_provider = email_provider_label()
    email_ready = email_provider in {"resend", "sendgrid", "smtp"} and email_from_address() != "no-reply@errorsweep.local"
    billing_provider = billing_provider_label()
    billing_webhook_ready = bool(secret("STRIPE_WEBHOOK_SECRET", "") or secret("RAZORPAY_WEBHOOK_SECRET", "") or secret("ERRORSWEEP_BILLING_WEBHOOK_SECRET", ""))
    storage_health = object_storage_status() if object_storage_status is not None else {"provider": "local", "configured": False, "mode": "local_fallback"}
    async_health = async_backend_status() if async_backend_status is not None else {"provider": "local", "ready": False, "mode": "local_inline"}
    rows = [
        {
            "Requirement": "Billing & subscriptions",
            "Current status": "Mandates + webhook events ready" if st.session_state.get("billing_events") else "Mandate checkout foundation available",
            "Launch gate": "Ready to test provider" if billing_provider_ready(billing_provider) and billing_webhook_ready else "Needs live provider + webhook secret",
            "Next action": "Set ERRORSWEEP_BILLING_PROVIDER, Stripe/Razorpay keys, monthly mandate links, and webhook secret; then test webhook reconciliation.",
        },
        {
            "Requirement": "Async task queue",
            "Current status": f"{async_health.get('provider', 'local')} / {async_health.get('mode', 'local_inline')}",
            "Launch gate": "External worker configured" if async_health.get("mode") == "external" and async_health.get("ready") else "Needs worker service",
            "Next action": "Set ERRORSWEEP_ASYNC_WORKER_URL or REDIS_URL/CELERY_BROKER_URL, then run QA/Pro jobs through the worker before high-volume public launch.",
        },
        {
            "Requirement": "Cloud object storage",
            "Current status": f"{storage_health.get('provider', 'local')} / {storage_health.get('mode', 'local_fallback')}",
            "Launch gate": "Storage bucket configured" if storage_health.get("configured") and storage_health.get("provider") != "local" else "Needs storage bucket",
            "Next action": "Set ERRORSWEEP_OBJECT_STORAGE_PROVIDER plus Supabase Storage, S3, or GCS bucket secrets for multi-instance deployments.",
        },
        {
            "Requirement": "Authentication & onboarding",
            "Current status": "Password hashing, email verification, and reset links ready",
            "Launch gate": "Needs SSO provider" if not secret("ERRORSWEEP_ENTERPRISE_SSO_ENABLED") else "SSO flag configured",
            "Next action": "Configure OAuth/SAML for enterprise tenants and test verification/reset delivery through the email provider.",
        },
        {
            "Requirement": "Legal & compliance",
            "Current status": "Draft policy pages and consent banner ready",
            "Launch gate": "Needs legal approval" if safe_text(secret("ERRORSWEEP_LEGAL_REVIEWED", "")).lower() not in {"1", "true", "yes"} else "Legal flag approved",
            "Next action": "Replace draft Terms/Privacy/Security with lawyer-reviewed Terms, DPA, Cookie Notice, and privacy policy.",
        },
        {
            "Requirement": "CDN / WAF / SSL",
            "Current status": "App security controls ready",
            "Launch gate": "Needs deployment edge" if not secret("ERRORSWEEP_WAF_PROVIDER") else "Edge provider configured",
            "Next action": "Deploy behind Cloudflare/AWS CloudFront or equivalent with HTTPS, rate limits, and WAF rules.",
        },
        {
            "Requirement": "Transactional email",
            "Current status": f"Outbox + dispatch foundation ({email_provider})",
            "Launch gate": "Needs provider secrets/from address" if not email_ready else "Ready to send",
            "Next action": "Configure Resend, SendGrid, or SMTP credentials and verify the sender domain.",
        },
        {
            "Requirement": "Production persistence",
            "Current status": health.get("storage_mode", "unknown"),
            "Launch gate": "Ready" if supabase_ready else "Needs Supabase schema/secrets",
            "Next action": "Run the Supabase release schema and configure SUPABASE_URL + SUPABASE_SERVICE_ROLE_KEY.",
        },
    ]
    return rows


def split_text_lines(text: str) -> List[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def parse_timecode(value: str) -> float:
    value = value.strip().replace(",", ".")
    parts = value.split(":")
    try:
        if len(parts) == 3:
            h, m, s = parts
            return int(h) * 3600 + int(m) * 60 + float(s)
        if len(parts) == 2:
            m, s = parts
            return int(m) * 60 + float(s)
        return float(value)
    except Exception as exc:
        LOGGER.debug("Unable to parse timecode %s: %s", value, exc)
        return 0.0


def format_time(seconds: float, comma: bool = True) -> str:
    seconds = max(float(seconds), 0.0)
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds % 60
    if comma:
        return f"{h:02d}:{m:02d}:{int(s):02d},{int((s % 1) * 1000):03d}"
    return f"{h:02d}:{m:02d}:{int(s):02d}.{int((s % 1) * 1000):03d}"


def parse_srt_or_vtt(text: str) -> List[Dict[str, Any]]:
    text = text.replace("\ufeff", "").replace("\r\n", "\n").replace("\r", "\n")
    blocks = re.split(r"\n\s*\n", text.strip())
    rows = []
    for block in blocks:
        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue
        idx = 0
        if re.fullmatch(r"\d+", lines[0]):
            idx = 1
        if idx >= len(lines):
            continue
        if "-->" not in lines[idx]:
            # Maybe plain text block.
            content = " ".join(lines[idx:])
            if content:
                n = len(rows)
                rows.append({
                    "id": n + 1, "start": n * 4.0, "end": n * 4.0 + 3.0,
                    "source": content, "target": "", "status": "Untranslated", "match": ""
                })
            continue
        start_s, end_s = [x.strip().split(" ")[0] for x in lines[idx].split("-->", 1)]
        content = " ".join(lines[idx + 1:]).strip()
        n = len(rows)
        rows.append({
            "id": n + 1,
            "start": parse_timecode(start_s),
            "end": parse_timecode(end_s),
            "source": content,
            "target": "",
            "status": "Untranslated",
            "match": "",
        })
    return rows


def parse_uploaded_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            LOGGER.warning("Unable to decode uploaded text with %s: %s", enc, exc)
            continue
    return data.decode("utf-8", errors="replace")


def inspect_rules_zip(uploaded_file) -> Dict[str, Any]:
    if uploaded_file is None:
        return {"ok": True, "warnings": []}

    warnings: List[str] = []
    try:
        data = uploaded_file.getvalue()
        size = len(data)
        if size > RULE_ZIP_MAX_BYTES:
            warnings.append(
                f"Rules ZIP is {size / (1024 * 1024):.1f} MB. "
                f"Recommended maximum is {RULE_ZIP_MAX_BYTES / (1024 * 1024):.0f} MB."
            )
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            members = [
                info for info in zf.infolist()
                if not info.is_dir() and not info.filename.startswith("__MACOSX/")
            ]
            if len(members) > RULE_ZIP_MAX_FILES:
                warnings.append(
                    f"Rules ZIP contains {len(members)} files. "
                    f"Recommended maximum is {RULE_ZIP_MAX_FILES} files."
                )
            total_uncompressed = sum(info.file_size for info in members)
            if total_uncompressed > RULE_ZIP_MAX_BYTES * 4:
                warnings.append(
                    f"Rules ZIP expands to {total_uncompressed / (1024 * 1024):.1f} MB. "
                    "Large rule packs may slow QA; split the pack by client/domain."
                )
            return {
                "ok": not warnings,
                "warnings": warnings,
                "files": len(members),
                "bytes": size,
                "expanded_bytes": total_uncompressed,
            }
    except zipfile.BadZipFile:
        return {"ok": False, "warnings": ["Rules upload is not a valid ZIP file."]}
    except Exception as exc:
        LOGGER.warning("Unable to inspect rules ZIP: %s", exc)
        return {"ok": False, "warnings": [f"Could not inspect rules ZIP: {exc}"]}


def render_rules_zip_warning(uploaded_file) -> None:
    report = inspect_rules_zip(uploaded_file)
    for warning in report.get("warnings", []):
        st.warning(warning)


def parse_rules_zip(uploaded_file) -> Dict[str, Any]:
    if uploaded_file is None:
        return {}

    report = inspect_rules_zip(uploaded_file)
    chunks: List[Dict[str, str]] = []
    if any("not a valid ZIP" in warning for warning in report.get("warnings", [])):
        return {"chunks": chunks, "warnings": report.get("warnings", [])}

    try:
        with zipfile.ZipFile(io.BytesIO(uploaded_file.getvalue())) as zf:
            for info in zf.infolist():
                if info.is_dir() or info.filename.startswith("__MACOSX/"):
                    continue
                suffix = Path(info.filename).suffix.lower()
                data = zf.read(info)
                text = ""
                if suffix in {".txt", ".md", ".csv", ".tsv"}:
                    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
                        try:
                            text = data.decode(enc)
                            break
                        except UnicodeDecodeError:
                            continue
                elif suffix == ".docx":
                    doc = Document(io.BytesIO(data))
                    text = "\n".join(p.text for p in doc.paragraphs if p.text)
                elif suffix == ".xlsx":
                    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
                    lines = []
                    for ws in wb.worksheets:
                        for row in ws.iter_rows(values_only=True):
                            values = [safe_text(cell) for cell in row if safe_text(cell)]
                            if values:
                                lines.append(" | ".join(values))
                    text = "\n".join(lines)
                if text.strip():
                    chunks.append({"source": info.filename, "text": text[:200_000]})
    except Exception as exc:
        LOGGER.warning("Unable to parse rules ZIP: %s", exc)
        return {"chunks": chunks, "warnings": [f"Could not parse rules ZIP: {exc}"]}
    return {"chunks": chunks, "warnings": report.get("warnings", [])}


def _split_rule_terms(value: str) -> List[str]:
    value = safe_text(value)
    if not value:
        return []
    parts = re.split(r"[;,]|\t|\|", value)
    return [part.strip().strip("-•* ") for part in parts if part.strip().strip("-•* ")]


def enrich_rules_from_chunks(rules: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    """Extract high-signal glossary, DNT, and instruction hints from uploaded rule files."""
    enriched = dict(rules or {})
    glossary = list(enriched.get("glossary") or [])
    dnt = list(enriched.get("dnt") or [])
    instructions = list(enriched.get("instructions") or [])

    for chunk in enriched.get("chunks", []):
        source_name = safe_text(chunk.get("source", "Rules ZIP"))
        lower_name = source_name.lower()
        for raw_line in safe_text(chunk.get("text", "")).splitlines():
            line = raw_line.strip()
            if not line or len(line) > 700:
                continue
            lower = line.lower()

            if any(token in lower_name for token in ["instruction", "style", "guideline", "brief"]) or lower.startswith(("instruction", "style", "tone", "guideline", "rule:")):
                if len(instructions) < 80:
                    instructions.append({"text": line, "source": source_name})

            dnt_match = re.match(r"^(?:dnt|do\s*not\s*translate|locked\s*term|keep(?:\s+unchanged)?)\s*[:=-]\s*(.+)$", line, flags=re.I)
            if dnt_match:
                for term in _split_rule_terms(dnt_match.group(1)):
                    dnt.append({"term": term, "source": source_name})
                continue

            glossary_match = re.match(r"^(?:glossary|term|terminology)\s*[:=-]\s*(.+)$", line, flags=re.I)
            candidate = glossary_match.group(1).strip() if glossary_match else line
            if "->" in candidate or "=>" in candidate:
                sep = "->" if "->" in candidate else "=>"
                src, tgt = [part.strip() for part in candidate.split(sep, 1)]
                src = re.sub(r"^(?:source|src|english|term)\s*[:=]\s*", "", src, flags=re.I).strip()
                tgt = re.sub(r"^(?:target|tgt|translation|preferred)\s*[:=]\s*", "", tgt, flags=re.I).strip()
                if src and tgt and src.lower() != tgt.lower():
                    glossary.append({"source_term": src, "target_term": tgt, "source": source_name})
                elif src:
                    dnt.append({"term": src, "source": source_name})
            elif any(token in lower_name for token in ["dnt", "do-not-translate", "locked"]):
                for term in _split_rule_terms(line):
                    if len(term) <= 120:
                        dnt.append({"term": term, "source": source_name})

    def dedupe_dicts(items: List[Dict[str, Any]], keys: Tuple[str, ...]) -> List[Dict[str, Any]]:
        seen = set()
        output = []
        for item in items:
            key = tuple(safe_text(item.get(k, "")).lower() for k in keys)
            if not any(key) or key in seen:
                continue
            seen.add(key)
            output.append(item)
        return output

    enriched["glossary"] = dedupe_dicts(glossary, ("source_term", "target_term"))
    enriched["dnt"] = dedupe_dicts(dnt, ("term",))
    enriched["instructions"] = dedupe_dicts(instructions, ("text",))[:120]
    return enriched


def workspace_rules(uploaded_rules=None) -> Dict[str, Any]:
    """Merge uploaded rules with saved workspace memory so every workflow sees one rule pack."""
    rules = enrich_rules_from_chunks(parse_rules_zip(uploaded_rules) if uploaded_rules is not None else {})
    glossary = list(rules.get("glossary") or [])
    dnt = list(rules.get("dnt") or [])
    instructions = list(rules.get("instructions") or [])

    for item in st.session_state.get("glossary", []):
        source_term = safe_text(item.get("source") or item.get("source_term"))
        target_term = safe_text(item.get("target") or item.get("target_term"))
        notes = safe_text(item.get("notes", ""))
        if source_term and target_term:
            glossary.append({"source_term": source_term, "target_term": target_term, "source": "Saved Glossary"})
        if source_term and ("dnt" in notes.lower() or source_term == target_term):
            dnt.append({"term": source_term, "source": "Saved Glossary"})

    for term in st.session_state.get("dnt", []):
        if safe_text(term):
            dnt.append({"term": safe_text(term), "source": "Saved DNT"})

    if st.session_state.get("tm"):
        instructions.append({"text": "Prefer saved translation memory matches where applicable.", "source": "Saved TM"})

    for item in st.session_state.get("rule_instructions", []):
        text = safe_text(item.get("text") if isinstance(item, dict) else item)
        source = safe_text(item.get("source", "Saved Instructions")) if isinstance(item, dict) else "Saved Instructions"
        if text:
            instructions.append({"text": text, "source": source or "Saved Instructions"})

    merged = dict(rules)
    merged["glossary"] = enrich_rules_from_chunks({"glossary": glossary}).get("glossary", [])
    merged["dnt"] = enrich_rules_from_chunks({"dnt": dnt}).get("dnt", [])
    merged["instructions"] = enrich_rules_from_chunks({"instructions": instructions}).get("instructions", [])
    return merged


def rules_summary_for_ai(rules: Optional[Dict[str, Any]], max_items: int = 24) -> str:
    rules = rules or {}
    lines: List[str] = []
    dnt = [safe_text(item.get("term")) for item in rules.get("dnt", []) if safe_text(item.get("term"))]
    glossary = [
        (safe_text(item.get("source_term")), safe_text(item.get("target_term")))
        for item in rules.get("glossary", [])
        if safe_text(item.get("source_term")) and safe_text(item.get("target_term"))
    ]
    instructions = [safe_text(item.get("text")) for item in rules.get("instructions", []) if safe_text(item.get("text"))]
    if dnt:
        lines.append("Do-not-translate / locked terms: " + "; ".join(dnt[:max_items]))
    if glossary:
        pairs = [f"{src} => {tgt}" for src, tgt in glossary[:max_items]]
        lines.append("Required glossary terms: " + "; ".join(pairs))
    if instructions:
        lines.append("Client instructions: " + " | ".join(instructions[:12]))
    return "\n".join(lines)[:6000]


def protected_terms_from_rules(rules: Optional[Dict[str, Any]]) -> List[str]:
    terms = []
    for item in (rules or {}).get("dnt", []):
        if safe_text(item.get("term")):
            terms.append(safe_text(item.get("term")))
    for item in (rules or {}).get("glossary", []):
        src = safe_text(item.get("source_term"))
        tgt = safe_text(item.get("target_term"))
        if src and src == tgt:
            terms.append(src)
    return sorted(set(terms), key=len, reverse=True)


def _norm_header(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", " ", safe_text(value).lower()).strip()


def _header_contains(header: str, *needles: str) -> bool:
    return any(needle in header for needle in needles)


def _score_header_row(values: List[Any]) -> int:
    headers = [_norm_header(v) for v in values]
    score = 0
    if any(h in {"source", "source text", "source string", "english"} for h in headers):
        score += 4
    if any(("translation" in h and "language" not in h) or h in {"target", "target text", "translated text"} for h in headers):
        score += 4
    if any(h in {"item no", "segment id", "id", "key"} for h in headers):
        score += 1
    if any("reviewer" in h or "error" in h or "severity" in h for h in headers):
        score += 1
    return score


def _find_table_header(data: List[Tuple[Any, ...]]) -> Tuple[int, List[str]]:
    best_idx = 0
    best_score = -1
    for idx, row in enumerate(data[:50]):
        score = _score_header_row(list(row))
        if score > best_score:
            best_idx = idx
            best_score = score
    headers = [_norm_header(v) for v in data[best_idx]] if data else []
    return best_idx, headers


def _find_first_header(headers: List[str], candidates: List[str], blocked: Optional[List[str]] = None) -> Optional[int]:
    blocked = blocked or []
    for wanted in candidates:
        for idx, header in enumerate(headers):
            if not header or any(bad in header for bad in blocked):
                continue
            if header == wanted or wanted in header:
                return idx
    return None


def _default_source_index(headers: List[str]) -> Optional[int]:
    return _find_first_header(
        headers,
        ["source text", "source", "english", "source string"],
        blocked=["source language"],
    )


def _target_candidate_indices(headers: List[str], mode: str) -> List[int]:
    if mode == "translator":
        wanted = [
            "original translation",
            "translated text",
            "machine translation",
            "target text",
            "target",
            "translation",
        ]
        blocked = ["suggested", "reviewer", "comment", "response", "error"]
    elif mode == "reviewer":
        wanted = [
            "suggested translation",
            "final translation",
            "reviewed translation",
            "corrected translation",
            "reviewer translation",
            "target text",
            "target",
            "original translation",
            "translated text",
            "translation",
        ]
        blocked = ["comment", "response", "error"]
    elif mode == "source":
        return []
    else:
        wanted = [
            "target text",
            "target",
            "final translation",
            "suggested translation",
            "original translation",
            "translated text",
            "translation",
        ]
        blocked = ["comment", "response", "error"]

    found: List[int] = []
    for wanted_header in wanted:
        for idx, header in enumerate(headers):
            if idx in found or not header:
                continue
            if any(bad in header for bad in blocked):
                continue
            if header == wanted_header or wanted_header in header:
                found.append(idx)
    return found


def _row_value(row: Tuple[Any, ...], idx: Optional[int]) -> str:
    if idx is None or idx >= len(row):
        return ""
    return safe_text(row[idx])


def _column_label(headers: List[str], idx: int) -> str:
    label = safe_text(headers[idx]) if idx < len(headers) else ""
    return f"{idx + 1}. {label or f'Column {idx + 1}'}"


def _column_choices(headers: List[str]) -> List[str]:
    return [_column_label(headers, idx) for idx in range(len(headers))]


def _choice_index(choice: str) -> Optional[int]:
    match = re.match(r"\s*(\d+)\.", safe_text(choice))
    return int(match.group(1)) - 1 if match else None


def _looks_like_scorecard_label(text: str) -> bool:
    value = safe_text(text)
    if not value:
        return False
    return bool(re.fullmatch(r"(?:text|part|section)\s*\d+(?:\s*\(\d+\))?", value, flags=re.I))


def _extract_rows_from_table(
    data: List[Tuple[Any, ...]],
    mode: str = "review",
    source_idx: Optional[int] = None,
    target_idx: Optional[int] = None,
    header_idx: Optional[int] = None,
) -> List[Dict[str, Any]]:
    if not data:
        return []
    detected_header_idx, headers = _find_table_header(data)
    header_idx = detected_header_idx if header_idx is None else header_idx
    if header_idx != detected_header_idx:
        headers = [_norm_header(v) for v in data[header_idx]] if header_idx < len(data) else headers
    if source_idx is None:
        source_idx = _default_source_index(headers)
    if source_idx is None:
        source_idx = 0
    target_indices = [target_idx] if target_idx is not None else _target_candidate_indices(headers, mode)
    if not target_indices and mode != "source":
        fallback = 1 if len(headers) > 1 else None
        target_indices = [fallback] if fallback is not None else []

    rows: List[Dict[str, Any]] = []
    for row in data[header_idx + 1:]:
        source = _row_value(row, source_idx)
        target = ""
        for idx in target_indices:
            target = _row_value(row, idx)
            if target:
                break
        if _looks_like_scorecard_label(source) and (not target or _looks_like_scorecard_label(target)):
            continue
        if not source and not target:
            continue
        if source.lower() in {"source", "source text"} and target.lower() in {"target", "translation", "original translation"}:
            continue
        rows.append({
            "id": len(rows) + 1,
            "source": source,
            "target": target,
            "status": "Existing" if target else "Untranslated",
            "match": "",
        })
    return rows


def _scorecard_table_options(uploaded_file) -> List[Dict[str, Any]]:
    if uploaded_file is None:
        return []
    name = uploaded_file.name.lower()
    options: List[Dict[str, Any]] = []
    data_bytes = uploaded_file.getvalue()
    try:
        if name.endswith(".xlsx"):
            wb = load_workbook(io.BytesIO(data_bytes), data_only=True, read_only=True)
            for ws in wb.worksheets:
                data = list(ws.iter_rows(values_only=True))
                if not data:
                    continue
                header_idx, norm_headers = _find_table_header(data)
                headers = [safe_text(v) for v in data[header_idx]]
                if not any(headers):
                    continue
                score = _score_header_row(list(data[header_idx]))
                options.append({
                    "label": f"Sheet: {ws.title} ({len(data) - header_idx - 1} rows)",
                    "data": data,
                    "header_idx": header_idx,
                    "headers": headers,
                    "norm_headers": norm_headers,
                    "score": score,
                })
        elif name.endswith(".docx"):
            docx_options: List[Dict[str, Any]] = []
            for idx, table in enumerate(_docx_table_rows(data_bytes), start=1):
                if not table:
                    continue
                header_idx, norm_headers = _find_table_header(table)
                headers = [safe_text(v) for v in table[header_idx]]
                if not any(headers):
                    continue
                score = _score_header_row(list(table[header_idx]))
                docx_options.append({
                    "label": f"Table {idx} ({len(table) - header_idx - 1} rows)",
                    "data": table,
                    "header_idx": header_idx,
                    "headers": headers,
                    "norm_headers": norm_headers,
                    "score": score,
                })
            docx_options = sorted(docx_options, key=lambda opt: (opt.get("score", 0), len(opt.get("data", []))), reverse=True)
            if len(docx_options) > 1:
                best = docx_options[0]
                options.append({
                    "label": f"All detected DOCX tables ({sum(max(0, len(opt.get('data', [])) - int(opt.get('header_idx', 0)) - 1) for opt in docx_options)} rows)",
                    "tables": docx_options,
                    "data": best.get("data", []),
                    "header_idx": best.get("header_idx", 0),
                    "headers": best.get("headers", []),
                    "norm_headers": best.get("norm_headers", []),
                    "score": best.get("score", 0) + 1,
                })
            options.extend(docx_options)
        elif name.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(data_bytes))
            headers = [safe_text(c) for c in df.columns]
            if not any(headers):
                return options
            data = [tuple(headers)] + [tuple(row) for row in df.itertuples(index=False, name=None)]
            norm_headers = [_norm_header(h) for h in headers]
            options.append({
                "label": f"CSV ({len(df)} rows)",
                "data": data,
                "header_idx": 0,
                "headers": headers,
                "norm_headers": norm_headers,
                "score": _score_header_row(headers),
            })
    except Exception as exc:
        LOGGER.warning("Unable to inspect scorecard file %s: %s", getattr(uploaded_file, "name", ""), exc)
    return sorted(options, key=lambda opt: (opt.get("score", 0), len(opt.get("data", []))), reverse=True)


def _default_table_index(options: List[Dict[str, Any]]) -> int:
    return 0 if options else -1


def _default_column_choice(options: List[Dict[str, Any]], table_idx: int, mode: str, column_kind: str) -> int:
    if table_idx < 0 or table_idx >= len(options):
        return 0
    headers = options[table_idx].get("headers", [])
    norm_headers = options[table_idx].get("norm_headers", [])
    if not headers:
        return 0
    if column_kind == "source":
        idx = _default_source_index(norm_headers)
    else:
        candidates = _target_candidate_indices(norm_headers, mode)
        idx = candidates[0] if candidates else None
    if idx is None:
        idx = 0 if column_kind == "source" else min(1, len(headers) - 1)
    return max(0, min(idx, len(headers) - 1))


def _rows_from_mapping(options: List[Dict[str, Any]], table_idx: int, mode: str, source_col: Optional[int], target_col: Optional[int]) -> List[Dict[str, Any]]:
    if table_idx < 0 or table_idx >= len(options):
        return []
    opt = options[table_idx]
    if opt.get("tables"):
        rows: List[Dict[str, Any]] = []
        for table_opt in opt.get("tables", []):
            for item in _extract_rows_from_table(
                table_opt.get("data", []),
                mode=mode,
                source_idx=source_col,
                target_idx=target_col,
                header_idx=table_opt.get("header_idx"),
            ):
                item["id"] = len(rows) + 1
                rows.append(item)
        return rows
    return _extract_rows_from_table(
        opt.get("data", []),
        mode=mode,
        source_idx=source_col,
        target_idx=target_col,
        header_idx=opt.get("header_idx"),
    )


def render_scorecard_mapping_controls(label: str, uploaded_file, mode: str, key_prefix: str) -> Tuple[List[Dict[str, Any]], Dict[str, Optional[int]]]:
    options = _scorecard_table_options(uploaded_file)
    mapping: Dict[str, Optional[int]] = {"table_idx": None, "source_col": None, "target_col": None}
    if not options:
        st.warning(f"No structured table could be detected in {label}. ErrorSweep will use fallback extraction.")
        return options, mapping

    table_labels = [opt["label"] for opt in options]
    table_idx = st.selectbox(
        f"{label} table",
        range(len(table_labels)),
        index=_default_table_index(options),
        format_func=lambda idx: table_labels[idx],
        key=f"{key_prefix}_table",
    )
    headers = options[table_idx].get("headers", [])
    choices = _column_choices(headers)
    source_default = _default_column_choice(options, table_idx, mode, "source")
    source_choice = st.selectbox(
        f"{label} source column",
        choices,
        index=source_default,
        key=f"{key_prefix}_source_col",
    )
    target_choice = None
    if mode != "source":
        target_choices = ["Auto detected target"] + choices
        target_choice = st.selectbox(
            f"{label} target column",
            target_choices,
            index=0,
            key=f"{key_prefix}_target_col",
        )
    mapping = {
        "table_idx": int(table_idx),
        "source_col": _choice_index(source_choice),
        "target_col": _choice_index(target_choice) if target_choice and target_choice != "Auto detected target" else None,
    }
    preview_rows = _rows_from_mapping(options, int(table_idx), mode, mapping["source_col"], mapping["target_col"])[:5]
    if preview_rows:
        preview = [{
            "source": clean_scorecard_source_text(row.get("source", ""))[:220],
            "target": safe_text(row.get("target", ""))[:220],
        } for row in preview_rows]
        st.dataframe(pd.DataFrame(preview), use_container_width=True, hide_index=True)
    return options, mapping


def _docx_cell_text(cell: Any, ns: Dict[str, str]) -> str:
    paragraphs = []
    for paragraph in cell.findall(".//w:p", ns):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns))
        if text.strip():
            paragraphs.append(text.strip())
    if paragraphs:
        return "\n".join(paragraphs)
    return " ".join(node.text or "" for node in cell.findall(".//w:t", ns)).strip()


def _docx_table_rows(data: bytes) -> List[List[Tuple[str, ...]]]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    tables: List[List[Tuple[str, ...]]] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            document_xml = archive.read("word/document.xml")
        if b"<!DOCTYPE" in document_xml.upper() or b"<!ENTITY" in document_xml.upper():
            raise ValueError("DOCX XML contains prohibited DTD/entity declarations.")
        if "defusedxml" in getattr(ET, "__name__", ""):
            root = ET.fromstring(document_xml, forbid_dtd=True, forbid_entities=True)
        else:
            root = ET.fromstring(document_xml)
    except Exception as exc:
        LOGGER.warning("Unable to parse DOCX XML tables: %s", exc)
        return tables

    for table in root.findall(".//w:tbl", ns):
        table_rows: List[Tuple[str, ...]] = []
        for row in table.findall("./w:tr", ns):
            cells = tuple(_docx_cell_text(cell, ns) for cell in row.findall("./w:tc", ns))
            if any(safe_text(cell) for cell in cells):
                table_rows.append(cells)
        if table_rows:
            tables.append(table_rows)
    return tables


def extract_rows_from_upload(uploaded_file, mode: str = "review") -> List[Dict[str, Any]]:
    if uploaded_file is None:
        return []
    name = uploaded_file.name.lower()
    rows: List[Dict[str, Any]] = []
    try:
        if name.endswith(".xlsx"):
            wb = load_workbook(io.BytesIO(uploaded_file.getvalue()), data_only=True)
            best_rows = []
            best_rank = (-1, -1)
            for ws in wb.worksheets:
                data = list(ws.iter_rows(values_only=True))
                extracted = _extract_rows_from_table(data, mode=mode)
                header_idx, headers = _find_table_header(data)
                header_score = _score_header_row(list(data[header_idx])) if data else 0
                rank = (header_score, len(extracted))
                if extracted and rank > best_rank:
                    best_rank = rank
                    best_rows = extracted
            rows.extend(best_rows)
        elif name.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(uploaded_file.getvalue()))
            cols = [c.lower() for c in df.columns.astype(str)]
            src_col = df.columns[0]
            tgt_col = df.columns[1] if len(df.columns) > 1 else None
            for c in df.columns:
                cl = str(c).lower()
                if "source" in cl or "english" in cl:
                    src_col = c
                if "target" in cl or "translation" in cl:
                    tgt_col = c
            for _, row in df.iterrows():
                src = safe_text(row.get(src_col, ""))
                tgt = safe_text(row.get(tgt_col, "")) if tgt_col is not None else ""
                if src or tgt:
                    rows.append({"id": len(rows)+1, "source": src, "target": tgt, "status": "Existing" if tgt else "Untranslated", "match": ""})
        elif name.endswith(".docx"):
            data = uploaded_file.getvalue()
            for table in _docx_table_rows(data):
                extracted = _extract_rows_from_table(table, mode=mode)
                for item in extracted:
                    item["id"] = len(rows) + 1
                    rows.append(item)
            if not rows:
                doc = Document(io.BytesIO(data))
                for p in doc.paragraphs:
                    txt = safe_text(p.text)
                    if txt:
                        rows.append({"id": len(rows)+1, "source": txt, "target": "", "status": "Untranslated", "match": ""})
        else:
            text = parse_uploaded_text(uploaded_file)
            if name.endswith((".srt", ".vtt")):
                rows = parse_srt_or_vtt(text)
            else:
                lines = split_text_lines(text)
                for line in lines:
                    rows.append({"id": len(rows)+1, "source": line, "target": "", "status": "Untranslated", "match": ""})
    except Exception as exc:
        st.error(f"Could not parse uploaded file: {exc}")
    return rows


def rows_to_csv(rows: List[Dict[str, Any]]) -> bytes:
    if not rows:
        return b""
    return pd.DataFrame(rows).to_csv(index=False).encode("utf-8-sig")


def rows_to_srt(rows: List[Dict[str, Any]], use_target: bool = True) -> bytes:
    out = []
    for i, row in enumerate(rows, start=1):
        text = safe_text(row.get("target" if use_target else "source", ""))
        if not text:
            text = safe_text(row.get("source", ""))
        out.append(str(i))
        out.append(f"{format_time(row.get('start', (i-1)*4), comma=True)} --> {format_time(row.get('end', (i-1)*4+3), comma=True)}")
        out.append(text)
        out.append("")
    return "\n".join(out).encode("utf-8")


@st.cache_resource
def get_review_session_store() -> Dict[str, Dict[str, Any]]:
    """Small in-memory review-session store.

    Streamlit reruns can cause a button on the Pro page to re-render before the
    hidden editor page has read session_state. We store the Pro review rows under
    a stable session id and pass that id through the query string, so the editor
    can restore rows reliably instead of opening as a blank page.
    """
    return {}


def save_review_session_to_store(rows: List[Dict[str, Any]], title: str, target_language: str, file_name: str, rules: Optional[Dict[str, Any]] = None) -> str:
    """Save a Pro review job for separate-tab editor use.

    v42 release hardening:
    - session memory for current tab
    - local JSON fallback for development
    - Supabase persistence when SUPABASE_URL + SUPABASE_SERVICE_ROLE_KEY are configured
    """
    session_id = uuid.uuid4().hex
    user = current_user() or {}
    metadata = {
        "title": title,
        "target_language": target_language,
        "file_name": file_name,
        "created": now_stamp() if "now_stamp" in globals() else "",
        "source": "ErrorSweep Pro",
        "workspace": user.get("workspace", "Demo Workspace"),
        "user_email": user.get("email", ""),
        "status": "draft",
        "rules": rules or {},
    }
    payload = {
        "job_id": session_id,
        "rows": rows,
        "metadata": metadata,
        "rules": rules or {},
        "title": title,
        "target_language": target_language,
        "file_name": file_name,
        "created": metadata["created"],
        "job_type": "cat",
    }
    get_review_session_store()[session_id] = payload

    # Existing v41 local job store fallback.
    if save_editor_job is not None:
        try:
            save_editor_job("cat", rows, metadata=metadata, job_id=session_id)
        except Exception as exc:
            LOGGER.warning("Unable to save CAT job to local editor store: %s", exc)

    # v42 production persistence.
    if save_persistent_editor_job is not None:
        try:
            save_persistent_editor_job("cat", rows, metadata=metadata, job_id=session_id, user=user)
        except Exception as exc:
            LOGGER.warning("Unable to save CAT job to persistent store: %s", exc)

    st.session_state["active_review_session_id"] = session_id

    # v43 owner-console handoff: keep visible current/recent job details even if
    # Supabase/local persistence is not immediately readable in this Streamlit run.
    owner_job_record = {
        "id": session_id,
        "job_type": "cat",
        "workspace": metadata.get("workspace", ""),
        "user_email": metadata.get("user_email", ""),
        "file_name": file_name,
        "target_language": target_language,
        "status": metadata.get("status", "draft"),
        "row_count": len(rows or []),
        "created": metadata.get("created", ""),
        "updated_at": metadata.get("created", ""),
        "source": metadata.get("source", "ErrorSweep Pro"),
    }
    st.session_state["last_pro_task_details"] = owner_job_record
    st.session_state.setdefault("owner_recent_editor_jobs", [])
    st.session_state["owner_recent_editor_jobs"] = [
        owner_job_record,
        *[j for j in st.session_state["owner_recent_editor_jobs"] if j.get("id") != session_id],
    ][:25]
    return session_id


def media_preview_root() -> Path:
    configured = os.getenv("ERRORSWEEP_MEDIA_PREVIEW_DIR", "").strip()
    root = Path(configured) if configured else Path(tempfile.gettempdir()) / "errorsweep_media_previews"
    root.mkdir(parents=True, exist_ok=True)
    return root


def cleanup_media_preview_files(ttl_seconds: int = MEDIA_PREVIEW_TTL_SECONDS) -> int:
    """Remove stale local media preview files so uploaded videos do not fill disk."""
    if ttl_seconds <= 0:
        return 0
    root = media_preview_root()
    now = time.time()
    removed = 0
    for path in root.iterdir():
        try:
            if not path.is_file():
                continue
            if now - path.stat().st_mtime <= ttl_seconds:
                continue
            path.unlink(missing_ok=True)
            removed += 1
        except Exception as exc:
            LOGGER.warning("Unable to clean media preview file %s: %s", path, exc)
    return removed


def maybe_cleanup_media_preview_files(
    ttl_seconds: int = MEDIA_PREVIEW_TTL_SECONDS,
    interval_seconds: int = 60 * 30,
) -> int:
    """Throttle media preview cleanup during normal Streamlit reruns."""
    now = time.time()
    key = "_media_preview_last_cleanup_at"
    last_cleanup = float(st.session_state.get(key, 0) or 0)
    if interval_seconds > 0 and now - last_cleanup < interval_seconds:
        return 0
    st.session_state[key] = now
    try:
        return cleanup_media_preview_files(ttl_seconds=ttl_seconds)
    except Exception as exc:
        LOGGER.warning("Periodic media preview cleanup failed: %s", exc)
        return 0


def stream_uploaded_file_to_path(uploaded_file: Any, path: Path, chunk_size: int = 1024 * 1024) -> int:
    """Write a Streamlit UploadedFile to disk in chunks without calling getvalue()."""
    path.parent.mkdir(parents=True, exist_ok=True)
    written = 0
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    with path.open("wb") as target:
        while True:
            chunk = uploaded_file.read(chunk_size)
            if not chunk:
                break
            target.write(chunk)
            written += len(chunk)
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    return written


def hash_file(path: Path, chunk_size: int = 1024 * 1024) -> Tuple[str, int]:
    digest = hashlib.sha256()
    size = 0
    with Path(path).open("rb") as handle:
        while True:
            chunk = handle.read(chunk_size)
            if not chunk:
                break
            digest.update(chunk)
            size += len(chunk)
    return digest.hexdigest(), size


def job_attachment_root() -> Path:
    configured = os.getenv("ERRORSWEEP_JOB_ATTACHMENT_DIR", "").strip()
    root = Path(configured) if configured else Path(tempfile.gettempdir()) / "errorsweep_job_attachments"
    root.mkdir(parents=True, exist_ok=True)
    return root


def store_file_object(
    path: Path,
    workspace: str,
    purpose: str,
    object_id: str,
    file_name: str,
    mime_type: str = "application/octet-stream",
) -> Dict[str, Any]:
    """Store a local file path through the configured object storage provider."""
    key = ""
    if build_object_key is not None:
        key = build_object_key(workspace, purpose, object_id, file_name)
    if put_object_file is None or not key:
        return {
            "storage_provider": "local",
            "storage_bucket": "local",
            "storage_key": str(path),
            "local_path": str(path),
            "public_url": "",
            "status": "stored",
        }
    try:
        return put_object_file(path, key, content_type=mime_type)
    except Exception as exc:
        LOGGER.warning("Object storage upload failed for %s; using local fallback: %s", file_name, exc)
        return {
            "storage_provider": "local",
            "storage_bucket": "local",
            "storage_key": str(path),
            "local_path": str(path),
            "public_url": "",
            "status": "stored_local_fallback",
        }


def save_job_attachment_files(job_id: str, uploaded_files: List[Any], purpose: str = "job_assignment") -> List[Dict[str, Any]]:
    """Persist manual job assignment uploads and create file manifest records."""
    manifests: List[Dict[str, Any]] = []
    user = current_user() or {}
    workspace = user.get("workspace", "Demo Workspace")
    safe_job_id = re.sub(r"[^A-Za-z0-9_.-]+", "_", safe_text(job_id) or uuid.uuid4().hex)
    for uploaded in uploaded_files or []:
        original_name = safe_text(getattr(uploaded, "name", "attachment")) or "attachment"
        safe_name = re.sub(r"[^A-Za-z0-9_. -]+", "_", original_name).strip(" .") or "attachment"
        target_dir = job_attachment_root() / safe_job_id
        target_path = target_dir / safe_name
        if target_path.exists():
            target_path = target_dir / f"{uuid.uuid4().hex[:8]}_{safe_name}"
        tmp_path = target_path.with_suffix(target_path.suffix + ".tmp")
        sha = hashlib.sha256()
        written = 0
        try:
            target_dir.mkdir(parents=True, exist_ok=True)
            try:
                uploaded.seek(0)
            except Exception:
                pass
            with tmp_path.open("wb") as handle:
                while True:
                    chunk = uploaded.read(1024 * 1024)
                    if not chunk:
                        break
                    handle.write(chunk)
                    sha.update(chunk)
                    written += len(chunk)
            os.replace(tmp_path, target_path)
            try:
                uploaded.seek(0)
            except Exception:
                pass
            mime_type = safe_text(getattr(uploaded, "type", "")) or "application/octet-stream"
            storage_result = store_file_object(target_path, workspace, purpose, safe_job_id, original_name, mime_type)
            manifest = persist_saas_record("files", {
                "workspace": workspace,
                "user_email": user.get("email", ""),
                "file_name": original_name,
                "purpose": purpose,
                "mime_type": mime_type,
                "size_bytes": written,
                "sha256": sha.hexdigest(),
                "storage_key": storage_result.get("storage_key", str(target_path)),
                "storage_provider": storage_result.get("storage_provider", "local"),
                "storage_bucket": storage_result.get("storage_bucket", ""),
                "public_url": storage_result.get("public_url", ""),
                "local_path": storage_result.get("local_path", str(target_path)),
                "status": storage_result.get("status", "stored"),
                "created_at": now_stamp(),
                "updated_at": now_stamp(),
            })
            manifests.append({
                "id": manifest.get("id", ""),
                "file_name": original_name,
                "size_bytes": written,
                "sha256": sha.hexdigest(),
                "storage_key": manifest.get("storage_key", ""),
                "storage_provider": manifest.get("storage_provider", "local"),
                "storage_bucket": manifest.get("storage_bucket", ""),
                "public_url": manifest.get("public_url", ""),
                "mime_type": manifest.get("mime_type", ""),
            })
            st.session_state.setdefault("files", [])
            st.session_state.files.insert(0, manifest)
            trim_session_list("files")
            if manifest.get("storage_provider") not in {"", "local"} and safe_text(manifest.get("status")) == "stored":
                try:
                    target_path.unlink(missing_ok=True)
                except Exception as exc:
                    LOGGER.warning("Unable to remove local staging attachment %s: %s", target_path, exc)
        except Exception as exc:
            LOGGER.warning("Unable to save job attachment %s: %s", original_name, exc)
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass
    return manifests


def save_media_preview_file(job_id: str, video_file) -> Dict[str, str]:
    if video_file is None:
        return {}
    try:
        cleanup_media_preview_files()
        original_name = safe_text(getattr(video_file, "name", "")) or "media"
        suffix = Path(original_name).suffix.lower()
        if suffix not in {".mp4", ".mov", ".m4v", ".webm", ".mp3", ".wav", ".m4a"}:
            suffix = ".bin"
        path = media_preview_root() / f"{job_id}{suffix}"
        tmp_path = path.with_suffix(path.suffix + ".tmp")

        stream_uploaded_file_to_path(video_file, tmp_path)
        os.replace(tmp_path, path)
        user = current_user() or {}
        workspace = user.get("workspace", "Demo Workspace")
        mime_type = safe_text(getattr(video_file, "type", "")) or "video/mp4"
        storage_result = store_file_object(path, workspace, "media_preview", job_id, original_name, mime_type)
        try:
            sha, size_bytes = hash_file(path)
        except Exception as exc:
            LOGGER.warning("Unable to hash media preview file %s: %s", path, exc)
            sha = ""
            size_bytes = int(path.stat().st_size) if path.exists() else 0
        manifest = persist_saas_record("files", {
            "workspace": workspace,
            "user_email": user.get("email", ""),
            "file_name": original_name,
            "purpose": "media_preview",
            "mime_type": mime_type,
            "size_bytes": size_bytes,
            "sha256": sha,
            "storage_key": storage_result.get("storage_key", str(path)),
            "storage_provider": storage_result.get("storage_provider", "local"),
            "storage_bucket": storage_result.get("storage_bucket", ""),
            "public_url": storage_result.get("public_url", ""),
            "local_path": storage_result.get("local_path", str(path)),
            "status": storage_result.get("status", "stored"),
            "expires_at": datetime.fromtimestamp(time.time() + MEDIA_PREVIEW_TTL_SECONDS, timezone.utc).isoformat() if MEDIA_PREVIEW_TTL_SECONDS > 0 else None,
            "created_at": now_stamp(),
            "updated_at": now_stamp(),
        })
        st.session_state.setdefault("files", [])
        st.session_state.files.insert(0, manifest)
        trim_session_list("files")
        return {
            "media_preview_path": str(path),
            "media_preview_type": mime_type,
            "media_preview_name": original_name,
            "media_preview_storage_key": safe_text(storage_result.get("storage_key", "")),
            "media_preview_storage_provider": safe_text(storage_result.get("storage_provider", "local")),
            "media_preview_storage_bucket": safe_text(storage_result.get("storage_bucket", "")),
            "media_preview_public_url": safe_text(storage_result.get("public_url", "")),
        }
    except Exception as exc:
        LOGGER.warning("Unable to save media preview file for job %s: %s", job_id, exc)
        return {}


def read_media_preview_bytes(metadata: Dict[str, Any]) -> Tuple[Optional[Any], str, str]:
    path = safe_text(metadata.get("media_preview_path", ""))
    mime = safe_text(metadata.get("media_preview_type", metadata.get("video_type", ""))) or "video/mp4"
    name = safe_text(metadata.get("media_preview_name", metadata.get("video_name", "")))
    if path:
        try:
            media_path = Path(path)
            root = media_preview_root().resolve()
            resolved = media_path.resolve()
            if resolved.exists() and (root in resolved.parents or resolved == root):
                return str(resolved), mime, name or resolved.name
            if resolved.exists():
                LOGGER.warning("Rejected media preview path outside preview root: %s", resolved)
        except Exception as exc:
            LOGGER.warning("Unable to read media preview file %s: %s", path, exc)
    public_url = safe_text(metadata.get("media_preview_public_url", ""))
    if public_url:
        return public_url, mime, name
    storage_key = safe_text(metadata.get("media_preview_storage_key", ""))
    storage_provider = safe_text(metadata.get("media_preview_storage_provider", ""))
    if storage_key and storage_provider == "local":
        try:
            local_path = Path(storage_key)
            if local_path.exists():
                return str(local_path), mime, name or local_path.name
        except Exception as exc:
            LOGGER.warning("Unable to read local object storage preview %s: %s", storage_key, exc)
    if storage_key and storage_provider != "local" and signed_url_for_key is not None:
        try:
            url = signed_url_for_key(storage_key)
            if url:
                return url, mime, name
        except Exception as exc:
            LOGGER.warning("Unable to create signed media preview URL for %s: %s", storage_key, exc)
    session_bytes = st.session_state.get("subtitle_video_bytes")
    if session_bytes:
        return session_bytes, safe_text(st.session_state.get("subtitle_video_type", mime)) or mime, safe_text(st.session_state.get("subtitle_video_name", name))
    return None, mime, name


def render_media_preview(media_source: Optional[Any], mime: str, name: str = "") -> None:
    if not media_source:
        st.warning("Media preview is not available for this editor job. Go back to setup and recreate the media job.")
        return
    st.caption(name or "Media preview")
    if str(mime or "").lower().startswith("audio/"):
        st.audio(media_source, format=mime)
    else:
        st.video(media_source, format=mime)


def save_media_session_to_store(workflow: str, rows: List[Dict[str, Any]], video_file=None, target_language: str = "") -> str:
    """Save subtitle/transcription rows as an external media editor job.

    This fixes the workflow where video upload created rows but did not open a
    separate editor like the CAT review editor. A local media preview copy is
    saved for browser playback; rows and metadata are stored safely.
    """
    job_id = uuid.uuid4().hex
    user = current_user() or {}
    video_name = getattr(video_file, "name", "") if video_file is not None else ""
    video_type = getattr(video_file, "type", "") if video_file is not None else ""
    media_preview = save_media_preview_file(job_id, video_file)
    metadata = {
        "title": f"ErrorSweep {workflow} Editor",
        "target_language": target_language or st.session_state.get("subtitle_target_language", ""),
        "file_name": video_name or f"{workflow.lower()}_job",
        "video_name": video_name,
        "video_type": video_type,
        **media_preview,
        "created": now_stamp() if "now_stamp" in globals() else "",
        "source": "Subtitle / Transcription Editor",
        "workspace": user.get("workspace", "Demo Workspace"),
        "user_email": user.get("email", ""),
        "status": "draft",
        "workflow": workflow,
    }
    payload = {
        "job_id": job_id,
        "rows": rows or [],
        "metadata": metadata,
        "title": metadata["title"],
        "target_language": metadata.get("target_language", ""),
        "file_name": metadata.get("file_name", ""),
        "created": metadata["created"],
        "job_type": "media",
    }
    get_review_session_store()[job_id] = payload

    if save_editor_job is not None:
        try:
            save_editor_job("media", rows or [], metadata=metadata, job_id=job_id)
        except Exception as exc:
            LOGGER.warning("Unable to save media job to local editor store: %s", exc)
    if save_persistent_editor_job is not None:
        try:
            save_persistent_editor_job("media", rows or [], metadata=metadata, job_id=job_id, user=user)
        except Exception as exc:
            LOGGER.warning("Unable to save media job to persistent store: %s", exc)

    st.session_state["last_media_editor_job_id"] = job_id
    owner_job_record = {
        "id": job_id,
        "job_type": "media",
        "workspace": metadata.get("workspace", ""),
        "user_email": metadata.get("user_email", ""),
        "file_name": metadata.get("file_name", ""),
        "target_language": metadata.get("target_language", ""),
        "status": "draft",
        "row_count": len(rows or []),
        "created": metadata.get("created", ""),
        "updated_at": metadata.get("created", ""),
        "source": metadata.get("source", ""),
    }
    st.session_state.setdefault("owner_recent_editor_jobs", [])
    st.session_state["owner_recent_editor_jobs"] = [
        owner_job_record,
        *[j for j in st.session_state["owner_recent_editor_jobs"] if j.get("id") != job_id],
    ][:25]
    return job_id


def load_review_session_from_store(session_id: str) -> bool:
    """Load review rows from memory, Supabase persistence, or local fallback."""
    if not session_id:
        return False
    payload = get_review_session_store().get(session_id)

    if not payload and load_persistent_editor_job is not None:
        try:
            payload = load_persistent_editor_job(session_id)
        except Exception as exc:
            LOGGER.warning("Unable to load job %s from persistent store: %s", session_id, exc)
            payload = None

    if not payload and load_editor_job is not None:
        try:
            payload = load_editor_job(session_id)
        except Exception as exc:
            LOGGER.warning("Unable to load job %s from local editor store: %s", session_id, exc)
            payload = None

    if not payload:
        return False
    rows = payload.get("rows") or []
    if not rows:
        return False
    metadata = payload.get("metadata") or payload
    st.session_state.review_segments = rows
    st.session_state.last_pro_review_segments = rows
    st.session_state.latest_human_review_segments = rows
    st.session_state.pro_post_editing_ready = True
    st.session_state.selected_review_index = min(int(st.session_state.get("selected_review_index", 0) or 0), max(len(rows)-1, 0))
    st.session_state.review_workspace_title = metadata.get("title") or payload.get("title") or "ErrorSweep Pro"
    st.session_state.review_workspace_language = metadata.get("target_language") or payload.get("target_language") or ""
    st.session_state.review_workspace_file_name = metadata.get("file_name") or payload.get("file_name") or ""
    st.session_state.active_review_session_id = session_id
    return True


def prepare_human_review_session(rows: List[Dict[str, Any]], source: str = "ErrorSweep Pro", target_language: str = "", file_name: str = "", rules: Optional[Dict[str, Any]] = None) -> None:
    """Store rows in a durable Human Review session before opening the editor.

    This fixes the blank-page issue: the review workspace reads from
    st.session_state.review_segments, so Pro must always seed that state
    before routing to the dedicated Human Review Workspace page.
    """
    prepared = []
    for i, row in enumerate(rows, start=1):
        src = safe_text(row.get("source", ""))
        tgt = repair_localization_translation(src, safe_text(row.get("target", row.get("translation", ""))))
        status = safe_text(row.get("status", "MT" if tgt else "Needs Review"))
        match = safe_text(row.get("match", "MT" if tgt else "Untranslated"))
        prepared.append({
            "id": row.get("id", i),
            "location": row.get("location", f"Segment {i}"),
            "source": src,
            "target": tgt,
            "status": status,
            "match": match,
            "language": target_language,
            "file_name": file_name,
            "source_workflow": source,
            "notes": row.get("notes", ""),
            "start": row.get("start", ""),
            "end": row.get("end", ""),
        })
    # Store in more than one session key. Some Streamlit reruns can make a hidden
    # route render before the editor reads review_segments; these backup keys let
    # the workspace restore itself instead of opening as a blank page.
    st.session_state.review_segments = prepared
    st.session_state.last_pro_review_segments = prepared
    st.session_state.latest_human_review_segments = prepared
    st.session_state.pro_review_rows = prepared
    st.session_state.pro_post_edit_rows = prepared
    st.session_state.pro_post_edit_language = target_language
    st.session_state.pro_post_edit_file_name = file_name
    st.session_state.pro_post_editing_ready = True
    st.session_state.selected_review_index = 0
    st.session_state.review_workspace_title = source
    st.session_state.review_workspace_language = target_language
    st.session_state.review_workspace_file_name = file_name
    st.session_state.review_workspace_created = now_stamp() if "now_stamp" in globals() else ""
    session_id = save_review_session_to_store(prepared, source, target_language, file_name, rules=rules)
    query_set("review_id", session_id)


def restore_human_review_session_from_cache() -> bool:
    """Restore Pro review rows if the dedicated page is opened after a rerun.

    This avoids the confusing blank-page experience after clicking
    "Open Human Review workspace". The restore order is:
    1. current review_segments
    2. review_id query/session store
    3. backup session_state lists
    """
    if st.session_state.get("review_segments"):
        return True

    session_id = st.session_state.get("active_review_session_id") or query_get("review_id")
    if session_id and load_review_session_from_store(str(session_id)):
        return True

    for key in ("last_pro_review_segments", "latest_human_review_segments", "pro_review_rows", "pro_post_edit_rows"):
        cached = st.session_state.get(key)
        if isinstance(cached, list) and cached:
            st.session_state.review_segments = cached
            st.session_state.selected_review_index = 0
            st.session_state.pro_post_editing_ready = True
            return True
    return False


def go_to_human_review_workspace() -> None:
    """Route to the dedicated Pro post-editing workspace safely."""
    restore_human_review_session_from_cache()
    session_id = st.session_state.get("active_review_session_id")
    if session_id:
        query_set("review_id", str(session_id))
    open_page("Human Review Workspace")


def current_session_token_for_links() -> str:
    token = query_get("es_session")
    if token:
        return token
    user = current_user() or {}
    if not user:
        return ""
    payload = {**user, "exp": int(time.time()) + SESSION_TTL_SECONDS}
    try:
        return sign_payload(payload)
    except Exception as exc:
        LOGGER.debug("Unable to sign editor session link: %s", exc)
        return ""


def external_editor_url(editor_type: str, job_id: str) -> str:
    parts = []
    token = current_session_token_for_links()
    if token:
        parts.append(f"es_session={quote(token)}")
    parts.append(f"es_editor={quote(editor_type)}")
    parts.append(f"job_id={quote(str(job_id))}")
    return "?" + "&".join(parts)


def render_external_editor_link(label: str, editor_type: str, job_id: str) -> None:
    url = external_editor_url(editor_type, job_id)
    st.markdown(
        f"""
        <a href="{url}" target="_blank" style="
            display:flex; align-items:center; justify-content:center; width:100%;
            padding: 0.78rem 1rem; border-radius:14px; text-decoration:none;
            background: linear-gradient(90deg,#00d985,#34bdf6); color:#061018;
            font-weight:900; box-shadow:0 12px 30px rgba(52,189,246,.25);
        ">{escape(label)} ↗</a>
        """,
        unsafe_allow_html=True,
    )


def load_external_editor_payload(job_id: str) -> Optional[Dict[str, Any]]:
    if not job_id:
        return None

    # v42 production persistence first, so new browser tabs survive Streamlit restarts.
    if load_persistent_editor_job is not None:
        try:
            payload = load_persistent_editor_job(job_id)
            if payload:
                return payload
        except Exception as exc:
            LOGGER.warning("Unable to load external editor job %s from persistent store: %s", job_id, exc)

    # v41 local JSON fallback.
    if load_editor_job is not None:
        try:
            payload = load_editor_job(job_id)
            if payload:
                return payload
        except Exception as exc:
            LOGGER.warning("Unable to load external editor job %s from local store: %s", job_id, exc)
    return get_review_session_store().get(job_id)


def save_external_editor_payload(job_id: str, payload: Dict[str, Any]) -> None:
    if not job_id or not payload:
        return
    rows = payload.get("rows") or []
    metadata = payload.get("metadata") or {
        "title": payload.get("title", "ErrorSweep CAT"),
        "target_language": payload.get("target_language", ""),
        "file_name": payload.get("file_name", ""),
    }
    get_review_session_store()[job_id] = {
        "rows": rows,
        "metadata": metadata,
        "rules": payload.get("rules") or metadata.get("rules") or {},
        "title": metadata.get("title", "ErrorSweep CAT"),
        "target_language": metadata.get("target_language", ""),
        "file_name": metadata.get("file_name", ""),
        "created": metadata.get("created", ""),
        "job_type": payload.get("job_type", "cat"),
    }

    if update_editor_job is not None:
        try:
            update_editor_job(job_id, rows=rows, metadata=metadata)
        except Exception as exc:
            LOGGER.warning("Unable to update external editor job %s in local store: %s", job_id, exc)

    # v42 persistent save/update.
    if update_persistent_editor_job is not None:
        try:
            update_persistent_editor_job(job_id, rows=rows, metadata=metadata, status=metadata.get("status", "draft"))
        except Exception as exc:
            LOGGER.warning("Unable to update external editor job %s in persistent store: %s", job_id, exc)


def render_external_cat_editor(job_id: str) -> None:
    payload = load_external_editor_payload(job_id)
    if not payload:
        st.error("Editor job not found or expired. Please go back to ErrorSweep Pro and open the editor again.")
        return
    rows = payload.get("rows") or []
    metadata = payload.get("metadata") or payload
    if not rows:
        st.error("This editor job has no rows. Please rerun Pro translation and open the editor again.")
        return

    st.markdown(
        """
        <style>
        .block-container { max-width: 100vw !important; padding: .15rem .25rem .25rem .25rem !important; }
        .es-editor-shell { border: 1px solid rgba(148,163,184,.25); background:#080a12; min-height: calc(100vh - 10px); overflow:hidden; }
        .es-editor-top { height: 50px; display:flex; align-items:center; justify-content:space-between; padding:0 14px; background:#242a2f; border-bottom:1px solid rgba(255,255,255,.08); }
        .es-editor-brand { display:flex; align-items:center; gap:10px; font-weight:900; color:#fff; }
        .es-editor-logo { width:32px; height:32px; border-radius:10px; display:flex; align-items:center; justify-content:center; background:linear-gradient(135deg,#00d985,#34bdf6,#8b5cf6); color:#061018; font-weight:1000; }
        .es-editor-pill { display:inline-flex; border-radius:999px; padding:4px 10px; font-size:12px; font-weight:900; background:rgba(255,255,255,.06); border:1px solid rgba(255,255,255,.14); color:#e5edff; }
        .es-editor-pill.green { background:rgba(0,217,133,.14); border-color:rgba(0,217,133,.35); color:#63ffc4; }
        .es-editor-tabs { height:36px; display:flex; align-items:center; gap:0; background:#eef2f7; color:#0f172a; border-bottom:1px solid #cbd5e1; }
        .es-editor-tab { height:36px; padding:0 22px; display:flex; align-items:center; font-weight:800; font-size:13px; border-right:1px solid #cbd5e1; }
        .es-editor-tab.active { background:#fff; border-bottom:2px solid #00d985; }
        .es-context-preview { height:96px; background:#e9eef5; border-bottom:1px solid #cbd5e1; display:flex; justify-content:center; }
        .es-context-phone { width:430px; background:#1f2027; display:flex; align-items:center; justify-content:center; }
        .es-context-highlight { padding:8px 54px; border-radius:999px; background:#4b5563; color:#7dd3fc; border:2px dashed #7dd3fc; font-weight:800; }
        .es-formatbar { height:32px; display:flex; align-items:center; gap:12px; background:#313940; color:#fff; border-bottom:1px solid rgba(255,255,255,.10); padding:0 10px; font-size:13px; }
        .es-side-card { background:#f5f6f8; color:#1f2937; border-left:1px solid #cbd5e1; height: calc(100vh - 254px); overflow-y:auto; }
        .es-side-head { display:flex; align-items:center; justify-content:space-between; padding:10px 14px; border-bottom:1px solid #cbd5e1; font-weight:900; }
        .es-side-section { padding:14px; border-bottom:1px solid #d8dee8; }
        .es-side-title { font-size:13px; font-weight:900; color:#4b5563; margin-bottom:8px; }
        .es-resource { display:grid; grid-template-columns:38px 1fr; border:1px solid #d1d5db; background:#fff; margin-bottom:8px; }
        .es-resource-code { background:#cbd5e1; display:flex; align-items:center; justify-content:center; font-weight:900; font-size:12px; }
        .es-resource-body { padding:8px; font-size:13px; }
        .es-muted { color:#64748b; font-size:12px; }
        div[data-testid="stDataEditor"] { border-radius:0 !important; border:0 !important; contain:paint !important; isolation:isolate !important; overflow:hidden !important; }
        div[data-testid="stDataEditor"] textarea, div[data-testid="stDataEditor"] input { font-size:14px !important; }
        div[data-testid="stDataEditor"] textarea:not(:focus), div[data-testid="stDataEditor"] input:not(:focus) { opacity:0 !important; pointer-events:none !important; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    title = safe_text(metadata.get("title", "ErrorSweep CAT")) or "ErrorSweep CAT"
    file_name = safe_text(metadata.get("file_name", "translation_job")) or "translation_job"
    language = safe_text(metadata.get("target_language", "Target")) or "Target"
    completion = compute_review_completion(rows)

    st.markdown(
        f"""
        <div class="es-editor-shell">
          <div class="es-editor-top">
            <div class="es-editor-brand"><div class="es-editor-logo">ES</div><div><div>{escape(file_name)}</div><div class="es-muted">{escape(title)} · Target: {escape(language)} · Job: {escape(job_id[:10])}</div></div></div>
            <div style="display:flex; align-items:center; gap:8px;"><span class="es-editor-pill green">Accepted</span><span class="es-editor-pill">TM</span><span class="es-editor-pill">TB</span><span class="es-editor-pill">MT</span></div>
          </div>
          <div class="es-editor-tabs"><div class="es-editor-tab active">Context</div><div class="es-editor-tab">Quality Checks</div><div class="es-editor-tab">Search TM</div><div class="es-editor-tab">Glossary</div><div style="margin-left:auto; padding-right:14px; font-size:13px;">Mode: <b>Highlight strings</b></div></div>
          <div class="es-context-preview"><div class="es-context-phone"><div class="es-context-highlight">Open account settings</div></div></div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    toolbar_cols = st.columns([0.45, 0.18, 0.13, 0.10, 0.14], gap="small")
    with toolbar_cols[0]:
        search = st.text_input("Search", placeholder="Search source and translations", label_visibility="collapsed", key=f"ext_cat_search_{job_id}")
    with toolbar_cols[1]:
        status_filter = st.selectbox("Status", ["All", "MT", "Needs Review", "Approved", "Untranslated", "Fuzzy 75%", "Fuzzy 85%", "100%", "101%", "Rejected", "Needs Rework"], label_visibility="collapsed", key=f"ext_cat_status_{job_id}")
    with toolbar_cols[2]:
        pending_only = st.checkbox("Pending", value=False, key=f"ext_cat_pending_{job_id}")
    with toolbar_cols[3]:
        st.metric("Rows", len(rows))
    with toolbar_cols[4]:
        st.metric("Approved", completion["approved"])

    st.markdown('<div class="es-formatbar"><b>B</b><i>I</i><u>U</u><span>↶</span><span>↷</span><span>BR</span><span>NBSP</span><span>✓</span></div>', unsafe_allow_html=True)

    filtered_indexes = []
    needle = safe_text(search).lower().strip()
    for i, r in enumerate(rows):
        src = safe_text(r.get("source", ""))
        tgt = safe_text(r.get("target", ""))
        status = safe_text(r.get("status", "Needs Review")) or "Needs Review"
        if needle and needle not in src.lower() and needle not in tgt.lower():
            continue
        if status_filter != "All" and status != status_filter:
            continue
        if pending_only and status in {"Approved", "100%", "101%"}:
            continue
        filtered_indexes.append(i)
    if not filtered_indexes:
        filtered_indexes = list(range(len(rows)))

    grid_rows = []
    for i in filtered_indexes:
        r = rows[i]
        status = safe_text(r.get("status", "Needs Review")) or "Needs Review"
        sensitive_kinds = detect_sensitive_text(f"{r.get('source', '')} {r.get('target', '')}")
        grid_rows.append({
            "No": i + 1,
            "Source (EN)": safe_text(r.get("source", "")),
            "Target": repair_localization_translation(r.get("source", ""), r.get("target", "")),
            "Match": safe_text(r.get("match", "MT")) or "MT",
            "Status": status,
            "QA": "PII" if sensitive_kinds else ("OK" if status in {"Approved", "100%", "101%"} else "Open"),
            "Notes": safe_text(r.get("notes", "")),
            "Location": safe_text(r.get("location", f"Segment {i+1}")),
        })

    grid_col, side_col = st.columns([0.81, 0.19], gap="small")
    with grid_col:
        edited_df = st.data_editor(
            pd.DataFrame(grid_rows),
            use_container_width=True,
            hide_index=True,
            height=620,
            num_rows="fixed",
            disabled=["No", "Source (EN)", "Match", "QA", "Location"],
            column_order=["No", "Source (EN)", "Target", "Match", "QA", "Status", "Notes", "Location"],
            column_config={
                "No": st.column_config.NumberColumn("#", width="small"),
                "Source (EN)": st.column_config.TextColumn("Source (EN)", width="large"),
                "Target": st.column_config.TextColumn("Target", width="large"),
                "Match": st.column_config.TextColumn("Match", width="small"),
                "QA": st.column_config.TextColumn("QA", width="small"),
                "Status": st.column_config.SelectboxColumn("Status", options=["MT", "Fuzzy 75%", "Fuzzy 85%", "100%", "101%", "Needs Review", "Approved", "Rejected", "Needs Rework", "Untranslated"], width="medium"),
                "Notes": st.column_config.TextColumn("Notes", width="medium"),
                "Location": st.column_config.TextColumn("Location", width="medium"),
            },
            key=f"external_cat_grid_{job_id}",
        )

        current_rows = [dict(r) for r in rows]
        for _, erow in edited_df.iterrows():
            idx = int(erow["No"]) - 1
            if 0 <= idx < len(current_rows):
                current_rows[idx]["target"] = safe_text(erow.get("Target", ""))
                current_rows[idx]["translation"] = safe_text(erow.get("Target", ""))
                current_rows[idx]["status"] = safe_text(erow.get("Status", "")) or "Needs Review"
                current_rows[idx]["notes"] = safe_text(erow.get("Notes", ""))
                current_rows[idx].setdefault("location", f"Segment {idx+1}")
        gate_rules = payload.get("rules") or metadata.get("rules") or workspace_rules()
        gate_findings = delivery_quality_findings(current_rows, language, "Human Review", gate_rules)
        gate_summary = render_delivery_gate(current_rows, gate_findings, "Pre-delivery quality gate")

        action_cols = st.columns([1, 1, 1, 1, 1])
        if action_cols[0].button("Save Page", type="primary", use_container_width=True, key=f"ext_cat_save_{job_id}"):
            for _, erow in edited_df.iterrows():
                idx = int(erow["No"]) - 1
                if 0 <= idx < len(rows):
                    rows[idx]["target"] = safe_text(erow.get("Target", ""))
                    rows[idx]["status"] = safe_text(erow.get("Status", "")) or "Needs Review"
                    rows[idx]["notes"] = safe_text(erow.get("Notes", ""))
            payload["rows"] = rows
            save_external_editor_payload(job_id, payload)
            st.success("Saved editor changes.")
        if action_cols[1].button("Approve visible", use_container_width=True, key=f"ext_cat_approve_{job_id}"):
            for _, erow in edited_df.iterrows():
                idx = int(erow["No"]) - 1
                if 0 <= idx < len(rows):
                    rows[idx]["target"] = safe_text(erow.get("Target", ""))
                    rows[idx]["status"] = "Approved" if safe_text(erow.get("Target", "")).strip() else "Needs Review"
                    rows[idx]["notes"] = safe_text(erow.get("Notes", ""))
            payload["rows"] = rows
            save_external_editor_payload(job_id, payload)
            st.success("Visible rows approved.")
        if action_cols[2].button("Submit", use_container_width=True, key=f"ext_cat_submit_{job_id}"):
            if not gate_summary.get("ready"):
                st.error("Submit blocked by the pre-delivery quality gate. Resolve blocking QA issues first.")
                st.stop()
            for r in rows:
                if safe_text(r.get("target", "")).strip() and safe_text(r.get("status", "")) not in {"Rejected", "Needs Rework"}:
                    r["status"] = "Approved"
            payload["rows"] = rows
            payload.setdefault("metadata", metadata)["submitted_at"] = now_stamp()
            save_external_editor_payload(job_id, payload)
            st.success("Submitted reviewed job.")
        if action_cols[3].button("Refresh", use_container_width=True, key=f"ext_cat_refresh_{job_id}"):
            st.rerun()
        if action_cols[4].button("Back to Pro", use_container_width=True, key=f"ext_cat_back_{job_id}"):
            query_clear("es_editor")
            query_clear("job_id")
            open_page("ErrorSweep Pro")

        dl1, dl2, dl3 = st.columns(3)
        base = re.sub(r"\.[^.]+$", "", file_name) or "reviewed_translation"
        dl1.download_button("Download reviewed Excel", build_reviewed_translation_workbook(current_rows), file_name=f"{base}_reviewed.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, disabled=not gate_summary.get("ready"))
        dl2.download_button("Download reviewed CSV", rows_to_csv(current_rows), file_name=f"{base}_reviewed.csv", mime="text/csv", use_container_width=True, disabled=not gate_summary.get("ready"))
        dl3.download_button("Download target text", build_reviewed_plain_text(current_rows), file_name=f"{base}_target.txt", mime="text/plain", use_container_width=True, disabled=not gate_summary.get("ready"))
        if not gate_summary.get("ready"):
            st.download_button("Download QA gate findings", rows_to_csv(gate_findings), file_name=f"{base}_qa_gate_findings.csv", mime="text/csv", use_container_width=True)

    with side_col:
        st.markdown('<div class="es-side-card">', unsafe_allow_html=True)
        st.markdown('<div class="es-side-head"><span>Additional Details</span><span class="es-editor-pill" style="background:#f97316;color:white;">1</span></div>', unsafe_allow_html=True)
        selected_no = int(edited_df.iloc[0]["No"]) if not edited_df.empty else 1
        selected_idx = max(0, min(selected_no - 1, len(rows) - 1))
        selected = rows[selected_idx]
        matches = compute_matches(safe_text(selected.get("source", "")))
        st.markdown('<div class="es-side-section"><div class="es-side-title">Language Resources</div>', unsafe_allow_html=True)
        if matches["glossary"]:
            for g in matches["glossary"][:5]:
                st.markdown(f'<div class="es-resource"><div class="es-resource-code">GT</div><div class="es-resource-body"><b>{escape(g.get("target", ""))}</b><br><span class="es-muted">{escape(g.get("source", ""))}</span></div></div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="es-resource"><div class="es-resource-code">GT</div><div class="es-resource-body"><b>Glossary</b><br><span class="es-muted">No glossary hit.</span></div></div>', unsafe_allow_html=True)
        if matches["tm"]:
            for m in matches["tm"][:3]:
                st.markdown(f'<div class="es-resource"><div class="es-resource-code">TM</div><div class="es-resource-body"><b>{escape(m.get("type", "TM"))}</b><br><span class="es-muted">{escape(m.get("target", ""))}</span></div></div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown('<div class="es-side-section"><div class="es-side-title">Quality Checks</div><div class="es-resource"><div class="es-resource-code">QA</div><div class="es-resource-body">Check placeholders, glossary, DNT, punctuation, and number consistency before submit.</div></div></div>', unsafe_allow_html=True)
        st.markdown(f'<div class="es-side-section"><div class="es-side-title">Selected Row</div><div class="es-resource"><div class="es-resource-code">#</div><div class="es-resource-body"><b>{selected_idx+1}</b><br><span class="es-muted">{escape(safe_text(selected.get("location", "")))}</span></div></div></div>', unsafe_allow_html=True)
        st.markdown('<div class="es-side-section"><div class="es-side-title">Issues</div><button style="width:100%;border:1px solid #94a3b8;background:white;padding:8px;font-weight:800;">Open New Issue</button><div class="es-muted" style="margin-top:8px;">View related source issues (0)</div></div>', unsafe_allow_html=True)
        st.markdown('<div class="es-side-section"><div class="es-side-title">History</div><div class="es-muted">Saved changes are stored under this job_id.</div></div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)


def render_external_media_editor(job_id: str) -> None:
    payload = load_external_editor_payload(job_id)
    if not payload:
        st.error("Media editor job not found or expired.")
        return
    rows = payload.get("rows") or []
    metadata = payload.get("metadata") or payload
    workflow = metadata.get("workflow") or metadata.get("title", "Media Editor")
    file_name = metadata.get("file_name", "media_job")
    target_language = metadata.get("target_language", "")

    st.markdown(
        """
        <style>
        [data-testid="stSidebar"] {display:none !important;}
        [data-testid="stHeader"] {display:none !important;}
        .block-container {max-width: 100% !important; padding: 0.6rem 0.8rem 1rem !important;}
        .es-media-top {display:flex;justify-content:space-between;align-items:center;background:#242a2f;color:white;border:1px solid #334155;border-radius:10px;padding:10px 14px;margin-bottom:10px;}
        .es-media-title {font-weight:900;font-size:15px;}
        .es-media-sub {font-size:12px;color:#cbd5e1;margin-top:2px;}
        .es-media-pill {background:#0ea5e9;color:#061018;border-radius:999px;padding:5px 10px;font-weight:900;font-size:12px;}
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(
        f"""
        <div class="es-media-top">
          <div><div class="es-media-title">ErrorSweep Media Editor</div><div class="es-media-sub">{escape(str(workflow))} · {escape(str(file_name))} · {escape(str(target_language))}</div></div>
          <div class="es-media-pill">Separate Editor</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if not rows:
        st.warning("No media rows found for this editor job.")
        return

    media_bytes, media_type, media_name = read_media_preview_bytes(metadata)
    preview_col, job_col = st.columns([0.44, 0.56], gap="large")
    with preview_col:
        render_media_preview(media_bytes, media_type, media_name or file_name)
        render_waveform_preview(rows, 0)
    with job_col:
        st.caption("Edit timing and transcript/subtitle text directly in the grid. Keep the media preview visible while reviewing timing and transcript/subtitle text.")
        st.metric("Rows", len(rows))
        st.metric("Approved", sum(1 for r in rows if safe_text(r.get("status", "")) == "Approved"))
        timing_issues = validate_timing_rows(rows)
        if timing_issues:
            st.warning(f"{len(timing_issues)} timing issue(s) need review.")

    df = pd.DataFrame(rows)
    wanted = [c for c in ["id", "start", "end", "source", "target", "status", "match"] if c in df.columns]
    if wanted:
        df = df[wanted]

    render_segment_timeline(rows, 0)

    column_config = {}
    if "target" in df.columns:
        column_config["target"] = st.column_config.TextColumn("Target Subtitle / Transcript", width="large")
    if "source" in df.columns:
        column_config["source"] = st.column_config.TextColumn("Source / Speaker Note", width="large", disabled=True)
    if "start" in df.columns:
        column_config["start"] = st.column_config.NumberColumn("Start", format="%.3f")
    if "end" in df.columns:
        column_config["end"] = st.column_config.NumberColumn("End", format="%.3f")

    edited = st.data_editor(
        df,
        use_container_width=True,
        hide_index=True,
        height=470,
        num_rows="dynamic",
        column_config=column_config,
        disabled=[c for c in df.columns if c not in {"target", "start", "end", "status", "match"}],
        key=f"external_media_grid_{job_id}",
    )

    c1, c2, c3, c4 = st.columns(4)
    if c1.button("Save Draft", use_container_width=True, key=f"media_save_{job_id}"):
        new_rows = edited.to_dict(orient="records")
        payload["rows"] = new_rows
        save_external_editor_payload(job_id, payload)
        st.success("Media editor draft saved.")
    c2.download_button("Download CSV", rows_to_csv(edited.to_dict(orient="records")), file_name=f"{re.sub(r'[^A-Za-z0-9_-]+','_', str(file_name))}_media_editor.csv", mime="text/csv", use_container_width=True)
    c3.download_button("Download SRT", rows_to_srt(edited.to_dict(orient="records"), use_target=True), file_name=f"{re.sub(r'[^A-Za-z0-9_-]+','_', str(file_name))}_output.srt", mime="text/plain", use_container_width=True)
    if c4.button("Back", use_container_width=True, key=f"media_back_{job_id}"):
        query_clear("es_editor")
        query_clear("job_id")
        open_page("Subtitle / Transcription Editor")


def render_external_editor_router() -> bool:
    editor_type = query_get("es_editor")
    if not editor_type:
        return False
    job_id = query_get("job_id")
    if editor_type == "cat":
        render_external_cat_editor(job_id)
        return True
    if editor_type == "media":
        render_external_media_editor(job_id)
        return True
    st.error("Unknown editor route.")
    return True



def build_reviewed_translation_workbook(rows: List[Dict[str, Any]]) -> bytes:
    """Excel output from Human Review. Always safe and easy for clients to review."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Reviewed Translation"
    headers = ["Segment ID", "Location", "Source Text", "Final Translation", "Status", "Match", "Language", "Notes"]
    ws.append(headers)
    for row in rows:
        source_text = safe_text(row.get("source", ""))
        target_text = repair_localization_translation(source_text, row.get("target", ""))
        ws.append([
            row.get("id", ""),
            row.get("location", ""),
            source_text,
            target_text,
            row.get("status", ""),
            row.get("match", ""),
            row.get("language", ""),
            row.get("notes", ""),
        ])
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        cell.alignment = Alignment(wrap_text=True, vertical="top")
    widths = {"A": 12, "B": 22, "C": 55, "D": 55, "E": 18, "F": 14, "G": 16, "H": 28}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_reviewed_plain_text(rows: List[Dict[str, Any]]) -> bytes:
    return "\n".join(repair_localization_translation(r.get("source", ""), r.get("target", "")) for r in rows).encode("utf-8-sig")


def compute_review_completion(rows: List[Dict[str, Any]]) -> Dict[str, int]:
    total = len(rows)
    translated = sum(1 for r in rows if safe_text(r.get("target", "")).strip())
    approved = sum(1 for r in rows if safe_text(r.get("status", "")) == "Approved")
    needs_review = sum(1 for r in rows if "Review" in safe_text(r.get("status", "")) or not safe_text(r.get("target", "")).strip())
    return {"total": total, "translated": translated, "approved": approved, "needs_review": needs_review}


def compute_matches(source: str) -> Dict[str, List[Dict[str, str]]]:
    source_l = source.lower()
    tm_hits = []
    for item in st.session_state.tm:
        if item.get("source", "").lower() == source_l:
            tm_hits.append({"type": "TM 100%", "source": item.get("source",""), "target": item.get("target","")})
        elif source_l and item.get("source", "").lower() in source_l:
            tm_hits.append({"type": "TM fuzzy", "source": item.get("source",""), "target": item.get("target","")})

    gloss_hits = []
    for item in st.session_state.glossary:
        if item.get("source","").lower() in source_l:
            gloss_hits.append(item)

    dnt_hits = [{"term": term} for term in st.session_state.dnt if term.lower() in source_l]
    return {"tm": tm_hits[:5], "glossary": gloss_hits[:8], "dnt": dnt_hits[:8]}


def openai_client() -> Optional[OpenAI]:
    key = secret("OPENAI_API_KEY")
    if not key:
        return None
    return OpenAI(api_key=key)


def current_ai_route_label() -> str:
    """Small safe label for UI. Do not expose provider names, URLs, or tokens."""
    if st.session_state.get("byo_openai_api_key"):
        provider = safe_text(st.session_state.get("byo_ai_provider", "Custom API"))
        model = safe_text(st.session_state.get("byo_openai_model", ""))
        return f"BYO {provider} active" + (f" ({model})" if model else "")
    if current_builtin_engine_label is not None:
        try:
            return current_builtin_engine_label()
        except Exception as exc:
            LOGGER.warning("Unable to read built-in translation engine label: %s", exc)
    return "Translation engine not configured"


def log_ai_usage_event(usage: Dict[str, Any], purpose: str, segment_count: int = 0) -> None:
    """Log AI/translation usage in session and, when configured, Supabase.

    Persists self-hosted MT usage and BYO/Managed AI events for owner reporting.
    """
    st.session_state.setdefault("ai_usage_events", [])
    user = current_user() or {}
    record = {
        "time": now_stamp() if "now_stamp" in globals() else datetime.now().strftime("%Y-%m-%d %H:%M"),
        "user_email": user.get("email", ""),
        "workspace": user.get("workspace", "Demo Workspace"),
        "purpose": purpose,
        "provider": usage.get("provider", usage.get("engine", "unknown")),
        "model": usage.get("model", usage.get("engine", "")),
        "managed": usage.get("managed", False),
        "billable": usage.get("billable", False),
        "usage_kind": usage.get("usage_kind", "ai_route"),
        "input_tokens": usage.get("input_tokens", 0),
        "output_tokens": usage.get("output_tokens", 0),
        "total_tokens": usage.get("total_tokens", 0),
        "characters": usage.get("characters", 0),
        "requests": usage.get("requests", 0),
        "success": usage.get("success", False),
        "error": usage.get("error", ""),
        "segments": segment_count,
    }
    st.session_state.ai_usage_events.insert(0, record)
    trim_session_list("ai_usage_events")

    if log_persistent_usage_event is not None:
        try:
            log_persistent_usage_event(
                record,
                purpose=purpose,
                segment_count=segment_count,
                user=user,
                metadata={
                    "app_version": APP_VERSION,
                    "billable": record.get("billable", False),
                    "usage_kind": record.get("usage_kind", "ai_route"),
                },
            )
        except Exception as exc:
            LOGGER.error("Unable to persist usage event: %s", exc)


def repair_localization_translation(source: str, translation: str) -> str:
    """Restore non-translatable localization tokens after MT/AI output.

    MT engines may alter protection markers by case (PH000Token) or drop visual
    prefixes. This final pass is source-driven and conservative: it only restores
    placeholders, tags, URLs/emails, emoji/icon prefixes, and wrapper brackets.
    """
    source = safe_text(source)
    out = safe_text(translation)
    if not source:
        return out

    out = re.sub(r"^\s*[ÂÃ]\s+", "", out).strip()

    protected = LOCALIZATION_PROTECTED_RE.findall(source)
    for idx, token in enumerate(protected):
        marker_re = re.compile(
            rf"(?i)[_\s]*(?:P\s*H\s*{idx:03d}\s*T\s*O\s*K\s*E\s*N|E\s*S\s*P\s*H\s*{idx}|Z\s*X\s*P\s*H\s*{idx}\s*Z\s*X|ZXPH\s*{idx}\s*ZX)[_\s]*"
        )
        out = marker_re.sub(token, out)
        if token and token not in out:
            out = f"{out.rstrip()} {token}".strip()

    prefix_match = LOCALIZATION_VISUAL_PREFIX_RE.match(source)
    if prefix_match:
        prefix = prefix_match.group(1)
        if prefix.strip() and not out.startswith(prefix.strip()):
            out = prefix + out.lstrip()

    for symbol in LOCALIZATION_EMOJI_RE.findall(source):
        if symbol and symbol not in out:
            out = f"{symbol} {out}".strip()

    src = source.strip()
    stripped = out.strip()
    if src.startswith("[") and src.endswith("]") and stripped and not (stripped.startswith("[") and stripped.endswith("]")):
        out = f"[{stripped.strip('[] ')}]"

    return out.strip()


def repair_translation_batch(sources: List[str], translations: List[str]) -> List[str]:
    repaired = []
    for idx, source in enumerate(sources):
        translation = translations[idx] if idx < len(translations) else ""
        repaired.append(repair_localization_translation(source, translation))
    return repaired


def call_main_api_translate(texts: List[str], target_language: str, domain: str, rules: Optional[Dict[str, Any]] = None) -> List[str]:
    """Translate through the v30 two-phase backend.

    User has BYO key:
        Use the selected OpenAI-compatible provider/base URL/model from Account.

    User has no key:
        Use commercial-safe self-hosted MT: IndicTrans2 for Indian languages, OPUS-MT for supported global pairs.
    """
    if not texts:
        return []

    user_key = str(st.session_state.get("byo_openai_api_key", "") or "").strip()
    rules_text = rules_summary_for_ai(rules)

    # ----------------------------------------------------------
    # BYO KEY PATH - user can provide OpenAI or any OpenAI-compatible API route.
    # ----------------------------------------------------------
    if user_key:
        if ai_json_items is None or select_ai_route is None:
            st.error("BYO-key AI router file is missing. Add managed_ai_router.py beside app.py.")
            return ["" for _ in texts]

        try:
            route = select_ai_route(user_openai_key=user_key, purpose="translate")
        except Exception as exc:
            st.error(f"User AI route is not configured: {exc}")
            return ["" for _ in texts]

        system_prompt = f"""
You are ErrorSweep AI, a professional localization translator.
Return JSON only. Do not use markdown.

Task:
Translate source strings into {target_language} for a {domain} localization project.

Hard rules:
1. Preserve placeholders exactly: {{{{email}}}}, {{{{password}}}}, %s, %d, <tags>, URLs, emails.
2. Preserve numbers, units, emoji/icons, bullets, and product names.
3. Obey uploaded rules ZIP and saved Memory & Rules. Preserve DNT/client locked terms if they appear.
4. Square bracket UI labels may be localized inside brackets, but keep the bracket structure.
5. Do not leave translations blank.
6. Return a JSON object with key "items".

Client rules:
{rules_text or "No uploaded or saved client rules were provided."}

Output shape:
{{
  "items": [
    {{"index": 0, "translation": "translated text"}}
  ]
}}
"""
        payload = {
            "target_language": target_language,
            "domain": domain,
            "rules_summary": rules_text,
            "texts": [{"index": i, "source": text} for i, text in enumerate(texts)],
        }

        items, usage = ai_json_items(
            system_prompt=system_prompt,
            user_prompt=json.dumps(payload, ensure_ascii=False),
            route=route,
            temperature=0.0,
            max_tokens=4500,
        )
        log_ai_usage_event(usage, "translate", len(texts))

        result = [""] * len(texts)
        for item in items:
            try:
                idx = int(item.get("index", 0))
            except Exception as exc:
                LOGGER.debug("Skipping malformed translation item index: %s", exc)
                continue
            if 0 <= idx < len(result):
                result[idx] = safe_text(item.get("translation", ""))

        if not any(result) and usage.get("error"):
            st.error(f"Translation service error: {usage.get('error')}")
        return repair_translation_batch(texts, result)

    # ----------------------------------------------------------
    # NO USER KEY PATH: commercial-safe self-hosted MT.
    # IndicTrans2 for Indian languages, MADLAD-400 when enabled, and OPUS-MT fallback.
    # ----------------------------------------------------------
    if builtin_translate_batch is None:
        st.error("Built-in translation router is missing. Add translator_router.py and selfhosted_mt_clients.py beside app.py.")
        return ["" for _ in texts]

    try:
        translations, usage = builtin_translate_batch(
            source_language="English",
            target_language=target_language,
            texts=texts,
            user_api_key="",
            protected_terms=protected_terms_from_rules(rules),
            metadata={
                "domain": domain,
                "rules_glossary_terms": len((rules or {}).get("glossary", [])),
                "rules_dnt_terms": len((rules or {}).get("dnt", [])),
                "rules_instruction_count": len((rules or {}).get("instructions", [])),
            },
        )
        # Store usage in the same owner-console log format.
        log_ai_usage_event({
            "provider": usage.get("provider", "built_in_translation"),
            "model": usage.get("engine", "built_in_translation"),
            "managed": True,
            "input_tokens": 0,
            "output_tokens": 0,
            "total_tokens": 0,
            "success": usage.get("success", True),
            "error": usage.get("error", ""),
            "characters": usage.get("characters", 0),
            "requests": usage.get("requests", 0),
        }, "translate", len(texts))
        return repair_translation_batch(texts, [safe_text(t) for t in translations])
    except Exception as exc:
        st.error(f"Translation service error: {str(exc)}")
        return ["" for _ in texts]



def generate_transcription_rows_from_video(video_file, locale: str = "en-US", prompt: str = "") -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    """Generate transcript rows from uploaded video/audio.

    v32 policy:
    - Auto transcription is available only when the user provides a BYO OpenAI-compatible API key.
    - Built-in self-hosted MT is text-only in ErrorSweep and is NOT used for transcription.
    - If no user API key is available, return blank starter rows for manual human editing.
    """
    if video_file is None:
        return default_subtitle_segments(10, transcription=True), {"success": False, "error": "No video uploaded."}

    user_key = str(st.session_state.get("byo_openai_api_key", "") or "").strip()
    if not user_key:
        return default_subtitle_segments(10, transcription=True), {
            "success": False,
            "provider": "manual_transcription",
            "engine": "manual_editor",
            "error": "No user API key available. Blank rows were created for manual transcription.",
        }

    if transcribe_media_file_to_rows is None:
        return default_subtitle_segments(10, transcription=True), {"success": False, "error": "speech_transcription.py is missing."}

    media_record = save_media_preview_file(f"transcribe_{uuid.uuid4().hex}", video_file)
    media_path = safe_text(media_record.get("media_preview_path"))
    if not media_path:
        return default_subtitle_segments(10, transcription=True), {
            "success": False,
            "provider": "manual_transcription",
            "engine": "manual_editor",
            "error": "Unable to save media file for transcription.",
        }

    rows, usage = transcribe_media_file_to_rows(
        media_path=media_path,
        filename=getattr(video_file, "name", "video.mp4"),
        mime_type=getattr(video_file, "type", "video/mp4") or "video/mp4",
        user_openai_key=user_key,
        locale=locale,
        prompt=prompt,
    )
    st.session_state.setdefault("ai_usage_events", [])
    user = current_user() or {}
    st.session_state.ai_usage_events.insert(0, {
        "time": now_stamp(),
        "user_email": user.get("email", ""),
        "workspace": user.get("workspace", "Demo Workspace"),
        "purpose": "transcription",
        "provider": usage.get("provider", "user_speech"),
        "model": usage.get("engine", "speech"),
        "managed": False,
        "billable": False,
        "usage_kind": "ai_route",
        "input_tokens": 0,
        "output_tokens": 0,
        "total_tokens": 0,
        "characters": sum(len(safe_text(row.get("source", ""))) for row in rows),
        "success": usage.get("success", False),
        "error": usage.get("error", ""),
        "segments": len(rows),
    })
    trim_session_list("ai_usage_events")
    return rows, usage


def translate_subtitle_sources(rows: List[Dict[str, Any]], target_language: str, domain: str = "Subtitling", rules: Optional[Dict[str, Any]] = None) -> Tuple[List[Dict[str, Any]], int]:
    """Translate source rows into target subtitle rows using BYO key or built-in self-hosted MT."""
    source_texts = [safe_text(r.get("source", "")) for r in rows]
    translations = call_main_api_translate(source_texts, target_language, domain, rules=rules)
    missing = 0
    for row, trans in zip(rows, translations):
        row["target"] = safe_text(trans)
        if row["target"]:
            row["status"] = "MT"
            row["match"] = "MT"
        else:
            row["status"] = "Needs Review"
            row["match"] = "Untranslated"
            missing += 1
    return rows, missing

def call_main_api_qa(rows: List[Dict[str, Any]], domain: str, strictness: str = "Standard", rules: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    """Optional Managed AI QA route. Deterministic checks still run first elsewhere."""
    if not rows:
        return []
    if ai_json_items is None or select_ai_route is None:
        return []
    try:
        route = select_ai_route(
            user_openai_key=st.session_state.get("byo_openai_api_key", ""),
            purpose="qa",
        )
    except Exception as exc:
        LOGGER.warning("Unable to select QA AI route: %s", exc)
        return []

    rules_text = rules_summary_for_ai(rules)
    system_prompt = f"""
You are ErrorSweep AI, a conservative localization QA reviewer.
Return JSON only. Do not invent errors.
Flag only real issues supported by source/target evidence.
DNT/placeholder/number damage is severe. Empty target is Critical.
Apply uploaded rules ZIP and saved Memory & Rules when judging glossary, DNT, style, and client instructions.

Client rules:
{rules_text or "No uploaded or saved client rules were provided."}

Output shape:
{{"items":[{{"id":1,"issue":"short issue","severity":"Minor|Major|Critical","suggestion":"fix","reason":"why"}}]}}
"""
    payload = {
        "domain": domain,
        "strictness": strictness,
        "rules_summary": rules_text,
        "segments": [
            {"id": r.get("id"), "source": r.get("source", ""), "target": r.get("target", "")}
            for r in rows[:80]
        ],
    }
    items, usage = ai_json_items(
        system_prompt=system_prompt,
        user_prompt=json.dumps(payload, ensure_ascii=False),
        route=route,
        temperature=0.0,
        max_tokens=3000,
    )
    log_ai_usage_event(usage, "qa", len(rows[:80]))
    return items

def auto_detect_domain(text_sample: str) -> str:
    t = text_sample.lower()
    if any(x in t for x in ["button", "dashboard", "settings", "login", "password", "screen", "menu", "tooltip"]):
        return "Software UI"
    if any(x in t for x in ["subtitle", "-->", "caption", "dialogue", "scene"]):
        return "Subtitling"
    if any(x in t for x in ["invoice", "payment", "bank", "account", "revenue"]):
        return "Finance"
    if any(x in t for x in ["agreement", "clause", "contract", "legal"]):
        return "Legal"
    if any(x in t for x in ["course", "lesson", "quiz", "module"]):
        return "E-learning"
    if any(x in t for x in ["campaign", "brand", "copy", "ad", "landing page"]):
        return "Marketing"
    return "General"


def detect_sensitive_text(text: str) -> List[str]:
    value = safe_text(text)
    if not value:
        return []
    hits = []
    if re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", value):
        hits.append("email")
    if re.search(r"\b(?:\+?\d[\d .()-]{7,}\d)\b", value):
        hits.append("phone/number")
    if re.search(r"\b(?:api[_-]?key|secret|token|password|passwd|pwd)\b\s*[:=]", value, re.I):
        hits.append("credential-like text")
    if not hits and SENSITIVE_TEXT_RE.search(value):
        hits.append("sensitive pattern")
    return hits


def sensitive_rows_summary(rows: List[Dict[str, Any]], limit: int = 8) -> Dict[str, Any]:
    matches = []
    for idx, row in enumerate(rows or []):
        joined = " ".join([safe_text(row.get("source", "")), safe_text(row.get("target", ""))])
        kinds = detect_sensitive_text(joined)
        if kinds:
            matches.append({"row": idx + 1, "kinds": sorted(set(kinds))})
        if len(matches) >= limit:
            break
    return {"count": len(matches), "matches": matches}


def render_privacy_route_notice(context: str = "translation") -> None:
    if st.session_state.get("byo_openai_api_key"):
        st.warning(f"{context}: BYO external AI key is active. Confirm client approval before sending confidential or NDA text.")
    else:
        st.success(f"{context}: built-in self-hosted MT route is active for no-key workflows.")


def run_global_qa_for_row(row: Dict[str, Any], target_language: str, domain: str, rules: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    if deterministic_checks_v2 is None:
        return []
    segment = {
        "id": row.get("id", ""),
        "location": row.get("location", f"Segment {row.get('id', '')}"),
        "source": row.get("source", ""),
        "translation": row.get("target", row.get("translation", "")),
        "mode": domain,
    }
    try:
        return deterministic_checks_v2(
            segment,
            rules=rules or {},
            target_language=target_language,
            domain=domain,
            enable_language_tool=False,
            language_tool_mode="local",
        )
    except Exception as exc:
        LOGGER.warning("Global QA failed for row %s: %s", row.get("id", ""), exc)
        return []


def summarize_qa_findings(findings: List[Dict[str, Any]]) -> str:
    if not findings:
        return ""
    top = findings[0]
    return f"{top.get('Severity', 'Review')}: {top.get('Error Type', 'QA')}"


def qa_severity_rank(severity: str) -> int:
    order = {"Critical": 4, "Major": 3, "Minor": 2, "Review": 1, "Info": 0, "Pass": -1}
    return order.get(safe_text(severity).title(), 1)


def qa_overall_status(findings: List[Dict[str, Any]]) -> str:
    if not findings:
        return "Pass"
    max_rank = max(qa_severity_rank(f.get("Severity", "Review")) for f in findings)
    if max_rank >= 4:
        return "Critical"
    if max_rank >= 3:
        return "Major"
    return "Needs Review"


def dedupe_qa_findings(findings: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    unique: List[Dict[str, Any]] = []
    seen = set()
    for finding in findings or []:
        key = (
            safe_text(finding.get("Location") or finding.get("Segment ID")),
            safe_text(finding.get("Rule ID")),
            safe_text(finding.get("Error Type")),
            safe_text(finding.get("Wrong Part") or finding.get("Explanation")),
        )
        if key in seen:
            continue
        seen.add(key)
        unique.append(finding)
    return unique


def delivery_quality_findings(
    rows: List[Dict[str, Any]],
    target_language: str,
    domain: str,
    rules: Optional[Dict[str, Any]] = None,
) -> List[Dict[str, Any]]:
    """Pre-delivery gate for objective localization integrity checks.

    This deliberately duplicates a few core checks outside the AI path so that
    exports still catch placeholders, emoji, DNT, and ZWNJ issues when no API key
    is configured.
    """
    all_findings: List[Dict[str, Any]] = []
    rules = rules or workspace_rules()
    dnt_terms = [safe_text(item.get("term")) for item in rules.get("dnt", []) if safe_text(item.get("term"))]
    zwnj = "\u200c"

    for idx, original in enumerate(rows or [], start=1):
        row = dict(original or {})
        row.setdefault("id", idx)
        row.setdefault("location", f"Segment {idx}")
        row["target"] = safe_text(row.get("target", row.get("translation", "")))
        row["source"] = safe_text(row.get("source", ""))
        row_findings: List[Dict[str, Any]] = []

        if not row["target"].strip():
            row_findings.append(qa_manual_finding(row, "Major", "Completeness", "Target translation is blank.", "Add the missing translation."))

        protected = LOCALIZATION_PROTECTED_RE.findall(row["source"])
        for token in protected:
            if token and token not in row["target"]:
                row_findings.append(qa_manual_finding(row, "Major", "Protected Token", f"Missing protected token {token}", f"Keep {token} unchanged in the target."))

        for symbol in LOCALIZATION_EMOJI_RE.findall(row["source"]):
            if symbol and symbol not in row["target"]:
                row_findings.append(qa_manual_finding(row, "Major", "Formatting", f"Missing source emoji/icon {symbol}", "Keep source emoji/icon characters unless client rules say to remove them."))

        for term in dnt_terms:
            if term and term.lower() in row["source"].lower() and term not in row["target"]:
                row_findings.append(qa_manual_finding(row, "Major", "DNT", f"DNT term changed or missing: {term}", f"Keep {term} unchanged."))

        source_zwnj = row["source"].count(zwnj)
        target_zwnj = row["target"].count(zwnj)
        if source_zwnj and target_zwnj < source_zwnj:
            row_findings.append(qa_manual_finding(row, "Major", "ZWNJ", "Zero Width Non-Joiner is missing from target.", "Preserve ZWNJ characters where the source requires them."))
        if target_zwnj and re.search(rf"(^|[\s\.,;:!?]){zwnj}|{zwnj}($|[\s\.,;:!?])", row["target"]):
            row_findings.append(qa_manual_finding(row, "Major", "ZWNJ", "Zero Width Non-Joiner is placed next to whitespace, punctuation, or a boundary.", "Keep ZWNJ only between characters that must not join."))

        row_findings.extend(run_global_qa_for_row(row, target_language=target_language, domain=domain, rules=rules))
        for finding in row_findings:
            finding.setdefault("Location", row.get("location", f"Segment {idx}"))
            finding.setdefault("Segment ID", row.get("id", idx))
        all_findings.extend(row_findings)

    return dedupe_qa_findings(all_findings)


def delivery_gate_summary(rows: List[Dict[str, Any]], findings: List[Dict[str, Any]]) -> Dict[str, Any]:
    status_by_location: Dict[str, List[Dict[str, Any]]] = {}
    for finding in findings or []:
        location = safe_text(finding.get("Location") or finding.get("Segment ID"))
        if location:
            status_by_location.setdefault(location, []).append(finding)
    gate_rows = []
    for idx, row in enumerate(rows or [], start=1):
        location = safe_text(row.get("location") or f"Segment {row.get('id', idx)}")
        row_findings = status_by_location.get(location, [])
        gate_rows.append({
            **row,
            "id": row.get("id", idx),
            "status": qa_overall_status(row_findings),
            "issues": summarize_qa_findings(row_findings),
            "match": row.get("match", ""),
        })
    _, _, summary = professional_qa_rows(gate_rows, findings)
    blocking = [f for f in findings if qa_severity_rank(f.get("Severity", "Review")) >= 3]
    summary["blocking_findings"] = len(blocking)
    summary["ready"] = summary["result"] == "PASS" and not blocking
    return summary


def render_delivery_gate(rows: List[Dict[str, Any]], findings: List[Dict[str, Any]], title: str = "Delivery readiness") -> Dict[str, Any]:
    summary = delivery_gate_summary(rows, findings)
    st.markdown(f"### {title}")
    metrics([
        ("Gate", "PASS" if summary["ready"] else "FAIL", f"{summary['qa_score']}%"),
        ("Blocking", summary["blocking_findings"], "Critical/Major"),
        ("Findings", summary["total_findings"], "total"),
        ("Threshold", "95%", "minimum pass"),
    ])
    if summary["ready"]:
        st.success("Delivery gate passed. Objective localization checks did not find blocking issues.")
    else:
        st.error("Delivery gate failed. Fix Critical/Major issues or keep the file in Human Review before delivery.")
        preview = [f for f in findings if qa_severity_rank(f.get("Severity", "Review")) >= 3][:12]
        if preview:
            cols = ["Location", "Source Text", "Translation", "Error Type", "Severity", "Wrong Part", "Suggestion", "Rule ID"]
            st.dataframe(pd.DataFrame(preview)[[c for c in cols if c in pd.DataFrame(preview).columns]], use_container_width=True, hide_index=True)
    return summary


def qa_manual_finding(row: Dict[str, Any], severity: str, category: str, issue: str, suggestion: str = "") -> Dict[str, Any]:
    return {
        "Sheet": "",
        "Location": safe_text(row.get("location") or f"Segment {row.get('id', '')}"),
        "Mode": "Manual QA",
        "Source Text": safe_text(row.get("source", "")),
        "Translation": safe_text(row.get("target", row.get("translation", ""))),
        "Error Type": category,
        "Severity": severity,
        "Wrong Part": issue,
        "Suggestion": suggestion,
        "Explanation": issue,
        "Check Source": "ErrorSweep QA",
        "Rule Source": "Built-in QA",
        "Confidence": "High",
        "Rule ID": f"manual.{re.sub(r'[^a-z0-9]+', '_', category.lower()).strip('_')}",
        "Autofix Possible": "No",
        "Priority": 90 if severity == "Critical" else 70 if severity == "Major" else 50,
    }


def professional_qa_rows(segment_rows: List[Dict[str, Any]], detailed_findings: List[Dict[str, Any]]) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], Dict[str, Any]]:
    finding_rows: List[Dict[str, Any]] = []
    severity_counts = {"Critical": 0, "Major": 0, "Minor": 0, "Review": 0}
    for idx, finding in enumerate(detailed_findings, start=1):
        severity = safe_text(finding.get("Severity") or "Review").title()
        if severity not in severity_counts:
            severity = "Review"
        severity_counts[severity] += 1
        finding_rows.append({
            "Finding ID": idx,
            "Segment ID": safe_text(finding.get("Segment ID") or finding.get("Location") or ""),
            "Source Text": safe_text(finding.get("Source Text") or ""),
            "Target Text": safe_text(finding.get("Translation") or ""),
            "Suggested Target": safe_text(finding.get("Suggestion") or ""),
            "Error Category": safe_text(finding.get("Error Type") or "QA"),
            "Severity": severity,
            "Issue": safe_text(finding.get("Wrong Part") or finding.get("Explanation") or ""),
            "Explanation": safe_text(finding.get("Explanation") or ""),
            "Check": safe_text(finding.get("Rule Source") or finding.get("Check Source") or "Rule Engine"),
            "Confidence": safe_text(finding.get("Confidence") or ""),
            "Rule ID": safe_text(finding.get("Rule ID") or ""),
        })

    segment_overview = []
    failed_segments = 0
    for row in segment_rows:
        status = safe_text(row.get("status", "Pass"))
        if status != "Pass":
            failed_segments += 1
        segment_overview.append({
            "Segment ID": row.get("id", ""),
            "Status": status,
            "Source Text": row.get("source", ""),
            "Target Text": row.get("target", ""),
            "Issue Summary": row.get("issues", ""),
            "Match": row.get("match", ""),
        })

    total_segments = len(segment_rows)
    pass_segments = max(0, total_segments - failed_segments)
    qa_score = round((pass_segments / total_segments) * 100, 2) if total_segments else 100.0
    summary = {
        "total_segments": total_segments,
        "pass_segments": pass_segments,
        "review_segments": failed_segments,
        "total_findings": len(finding_rows),
        "qa_score": qa_score,
        "result": "PASS" if qa_score >= 95 and severity_counts["Critical"] == 0 else "FAIL",
        "severity_counts": severity_counts,
    }
    return finding_rows, segment_overview, summary


def create_qa_excel_report(segment_rows: List[Dict[str, Any]], detailed_findings: List[Dict[str, Any]]) -> bytes:
    finding_rows, segment_overview, summary = professional_qa_rows(segment_rows, detailed_findings)
    wb = Workbook()
    ws_summary = wb.active
    ws_summary.title = "Summary"
    ws_findings = wb.create_sheet("QA Findings")
    ws_segments = wb.create_sheet("Segment Overview")
    ws_instructions = wb.create_sheet("Review Notes")

    dark = "1F2937"
    blue = "D9EAF7"
    green = "D9EAD3"
    yellow = "FFF2CC"
    red = "F4CCCC"
    orange = "FCE4D6"
    border = Side(style="thin", color="A6A6A6")

    ws_summary.merge_cells("A1:D1")
    ws_summary["A1"] = "ErrorSweep QA Report"
    ws_summary["A1"].font = Font(bold=True, size=16, color="FFFFFF")
    ws_summary["A1"].fill = PatternFill("solid", fgColor=dark)
    ws_summary["A1"].alignment = Alignment(horizontal="center")
    summary_rows = [
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Result", summary["result"]),
        ("QA Score", summary["qa_score"]),
        ("Total Segments", summary["total_segments"]),
        ("Passed Segments", summary["pass_segments"]),
        ("Segments Needing Review", summary["review_segments"]),
        ("Total Findings", summary["total_findings"]),
        ("Critical", summary["severity_counts"]["Critical"]),
        ("Major", summary["severity_counts"]["Major"]),
        ("Minor", summary["severity_counts"]["Minor"]),
        ("Review", summary["severity_counts"]["Review"]),
    ]
    for idx, (label, value) in enumerate(summary_rows, start=3):
        ws_summary.cell(idx, 1, label)
        ws_summary.cell(idx, 2, value)
        ws_summary.cell(idx, 1).font = Font(bold=True)
        ws_summary.cell(idx, 1).fill = PatternFill("solid", fgColor=blue)
        if label == "Result":
            ws_summary.cell(idx, 2).fill = PatternFill("solid", fgColor=green if value == "PASS" else red)
            ws_summary.cell(idx, 2).font = Font(bold=True)
    ws_summary["D3"] = "Pass rule"
    ws_summary["E3"] = "PASS requires QA score >= 95 and zero Critical findings."
    ws_summary["D3"].font = Font(bold=True)
    ws_summary["D3"].fill = PatternFill("solid", fgColor=blue)
    apply_widths(ws_summary, {"A": 28, "B": 20, "D": 18, "E": 62})

    finding_headers = [
        "Finding ID", "Segment ID", "Source Text", "Target Text", "Suggested Target",
        "Error Category", "Severity", "Issue", "Explanation", "Check", "Confidence", "Rule ID",
    ]
    for col_idx, header in enumerate(finding_headers, start=1):
        cell = ws_findings.cell(1, col_idx, header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor=blue)
    for row_idx, row in enumerate(finding_rows, start=2):
        for col_idx, header in enumerate(finding_headers, start=1):
            cell = ws_findings.cell(row_idx, col_idx, row.get(header, ""))
            if header == "Severity":
                color = red if row.get(header) == "Critical" else orange if row.get(header) == "Major" else yellow if row.get(header) == "Minor" else blue
                cell.fill = PatternFill("solid", fgColor=color)
                cell.font = Font(bold=True)
    apply_widths(ws_findings, {"A": 12, "B": 18, "C": 48, "D": 48, "E": 42, "F": 20, "G": 14, "H": 34, "I": 42, "J": 18, "K": 16, "L": 24})
    ws_findings.freeze_panes = "A2"

    segment_headers = ["Segment ID", "Status", "Source Text", "Target Text", "Issue Summary", "Match"]
    for col_idx, header in enumerate(segment_headers, start=1):
        cell = ws_segments.cell(1, col_idx, header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor=blue)
    for row_idx, row in enumerate(segment_overview, start=2):
        for col_idx, header in enumerate(segment_headers, start=1):
            cell = ws_segments.cell(row_idx, col_idx, row.get(header, ""))
            if header == "Status":
                cell.fill = PatternFill("solid", fgColor=green if row.get(header) == "Pass" else yellow)
    apply_widths(ws_segments, {"A": 12, "B": 18, "C": 52, "D": 52, "E": 46, "F": 16})
    ws_segments.freeze_panes = "A2"

    notes = [
        ("Purpose", "Use this workbook as the client-facing QA handoff. The Findings sheet contains one issue per row."),
        ("Severity", "Critical and Major issues should be reviewed first. Minor issues are lower-risk language or formatting corrections."),
        ("Status", "Segment Overview shows all uploaded segments, including clean pass rows."),
        ("Threshold", "PASS requires QA score >= 95 and zero Critical findings."),
    ]
    ws_instructions["A1"] = "Review Notes"
    ws_instructions["A1"].font = Font(bold=True, size=14)
    for idx, (label, text) in enumerate(notes, start=3):
        ws_instructions.cell(idx, 1, label)
        ws_instructions.cell(idx, 2, text)
        ws_instructions.cell(idx, 1).font = Font(bold=True)
        ws_instructions.cell(idx, 1).fill = PatternFill("solid", fgColor=blue)
    apply_widths(ws_instructions, {"A": 18, "B": 90})

    for sheet in wb.worksheets:
        style_sheet_base(sheet)
        for row in sheet.iter_rows():
            for cell in row:
                cell.border = Border(left=border, right=border, top=border, bottom=border)
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()


def validate_timing_rows(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    issues = []
    previous_end = None
    for idx, row in enumerate(rows or []):
        start = float(row.get("start", 0) or 0)
        end = float(row.get("end", 0) or 0)
        if end <= start:
            issues.append({"row": idx + 1, "issue": "End time must be after start time."})
        if previous_end is not None and start < previous_end:
            issues.append({"row": idx + 1, "issue": "Segment overlaps the previous segment."})
        previous_end = max(previous_end or 0, end)
    return issues


# ==========================================================
# Login
# ==========================================================

def render_landing_page() -> None:
    local_status = current_ai_route_label()
    st.html(
        dedent(f"""
        <div class="es-lp">
          <header class="es-lp-nav">
            <div class="es-lp-nav-inner">
              <div class="es-lp-brand">
                <div class="es-lp-logo">ES</div>
                <div>
                  <div class="es-lp-brand-name">ErrorSweep</div>
                  <div class="es-lp-brand-sub">Nawin Corp</div>
                </div>
              </div>
              <div class="es-lp-links">
                <a class="es-lp-link" href="#solutions">Solutions</a>
                <a class="es-lp-link" href="#product">Product</a>
                <a class="es-lp-link" href="#resources">Resources</a>
                <a class="es-lp-link" href="#developers">Developers</a>
                <a class="es-lp-link" href="#pricing">Pricing</a>
              </div>
              <div class="es-lp-actions">
                <a class="es-lp-btn" href="{public_page_link('login')}" target="_self">Login</a>
                <a class="es-lp-btn primary" href="{public_page_link('signup')}" target="_self">Sign up</a>
              </div>
            </div>
          </header>

          <main>
            <section class="es-lp-hero">
              <div class="es-lp-grid-bg"></div>
              <div class="es-lp-inner">
                <div class="es-lp-hero-content">
                  <div class="es-lp-badge"><span class="es-lp-dot"></span>{escape(local_status)}</div>
                  <h1 class="es-lp-title">AI localization that scales your growth, not your overhead.</h1>
                  <p class="es-lp-copy">
                    Simplify localization every step of the way. Focus on growth while our AI automates translation quality checks,
                    identifies linguistic bugs, and drives issue resolution without slowing you down.
                  </p>
                  <div class="es-lp-hero-actions">
                    <a class="es-lp-btn primary" href="{public_page_link('signup')}" target="_self">Start for free</a>
                    <a class="es-lp-btn" href="{public_page_link('login')}" target="_self">Login</a>
                  </div>
                </div>
              </div>
            </section>

            <section id="product" class="es-lp-product">
              <div class="es-lp-inner">
                <div class="es-lp-glass">
                  <div class="es-lp-glow-a"></div>
                  <div class="es-lp-glow-b"></div>
                  <div class="es-lp-window">
                    <div class="es-lp-window-top">
                      <div class="es-lp-traffic"><span></span><span></span><span></span></div>
                      <div class="es-lp-window-label">ErrorSweep Localization Workflow</div>
                    </div>

                    <div class="es-lp-workflow">
                      <div class="es-lp-panel">
                        <div class="es-lp-panel-head">
                          <div>
                            <div class="es-lp-eyebrow">Active QA job</div>
                            <h3>Mobile App UI - French</h3>
                          </div>
                          <div class="es-lp-status">Review ready</div>
                        </div>
                        <div class="es-lp-task">
                          <div class="es-lp-task-row"><span>Placeholder integrity</span><span class="es-lp-green">Passed</span></div>
                          <div class="es-lp-progress"><span style="width:92%;"></span></div>
                        </div>
                        <div class="es-lp-task">
                          <div class="es-lp-task-row"><span>Linguistic issues</span><span class="es-lp-yellow">12 findings</span></div>
                          <div class="es-lp-findings">
                            <div class="es-lp-finding red">2 Critical</div>
                            <div class="es-lp-finding yellow">6 Major</div>
                            <div class="es-lp-finding blue">4 Minor</div>
                          </div>
                        </div>
                        <div class="es-lp-task">
                          <div class="es-lp-task-row"><span>Human Review</span><span class="es-lp-violet">Assigned to Reviewer</span></div>
                        </div>
                      </div>

                      <div>
                        <div class="es-lp-panel">
                          <div class="es-lp-eyebrow es-lp-sky">Translation Quality Index</div>
                          <div class="es-lp-score">
                            <div class="es-lp-score-number">A</div>
                            <div class="es-lp-muted">review-ready score</div>
                          </div>
                          <div class="es-lp-mini-grid">
                            <div class="es-lp-mini">Glossary<br><span class="es-lp-green">Aligned</span></div>
                            <div class="es-lp-mini">DNT<br><span class="es-lp-green">Locked</span></div>
                          </div>
                        </div>
                        <div class="es-lp-panel es-lp-pipeline">
                          <div class="es-lp-eyebrow">Pipeline</div>
                          <div class="es-lp-pipeline-row"><span>Extract strings</span><span class="es-lp-green">Done</span></div>
                          <div class="es-lp-pipeline-row"><span>Run AI QA</span><span class="es-lp-green">Done</span></div>
                          <div class="es-lp-pipeline-row"><span>Review scorecard</span><span class="es-lp-yellow">In progress</span></div>
                        </div>
                      </div>
                    </div>
                  </div>
                </div>
              </div>
            </section>
          </main>
        </div>
        """).strip(),
    )

    st.html(
        dedent(f"""
        <div class="es-lp">
          <section class="es-lp-social">
            <div class="es-lp-inner">
              <p class="es-lp-social-title">Built for localization teams who want fewer spreadsheet handoffs and cleaner multilingual releases.</p>
              <div class="es-lp-logo-grid">
                <div class="es-lp-logo-card">ACME</div>
                <div class="es-lp-logo-card">GLOBEX</div>
                <div class="es-lp-logo-card">NOVA APP</div>
                <div class="es-lp-logo-card">ATLAS</div>
                <div class="es-lp-logo-card">BRIGHTTEL</div>
                <div class="es-lp-logo-card">TASKFLOW</div>
              </div>
              <div class="es-lp-awards">
                <span class="es-lp-award es-lp-yellow">QA Workflow</span>
                <span class="es-lp-award es-lp-sky">Human Review</span>
                <span class="es-lp-award es-lp-green">No-Key MT</span>
                <span class="es-lp-award es-lp-violet">Scorecards</span>
              </div>
            </div>
          </section>

          <section class="es-lp-section">
            <div class="es-lp-inner">
              <div class="es-lp-stats">
                <div class="es-lp-stat"><div class="es-lp-stat-number es-lp-green">QA</div><p>Placeholder, terminology, formatting, and DNT checks</p></div>
                <div class="es-lp-stat"><div class="es-lp-stat-number es-lp-sky">MT</div><p>Built-in self-hosted translation route for no-key workflows</p></div>
                <div class="es-lp-stat"><div class="es-lp-stat-number es-lp-violet">LQA</div><p>Review-ready editor jobs and Excel scorecards</p></div>
              </div>
            </div>
          </section>

          <section id="solutions" class="es-lp-section">
            <div class="es-lp-inner">
              <div class="es-lp-section-title">
                <div class="es-lp-eyebrow">Solutions</div>
                <h2>Shape your localization strategy to fit your needs</h2>
              </div>
              <div class="es-lp-card-grid">
                <article class="es-lp-card">
                  <div class="es-lp-icon">DEV</div>
                  <h3>Developers</h3>
                  <p>Ship multilingual product updates without spreadsheet handoffs. Connect strings, rules, and QA directly to release workflows.</p>
                  <span class="es-lp-card-link">Explore developer workflows -></span>
                </article>
                <article class="es-lp-card">
                  <div class="es-lp-icon">OPS</div>
                  <h3>Localization Managers</h3>
                  <p>Track jobs, linguists, scorecards, client rules, TMs, and QA findings from one control center.</p>
                  <span class="es-lp-card-link es-lp-sky">Build scalable operations -></span>
                </article>
                <article class="es-lp-card">
                  <div class="es-lp-icon">UX</div>
                  <h3>Product Teams</h3>
                  <p>Find localization bugs before they reach users. Prioritize high-impact issues with quality scores and review status.</p>
                  <span class="es-lp-card-link es-lp-violet">Protect user experience -></span>
                </article>
              </div>
            </div>
          </section>

          <section id="developers" class="es-lp-section">
            <div class="es-lp-inner">
              <div class="es-lp-feature-grid">
                <div class="es-lp-feature">
                  <div class="es-lp-eyebrow">Collaboration hub</div>
                  <h2>One place for every team, every language, every project.</h2>
                  <p>Centralize jobs, source files, translations, reviewer feedback, terminology, and vendor scorecards. Give developers, translators, reviewers, and clients the right view without slowing delivery.</p>
                  <ul class="es-lp-list">
                    <li><span class="es-lp-green">OK</span> Role-based workspaces for owners, PMs, translators, reviewers, and clients</li>
                    <li><span class="es-lp-green">OK</span> Human review editor with TM, glossary, DNT, and QA context</li>
                    <li><span class="es-lp-green">OK</span> Excel-first scorecards for vendor and translator quality</li>
                  </ul>
                </div>
                <div class="es-lp-mock">
                  <div class="es-lp-mock-inner">
                    <div class="es-lp-task-row"><span class="es-lp-muted">Workspace</span><span class="es-lp-status">Live</span></div>
                    <div class="es-lp-mock-row"><b>Project:</b> Mobile App UI <span style="float:right;" class="es-lp-green">12 languages</span></div>
                    <div class="es-lp-mock-row"><b>Reviewer:</b> Assigned <span style="float:right;" class="es-lp-sky">68 rows</span></div>
                    <div class="es-lp-mock-row"><b>Client rules:</b> Active <span style="float:right;" class="es-lp-violet">Glossary + DNT</span></div>
                  </div>
                </div>
              </div>

              <div class="es-lp-feature-grid">
                <div class="es-lp-mock">
                  <div class="es-lp-mock-inner">
                    <div class="es-lp-three">
                      <div class="es-lp-count" style="background:rgba(239,68,68,.10);"><strong class="es-lp-violet">3</strong><span class="es-lp-muted">Critical</span></div>
                      <div class="es-lp-count" style="background:rgba(234,179,8,.10);"><strong class="es-lp-yellow">9</strong><span class="es-lp-muted">Major</span></div>
                      <div class="es-lp-count" style="background:rgba(17,245,181,.10);"><strong class="es-lp-green">A</strong><span class="es-lp-muted">TQI</span></div>
                    </div>
                    <div class="es-lp-issue">
                      <b>Detected issue</b>
                      <p>Placeholder <span class="es-lp-code">{{email}}</span> was removed from target string.</p>
                      <a class="es-lp-btn primary" href="{public_page_link('signup')}" target="_self">Send to Human Review</a>
                    </div>
                  </div>
                </div>
                <div class="es-lp-feature">
                  <div class="es-lp-eyebrow es-lp-sky">Translation quality</div>
                  <h2>Redefine translation quality with TQI.</h2>
                  <p>Automatically validate string health, linguistic accuracy, placeholders, formatting, DNT terms, and terminology. Route risky segments to reviewers and keep production-ready content moving.</p>
                  <ul class="es-lp-list">
                    <li><span class="es-lp-sky">OK</span> Automated QA checks for placeholders, numbers, tags, and language issues</li>
                    <li><span class="es-lp-sky">OK</span> Scorecard output for PMs, reviewers, and vendor evaluation</li>
                    <li><span class="es-lp-sky">OK</span> Human-in-the-loop workflows for uncertain translations</li>
                  </ul>
                </div>
              </div>
            </div>
          </section>

          <section id="pricing" class="es-lp-section">
            <div class="es-lp-inner">
              <div class="es-lp-cta">
                <h2>Build a localization operation your product team actually enjoys using.</h2>
                <p>Start with AI QA, translation review, scorecards, and managed workflows, then scale into full localization operations.</p>
                <div class="es-lp-hero-actions">
                  <a class="es-lp-btn primary" href="{public_page_link('signup')}" target="_self">Start for free</a>
                  <a class="es-lp-btn" href="{public_page_link('login')}" target="_self">Login</a>
                </div>
              </div>
            </div>
          </section>

          <footer id="resources" class="es-lp-footer">
            <div class="es-lp-inner es-lp-footer-row">
              <span>Copyright 2026 ErrorSweep by Nawin Corp. All rights reserved.</span>
              <span class="es-lp-footer-links">
                <a class="es-lp-link" href="{public_page_link('security')}" target="_self">Security</a>
                <a class="es-lp-link" href="#developers">Developers</a>
                <a class="es-lp-link" href="{public_page_link('terms')}" target="_self">Terms</a>
                <a class="es-lp-link" href="{public_page_link('privacy')}" target="_self">Privacy</a>
              </span>
            </div>
          </footer>
        </div>
        """).strip(),
    )

    st.html(
        dedent("""
        <style>
        @media (max-width: 1100px) {
          .es-lp-links { display: none; }
          .es-lp-workflow, .es-lp-feature-grid { grid-template-columns: 1fr; }
          .es-lp-logo-grid { grid-template-columns: repeat(3, minmax(0, 1fr)); }
          .es-lp-stats, .es-lp-card-grid { grid-template-columns: 1fr; }
        }
        @media (max-width: 640px) {
          .es-lp-nav-inner { align-items: flex-start; }
          .es-lp-actions { flex-direction: column; align-items: stretch; }
          .es-lp-btn { width: 100%; }
          .es-lp-title { font-size: 44px; }
          .es-lp-copy { font-size: 17px; }
          .es-lp-window { padding: 16px; }
          .es-lp-panel-head, .es-lp-task-row, .es-lp-pipeline-row { flex-direction: column; gap: 8px; }
          .es-lp-findings, .es-lp-mini-grid, .es-lp-three { grid-template-columns: 1fr; }
          .es-lp-logo-grid { grid-template-columns: repeat(2, minmax(0, 1fr)); }
          .es-lp-cta { padding: 28px; }
        }
        </style>
        """).strip(),
    )


def render_login() -> None:
    st.html(
        dedent(f"""
        <div class="es-auth-shell">
          <div class="es-lp-brand">
            <div class="es-lp-logo">ES</div>
            <div>
              <div class="es-lp-brand-name">ErrorSweep</div>
              <div class="es-lp-brand-sub">Nawin Corp</div>
            </div>
          </div>
          <div class="es-auth-links">
            <a class="es-lp-btn" href="{public_page_link('landing')}" target="_self">Back to landing</a>
            <a class="es-lp-btn primary" href="{public_page_link('signup')}" target="_self">Sign up</a>
          </div>
        </div>
        """).strip(),
    )
    st.markdown("## Login to ErrorSweep")

    tabs = st.tabs(["Platform owner", "Workspace user", "Demo access"])

    with tabs[0]:
        st.markdown("### Platform Owner Login")
        owner_user = secret("ERRORSWEEP_OWNER_USERNAME", "owner@errorsweep.local")
        owner_is_configured = password_configured("ERRORSWEEP_OWNER_PASSWORD_HASH", "ERRORSWEEP_OWNER_PASSWORD")
        with st.form("owner_login"):
            email = st.text_input("Owner email", value=owner_user if not owner_is_configured else "")
            password = st.text_input("Owner password", type="password")
            accepted = st.checkbox(COMPLIANCE_ACK_LABEL, key="owner_compliance_ack")
            submitted = st.form_submit_button("Sign in as Platform Owner", use_container_width=True)
        sso1, sso2 = st.columns(2)
        sso1.button("Continue with Enterprise SSO", disabled=True, use_container_width=True, key="owner_entra_sso")
        sso2.button("Continue with Okta", disabled=True, use_container_width=True, key="owner_okta_sso")
        if submitted:
            if not accepted:
                st.error("Please accept the workspace compliance terms before signing in.")
            elif owner_is_configured and hmac.compare_digest(email.strip(), owner_user.strip()) and verify_login_password(password, "ERRORSWEEP_OWNER_PASSWORD_HASH", "ERRORSWEEP_OWNER_PASSWORD"):
                login_user(email, "Platform Owner", "owner", "Platform")
                add_audit("Owner sign-in", email)
                st.rerun()
            elif not owner_is_configured:
                st.warning("Owner password hash is not configured. Use Demo access while building.")
            else:
                st.error("Invalid owner credentials.")

    with tabs[1]:
        st.markdown("### Workspace User Login")
        user_name = secret("ERRORSWEEP_USER_USERNAME", "user@errorsweep.local")
        user_is_configured = password_configured("ERRORSWEEP_USER_PASSWORD_HASH", "ERRORSWEEP_USER_PASSWORD")
        default_role = secret("ERRORSWEEP_DEFAULT_USER_ROLE", "Workspace Owner")
        with st.form("user_login"):
            email = st.text_input("User email", value=user_name if not user_is_configured else "")
            password = st.text_input("User password", type="password")
            accepted = st.checkbox(COMPLIANCE_ACK_LABEL, key="user_compliance_ack")
            submitted = st.form_submit_button("Sign in", use_container_width=True)
        sso1, sso2 = st.columns(2)
        sso1.button("Continue with Enterprise SSO", disabled=True, use_container_width=True, key="user_entra_sso")
        sso2.button("Continue with Okta", disabled=True, use_container_width=True, key="user_okta_sso")
        if submitted:
            if not accepted:
                st.error("Please accept the workspace compliance terms before signing in.")
            elif user_is_configured and hmac.compare_digest(email.strip(), user_name.strip()) and verify_login_password(password, "ERRORSWEEP_USER_PASSWORD_HASH", "ERRORSWEEP_USER_PASSWORD"):
                login_user(email, default_role, "workspace", secret("ERRORSWEEP_ORG_NAME", "Demo Workspace"))
                add_audit("Workspace user sign-in", email)
                st.rerun()
            else:
                stored_users = load_saas_records("users", include_all_workspaces=True, limit=1000)
                matched = next((u for u in stored_users if hmac.compare_digest(str(u.get("email", "")).strip().lower(), email.strip().lower())), None)
                if matched and safe_text(matched.get("status", "Active")).lower() not in {"active", "invited"}:
                    st.error("This workspace account is not active. Contact your workspace owner.")
                elif matched and verify_password(password, str(matched.get("password_hash", ""))):
                    if is_production_mode() and not bool(matched.get("email_verified")):
                        verify_url = queue_verification_email(email.strip(), matched.get("workspace", "Demo Workspace"))
                        st.warning("Please verify your email before signing in. A verification link was added to the notification outbox.")
                        if not email_provider_label() or email_provider_label() == "not_configured":
                            st.caption(f"Local verification link: {verify_url}")
                        return
                    login_user(email.strip(), matched.get("role", default_role) or default_role, "workspace", matched.get("workspace", "Demo Workspace") or "Demo Workspace")
                    add_audit("Workspace user sign-in", email.strip())
                    st.rerun()
                elif not user_is_configured and not matched:
                    st.warning("No matching workspace account found. Create one from Sign up, or use Demo access while building.")
                else:
                    st.error("Invalid workspace credentials.")
        with st.expander("Forgot password?", expanded=False):
            reset_email = st.text_input("Account email", key="password_reset_email")
            if st.button("Send password reset link", use_container_width=True, key="send_password_reset"):
                matched = next(
                    (
                        u for u in load_saas_records("users", include_all_workspaces=True, limit=1000)
                        if safe_text(u.get("email")).lower() == safe_text(reset_email).lower()
                    ),
                    None,
                )
                workspace = safe_text((matched or {}).get("workspace") or "Demo Workspace")
                if safe_text(reset_email):
                    reset_url = queue_password_reset_email(reset_email.strip(), workspace)
                else:
                    reset_url = ""
                st.success("If that account exists, a password reset link has been added to the notification outbox.")
                if reset_url and email_provider_label() == "not_configured":
                    st.caption(f"Local reset link: {reset_url}")

    with tabs[2]:
        st.markdown("### Demo Access")
        if is_production_mode():
            st.warning("Demo access is disabled in production.")
            return
        demo_role = st.selectbox(
            "Preview as",
            ["Platform Owner", "Workspace Owner", "Workspace Admin", "Project Manager", "Translator", "Reviewer", "Client Viewer", "Billing Admin", "User"],
        )
        accepted = st.checkbox(COMPLIANCE_ACK_LABEL, key="demo_compliance_ack")
        if st.button("Enter demo workspace", use_container_width=True):
            if not accepted:
                st.error("Please accept the workspace compliance terms before entering the demo workspace.")
            else:
                account_type = "owner" if demo_role == "Platform Owner" else "workspace"
                login_user(f"{demo_role.lower().replace(' ', '_')}@errorsweep.local", demo_role, account_type, "Demo Workspace")
                add_audit("Demo login", demo_role)
                st.rerun()


def render_signup() -> None:
    st.html(
        dedent(f"""
        <div class="es-auth-shell">
          <div class="es-lp-brand">
            <div class="es-lp-logo">ES</div>
            <div>
              <div class="es-lp-brand-name">ErrorSweep</div>
              <div class="es-lp-brand-sub">Nawin Corp</div>
            </div>
          </div>
          <div class="es-auth-links">
            <a class="es-lp-btn" href="{public_page_link('landing')}" target="_self">Back to landing</a>
            <a class="es-lp-btn primary" href="{public_page_link('login')}" target="_self">Login</a>
          </div>
        </div>
        """).strip(),
    )
    st.markdown("## Create your ErrorSweep workspace")
    st.caption("This creates a trial workspace in the current app session.")

    with st.form("signup_form"):
        name = st.text_input("Full name")
        email = st.text_input("Work email")
        workspace = st.text_input("Company / workspace", value="Demo Workspace")
        password = st.text_input("Password", type="password")
        accepted = st.checkbox(COMPLIANCE_ACK_LABEL, key="signup_compliance_ack")
        submitted = st.form_submit_button("Create workspace", use_container_width=True)

    if submitted:
        if not name.strip() or not email.strip() or not workspace.strip() or not password:
            st.error("Please enter your name, email, workspace, and password.")
            return
        if not accepted:
            st.error("Please accept the workspace compliance terms before creating a workspace.")
            return
        existing_user = next(
            (
                u for u in load_saas_records("users", include_all_workspaces=True, limit=1000)
                if safe_text(u.get("email")).lower() == email.strip().lower()
            ),
            None,
        )
        if existing_user:
            st.error("An account with this email already exists. Use login or password reset.")
            return
        user_record = persist_saas_record("users", {
            "email": email.strip(),
            "workspace": workspace.strip(),
            "role": "Workspace Owner",
            "plan": "Trial",
            "status": "Active",
            "password_hash": hash_password(password),
            "email_verified": False,
            "verified_at": None,
        })
        st.session_state.users.append(user_record)
        trim_session_list("users")
        if not any(w.get("workspace") == workspace.strip() for w in st.session_state.workspaces):
            workspace_record = persist_saas_record("workspaces", {
                "workspace": workspace.strip(),
                "owner": email.strip(),
                "plan": "Trial",
                "status": "Active",
                "users": 1,
                "jobs": 0,
            })
            st.session_state.workspaces.append(workspace_record)
            trim_session_list("workspaces")
        add_audit("Trial workspace signup", email.strip())
        verify_url = queue_verification_email(email.strip(), workspace.strip(), name.strip())
        queue_email_notification(
            email.strip(),
            "Welcome to ErrorSweep",
            f"Your trial workspace '{workspace.strip()}' is ready. Verify your email to secure the account: {verify_url}",
            "signup.welcome",
            metadata={"workspace": workspace.strip(), "name": name.strip(), "verify_url": verify_url},
            workspace=workspace.strip(),
        )
        if is_production_mode():
            st.success("Workspace created. Please verify your email before signing in.")
            if email_provider_label() == "not_configured":
                st.caption(f"Local verification link: {verify_url}")
            return
        login_user(email.strip(), "Workspace Owner", "workspace", workspace.strip())
        st.rerun()


def render_public_document(kind: str) -> None:
    docs = {
        "terms": {
            "title": "Terms of Service",
            "body": [
                "ErrorSweep is provided for authorized localization QA, translation review, subtitle, transcription, and scorecard workflows.",
                "Users must only upload content they are authorized to process, including client, confidential, copyrighted, or regulated material.",
                "Accounts are workspace-scoped. Do not share credentials or attempt to access another workspace's data.",
                "Trial and local development records are provided for evaluation and may be reset during product setup.",
            ],
        },
        "privacy": {
            "title": "Privacy Policy",
            "body": [
                "ErrorSweep stores account, workspace, job, audit, and usage records needed to operate the SaaS workflow.",
                "Uploaded localization content is processed for translation QA, human review, media editing, and export generation.",
                "When Supabase is not configured, local development uses JSON fallback storage on this machine.",
                "External AI or grammar routes should be used only when the customer has approved that processing path.",
            ],
        },
        "security": {
            "title": "Security",
            "body": [
                "Passwords are stored as PBKDF2 hashes, and production deployments require a custom session secret.",
                "Editor jobs and SaaS records persist through Supabase when configured, with local atomic JSON fallback for development.",
                "Built-in no-key MT can run through local/self-hosted engines, reducing reliance on user API keys.",
                "Sensitive-text indicators help reviewers identify emails, phone-like values, and credential-like content before routing externally.",
            ],
        },
    }
    doc = docs.get(kind, docs["terms"])
    st.html(
        dedent(f"""
        <div class="es-auth-shell">
          <div class="es-lp-brand">
            <div class="es-lp-logo">ES</div>
            <div>
              <div class="es-lp-brand-name">ErrorSweep</div>
              <div class="es-lp-brand-sub">Nawin Corp</div>
            </div>
          </div>
          <div class="es-auth-links">
            <a class="es-lp-btn" href="{public_page_link('landing')}" target="_self">Back to landing</a>
            <a class="es-lp-btn primary" href="{public_page_link('signup')}" target="_self">Sign up</a>
          </div>
        </div>
        """).strip(),
    )
    st.markdown(f"## {doc['title']}")
    for item in doc["body"]:
        st.markdown(f"- {item}")
    st.caption("Draft policy text for local product testing. Replace with reviewed legal documents before public launch.")


def render_verify_email() -> None:
    st.html(
        dedent(f"""
        <div class="es-auth-shell">
          <div class="es-lp-brand">
            <div class="es-lp-logo">ES</div>
            <div>
              <div class="es-lp-brand-name">ErrorSweep</div>
              <div class="es-lp-brand-sub">Nawin Corp</div>
            </div>
          </div>
          <div class="es-auth-links">
            <a class="es-lp-btn" href="{public_page_link('landing')}" target="_self">Back to landing</a>
            <a class="es-lp-btn primary" href="{public_page_link('login')}" target="_self">Login</a>
          </div>
        </div>
        """).strip(),
    )
    st.markdown("## Verify your email")
    token = query_get("token")
    if not token:
        st.error("Verification token is missing.")
        return
    record = find_auth_token(token, "email_verification")
    if not record:
        st.error("This verification link is invalid or expired. Request a new link from the login screen.")
        return
    email = safe_text(record.get("email"))
    verified = update_stored_user(email, {"email_verified": True, "verified_at": now_stamp(), "status": "Active"})
    consume_auth_token(record)
    if verified:
        add_audit("Email verified", email)
        st.success("Email verified. You can now sign in.")
        st.link_button("Go to login", public_page_link("login"), use_container_width=True)
    else:
        st.error("The account for this verification link could not be found.")


def render_password_reset() -> None:
    st.html(
        dedent(f"""
        <div class="es-auth-shell">
          <div class="es-lp-brand">
            <div class="es-lp-logo">ES</div>
            <div>
              <div class="es-lp-brand-name">ErrorSweep</div>
              <div class="es-lp-brand-sub">Nawin Corp</div>
            </div>
          </div>
          <div class="es-auth-links">
            <a class="es-lp-btn" href="{public_page_link('landing')}" target="_self">Back to landing</a>
            <a class="es-lp-btn primary" href="{public_page_link('login')}" target="_self">Login</a>
          </div>
        </div>
        """).strip(),
    )
    st.markdown("## Reset password")
    token = query_get("token")
    if not token:
        st.error("Password reset token is missing. Request a new reset link from the login screen.")
        return
    record = find_auth_token(token, "password_reset")
    if not record:
        st.error("This reset link is invalid or expired. Request a new link from the login screen.")
        return
    with st.form("password_reset_form"):
        password = st.text_input("New password", type="password")
        confirm = st.text_input("Confirm password", type="password")
        submitted = st.form_submit_button("Reset password", use_container_width=True)
    if submitted:
        if len(password) < 8:
            st.error("Use at least 8 characters for the new password.")
            return
        if password != confirm:
            st.error("Passwords do not match.")
            return
        email = safe_text(record.get("email"))
        updated = update_stored_user(email, {"password_hash": hash_password(password), "status": "Active"})
        consume_auth_token(record)
        if updated:
            add_audit("Password reset completed", email)
            st.success("Password reset complete. You can now sign in.")
            st.link_button("Go to login", public_page_link("login"), use_container_width=True)
        else:
            st.error("The account for this reset link could not be found.")


def render_public_app() -> None:
    route = (query_get("public") or "landing").strip().lower()
    if route == "login":
        render_login()
    elif route == "signup":
        render_signup()
    elif route == "verify":
        render_verify_email()
    elif route == "reset":
        render_password_reset()
    elif route in {"terms", "privacy", "security"}:
        render_public_document(route)
    else:
        render_landing_page()


# ==========================================================
# Pages
# ==========================================================

def _legacy_page_dashboard_unused() -> None:
    hero("Dashboard", "Localization operations hub", "Manage projects, jobs, review tasks, scorecards, and translation memory from one workspace.")
    pending_review = sum(1 for r in st.session_state.review_segments if r.get("status") not in ("Approved", "Rejected"))
    spark_base = [
        len(st.session_state.projects),
        len(st.session_state.jobs),
        len(st.session_state.tm),
        pending_review,
        len(st.session_state.audit_logs),
        len(st.session_state.glossary),
        len(st.session_state.dnt),
    ]
    st.markdown(
        f"""
        <div class="es-bento">
          <div class="es-bento-card wide">
            <div class="es-metric-label">Mission Control</div>
            <div class="es-metric-value">{len(st.session_state.jobs)}</div>
            <div class="es-small">Jobs across QA, Pro, subtitle, transcription, and scorecard workflows.</div>
            {sparkline_svg(spark_base)}
          </div>
          <div class="es-bento-card">
            <div class="es-metric-label">Projects</div>
            <div class="es-metric-value">{len(st.session_state.projects)}</div>
            <div class="es-small">Client/product workspaces</div>
            {sparkline_svg([0, len(st.session_state.projects), len(st.session_state.projects) + len(st.session_state.glossary)])}
          </div>
          <div class="es-bento-card">
            <div class="es-metric-label">TM Entries</div>
            <div class="es-metric-value">{len(st.session_state.tm)}</div>
            <div class="es-small">Approved translations</div>
            {sparkline_svg([0, max(1, len(st.session_state.tm)//3), len(st.session_state.tm)])}
          </div>
          <div class="es-bento-card">
            <div class="es-metric-label">Pending Review</div>
            <div class="es-metric-value">{pending_review}</div>
            <div class="es-small">Segments requiring attention</div>
            {sparkline_svg([pending_review + 2, pending_review + 1, pending_review])}
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("### Recommended next steps")
    c1, c2, c3 = st.columns(3)
    with c1.container(border=True):
        st.markdown("### 📁 Create a project")
        st.caption("Set source/target languages, domain, and reusable rules.")
    with c2.container(border=True):
        st.markdown("### 🚀 Run QA or Pro")
        st.caption("Upload bilingual files or source files and route outputs to review.")
    with c3.container(border=True):
        st.markdown("### 🎬 Subtitle / Transcription")
        st.caption("Create subtitles, transcripts, and timing rows in a focused editor.")

    jobs_col, activity_col = st.columns([0.62, 0.38], gap="large")
    with jobs_col:
        st.markdown("### Recent jobs")
        if st.session_state.jobs:
            st.dataframe(pd.DataFrame(display_records(st.session_state.jobs)), use_container_width=True, hide_index=True)
        else:
            st.info("No jobs yet.")
    with activity_col:
        st.markdown("### Activity pulse")
        items = []
        for item in st.session_state.audit_logs[:8]:
            items.append(
                f'<div class="es-activity-item"><div class="es-small">{escape(format_local_time(item.get("time", "")))}</div>'
                f'<div><b>{escape(safe_text(item.get("action", "")))}</b><br><span class="es-small">{escape(safe_text(item.get("details", "")))}</span></div></div>'
            )
        activity_html = "".join(items) or '<span class="es-small">No activity yet.</span>'
        st.markdown(f'<div class="es-activity-drawer">{activity_html}</div>', unsafe_allow_html=True)

def page_dashboard() -> None:
    user = current_user() or {}
    pending_review = sum(1 for r in st.session_state.review_segments if r.get("status") not in ("Approved", "Rejected"))
    total_jobs = len(st.session_state.jobs)
    total_projects = len(st.session_state.projects)
    total_memory = len(st.session_state.tm)
    active_rules = len(st.session_state.glossary) + len(st.session_state.dnt)
    completed_jobs = sum(1 for j in st.session_state.jobs if safe_text(j.get("status")).lower() == "completed")
    tqi_score = max(62, min(99, 88 + completed_jobs * 2 - pending_review * 3 + min(active_rules, 8)))
    attention_items = dashboard_attention_items()
    spark_base = [
        total_projects,
        total_jobs,
        total_memory,
        pending_review,
        len(st.session_state.audit_logs),
        len(st.session_state.glossary),
        len(st.session_state.dnt),
    ]
    st.html(
        dedent(f"""
        <section class="es-personal-hero">
          <div class="es-hero-row">
            <div>
              <div class="es-kicker">Workspace command center</div>
              <div class="es-welcome-title">Good {escape('morning' if datetime.now().hour < 12 else 'afternoon' if datetime.now().hour < 17 else 'evening')}, {escape(first_name_from_user(user))}</div>
              <div class="es-hero-summary">
                You have <b>{pending_review}</b> segment(s) waiting for review, <b>{len(attention_items)}</b> priority item(s), and <b>{active_rules}</b> saved rule asset(s) ready for QA and translation.
              </div>
              <div class="es-fab-row">
                <a class="es-fab-action" href="{page_link('Projects')}" target="_self">New Project</a>
                <a class="es-fab-action" href="{page_link('ErrorSweep Pro')}" target="_self">Run Pro Translation</a>
                <a class="es-fab-action secondary" href="{page_link('Memory & Rules')}" target="_self">Upload Rules</a>
                <a class="es-fab-action secondary" href="{page_link('ErrorSweep QA')}" target="_self">Run QA</a>
              </div>
            </div>
            <div class="es-hero-orb">ErrorSweep<br/>Live</div>
          </div>
        </section>

        <div class="es-bento">
          <div class="es-bento-card wide">
            <div class="es-metric-label">Mission Control</div>
            <div class="es-metric-value">{total_jobs}</div>
            <div class="es-small">Jobs across QA, Pro, subtitle, transcription, and scorecard workflows.</div>
            {area_chart_svg(spark_base, "mission")}
          </div>
          <div class="es-bento-card">
            <div class="es-metric-label">Projects</div>
            <div class="es-metric-value">{total_projects}</div>
            <div class="es-small">Client/product workspaces</div>
            {area_chart_svg([0, total_projects, total_projects + len(st.session_state.glossary)], "projects")}
          </div>
          <div class="es-bento-card">
            <div class="es-metric-label">Language Memory</div>
            <div class="es-metric-value">{total_memory}</div>
            <div class="es-small">Approved translations</div>
            {area_chart_svg([0, max(1, total_memory//3), total_memory], "memory")}
          </div>
          <div class="es-bento-card">
            {radial_progress_svg(tqi_score, "TQI")}
          </div>
        </div>

        <div class="es-dashboard-grid">
          <section class="es-dashboard-panel">
            <div class="es-dashboard-title">
              <h3>Production workflow</h3>
              <span class="es-status-pill">Live workspace</span>
            </div>
            <div class="es-flow">
              <div class="es-flow-step">
                <span class="es-code-chip">01 Intake</span>
                <b>Create project</b>
                <div class="es-small">Source language, targets, domain, reusable client rules.</div>
              </div>
              <div class="es-flow-step">
                <span class="es-code-chip">02 Analyze</span>
                <b>Run QA or Pro</b>
                <div class="es-small">Detect placeholders, DNT, terminology, and linguistic risk.</div>
              </div>
              <div class="es-flow-step">
                <span class="es-code-chip">03 Review</span>
                <b>Human workspace</b>
                <div class="es-small">Edit segments with source, target, memory, and QA context.</div>
              </div>
              <div class="es-flow-step">
                <span class="es-code-chip">04 Deliver</span>
                <b>Reports</b>
                <div class="es-small">Export QA reports, subtitles, scorecards, and reviewed files.</div>
              </div>
            </div>
          </section>

          <section class="es-dashboard-panel es-orbit">
            <div class="es-dashboard-title">
              <h3>System readiness</h3>
              <span class="es-status-pill">Operational</span>
            </div>
            <div class="es-node"><strong>{active_rules} active rules</strong><span>Glossary and DNT terms available to QA and translation.</span></div>
            <div class="es-node"><strong>{len(st.session_state.audit_logs)} activity events</strong><span>Workspace actions are tracked for owner review.</span></div>
            <div class="es-node"><strong>{pending_review} review queue</strong><span>Segments waiting for human decision.</span></div>
            <div class="es-node"><strong>{escape(current_ai_route_label())}</strong><span>Current AI and translation route.</span></div>
          </section>
        </div>
        """).strip(),
    )

    st.markdown("")
    jobs_col, activity_col = st.columns([0.66, 0.34], gap="large")
    with jobs_col:
        if attention_items:
            attention_html = []
            for item in attention_items:
                severity = safe_text(item.get("severity", ""))
                attention_html.append(
                    f'<div class="es-attention-item {escape(severity)}">'
                    f'<span class="es-attention-dot"></span>'
                    f'<div><b>{escape(safe_text(item.get("title", "")))}</b><br><span class="es-small">{escape(safe_text(item.get("meta", "")))}</span></div>'
                    f'<span class="es-code-chip">{escape(severity or "open")}</span>'
                    f'</div>'
                )
            body_html = f'<div class="es-attention-list">{"".join(attention_html)}</div>'
        else:
            body_html = """
                <div class="es-empty-state">
                  <div>
                    <div class="es-empty-icon"></div>
                    <h3>No urgent work waiting</h3>
                    <div class="es-small">Start by creating a project, importing rules, or running QA on a bilingual file.</div>
                  </div>
                </div>
                """
        st.html(
            f"""
            <div class="es-dashboard-panel">
              <div class="es-dashboard-title"><h3>Needs attention</h3><span class="es-code-chip">Priority queue</span></div>
              {body_html}
            </div>
            """
        )
    with activity_col:
        items = []
        for item in st.session_state.audit_logs[:8]:
            items.append(
                f'<div class="es-activity-item"><div class="es-small">{escape(format_local_time(item.get("time", "")))}</div>'
                f'<div><b>{escape(safe_text(item.get("action", "")))}</b><br><span class="es-small">{escape(safe_text(item.get("details", "")))}</span></div></div>'
            )
        activity_html = "".join(items) or '<span class="es-small">No activity yet.</span>'
        st.html(
            f"""
            <div class="es-dashboard-panel">
              <div class="es-dashboard-title"><h3>Activity pulse</h3><span class="es-code-chip">Audit</span></div>
              <div class="es-activity-drawer">{activity_html}</div>
            </div>
            """
        )


def page_projects() -> None:
    hero("Projects", "Client and product workspaces", "Create language projects, attach rule packs, and keep memory scoped correctly.")
    with st.form("create_project", enter_to_submit=False):
        c1, c2 = st.columns(2)
        name = c1.text_input("Project name")
        client = c2.text_input("Client / workspace", value=(current_user() or {}).get("workspace", "Demo Workspace"))
        c3, c4 = st.columns(2)
        source_lang = c3.selectbox("Source language", LANGUAGE_CATALOG, index=LANGUAGE_CATALOG.index("English"))
        target_langs = c4.multiselect(
            "Target languages",
            LANGUAGE_CATALOG,
            default=["French", "Spanish", "German", "Italian", "Portuguese", "Telugu", "Hindi", "Tamil", "Malayalam"],
        )
        domain = st.selectbox("Default domain", ["Auto-detect", "Software UI", "Marketing", "Legal", "Medical", "E-learning", "Subtitling", "Gaming", "Finance", "General"])
        submitted = st.form_submit_button("Create project", use_container_width=True)
    if submitted:
        if not safe_text(name).strip():
            st.error("Please enter a project name before creating the project.")
            return
        if not safe_text(client).strip():
            st.error("Please enter a client/workspace before creating the project.")
            return
        if not target_langs:
            st.error("Please select at least one target language before creating the project.")
            return
        project_id = f"{safe_text(client).replace(' ', '')}-{safe_text(name).replace(' ', '')}-{datetime.now().strftime('%Y%m%d%H%M%S')}"
        record = persist_saas_record("projects", {
            "id": project_id,
            "created": now_stamp(),
            "project": name,
            "client": client,
            "workspace": client,
            "source": source_lang,
            "targets": ", ".join(target_langs),
            "domain": domain,
            "status": "Active",
            "job_count": 0,
        })
        st.session_state.projects.append(record)
        trim_session_list("projects")
        add_audit("Project created", name)
        st.success("Project created.")

    if st.session_state.projects:
        st.dataframe(pd.DataFrame(display_records(st.session_state.projects)), use_container_width=True, hide_index=True)
    else:
        st.info("No projects yet.")


def page_jobs() -> None:
    hero("Jobs", "Workflow queue", "Track uploads, translation, QA, Pro post-editing, scorecards, and delivery status.")
    render_task_queue_panel()
    if st.session_state.jobs:
        render_jobs_kanban(st.session_state.jobs)
        labels = [
            f"{idx + 1}. {safe_text(job.get('type', 'Job'))} - {safe_text(job.get('status', ''))} - {safe_text(job.get('language', job.get('workspace', '')))}"
            for idx, job in enumerate(st.session_state.jobs)
        ]
        with st.expander("Open job context panel", expanded=False):
            selected = st.selectbox("Job", labels, label_visibility="collapsed")
            job = st.session_state.jobs[labels.index(selected)]
            st.html(
                f"""
                <div class="es-flyout">
                  <span class="es-code-chip">Job context</span>
                  <h3>{escape(safe_text(job.get("type", "Job")))}</h3>
                  <div class="es-mini-table">
                    <div class="es-mini-row"><span>Status</span><b>{escape(safe_text(job.get("status", "")))}</b></div>
                    <div class="es-mini-row"><span>Project</span><b>{escape(safe_text(job.get("project", "")) or "Unlinked")}</b></div>
                    <div class="es-mini-row"><span>Language</span><b>{escape(safe_text(job.get("language", "")))}</b></div>
                    <div class="es-mini-row"><span>Attachments</span><b>{escape(safe_text(job.get("attachment_count", 0)))}</b></div>
                    <div class="es-mini-row"><span>Workspace</span><b>{escape(safe_text(job.get("workspace", "")))}</b></div>
                    <div class="es-mini-row"><span>Created</span><b>{escape(format_local_time(job.get("created", job.get("created_at", ""))))}</b></div>
                  </div>
                  <p class="es-small">{escape(safe_text(job.get("note", "")) or "No job note supplied.")}</p>
                </div>
                """
            )
        with st.expander("Raw job table", expanded=False):
            st.dataframe(pd.DataFrame(display_records(st.session_state.jobs)), use_container_width=True, hide_index=True)
    else:
        st.html(
            """
            <div class="es-empty-state">
              <div>
                <div class="es-empty-icon"></div>
                <h3>No workflow jobs yet</h3>
                <div class="es-small">Create a manual job, run QA, or start Pro translation to populate the pipeline.</div>
              </div>
            </div>
            """
        )
    st.markdown("### Create manual job")
    project_options = ["No project selected"]
    project_lookup: Dict[str, Dict[str, Any]] = {}
    for project in st.session_state.get("projects", []):
        label = f"{safe_text(project.get('project', 'Project'))} · {safe_text(project.get('client') or project.get('workspace', 'Workspace'))}"
        project_options.append(label)
        project_lookup[label] = project
    with st.form("manual_job", enter_to_submit=False):
        c1, c2, c3 = st.columns(3)
        job_type = c1.selectbox("Job type", ["QA", "Pro Translation", "Post-editing Review", "Subtitle Review", "Transcription", "Scorecard"])
        language = c2.selectbox("Target language", LANGUAGE_CATALOG, index=LANGUAGE_CATALOG.index("French"))
        assignee = c3.text_input("Assignee", value="reviewer@errorsweep.local")
        project_label = st.selectbox("Project context", project_options)
        assignment_files = st.file_uploader(
            "Assignment upload (optional)",
            accept_multiple_files=True,
            help="Upload any file or ZIP package that should be assigned with this job.",
        )
        note = st.text_area("Notes", height=80)
        submitted = st.form_submit_button("Create job", use_container_width=True)
    if submitted:
        if not safe_text(language).strip():
            st.error("Please select a target language before creating the job.")
            return
        if not safe_text(assignee).strip():
            st.error("Please enter an assignee before creating the job.")
            return
        selected_project = project_lookup.get(project_label, {})
        job_id = uuid.uuid4().hex
        attachment_manifests = save_job_attachment_files(job_id, assignment_files or [])
        record = persist_saas_record("jobs", {
            "id": job_id,
            "created": now_stamp(),
            "workspace": (current_user() or {}).get("workspace", "Demo Workspace"),
            "type": job_type,
            "language": language,
            "assignee": assignee,
            "status": "Draft",
            "note": note,
            "project_id": safe_text(selected_project.get("id", "")),
            "project": safe_text(selected_project.get("project", "")),
            "attachment_count": len(attachment_manifests),
            "attachments_json": attachment_manifests,
        })
        st.session_state.jobs.insert(0, record)
        if selected_project:
            selected_project["job_count"] = int(selected_project.get("job_count") or 0) + 1
        trim_session_list("jobs")
        add_audit("Manual job created", f"{job_type} assigned to {assignee}")
        queue_email_notification(
            assignee,
            "New ErrorSweep job assigned",
            f"A new {job_type} job has been assigned to you for {language}. Attachments: {len(attachment_manifests)}.",
            "job.assigned",
            metadata={"job_id": job_id, "job_type": job_type, "language": language, "attachment_count": len(attachment_manifests)},
        )
        st.success("Job created.")


def page_qa() -> None:
    hero("ErrorSweep QA", "Review existing translation", "Upload bilingual files, detect issues, and create review-ready findings.")
    render_stepper(["Upload bilingual file", "Select rules and strictness", "Run QA report"], active_idx=0)
    render_upload_dropzone("Drop bilingual content here", "Supports Excel, CSV, DOCX, TXT, SRT, and VTT. Rules ZIP can be attached for client-specific QA.", "XLSX / CSV / DOCX / SRT")
    file = st.file_uploader("Upload bilingual file", type=["xlsx", "csv", "docx", "txt", "srt", "vtt"], key="qa_file")
    rules = st.file_uploader("Upload rules ZIP (optional)", type=["zip"], key="qa_rules")
    render_rules_zip_warning(rules)
    c1, c2 = st.columns(2)
    strictness = c1.selectbox("Strictness", ["Lenient", "Standard", "Strict", "Very Strict"], index=2)
    domain = c2.selectbox("Domain", ["Auto-detect", "Software UI", "Marketing", "Legal", "Medical", "E-learning", "Subtitling", "General"])

    if st.button("Run QA", use_container_width=True, disabled=file is None):
        task = create_task_record(
            "qa",
            f"QA report: {getattr(file, 'name', 'uploaded file')}",
            metadata={"file_name": getattr(file, "name", ""), "strictness": strictness, "domain": domain},
        )
        update_task_record(task["id"], status="running", progress=5)
        rows = extract_rows_from_upload(file)
        for upload in (file, rules):
            try:
                if upload is not None and hasattr(upload, "seek"):
                    upload.seek(0)
            except Exception as exc:
                LOGGER.debug("Unable to rewind uploaded file before queueing: %s", exc)
        if not rows:
            update_task_record(task["id"], status="failed", progress=100, error="No segments found in uploaded file.")
            st.error("No segments found.")
            return
        allowed, usage_message, usage_details = check_workspace_usage_allowance(rows, "QA run")
        if not allowed:
            update_task_record(task["id"], status="failed", progress=100, error=usage_message, metadata_json={**(task.get("metadata_json") or {}), "usage_check": usage_details})
            st.error(usage_message)
            return
        if usage_message:
            st.warning(usage_message)
        if queue_external_workflow_if_configured(
            task,
            "qa",
            file,
            rules,
            parameters={
                "file_name": getattr(file, "name", ""),
                "strictness": strictness,
                "domain": domain,
                "estimated_usage": usage_details,
            },
        ):
            return
        update_task_record(task["id"], progress=15, total_units=len(rows), processed_units=0)
        client_rules = workspace_rules(rules)
        if client_rules.get("warnings"):
            for warning in client_rules.get("warnings", []):
                st.warning(warning)
        st.info(
            "Rules applied: "
            f"{len(client_rules.get('glossary', []))} glossary terms, "
            f"{len(client_rules.get('dnt', []))} DNT terms, "
            f"{len(client_rules.get('instructions', []))} instructions."
        )
        ai_by_id: Dict[str, List[Dict[str, Any]]] = {}
        if st.session_state.get("byo_openai_api_key"):
            update_task_record(task["id"], progress=25, processed_units=0)
            with st.spinner("Running optional AI QA against client rules..."):
                for item in call_main_api_qa(rows, domain=domain, strictness=strictness, rules=client_rules):
                    row_id = safe_text(item.get("id"))
                    if row_id:
                        ai_by_id.setdefault(row_id, []).append(item)
        update_task_record(task["id"], progress=40, processed_units=0)
        findings = []
        report_rows = []
        for idx, r in enumerate(rows, start=1):
            src = r.get("source", "")
            tgt = r.get("target", "")
            status = "Pass"
            row_findings = []
            if not tgt:
                row_findings.append(qa_manual_finding(r, "Major", "Completeness", "Target translation is blank.", "Add the missing translation."))
            for ph in re.findall(r"\{\{[^}]+\}\}", src):
                if ph not in tgt:
                    row_findings.append(qa_manual_finding(r, "Major", "Placeholder", f"Missing placeholder {ph}", f"Keep {ph} unchanged in the target."))
            if src and tgt and src == tgt:
                row_findings.append(qa_manual_finding(r, "Critical", "Untranslated Text", "Target is identical to source.", "Translate the source text or confirm this is intentionally unchanged."))
            global_findings = run_global_qa_for_row(r, target_language="Auto-detect", domain=domain, rules=client_rules)
            row_findings.extend(global_findings)
            for item in ai_by_id.get(safe_text(r.get("id")), []):
                severity = safe_text(item.get("severity") or "Review").title()
                if severity not in {"Minor", "Major", "Critical"}:
                    severity = "Review"
                ai_finding = qa_manual_finding(
                    r,
                    severity,
                    "AI QA",
                    safe_text(item.get("issue") or item.get("reason") or "AI reviewer flagged this segment."),
                    safe_text(item.get("suggestion") or ""),
                )
                ai_finding["Check Source"] = "ErrorSweep AI QA"
                ai_finding["Rule Source"] = "Uploaded/Saved Rules + AI"
                ai_finding["Rule ID"] = "ai.rules_qa"
                ai_finding["Explanation"] = safe_text(item.get("reason") or ai_finding["Explanation"])
                row_findings.append(ai_finding)
            report_rows.extend(row_findings)
            if row_findings:
                status = qa_overall_status(row_findings)
            issue_summary = "; ".join(
                f"{f.get('Severity', 'Review')}: {f.get('Error Type', 'QA')} - {f.get('Wrong Part') or f.get('Explanation', '')}"
                for f in row_findings[:4]
            )
            findings.append({**r, "status": status, "issues": issue_summary, "match": r.get("match") or ("MT" if tgt else "Untranslated")})
            if idx == len(rows) or idx % 25 == 0:
                update_task_record(
                    task["id"],
                    progress=min(85, 40 + int((idx / max(len(rows), 1)) * 45)),
                    processed_units=idx,
                    total_units=len(rows),
                )

        report_rows = dedupe_qa_findings(report_rows + delivery_quality_findings(findings, "Auto-detect", domain, client_rules))
        for row in findings:
            row_findings = [
                f for f in report_rows
                if safe_text(f.get("Location")) == safe_text(row.get("location") or f"Segment {row.get('id', '')}")
            ]
            row["status"] = qa_overall_status(row_findings)
            row["issues"] = "; ".join(
                f"{f.get('Severity', 'Review')}: {f.get('Error Type', 'QA')}"
                for f in row_findings[:4]
            )

        qa_job = persist_saas_record("jobs", {
            "created": now_stamp(),
            "workspace": (current_user() or {}).get("workspace", "Demo Workspace"),
            "type": "QA",
            "language": "",
            "status": "Completed",
            "segments": len(findings),
        })
        st.session_state.jobs.insert(0, qa_job)
        trim_session_list("jobs")
        add_audit("QA run", f"{len(findings)} segments")
        professional_findings, segment_overview, qa_summary = professional_qa_rows(findings, report_rows)
        update_task_record(
            task["id"],
            status="completed",
            progress=100,
            processed_units=len(findings),
            total_units=len(findings),
            result_ref=safe_text(qa_job.get("id", "")),
            metadata_json={"score": qa_summary["qa_score"], "result": qa_summary["result"], "findings": qa_summary["total_findings"]},
        )
        record_billable_workflow_usage("qa_workflow", findings, provider="errorsweep_qa", model="deterministic_ai_rules")
        queue_email_notification(
            (current_user() or {}).get("email", ""),
            "ErrorSweep QA completed",
            f"QA completed for {len(findings)} segment(s). Score: {qa_summary['qa_score']} ({qa_summary['result']}).",
            "qa.completed",
            metadata={"segments": len(findings), "score": qa_summary["qa_score"], "result": qa_summary["result"]},
        )
        st.success("QA completed. Download the professional QA workbook below. Post-editing Human Review is available after Pro translation runs only.")
        metrics([
            ("QA Score", qa_summary["qa_score"], qa_summary["result"]),
            ("Segments", qa_summary["total_segments"], "checked"),
            ("Findings", qa_summary["total_findings"], "issues"),
            ("Needs Review", qa_summary["review_segments"], "segments"),
        ])
        render_delivery_gate(findings, report_rows, "QA delivery gate")
        st.dataframe(pd.DataFrame(segment_overview), use_container_width=True, hide_index=True)
        st.download_button(
            "Download Professional QA Excel",
            create_qa_excel_report(findings, report_rows),
            file_name="ErrorSweep_QA_Report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
        if professional_findings:
            st.markdown("### QA Findings")
            finding_preview_cols = [
                "Finding ID", "Segment ID", "Source Text", "Target Text", "Suggested Target",
                "Error Category", "Severity", "Issue", "Explanation", "Check", "Confidence", "Rule ID",
            ]
            st.dataframe(pd.DataFrame(professional_findings)[finding_preview_cols], use_container_width=True, hide_index=True)
            st.download_button("Download QA Findings CSV", rows_to_csv(professional_findings), file_name="errorsweep_qa_findings_professional.csv", mime="text/csv", use_container_width=True)
        st.download_button("Download Segment Overview CSV", rows_to_csv(segment_overview), file_name="errorsweep_qa_segment_overview.csv", mime="text/csv", use_container_width=True)


def page_pro() -> None:
    hero("ErrorSweep Pro", "Translate + QA + Human Review", "Translate first, then open a dedicated Human Review workspace for editing and approval.")
    st.caption(f"AI access: {current_ai_route_label()}")
    render_privacy_route_notice("Translation route")
    render_stepper(["Upload source", "Translate with rules", "Open Human Review"], active_idx=0)
    with st.expander("Built-in MT engine diagnostics", expanded=False):
        if builtin_engine_status is None:
            st.warning("Built-in MT router diagnostics are unavailable.")
        else:
            status_rows = builtin_engine_status()
            display_rows = []
            for row in status_rows:
                detail = row.get("detail", "")
                if isinstance(detail, dict):
                    detail = f"{detail.get('provider', row.get('engine'))} on {detail.get('device', 'unknown')}"
                display_rows.append({
                    "Engine": row.get("engine"),
                    "Role": row.get("priority"),
                    "Endpoint": row.get("endpoint"),
                    "Enabled": bool(row.get("enabled")),
                    "Ready": bool(row.get("ready")),
                    "Detail": safe_text(detail)[:180],
                })
            st.dataframe(pd.DataFrame(display_rows), use_container_width=True, hide_index=True)
            if st.button("Run no-key MT smoke tests", use_container_width=True):
                if smoke_test_builtin_engines is None:
                    st.warning("Smoke tests are unavailable.")
                else:
                    with st.spinner("Testing local MT engines..."):
                        smoke_rows = smoke_test_builtin_engines()
                    st.dataframe(pd.DataFrame(smoke_rows), use_container_width=True, hide_index=True)

    # If a Pro review session already exists, keep the action visible even after reruns.
    # This prevents the user from losing access to the post-editing workspace.
    if st.session_state.get("review_segments") or st.session_state.get("last_pro_review_segments"):
        restore_human_review_session_from_cache()
        st.success("A Pro Human Review session is ready.")
        if st.button("Open Human Review workspace", type="primary", use_container_width=True, key="open_existing_pro_review"):
            go_to_human_review_workspace()
    render_upload_dropzone("Drop source or bilingual content here", "ErrorSweep will translate, apply saved/uploaded rules, run QA, and prepare a Human Review workspace.", "XLSX / CSV / DOCX / SRT")
    uploaded = st.file_uploader("Upload source or bilingual file", type=["xlsx", "csv", "docx", "txt", "srt", "vtt"], key="pro_file")
    rules_zip = st.file_uploader("Upload rules ZIP (optional)", type=["zip"], key="pro_rules")
    render_rules_zip_warning(rules_zip)
    c1, c2, c3 = st.columns([1.2, 1.2, 1])
    target_language = c1.text_input("Target language", value="French")
    domain_choice = c2.selectbox("Domain", ["Auto-detect", "Software UI", "Marketing", "Legal", "Medical", "E-learning", "Subtitling", "Gaming", "Finance", "General"])
    threshold = c3.slider("Allow with review threshold", min_value=0, max_value=25, value=12, help="Incomplete rows below this percentage can be routed to Human Review.")

    if st.button("Run Translate + Review", use_container_width=True, disabled=uploaded is None):
        task = create_task_record(
            "pro_translation",
            f"Pro translation: {getattr(uploaded, 'name', 'uploaded file')} -> {target_language}",
            metadata={"file_name": getattr(uploaded, "name", ""), "target_language": target_language, "domain": domain_choice},
        )
        update_task_record(task["id"], status="running", progress=5)
        rows = extract_rows_from_upload(uploaded)
        for upload in (uploaded, rules_zip):
            try:
                if upload is not None and hasattr(upload, "seek"):
                    upload.seek(0)
            except Exception as exc:
                LOGGER.debug("Unable to rewind uploaded file before queueing: %s", exc)
        if not rows:
            update_task_record(task["id"], status="failed", progress=100, error="No segments found in uploaded file.")
            st.error("No segments found.")
            return
        allowed, usage_message, usage_details = check_workspace_usage_allowance(rows, "Pro translation")
        if not allowed:
            update_task_record(task["id"], status="failed", progress=100, error=usage_message, metadata_json={**(task.get("metadata_json") or {}), "usage_check": usage_details})
            st.error(usage_message)
            return
        if usage_message:
            st.warning(usage_message)
        if queue_external_workflow_if_configured(
            task,
            "pro_translation",
            uploaded,
            rules_zip,
            parameters={
                "file_name": getattr(uploaded, "name", ""),
                "target_language": target_language,
                "domain": domain_choice,
                "review_threshold": threshold,
                "estimated_usage": usage_details,
            },
        ):
            return
        update_task_record(task["id"], progress=12, total_units=len(rows), processed_units=0)
        sensitive = sensitive_rows_summary(rows)
        if sensitive["matches"]:
            labels = ", ".join(f"row {m['row']} ({', '.join(m['kinds'])})" for m in sensitive["matches"])
            st.warning(f"Sensitive text indicators found before translation: {labels}. Check client approval before using external BYO AI routes.")
        sample = " ".join([r.get("source", "") for r in rows[:20]])
        domain = auto_detect_domain(sample) if domain_choice == "Auto-detect" else domain_choice
        client_rules = workspace_rules(rules_zip)
        st.info(f"Detected domain: {domain}")
        st.info(
            "Rules applied: "
            f"{len(client_rules.get('glossary', []))} glossary terms, "
            f"{len(client_rules.get('dnt', []))} DNT terms, "
            f"{len(client_rules.get('instructions', []))} instructions."
        )

        source_texts = [r.get("source", "") or r.get("target", "") for r in rows]
        update_task_record(task["id"], progress=25, processed_units=0)
        with st.spinner("Translating with available AI route..."):
            translations = call_main_api_translate(source_texts, target_language, domain, rules=client_rules)
        update_task_record(task["id"], progress=55, processed_units=len(translations), total_units=len(rows))

        review_rows = []
        missing = 0
        for idx, (r, trans) in enumerate(zip(rows, translations), start=1):
            r["target"] = trans
            r["status"] = "MT"
            r["match"] = "MT"
            if not trans or trans.strip() == r.get("source", "").strip():
                missing += 1
                r["status"] = "Needs Review"
                r["match"] = "Untranslated"
            qa_findings = run_global_qa_for_row(r, target_language=target_language, domain=domain, rules=client_rules)
            if qa_findings:
                r["qa_findings"] = qa_findings[:10]
                r["qa_summary"] = summarize_qa_findings(qa_findings)
                if r["status"] in {"MT", "Existing"}:
                    r["status"] = "Needs Review"
            review_rows.append(r)
            if idx == len(rows) or idx % 25 == 0:
                update_task_record(
                    task["id"],
                    progress=min(78, 55 + int((idx / max(len(rows), 1)) * 23)),
                    processed_units=idx,
                    total_units=len(rows),
                )

        if st.session_state.get("byo_openai_api_key"):
            ai_by_id: Dict[str, List[Dict[str, Any]]] = {}
            update_task_record(task["id"], progress=82, processed_units=len(review_rows), total_units=len(rows))
            with st.spinner("Running optional AI QA against client rules..."):
                for item in call_main_api_qa(review_rows, domain=domain, strictness="Standard", rules=client_rules):
                    row_id = safe_text(item.get("id"))
                    if row_id:
                        ai_by_id.setdefault(row_id, []).append(item)
            for r in review_rows:
                ai_findings = []
                for item in ai_by_id.get(safe_text(r.get("id")), []):
                    severity = safe_text(item.get("severity") or "Review").title()
                    if severity not in {"Minor", "Major", "Critical"}:
                        severity = "Review"
                    ai_finding = qa_manual_finding(
                        r,
                        severity,
                        "AI QA",
                        safe_text(item.get("issue") or item.get("reason") or "AI reviewer flagged this segment."),
                        safe_text(item.get("suggestion") or ""),
                    )
                    ai_finding["Check Source"] = "ErrorSweep AI QA"
                    ai_finding["Rule Source"] = "Uploaded/Saved Rules + AI"
                    ai_finding["Rule ID"] = "ai.rules_qa"
                    ai_finding["Explanation"] = safe_text(item.get("reason") or ai_finding["Explanation"])
                    ai_findings.append(ai_finding)
                if ai_findings:
                    existing = list(r.get("qa_findings") or [])
                    r["qa_findings"] = (existing + ai_findings)[:10]
                    r["qa_summary"] = summarize_qa_findings(r["qa_findings"])
                    if r["status"] in {"MT", "Existing"}:
                        r["status"] = "Needs Review"

        gate_findings = delivery_quality_findings(review_rows, target_language, domain, client_rules)
        gate_summary = delivery_gate_summary(review_rows, gate_findings)
        findings_by_location: Dict[str, List[Dict[str, Any]]] = {}
        for finding in gate_findings:
            findings_by_location.setdefault(safe_text(finding.get("Location")), []).append(finding)
        for idx, r in enumerate(review_rows, start=1):
            location = safe_text(r.get("location") or f"Segment {r.get('id', idx)}")
            row_gate_findings = findings_by_location.get(location, [])
            if row_gate_findings:
                existing = list(r.get("qa_findings") or [])
                r["qa_findings"] = dedupe_qa_findings(existing + row_gate_findings)[:10]
                r["qa_summary"] = summarize_qa_findings(row_gate_findings)
                if r["status"] in {"MT", "Existing", "Approved"}:
                    r["status"] = "Needs Review"

        missing_rate = missing / max(len(review_rows), 1)
        # IMPORTANT: seed the dedicated Human Review workspace BEFORE any button click.
        # Without this, the Human Review page can open with no rows and look blank.
        prepare_human_review_session(
            review_rows,
            source="ErrorSweep Pro",
            target_language=target_language,
            file_name=getattr(uploaded, "name", "uploaded_file"),
            rules=client_rules,
        )
        status = "Completed" if missing == 0 else ("Needs Human Review" if missing_rate <= threshold / 100 else "Blocked")
        pro_job = persist_saas_record("jobs", {
            "created": now_stamp(),
            "workspace": (current_user() or {}).get("workspace", "Demo Workspace"),
            "type": "Pro Translation",
            "language": target_language,
            "status": status,
            "segments": len(review_rows),
            "missing": missing,
        })
        st.session_state.jobs.insert(0, pro_job)
        trim_session_list("jobs")
        add_audit("Pro translation run", f"{target_language}: {status}")
        update_task_record(
            task["id"],
            status="completed" if status != "Blocked" else "needs_review",
            progress=100,
            processed_units=len(review_rows),
            total_units=len(review_rows),
            result_ref=safe_text(pro_job.get("id", "")),
            metadata_json={
                "target_language": target_language,
                "status": status,
                "missing": missing,
                "review_job_id": st.session_state.get("active_review_session_id", ""),
            },
        )
        record_billable_workflow_usage("pro_translation", review_rows, provider="errorsweep_pro", model="translate_qa_review")
        queue_email_notification(
            (current_user() or {}).get("email", ""),
            "ErrorSweep Pro translation completed",
            f"Pro translation for {target_language} finished with status '{status}'. Segments: {len(review_rows)}. Missing/review rows: {missing}.",
            "pro.completed",
            metadata={
                "target_language": target_language,
                "status": status,
                "segments": len(review_rows),
                "missing": missing,
                "review_job_id": st.session_state.get("active_review_session_id", ""),
            },
        )

        if status == "Blocked":
            st.error(f"Translation incomplete: {missing}/{len(review_rows)} rows need review. Output blocked, but segments are available in Human Review.")
        elif status == "Needs Human Review":
            st.warning(f"Translation mostly completed. {missing}/{len(review_rows)} rows need Human Review.")
        else:
            st.success("Translation completed. Review is ready for approval.")

        render_delivery_gate(review_rows, gate_findings, "Draft delivery gate")
        st.dataframe(pd.DataFrame(review_rows), use_container_width=True, hide_index=True)

        # v41: External editor launcher. The CAT editor opens in a new browser tab
        # by job_id, so it feels like a professional editor window instead of a
        # normal Streamlit dashboard page.
        st.markdown("### Next step")
        cta1, cta2 = st.columns([1, 1])
        with cta1:
            review_job_id = st.session_state.get("active_review_session_id") or query_get("review_id")
            if review_job_id:
                render_external_editor_link("Open Human Review Editor", "cat", str(review_job_id))
            else:
                st.error("Review job was not created. Please rerun Pro translation.")
        with cta2:
            st.download_button("Download draft CSV", rows_to_csv(review_rows), "errorsweep_pro_draft_review_rows.csv", "text/csv", use_container_width=True)
        st.info("Human Review now opens in a separate full-window CAT editor. Target editing happens directly in the main grid; the right panel is only for TM, glossary, DNT, QA, issues, and history.")


# Pro post-editing and Subtitle/Transcription editors
# ==========================================================

def render_assist_panel(source: str) -> None:
    matches = compute_matches(source)
    st.markdown("#### Assist panel")
    st.markdown("##### TM matches")
    if matches["tm"]:
        for m in matches["tm"]:
            st.markdown(f'<div class="es-row-card"><span class="es-chip green">{escape(m["type"])}</span><br><b>{escape(m["source"])}</b><br><span class="es-small">{escape(m["target"])}</span></div>', unsafe_allow_html=True)
    else:
        st.caption("No TM match.")

    st.markdown("##### Glossary")
    if matches["glossary"]:
        for g in matches["glossary"]:
            st.markdown(f'<div class="es-row-card"><b>{escape(g.get("source",""))}</b> → {escape(g.get("target",""))}<br><span class="es-small">{escape(g.get("notes",""))}</span></div>', unsafe_allow_html=True)
    else:
        st.caption("No glossary hits.")

    st.markdown("##### DNT")
    if matches["dnt"]:
        for d in matches["dnt"]:
            st.markdown(f'<span class="es-chip amber">{escape(d["term"])}</span> ', unsafe_allow_html=True)
    else:
        st.caption("No DNT hits.")


def render_text_review_editor() -> None:
    """Full-width CAT-style Pro post-editing workspace.

    This version is intentionally closer to CAT tools such as Phrase/Memsource:
    a compact job bar, filter row, spreadsheet-like source/target grid, match
    score/status columns, and a right CAT panel. It opens as a dedicated page
    without the normal platform navigation so the grid can use the full screen.
    """
    st.markdown(
        """
        <style>
        .block-container {
            max-width: 100vw !important;
            padding: .45rem .75rem .65rem .75rem !important;
        }
        .es-cat-app-shell {
            min-height: calc(100vh - 16px);
            background: rgba(8, 10, 19, .98);
            border: 1px solid rgba(84,105,180,.20);
            border-radius: 16px;
            overflow: hidden;
        }
        .es-cat-topbar {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 14px;
            padding: 10px 14px;
            background: #24252b;
            border-bottom: 1px solid rgba(255,255,255,.09);
        }
        .es-cat-brandline {
            display: flex;
            align-items: center;
            gap: 10px;
            min-width: 0;
            color: #fff;
            font-weight: 800;
            font-size: 15px;
        }
        .es-cat-pill {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            border-radius: 999px;
            padding: 4px 10px;
            font-size: 12px;
            font-weight: 900;
            border: 1px solid rgba(255,255,255,.15);
            background: rgba(255,255,255,.06);
            color: #eaf2ff;
            white-space: nowrap;
        }
        .es-cat-pill.green { background: rgba(0,217,133,.13); border-color: rgba(0,217,133,.32); color: #66ffc4; }
        .es-cat-pill.amber { background: rgba(245,158,11,.12); border-color: rgba(245,158,11,.35); color: #ffd38a; }
        .es-cat-toolbar2 {
            display: flex;
            gap: 10px;
            align-items: center;
            padding: 7px 14px;
            background: #1f2026;
            border-bottom: 1px solid rgba(255,255,255,.08);
            color: rgba(255,255,255,.72);
            font-size: 14px;
        }
        .es-cat-tool-icon {
            width: 28px; height: 28px; border-radius: 7px;
            display: inline-flex; align-items: center; justify-content: center;
            background: rgba(255,255,255,.045); border: 1px solid rgba(255,255,255,.06);
        }
        .es-cat-metrics {
            display: grid;
            grid-template-columns: repeat(4, minmax(0, 1fr));
            gap: 8px;
            padding: 10px 14px 6px 14px;
            background: #0b0d17;
        }
        .es-cat-metric {
            background: rgba(18,21,38,.76);
            border: 1px solid rgba(84,105,180,.20);
            border-radius: 10px;
            padding: 8px 10px;
        }
        .es-cat-metric-label { font-family: Space Mono, monospace; color: #9aa7da; font-size: 10px; text-transform: uppercase; }
        .es-cat-metric-value { color:#fff; font-size:20px; font-weight:900; line-height: 1.1; }
        .es-cat-filterbar {
            display: grid;
            grid-template-columns: minmax(220px, 1fr) minmax(220px, 1fr) 180px 120px;
            gap: 10px;
            padding: 8px 14px 10px 14px;
            background: #0b0d17;
            border-bottom: 1px solid rgba(255,255,255,.08);
        }
        .es-cat-grid-title {
            display: grid;
            grid-template-columns: 56px minmax(270px, 1fr) minmax(270px, 1fr) 74px 90px 90px;
            gap: 0;
            padding: 7px 12px;
            background: #161820;
            border: 1px solid rgba(255,255,255,.08);
            border-bottom: none;
            border-radius: 12px 12px 0 0;
            font-family: Space Mono, monospace;
            color: #aeb8dc;
            text-transform: uppercase;
            font-size: 10px;
            font-weight: 700;
        }
        .es-cat-grid-card {
            border: 1px solid rgba(255,255,255,.08);
            border-radius: 0 0 12px 12px;
            overflow: hidden;
            background: #10121b;
        }
        .es-cat-side-card {
            border: 1px solid rgba(255,255,255,.10);
            background: #151721;
            border-radius: 12px;
            padding: 10px;
            height: 735px;
            overflow-y: auto;
        }
        .es-cat-seg-preview {
            background: #0e1018;
            border: 1px solid rgba(255,255,255,.08);
            border-radius: 10px;
            padding: 10px;
            margin-bottom: 10px;
        }
        .es-cat-assist-title {
            font-size: 13px;
            font-weight: 900;
            color: #fff;
            margin: 12px 0 5px 0;
        }
        .es-cat-assist-empty { color:#8d95bb; font-size:12px; }
        .es-cat-mini-row {
            border-bottom: 1px solid rgba(255,255,255,.06);
            padding: 7px 0;
            color: #dbe6ff;
            font-size: 12px;
        }
        div[data-testid="stDataFrame"], div[data-testid="stDataEditor"] {
            border-radius: 0 !important;
            border: none !important;
            contain: paint !important;
            isolation: isolate !important;
            overflow: hidden !important;
        }
        div[data-testid="stDataEditor"] [role="grid"] {
            font-size: 13px !important;
        }
        div[data-testid="stDataEditor"] textarea:not(:focus),
        div[data-testid="stDataEditor"] input:not(:focus),
        div[data-testid="stDataFrame"] textarea:not(:focus),
        div[data-testid="stDataFrame"] input:not(:focus) {
            opacity: 0 !important;
            pointer-events: none !important;
        }
        div[data-testid="stDataEditor"] textarea,
        div[data-testid="stDataEditor"] input {
            font-size: 13px !important;
        }
        @media (max-width: 1100px) {
            .es-cat-filterbar { grid-template-columns: 1fr; }
            .es-cat-metrics { grid-template-columns: repeat(2, minmax(0, 1fr)); }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    rows = st.session_state.get("review_segments", []) or []
    if not rows:
        st.info("No Pro translated rows are loaded. Run ErrorSweep Pro first, then click Open Human Review workspace.")
        return

    completion = compute_review_completion(rows)
    file_name = safe_text(st.session_state.get("review_workspace_file_name", "Current file")) or "Current file"
    language = safe_text(st.session_state.get("review_workspace_language", "Target")) or "Target"
    title = safe_text(st.session_state.get("review_workspace_title", "Pro Human Review")) or "Pro Human Review"

    # Top job bar like a CAT tool.
    st.markdown(
        f"""
        <div class="es-cat-app-shell">
          <div class="es-cat-topbar">
            <div class="es-cat-brandline">
              <span style="font-size:18px;">▣</span>
              <span>ErrorSweep CAT</span>
              <span style="color:#9aa0b9; font-weight:600;">/ {escape(title)} / {escape(file_name)} / {escape(language)}</span>
            </div>
            <div style="display:flex; gap:8px; align-items:center;">
              <span class="es-cat-pill green">Accepted</span>
              <span class="es-cat-pill">TM</span>
              <span class="es-cat-pill">TB</span>
              <span class="es-cat-pill amber">MT</span>
            </div>
          </div>
          <div class="es-cat-toolbar2">
            <span class="es-cat-tool-icon">B</span><span class="es-cat-tool-icon"><i>I</i></span><span class="es-cat-tool-icon">U</span>
            <span class="es-cat-tool-icon">⌘</span><span class="es-cat-tool-icon">✓</span><span class="es-cat-tool-icon">↶</span><span class="es-cat-tool-icon">↷</span>
            <span style="margin-left:auto; color:#8ea1dc; font-size:12px;">Post-editing workspace · source left · target right</span>
          </div>
          <div class="es-cat-metrics">
            <div class="es-cat-metric"><div class="es-cat-metric-label">Confirmed</div><div class="es-cat-metric-value">{completion['approved']}</div></div>
            <div class="es-cat-metric"><div class="es-cat-metric-label">Segments</div><div class="es-cat-metric-value">{completion['total']}</div></div>
            <div class="es-cat-metric"><div class="es-cat-metric-label">Translated</div><div class="es-cat-metric-value">{completion['translated']}</div></div>
            <div class="es-cat-metric"><div class="es-cat-metric-label">Needs Review</div><div class="es-cat-metric-value">{completion['needs_review']}</div></div>
          </div>
        """,
        unsafe_allow_html=True,
    )

    # Filter controls need Streamlit widgets, so they sit visually inside the shell.
    with st.container():
        f1, f2, f3, f4 = st.columns([1.25, 1.25, .7, .55])
        with f1:
            source_filter = st.text_input("Filter source", value="", placeholder="Filter source (en)", key="cat_v40_source_filter", label_visibility="collapsed")
        with f2:
            target_filter = st.text_input("Filter target", value="", placeholder="Filter target", key="cat_v40_target_filter", label_visibility="collapsed")
        with f3:
            status_options = ["All"] + sorted({safe_text(r.get("status", "Untranslated")) or "Untranslated" for r in rows})
            status_filter = st.selectbox("Status", status_options, key="cat_v40_status_filter", label_visibility="collapsed")
        with f4:
            pending_only = st.checkbox("Pending", value=False, key="cat_v40_pending_only")

    filtered_indexes: List[int] = []
    for i, r in enumerate(rows):
        src = safe_text(r.get("source", ""))
        tgt = safe_text(r.get("target", ""))
        status = safe_text(r.get("status", "Untranslated")) or "Untranslated"
        if source_filter and source_filter.lower() not in src.lower():
            continue
        if target_filter and target_filter.lower() not in tgt.lower():
            continue
        if status_filter != "All" and status != status_filter:
            continue
        if pending_only and status in {"Approved", "101%", "100%"}:
            continue
        filtered_indexes.append(i)

    if not filtered_indexes:
        st.warning("No segments match the current filters.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    current_idx = int(st.session_state.get("selected_review_index", filtered_indexes[0]) or filtered_indexes[0])
    if current_idx not in filtered_indexes:
        current_idx = filtered_indexes[0]
        st.session_state.selected_review_index = current_idx

    # Build the editable grid. Source and target appear side-by-side like Excel/CAT.
    grid_rows = []
    for i in filtered_indexes:
        r = rows[i]
        status = safe_text(r.get("status", "MT" if safe_text(r.get("target", "")) else "Needs Review"))
        match = safe_text(r.get("match", "MT" if safe_text(r.get("target", "")) else "Untranslated"))
        # Score column emulates Phrase/Memsource match percentage badges.
        if match in {"100%", "101%"}:
            score = match
        elif "Fuzzy" in match:
            score = match.replace("Fuzzy", "").strip() or "85%"
        elif status in {"Approved", "100%", "101%"}:
            score = "100"
        elif match == "MT":
            score = "MT"
        elif match == "Untranslated":
            score = "-"
        else:
            score = match or "MT"
        sensitive_kinds = detect_sensitive_text(f"{r.get('source', '')} {r.get('target', '')}")
        source_text = safe_text(r.get("source", ""))
        grid_rows.append({
            "No": i + 1,
            "Source": source_text,
            "Target": repair_localization_translation(source_text, r.get("target", "")),
            "Score": score,
            "Status": status,
            "QA": "✓" if status in {"Approved", "100%", "101%"} else "◯",
            "Notes": safe_text(r.get("notes", "")),
            "Location": safe_text(r.get("location", f"Segment {i+1}")),
        })
        if sensitive_kinds:
            grid_rows[-1]["QA"] = "PII"
        elif r.get("qa_findings"):
            grid_rows[-1]["QA"] = "QA"
        else:
            grid_rows[-1]["QA"] = "OK" if status in {"Approved", "100%", "101%"} else "Open"

    main_col, side_col = st.columns([4.25, 1.15], gap="small")
    with main_col:
        st.markdown(
            '<div class="es-cat-grid-title"><div>No</div><div>Source</div><div>Target</div><div>Score</div><div>Status</div><div>QA</div></div>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="es-cat-grid-card">', unsafe_allow_html=True)
        edited_df = st.data_editor(
            pd.DataFrame(grid_rows),
            use_container_width=True,
            hide_index=True,
            height=700,
            num_rows="fixed",
            column_order=["No", "Source", "Target", "Score", "Status", "QA", "Notes", "Location"],
            disabled=["No", "Source", "Score", "QA", "Location"],
            column_config={
                "No": st.column_config.NumberColumn("", width="small"),
                "Source": st.column_config.TextColumn("Source", width="large", help="Read-only source segment"),
                "Target": st.column_config.TextColumn("Target", width="large", help="Editable reviewed translation"),
                "Score": st.column_config.TextColumn("", width="small"),
                "Status": st.column_config.SelectboxColumn(
                    "Status",
                    width="medium",
                    options=["MT", "Fuzzy 75%", "Fuzzy 85%", "100%", "101%", "Needs Review", "Approved", "Rejected", "Needs Rework", "Untranslated"],
                ),
                "QA": st.column_config.TextColumn("", width="small"),
                "Notes": st.column_config.TextColumn("Notes", width="medium"),
                "Location": st.column_config.TextColumn("Location", width="medium"),
            },
            key="cat_v40_excel_grid",
        )
        st.markdown('</div>', unsafe_allow_html=True)

        b1, b2, b3, b4, b5 = st.columns([1, 1, 1, 1, 1])
        if b1.button("Save edits", type="primary", use_container_width=True):
            for _, erow in edited_df.iterrows():
                idx = int(erow["No"]) - 1
                if 0 <= idx < len(rows):
                    rows[idx]["target"] = safe_text(erow.get("Target", ""))
                    rows[idx]["status"] = safe_text(erow.get("Status", "")) or "Needs Review"
                    rows[idx]["notes"] = safe_text(erow.get("Notes", ""))
            st.session_state.review_segments = rows
            st.session_state.last_pro_review_segments = rows
            st.session_state.latest_human_review_segments = rows
            st.toast("Saved grid edits.")
            st.rerun()

        if b2.button("Approve visible", use_container_width=True):
            for _, erow in edited_df.iterrows():
                idx = int(erow["No"]) - 1
                if 0 <= idx < len(rows) and safe_text(erow.get("Target", "")).strip():
                    rows[idx]["target"] = safe_text(erow.get("Target", ""))
                    rows[idx]["status"] = "Approved"
                    rows[idx]["notes"] = safe_text(erow.get("Notes", ""))
            st.session_state.review_segments = rows
            st.toast("Visible translated rows approved.")
            st.rerun()

        if b3.button("Next pending", use_container_width=True):
            next_idx = None
            for i, r in enumerate(rows):
                if safe_text(r.get("status", "")) not in {"Approved", "101%", "100%"} or not safe_text(r.get("target", "")).strip():
                    next_idx = i
                    break
            if next_idx is not None:
                st.session_state.selected_review_index = next_idx
                st.rerun()
            else:
                st.success("All rows look complete.")

        if b4.button("Save to TM", use_container_width=True):
            saved = 0
            for r in rows:
                src = safe_text(r.get("source", ""))
                tgt = safe_text(r.get("target", ""))
                if src and tgt and safe_text(r.get("status", "")) in {"Approved", "100%", "101%"}:
                    st.session_state.tm.append({
                        "source": src,
                        "target": tgt,
                        "language": safe_text(st.session_state.get("review_workspace_language", "")),
                        "created": now_stamp(),
                        "approved_by": (current_user() or {}).get("email", ""),
                    })
                    trim_session_list("tm")
                    saved += 1
            st.success(f"Saved {saved} approved segment(s) to TM.")

        if b5.button("Back to Pro", use_container_width=True):
            open_page("ErrorSweep Pro")

        reviewed_base_name = safe_text(st.session_state.get("review_workspace_file_name", "human_review_output")) or "human_review_output"
        reviewed_base_name = re.sub(r"\.[^.]+$", "", reviewed_base_name)
        dl1, dl2, dl3 = st.columns(3)
        dl1.download_button(
            "Download reviewed Excel",
            build_reviewed_translation_workbook(rows),
            file_name=f"{reviewed_base_name}_reviewed_translation.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
        dl2.download_button(
            "Download reviewed CSV",
            rows_to_csv(rows),
            file_name=f"{reviewed_base_name}_reviewed_translation.csv",
            mime="text/csv",
            use_container_width=True,
        )
        dl3.download_button(
            "Download target text",
            build_reviewed_plain_text(rows),
            file_name=f"{reviewed_base_name}_target_text.txt",
            mime="text/plain",
            use_container_width=True,
        )

    with side_col:
        st.markdown('<div class="es-cat-side-card">', unsafe_allow_html=True)
        select_labels = [f"{i+1} · {safe_text(rows[i].get('source',''))[:42] or safe_text(rows[i].get('target',''))[:42]}" for i in filtered_indexes]
        selected_label = st.selectbox(
            "CAT",
            select_labels,
            index=filtered_indexes.index(current_idx) if current_idx in filtered_indexes else 0,
            key="cat_v40_focus_select",
            label_visibility="collapsed",
        )
        selected_idx = filtered_indexes[select_labels.index(selected_label)]
        st.session_state.selected_review_index = selected_idx
        focused = rows[selected_idx]
        focused_sensitive = detect_sensitive_text(f"{focused.get('source', '')} {focused.get('target', '')}")
        status = safe_text(focused.get("status", "Needs Review"))
        match = safe_text(focused.get("match", "MT"))
        chip_class = "green" if status in {"Approved", "100%", "101%"} else "amber" if "Review" in status or status in {"MT", "Untranslated", "Needs Rework"} else "red"
        st.markdown(
            f"""
            <div class="es-cat-seg-preview">
              <div style="display:flex; justify-content:space-between; gap:8px; align-items:center;">
                <div class="es-small">Segment {selected_idx + 1} / {len(rows)}</div>
                <div><span class="es-cat-pill {chip_class}">{escape(status)}</span></div>
              </div>
              <div style="margin-top:8px;"><span class="es-cat-pill">{escape(match)}</span></div>
              <div class="es-small" style="margin-top:8px;">{escape(safe_text(focused.get('location','')))}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown('<div class="es-cat-assist-title">Source</div>', unsafe_allow_html=True)
        if focused_sensitive:
            st.warning(f"Sensitive data indicator: {', '.join(focused_sensitive)}. Keep this row on approved local routes unless client consent exists.")
        st.markdown(highlight_localization_text(focused.get("source", "")), unsafe_allow_html=True)
        st.text_area("Source preview", value=safe_text(focused.get("source", "")), height=105, disabled=True, label_visibility="collapsed", key=f"v40_src_{selected_idx}")
        st.markdown('<div class="es-cat-assist-title">Target</div>', unsafe_allow_html=True)
        st.markdown(highlight_localization_text(focused.get("target", "")), unsafe_allow_html=True)
        focused_target = st.text_area("Target preview", value=safe_text(focused.get("target", "")), height=135, label_visibility="collapsed", key=f"v40_tgt_{selected_idx}")
        if safe_text(focused.get("target", "")) and focused_target != safe_text(focused.get("target", "")):
            st.markdown('<div class="es-cat-assist-title">Inline diff</div>', unsafe_allow_html=True)
            st.markdown(inline_diff_html(focused.get("target", ""), focused_target), unsafe_allow_html=True)
        focused_status = st.selectbox(
            "Status",
            ["MT", "Fuzzy 75%", "Fuzzy 85%", "100%", "101%", "Needs Review", "Approved", "Rejected", "Needs Rework", "Untranslated"],
            index=["MT", "Fuzzy 75%", "Fuzzy 85%", "100%", "101%", "Needs Review", "Approved", "Rejected", "Needs Rework", "Untranslated"].index(status) if status in ["MT", "Fuzzy 75%", "Fuzzy 85%", "100%", "101%", "Needs Review", "Approved", "Rejected", "Needs Rework", "Untranslated"] else 5,
            key=f"v40_status_{selected_idx}",
        )
        c1, c2 = st.columns(2)
        if c1.button("Save", key=f"v40_save_{selected_idx}", use_container_width=True):
            rows[selected_idx]["target"] = focused_target
            rows[selected_idx]["status"] = focused_status
            st.session_state.review_segments = rows
            st.toast("Saved selected segment.")
            st.rerun()
        if c2.button("Approve", key=f"v40_approve_{selected_idx}", use_container_width=True):
            rows[selected_idx]["target"] = focused_target
            rows[selected_idx]["status"] = "Approved"
            st.session_state.review_segments = rows
            st.toast("Approved selected segment.")
            st.rerun()

        st.markdown('<div class="es-cat-assist-title">TM matches</div>', unsafe_allow_html=True)
        matches = compute_matches(safe_text(focused.get("source", "")))
        if matches["tm"]:
            for m in matches["tm"][:7]:
                st.markdown(f'<div class="es-cat-mini-row"><b>{escape(m.get("type","TM"))}</b> · {escape(m.get("source",""))}<br><span class="es-small">{escape(m.get("target",""))}</span></div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="es-cat-assist-empty">No TM match.</div>', unsafe_allow_html=True)

        st.markdown('<div class="es-cat-assist-title">Glossary</div>', unsafe_allow_html=True)
        if matches["glossary"]:
            for g in matches["glossary"][:8]:
                st.markdown(f'<div class="es-cat-mini-row"><b>{escape(g.get("source",""))}</b> → {escape(g.get("target",""))}<br><span class="es-small">{escape(g.get("notes",""))}</span></div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="es-cat-assist-empty">No glossary hits.</div>', unsafe_allow_html=True)

        st.markdown('<div class="es-cat-assist-title">DNT</div>', unsafe_allow_html=True)
        if matches["dnt"]:
            for d in matches["dnt"][:12]:
                st.markdown(f'<span class="es-cat-pill amber">{escape(d.get("term",""))}</span> ', unsafe_allow_html=True)
        else:
            st.markdown('<div class="es-cat-assist-empty">No DNT hits.</div>', unsafe_allow_html=True)

        st.markdown('<div class="es-cat-assist-title">QA Details</div>', unsafe_allow_html=True)
        qa_findings = focused.get("qa_findings") or []
        if qa_findings:
            for finding in qa_findings[:6]:
                st.markdown(
                    f'<div class="es-cat-mini-row"><b>{escape(safe_text(finding.get("Severity", "Review")))}</b> '
                    f'{escape(safe_text(finding.get("Error Type", "QA")))}<br>'
                    f'<span class="es-small">{escape(safe_text(finding.get("Explanation", "")))}</span></div>',
                    unsafe_allow_html=True,
                )
        else:
            st.markdown('<div class="es-cat-assist-empty">No rule-engine findings for this row.</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


def default_subtitle_segments(count: int = 8, transcription: bool = False) -> List[Dict[str, Any]]:
    rows = []
    for i in range(count):
        rows.append({
            "id": i + 1,
            "start": round(i * 4.0, 3),
            "end": round(i * 4.0 + 3.5, 3),
            "source": "" if transcription else f"Source segment {i+1}",
            "target": "",
            "status": "Draft" if transcription else "Untranslated",
            "match": "",
        })
    return rows


def enter_subtitle_workspace(workflow: str, rows: List[Dict[str, Any]], video_file=None) -> None:
    """Open the dedicated subtitle/transcription workspace."""
    st.session_state.subtitle_workflow = workflow
    st.session_state.subtitle_segments = rows
    st.session_state.selected_subtitle_index = 0
    st.session_state.subtitle_editor_active = True
    if video_file is not None:
        preview = save_media_preview_file(f"session_{uuid.uuid4().hex}", video_file)
        st.session_state.subtitle_video_metadata = preview
        st.session_state.subtitle_video_bytes = None
        st.session_state.subtitle_video_name = preview.get("media_preview_name", getattr(video_file, "name", "uploaded_video"))
        st.session_state.subtitle_video_type = preview.get("media_preview_type", getattr(video_file, "type", "video/mp4") or "video/mp4")
    add_audit(f"{workflow} workspace opened", f"{len(rows)} rows")


def render_subtitle_transcription_setup() -> None:
    st.markdown("### Subtitle / Transcription Editor")
    st.caption("Create a dedicated editor workspace. Subtitling can use a source script. Transcription auto-generation is available only when the user provides an API key.")

    workflow = st.radio("Editor workflow", ["Subtitling", "Transcription"], horizontal=True, key="subtitle_workflow_picker")
    video = st.file_uploader("Upload video/audio", type=["mp4", "mov", "m4v", "webm", "mp3", "wav", "m4a"], key="subtitle_video_setup")
    user_key_available = bool(str(st.session_state.get("byo_openai_api_key", "") or "").strip())
    media_compliance_ack = st.checkbox(
        "I confirm I have rights or client authorization to process this media, including any confidential audio or copyrighted material.",
        key="media_compliance_ack",
    )

    if video:
        preview_col, info_col = st.columns([0.45, 0.55], gap="large")
        with preview_col:
            preview_key = hashlib.sha256(
                f"{getattr(video, 'name', 'media')}:{getattr(video, 'size', '')}".encode("utf-8")
            ).hexdigest()[:18]
            preview_record = save_media_preview_file(f"setup_{preview_key}", video)
            preview_source, preview_type, preview_name = read_media_preview_bytes(preview_record)
            render_media_preview(preview_source, preview_type or getattr(video, "type", "video/mp4"), preview_name or getattr(video, "name", "media"))
        with info_col:
            st.success("Video/audio loaded.")
            st.caption("The editor will open as a dedicated workspace page, separate from this setup screen.")
            if user_key_available:
                st.caption("Transcription route: user API key available.")
            else:
                st.caption("Transcription route: manual editing. No API key is available for speech-to-text.")
    else:
        st.info("Upload a video/audio file to begin.")

    if workflow == "Subtitling":
        source_file = st.file_uploader(
            "Upload English source subtitle/script (optional for subtitling)",
            type=["srt", "vtt", "txt", "csv", "xlsx", "docx"],
            key="subtitle_source_setup",
        )
        target_file = st.file_uploader(
            "Upload existing target subtitle file (optional)",
            type=["srt", "vtt", "txt", "csv", "xlsx", "docx"],
            key="subtitle_target_setup",
        )
        c1, c2, c3 = st.columns([1, 1, 1])
        subtitle_target_language = c1.text_input("Target subtitle language", value=st.session_state.get("subtitle_target_language", "French"), key="subtitle_target_lang_setup")
        speech_locale = c2.text_input("Source speech locale", value="en-US", help="Used only if no source file is uploaded and a user API key is available for transcription.")
        starter_rows = c3.number_input("Starter rows", min_value=1, max_value=200, value=10, key="subtitle_starter_rows")
        auto_generate = st.checkbox(
            "Generate draft target subtitles",
            value=True,
            help="If source rows exist, target subtitles use BYO API key or built-in self-hosted MT translation. If no source file is uploaded, speech-to-text requires a user API key; otherwise blank rows are created for manual editing.",
        )

        if st.button("Create subtitling workspace", use_container_width=True, disabled=video is None or not media_compliance_ack):
            st.session_state.subtitle_target_language = subtitle_target_language
            if source_file:
                rows = extract_rows_from_upload(source_file)
                for i, r in enumerate(rows):
                    r.setdefault("start", i * 4.0)
                    r.setdefault("end", i * 4.0 + 3.5)
                    r.setdefault("target", "")
                    r.setdefault("status", "Untranslated")
                    r.setdefault("match", "")
            else:
                if user_key_available:
                    with st.spinner("No source file uploaded. Transcribing source from video/audio using user API key..."):
                        transcript_rows, usage = generate_transcription_rows_from_video(video, locale=speech_locale)
                    rows = []
                    for i, tr in enumerate(transcript_rows):
                        rows.append({
                            "id": i + 1,
                            "start": tr.get("start", i * 3.5),
                            "end": tr.get("end", i * 3.5 + 3.0),
                            "source": tr.get("target", ""),
                            "target": "",
                            "status": "Transcribed Source" if tr.get("target") else "Untranslated",
                            "match": tr.get("match", "STT"),
                        })
                    if usage.get("error") and not usage.get("success"):
                        st.warning(f"Speech transcription was not available: {usage.get('error')}. Blank rows were created for manual subtitling.")
                else:
                    rows = default_subtitle_segments(int(starter_rows), transcription=False)
                    for r in rows:
                        r["source"] = ""
                        r["target"] = ""
                        r["status"] = "Draft"
                        r["match"] = "Manual"
                    st.info("No source file and no user API key were provided. Blank subtitle rows were created for manual source/target editing.")

            if target_file:
                target_rows = extract_rows_from_upload(target_file)
                for i, tr in enumerate(target_rows):
                    if i < len(rows):
                        rows[i]["target"] = tr.get("target") or tr.get("source") or ""
                        rows[i]["status"] = "Existing" if rows[i]["target"] else rows[i].get("status", "Untranslated")

            has_source_text = any(safe_text(r.get("source", "")) for r in rows)
            if auto_generate and rows and has_source_text:
                with st.spinner("Generating target subtitle draft..."):
                    rows, missing = translate_subtitle_sources(rows, subtitle_target_language, domain="Subtitling", rules=workspace_rules())
                if missing:
                    st.warning(f"Draft subtitles generated with {missing} untranslated row(s).")
                else:
                    st.success("Draft subtitles generated.")
            elif auto_generate and rows and not has_source_text:
                st.info("Draft target subtitles were skipped because there is no source text yet. Fill source rows manually, then translate/review later.")

            enter_subtitle_workspace("Subtitling", rows, video)
            job_id = save_media_session_to_store("Subtitling", rows, video_file=video, target_language=subtitle_target_language)
            st.success("Subtitling editor job created. Open it in the separate editor window below.")
            render_external_editor_link("Open Subtitle Editor", "media", job_id)
    else:
        st.caption("Transcription mode does not need a source file. Auto-transcription requires a user API key. Without a user key, blank transcript rows are created for human editing.")
        c1, c2 = st.columns([1, 1])
        speech_locale = c1.text_input("Speech locale", value="en-US", key="transcription_locale")
        starter_count = c2.number_input("Starter rows", min_value=1, max_value=200, value=10, key="transcription_starter_count")
        auto_transcribe = st.checkbox(
            "Auto-generate transcript using user API key",
            value=user_key_available,
            disabled=not user_key_available,
            help="Speech-to-text is available only when the user has added an API key in Account. Without a key, the editor opens with blank rows for manual transcription.",
        )
        if not user_key_available:
            st.info("No user API key found. The transcription workspace will open with blank rows for manual editing.")

        if st.button("Create transcription workspace", use_container_width=True, disabled=video is None or not media_compliance_ack):
            if auto_transcribe and user_key_available:
                with st.spinner("Creating transcript from video/audio using user API key..."):
                    rows, usage = generate_transcription_rows_from_video(video, locale=speech_locale)
                if usage.get("error") and not usage.get("success"):
                    st.warning(f"Auto-transcription was not available: {usage.get('error')}. Blank rows were created for manual transcription.")
                    rows = default_subtitle_segments(int(starter_count), transcription=True)
            else:
                rows = default_subtitle_segments(int(starter_count), transcription=True)
            if not rows:
                rows = default_subtitle_segments(int(starter_count), transcription=True)
            enter_subtitle_workspace("Transcription", rows, video)
            job_id = save_media_session_to_store("Transcription", rows, video_file=video, target_language="")
            st.success("Transcription editor job created. Open it in the separate editor window below.")
            render_external_editor_link("Open Transcription Editor", "media", job_id)

    if st.session_state.subtitle_segments:
        if st.button("Open existing subtitle/transcription workspace", use_container_width=True):
            st.session_state.subtitle_editor_active = True
            open_page("Subtitle Workspace" if st.session_state.get("subtitle_workflow") == "Subtitling" else "Transcription Workspace")


def render_focused_subtitle_workspace() -> None:
    workflow = st.session_state.get("subtitle_workflow", "Transcription")
    rows = st.session_state.subtitle_segments
    if not rows:
        st.session_state.subtitle_editor_active = False
        st.info("No editor rows available. Create a subtitle or transcription workspace first.")
        return

    top1, top2 = st.columns([0.78, 0.22])
    with top1:
        st.markdown(f"### {workflow} workspace")
        st.caption("Compact video on top, script writing in the middle, timing/text grid collapsed at the bottom.")
    with top2:
        if st.button("Back to setup", use_container_width=True):
            st.session_state.subtitle_editor_active = False
            open_page("Subtitle / Transcription Editor")

    video_source, video_type, video_name = read_media_preview_bytes(st.session_state.get("subtitle_video_metadata", {}))
    video_type = safe_text(video_type or st.session_state.get("subtitle_video_type", "video/mp4")) or "video/mp4"
    video_name = safe_text(video_name or st.session_state.get("subtitle_video_name", "media"))
    video_col, meta_col = st.columns([0.50, 0.50], gap="large")
    with video_col:
        render_media_preview(video_source, video_type, video_name)
        render_waveform_preview(rows, st.session_state.get("selected_subtitle_index", 0))
    with meta_col:
        st.markdown("#### Job notes")
        st.caption("Use the selected segment below to write transcript/subtitle text. Use the collapsed grid for detailed timing edits.")
        st.metric("Rows", len(rows))
        st.metric("Approved", sum(1 for r in rows if r.get("status") == "Approved"))
        timing_issues = validate_timing_rows(rows)
        if timing_issues:
            st.warning(f"{len(timing_issues)} timing issue(s) need review.")

    idx = min(st.session_state.selected_subtitle_index, len(rows) - 1)
    row = rows[idx]
    render_segment_timeline(rows, idx)

    list_col, editor_col, assist_col = st.columns([1.25, 2.25, 1.15], gap="large")

    with list_col:
        st.markdown("#### Segments")
        for i, seg in enumerate(rows):
            time_label = f"{format_time(seg.get('start',0))} → {format_time(seg.get('end',0))}"
            preview = seg.get("source") if workflow == "Subtitling" else seg.get("target")
            if not preview:
                preview = "Empty transcript row" if workflow == "Transcription" else "Empty subtitle row"
            status = seg.get("status", "Draft")
            st.caption(f"{i+1}. {time_label} · {status}")
            if st.button(preview[:80], key=f"focused_pick_{i}", use_container_width=True):
                st.session_state.selected_subtitle_index = i
                st.rerun()

    with editor_col:
        st.markdown(f"#### {workflow} segment {idx+1} / {len(rows)}")
        time_a, time_b = st.columns(2)
        start_val = time_a.number_input("Start", min_value=0.0, value=float(row.get("start", 0.0)), step=0.1, key=f"focus_start_{idx}")
        end_val = time_b.number_input("End", min_value=0.0, value=float(row.get("end", max(start_val + 2.0, 2.0))), step=0.1, key=f"focus_end_{idx}")
        rows[idx]["start"] = float(start_val)
        rows[idx]["end"] = float(max(end_val, start_val + 0.1))

        if workflow == "Subtitling":
            source_text = st.text_area("English source", value=row.get("source", ""), height=90, key=f"focus_source_{idx}")
            rows[idx]["source"] = source_text
            target_label = "Target subtitle"
        else:
            target_label = "Transcript text"

        target_text = st.text_area(target_label, value=row.get("target", ""), height=170, key=f"focus_target_{idx}")
        status = st.selectbox(
            "Status",
            ["Draft", "MT", "Fuzzy 75%", "Fuzzy 85%", "100%", "101%", "Needs Review", "Approved", "Rejected", "Needs Rework", "Untranslated"],
            index=["Draft", "MT", "Fuzzy 75%", "Fuzzy 85%", "100%", "101%", "Needs Review", "Approved", "Rejected", "Needs Rework", "Untranslated"].index(row.get("status", "Draft")) if row.get("status", "Draft") in ["Draft", "MT", "Fuzzy 75%", "Fuzzy 85%", "100%", "101%", "Needs Review", "Approved", "Rejected", "Needs Rework", "Untranslated"] else 0,
            key=f"focus_status_{idx}",
        )

        b1, b2, b3, b4 = st.columns(4)
        if b1.button("Save", key=f"focus_save_{idx}", use_container_width=True):
            rows[idx]["target"] = target_text
            rows[idx]["status"] = status
            st.success("Saved.")
        if b2.button("Approve", key=f"focus_approve_{idx}", use_container_width=True):
            rows[idx]["target"] = target_text
            rows[idx]["status"] = "Approved"
            st.success("Approved.")
        if b3.button("Split", key=f"focus_split_{idx}", use_container_width=True):
            start = float(rows[idx]["start"])
            end = float(rows[idx]["end"])
            mid = round((start + end) / 2, 3)
            rows[idx]["end"] = mid
            rows.insert(idx + 1, {**rows[idx], "id": len(rows) + 1, "start": mid, "end": end, "source": "", "target": "", "status": "Draft"})
            st.rerun()
        if b4.button("Next", key=f"focus_next_{idx}", use_container_width=True):
            rows[idx]["target"] = target_text
            rows[idx]["status"] = status
            st.session_state.selected_subtitle_index = min(idx + 1, len(rows) - 1)
            st.rerun()

    with assist_col:
        render_assist_panel(row.get("source", "") or row.get("target", ""))

    # No duration scale: compact timing/text grid only, hidden by default.
    with st.expander("Timing and text grid", expanded=bool(st.session_state.get("show_timing_grid", False))):
        grid_cols = ["id", "start", "end", "source", "target", "status", "match"] if workflow == "Subtitling" else ["id", "start", "end", "target", "status"]
        edited = st.data_editor(
            pd.DataFrame(rows)[grid_cols],
            use_container_width=True,
            hide_index=True,
            num_rows="dynamic",
            height=230,
            key="focused_subtitle_grid",
        )
        c1, c2, c3 = st.columns(3)
        if c1.button("Save grid", use_container_width=True):
            updated_rows = edited.to_dict("records")
            timing_issues = validate_timing_rows(updated_rows)
            if timing_issues:
                st.error("; ".join(f"Row {item['row']}: {item['issue']}" for item in timing_issues[:5]))
            else:
                st.session_state.subtitle_segments = updated_rows
                st.success("Grid saved.")
        c2.download_button("Download CSV", rows_to_csv(st.session_state.subtitle_segments), "subtitle_transcription_editor.csv", "text/csv", use_container_width=True)
        c3.download_button("Download SRT", rows_to_srt(st.session_state.subtitle_segments, use_target=True), "subtitle_transcription_output.srt", "text/plain", use_container_width=True)


def render_subtitle_transcription_editor() -> None:
    if st.session_state.get("subtitle_editor_active"):
        render_focused_subtitle_workspace()
    else:
        render_subtitle_transcription_setup()


def page_subtitle_transcription_editor() -> None:
    # Public/manual editor page. This page is only for subtitle and transcription work.
    # Pro post-editing Human Review is opened only from ErrorSweep Pro via the hidden
    # Human Review Workspace route.
    if st.session_state.get("subtitle_editor_active"):
        render_focused_subtitle_workspace()
        return
    hero("Subtitle / Transcription Editor", "Dedicated media localization workspace", "Create subtitles or transcripts. Pro post-editing opens separately only after a Pro translation run.")
    render_subtitle_transcription_editor()

# Backward-compatible alias for old references.
def page_human_review() -> None:
    page_subtitle_transcription_editor()


def page_human_review_workspace() -> None:
    """Dedicated CAT-style post-editing route opened from ErrorSweep Pro outputs only."""
    restore_human_review_session_from_cache()
    # Recovery guard: if Streamlit opened this hidden route after a rerun, restore
    # the Pro rows from the persistent Pro result cache. This prevents blank pages.
    if not st.session_state.get("review_segments") and st.session_state.get("pro_post_edit_rows"):
        prepare_human_review_session(
            st.session_state.get("pro_post_edit_rows", []),
            source="ErrorSweep Pro",
            target_language=st.session_state.get("pro_post_edit_language", ""),
            file_name=st.session_state.get("pro_post_edit_file_name", ""),
        )

    if not st.session_state.get("review_segments"):
        st.markdown(
            """
            <style>.block-container{max-width:100vw!important;padding:1rem!important;}</style>
            """,
            unsafe_allow_html=True,
        )
        st.warning("No Pro post-editing rows are loaded yet. Run ErrorSweep Pro first, then click Open Human Review workspace.")
        if st.button("Go to ErrorSweep Pro", type="primary", use_container_width=True):
            open_page("ErrorSweep Pro")
        return

    render_text_review_editor()



def page_subtitle_workspace() -> None:
    st.session_state.subtitle_editor_active = True
    st.session_state.subtitle_workflow = "Subtitling"
    render_focused_subtitle_workspace()


def page_transcription_workspace() -> None:
    st.session_state.subtitle_editor_active = True
    st.session_state.subtitle_workflow = "Transcription"
    render_focused_subtitle_workspace()



# ==========================================================
# SCORECARD EXCEL OUTPUT
# ==========================================================

ERROR_CATEGORIES = ["Accuracy", "Readability", "Style and Tone", "Grammar", "Country Standards"]
ERROR_SEVERITIES = ["Minor", "Major", "Critical"]
SEVERITY_POINTS = {"Minor": 1, "Major": 5, "Critical": 10}

CATEGORY_DESCRIPTIONS = {
    "Accuracy": "Translation does not accurately reflect the source meaning; omission, addition, mistranslation, wrong sense, or placeholder-critical meaning issue.",
    "Readability": "The natural flow of the sentence is compromised; structure is awkward, hard to understand, over-literal, or poorly segmented.",
    "Style and Tone": "The tone, register, or product style is not preserved; wording does not match the client style guide or expected UI tone.",
    "Grammar": "Grammar, spelling, punctuation, capitalization, spacing, or syntax issue in the target language.",
    "Country Standards": "Locale/country standard issue such as date/time/number format, unit, untranslated UI term, address/currency convention, or inappropriate locale adaptation.",
}

SEVERITY_DESCRIPTIONS = {
    "Minor": "Minor impact on meaning/readability. Overall meaning is accurate and understandable. Examples: typo, punctuation, minor grammar/style issue.",
    "Major": "Major impact on meaning/readability. Translation may be confusing, misleading, incomplete, or noticeably wrong.",
    "Critical": "Critical issue that can cause serious misunderstanding, offensive output, legal/compliance risk, or unusable delivery.",
}


def count_words(text: str) -> int:
    return len(re.findall(r"[\w\u0900-\u097F\u0C00-\u0C7F\u0B80-\u0BFF\u0600-\u06FF]+", safe_text(text)))


def extract_scorecard_placeholders(text: str) -> List[str]:
    return re.findall(r"\{\{[^{}]+\}\}|\{[^{}]+\}|%\$?\d*[sd]|%[sd]|<[^>]+>", safe_text(text))


def extract_scorecard_numbers(text: str) -> List[str]:
    return re.findall(r"\d+(?:[.,:]\d+)*", safe_text(text))


def normalized_compare_text(text: str) -> str:
    text = safe_text(text).lower()
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def clean_scorecard_source_text(text: str) -> str:
    """Remove non-translatable reviewer context from source cells."""
    cleaned_lines: List[str] = []
    for raw_line in safe_text(text).splitlines():
        line = raw_line.strip()
        if not line:
            continue
        context_match = re.search(r"\b(?:description|definition)\s*:", line, flags=re.I)
        if context_match:
            line = line[:context_match.start()].strip()
        if line:
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def scorecard_category_and_severity(source: str, translator: str, reviewer: str) -> Tuple[str, str, int, str]:
    """Heuristic categorization for translator-vs-reviewer scorecards.

    The user can edit the generated categories/severities inside the Excel output.
    This provides a practical first pass from reviewer changes.
    """
    src = safe_text(source)
    tr = safe_text(translator)
    rv = safe_text(reviewer)

    if tr.strip() == rv.strip():
        return "", "", 0, "No reviewer change."
    if normalized_compare_text(tr) == normalized_compare_text(rv):
        return "Grammar", "Minor", SEVERITY_POINTS["Minor"], "Reviewer changed spacing, casing, punctuation, or minor formatting only."

    if not tr and rv:
        return "Accuracy", "Critical", SEVERITY_POINTS["Critical"], "Translator target is blank while reviewer/final contains translation."
    if tr and not rv:
        return "Accuracy", "Major", SEVERITY_POINTS["Major"], "Reviewer/final target is blank or removed compared with translator output."

    tr_ph = extract_scorecard_placeholders(tr)
    rv_ph = extract_scorecard_placeholders(rv)
    if sorted(tr_ph) != sorted(rv_ph):
        return "Accuracy", "Major", SEVERITY_POINTS["Major"], "Placeholder/tag mismatch between translator and reviewer/final output."

    tr_nums = extract_scorecard_numbers(tr)
    rv_nums = extract_scorecard_numbers(rv)
    if sorted(tr_nums) != sorted(rv_nums):
        return "Country Standards", "Major", SEVERITY_POINTS["Major"], "Number, unit, or locale-sensitive value changed between translator and reviewer/final output."

    if src and normalized_compare_text(tr) == normalized_compare_text(src) and normalized_compare_text(rv) != normalized_compare_text(src):
        return "Accuracy", "Critical", SEVERITY_POINTS["Critical"], "Translator appears to have left source text untranslated."

    ratio = difflib.SequenceMatcher(None, normalized_compare_text(tr), normalized_compare_text(rv)).ratio()
    tr_compact = re.sub(r"[\s\W_]+", "", normalized_compare_text(tr))
    rv_compact = re.sub(r"[\s\W_]+", "", normalized_compare_text(rv))

    if tr_compact == rv_compact:
        return "Grammar", "Minor", SEVERITY_POINTS["Minor"], "Reviewer changed punctuation, spacing, casing, or minor formatting only."
    if ratio >= 0.92:
        return "Grammar", "Minor", SEVERITY_POINTS["Minor"], "Reviewer made a small language/formatting correction."
    if ratio >= 0.78:
        return "Style and Tone", "Minor", SEVERITY_POINTS["Minor"], "Reviewer made a style/tone or wording refinement."
    if ratio >= 0.55:
        return "Readability", "Major", SEVERITY_POINTS["Major"], "Reviewer substantially rewrote the segment for readability or clarity."

    return "Accuracy", "Major", SEVERITY_POINTS["Major"], "Reviewer/final translation differs substantially from translator output."


def build_scorecard_records(trans_rows: List[Dict[str, Any]], rev_rows: List[Dict[str, Any]], src_rows: List[Dict[str, Any]]) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    max_len = max(len(trans_rows), len(rev_rows), len(src_rows))
    records: List[Dict[str, Any]] = []
    category_counts = {cat: {sev: 0 for sev in ERROR_SEVERITIES} for cat in ERROR_CATEGORIES}
    severity_counts = {sev: 0 for sev in ERROR_SEVERITIES}
    total_penalty = 0
    changed_count = 0
    checked_words = 0

    for i in range(max_len):
        t = trans_rows[i] if i < len(trans_rows) else {}
        r = rev_rows[i] if i < len(rev_rows) else {}
        s = src_rows[i] if i < len(src_rows) else {}

        source = clean_scorecard_source_text(s.get("source") or t.get("source") or r.get("source") or "")
        translator = t.get("target") or t.get("translation") or t.get("source", "")
        reviewer = r.get("target") or r.get("translation") or r.get("source", "")
        source = safe_text(source)
        translator = safe_text(translator)
        reviewer = safe_text(reviewer)
        if not source:
            continue
        checked_words += count_words(source or translator or reviewer)

        changed_here = translator.strip() != reviewer.strip()
        category, severity, penalty, comment = scorecard_category_and_severity(source, translator, reviewer)
        if changed_here:
            changed_count += 1
            total_penalty += penalty
            if category in category_counts and severity in category_counts[category]:
                category_counts[category][severity] += 1
            if severity in severity_counts:
                severity_counts[severity] += 1

        records.append({
            "Item No.": len(records) + 1,
            "Source Text": source,
            "Original Translation": translator,
            "Suggested Translation": reviewer if changed_here else "",
            "Repeated Error": "",
            "Error Category": category,
            "Error Severity": severity,
            "Reviewer's Comment": comment if changed_here else "",
            "Agree? (Yes/No)": "",
            "Comment": "",
            "Reviewer's Response": "",
            "Final Error Category": "",
            "Final Error Severity": "",
            "Changed": "Yes" if changed_here else "No",
            "Penalty": penalty,
        })

    score = max(0, 100 - total_penalty)
    result = "PASS" if score >= 95 else "FAIL"
    summary = {
        "segments": len(records),
        "checked_words": checked_words,
        "changed_count": changed_count,
        "total_penalty": total_penalty,
        "score": score,
        "result": result,
        "category_counts": category_counts,
        "severity_counts": severity_counts,
    }
    return records, summary


def style_sheet_base(ws) -> None:
    ws.sheet_view.showGridLines = False
    thin = Side(style="thin", color="D7DEE8")
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)


def apply_widths(ws, widths: Dict[str, float]) -> None:
    for col, width in widths.items():
        ws.column_dimensions[col].width = width


def create_scorecard_excel(records: List[Dict[str, Any]], summary: Dict[str, Any], anonymized: bool = False) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "QA Eval Sheet"
    score_ws = wb.create_sheet("Quality Evaluation")
    instr_ws = wb.create_sheet("LQA Instructions")

    dark = "1F2937"
    blue = "D9EAF7"
    green = "D9EAD3"
    yellow = "FFF2CC"
    red = "F4CCCC"
    border = Side(style="thin", color="A6A6A6")

    # Sheet 1: QA Eval Sheet
    ws.merge_cells("B1:M1")
    ws["B1"] = "ErrorSweep Linguistic Review Form"
    ws["B1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws["B1"].alignment = Alignment(horizontal="center")
    ws["B1"].fill = PatternFill("solid", fgColor=dark)

    translator_label = "Translator-001" if anonymized else ""
    reviewer_label = "Reviewer-001" if anonymized else ""
    meta_rows = [
        ("B2", "Client", "C2", "", "D2", "Project ID", "E2", ""),
        ("B3", "Source language*", "C3", "", "D3", "Target language*", "E3", ""),
        ("B4", "Translator", "C4", translator_label, "D4", "Reviewer", "E4", reviewer_label),
        ("B5", "Date (mm/dd/yyyy)*", "C5", "", "D5", "Number of checked words*", "E5", summary.get("checked_words", 0)),
    ]
    for row in meta_rows:
        for label_cell, label, value_cell, value in [(row[0], row[1], row[2], row[3]), (row[4], row[5], row[6], row[7])]:
            ws[label_cell] = label
            ws[value_cell] = value
            ws[label_cell].font = Font(bold=True)
            ws[label_cell].fill = PatternFill("solid", fgColor=blue)

    headers = [
        "Item No.", "Source Text", "Original Translation", "Suggested Translation", "Repeated Error",
        "Error Category", "Error Severity", "Reviewer's Comment", "Agree? (Yes/No)", "Comment",
        "Reviewer's Response", "Error Category", "Error Severity"
    ]
    header_row = 7
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col_idx, value=header)
        cell.font = Font(bold=True, color="000000")
        cell.fill = PatternFill("solid", fgColor=blue)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for row_idx, rec in enumerate(records, start=header_row + 1):
        values = [
            rec.get("Item No."), rec.get("Source Text"), rec.get("Original Translation"), rec.get("Suggested Translation"),
            rec.get("Repeated Error"), rec.get("Error Category"), rec.get("Error Severity"), rec.get("Reviewer's Comment"),
            rec.get("Agree? (Yes/No)"), rec.get("Comment"), rec.get("Reviewer's Response"),
            rec.get("Final Error Category"), rec.get("Final Error Severity"),
        ]
        for col_idx, value in enumerate(values, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            if col_idx in (6, 12) and value:
                cell.fill = PatternFill("solid", fgColor=yellow)
            if col_idx in (7, 13) and value:
                if value == "Critical":
                    cell.fill = PatternFill("solid", fgColor=red)
                elif value == "Major":
                    cell.fill = PatternFill("solid", fgColor="FCE4D6")
                elif value == "Minor":
                    cell.fill = PatternFill("solid", fgColor=yellow)

    apply_widths(ws, {
        "A": 10, "B": 42, "C": 42, "D": 42, "E": 14, "F": 18, "G": 16, "H": 36,
        "I": 15, "J": 26, "K": 30, "L": 18, "M": 16,
    })
    ws.freeze_panes = "A8"

    # Sheet 2: Quality Evaluation
    score_ws.merge_cells("A1:D1")
    score_ws["A1"] = "Quality Evaluation Score Card"
    score_ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    score_ws["A1"].alignment = Alignment(horizontal="center")
    score_ws["A1"].fill = PatternFill("solid", fgColor=dark)

    score_ws["A3"] = "Number of checked words"
    score_ws["B3"] = summary.get("checked_words", 0)
    score_ws["A5"] = "Client"
    score_ws["B5"] = ""
    score_ws["A6"] = "Project ID"
    score_ws["B6"] = ""
    score_ws["A7"] = "Review date"
    score_ws["B7"] = ""
    score_ws["A8"] = "Source language"
    score_ws["B8"] = ""
    score_ws["A9"] = "Target language"
    score_ws["B9"] = ""

    score_ws["A12"] = "Error category"
    score_ws["B11"] = "Error severity"
    score_ws["B12"] = "Minor"
    score_ws["C12"] = "Major"
    score_ws["D12"] = "Critical"
    for cell in score_ws["A12:D12"][0]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor=blue)

    for idx, cat in enumerate(ERROR_CATEGORIES, start=13):
        score_ws.cell(row=idx, column=1, value=cat)
        for col_idx, sev in enumerate(ERROR_SEVERITIES, start=2):
            score_ws.cell(row=idx, column=col_idx, value=summary["category_counts"].get(cat, {}).get(sev, 0))

    total_row = 13 + len(ERROR_CATEGORIES) + 1
    score_ws.cell(total_row, 1, "Total errors")
    score_ws.cell(total_row, 2, summary["severity_counts"].get("Minor", 0))
    score_ws.cell(total_row, 3, summary["severity_counts"].get("Major", 0))
    score_ws.cell(total_row, 4, summary["severity_counts"].get("Critical", 0))
    for cell in score_ws[total_row]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor=green)

    score_ws["A22"] = "Penalty points"
    score_ws["B22"] = summary.get("total_penalty", 0)
    score_ws["A23"] = "LQA score"
    score_ws["B23"] = summary.get("score", 0)
    score_ws["A24"] = "Pass/Fail result"
    score_ws["B24"] = summary.get("result", "")
    score_ws["A25"] = "Changed segments"
    score_ws["B25"] = summary.get("changed_count", 0)
    score_ws["A26"] = "Compared segments"
    score_ws["B26"] = summary.get("segments", 0)
    for cell in ["A22", "A23", "A24", "A25", "A26"]:
        score_ws[cell].font = Font(bold=True)
        score_ws[cell].fill = PatternFill("solid", fgColor=blue)

    apply_widths(score_ws, {"A": 24, "B": 16, "C": 16, "D": 16})

    # Sheet 3: LQA Instructions
    instr_ws["A1"] = "Error categories"
    instr_ws["A1"].font = Font(bold=True, size=13)
    instr_ws["A2"] = "Category"
    instr_ws["B2"] = "Description"
    instr_ws["A2"].font = Font(bold=True)
    instr_ws["B2"].font = Font(bold=True)
    instr_ws["A2"].fill = PatternFill("solid", fgColor=blue)
    instr_ws["B2"].fill = PatternFill("solid", fgColor=blue)
    for idx, cat in enumerate(ERROR_CATEGORIES, start=3):
        instr_ws.cell(row=idx, column=1, value=cat)
        instr_ws.cell(row=idx, column=2, value=CATEGORY_DESCRIPTIONS[cat])

    sev_start = 3 + len(ERROR_CATEGORIES) + 2
    instr_ws.cell(row=sev_start, column=1, value="Error severities")
    instr_ws.cell(row=sev_start, column=1).font = Font(bold=True, size=13)
    instr_ws.cell(row=sev_start + 1, column=1, value="Severity")
    instr_ws.cell(row=sev_start + 1, column=2, value="Description")
    instr_ws.cell(row=sev_start + 1, column=1).font = Font(bold=True)
    instr_ws.cell(row=sev_start + 1, column=2).font = Font(bold=True)
    instr_ws.cell(row=sev_start + 1, column=1).fill = PatternFill("solid", fgColor=blue)
    instr_ws.cell(row=sev_start + 1, column=2).fill = PatternFill("solid", fgColor=blue)
    for offset, sev in enumerate(ERROR_SEVERITIES, start=sev_start + 2):
        instr_ws.cell(row=offset, column=1, value=sev)
        instr_ws.cell(row=offset, column=2, value=SEVERITY_DESCRIPTIONS[sev])

    apply_widths(instr_ws, {"A": 24, "B": 95})
    if anonymized:
        instr_ws.cell(row=sev_start + len(ERROR_SEVERITIES) + 3, column=1, value="Privacy")
        instr_ws.cell(row=sev_start + len(ERROR_SEVERITIES) + 3, column=2, value="Translator and reviewer identifiers were anonymized for external sharing.")

    for sheet in [ws, score_ws, instr_ws]:
        style_sheet_base(sheet)
        for row in sheet.iter_rows():
            for cell in row:
                cell.border = Border(left=border, right=border, top=border, bottom=border)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        for idx in range(1, sheet.max_row + 1):
            sheet.row_dimensions[idx].height = 28

    # Validation for editable category / severity columns in QA Eval Sheet.
    cat_validation = DataValidation(type="list", formula1='"Accuracy,Readability,Style and Tone,Grammar,Country Standards"', allow_blank=True)
    sev_validation = DataValidation(type="list", formula1='"Minor,Major,Critical"', allow_blank=True)
    yesno_validation = DataValidation(type="list", formula1='"Yes,No"', allow_blank=True)
    ws.add_data_validation(cat_validation)
    ws.add_data_validation(sev_validation)
    ws.add_data_validation(yesno_validation)
    last_row = max(header_row + 1, header_row + len(records))
    cat_validation.add(f"F{header_row+1}:F{last_row}")
    cat_validation.add(f"L{header_row+1}:L{last_row}")
    sev_validation.add(f"G{header_row+1}:G{last_row}")
    sev_validation.add(f"M{header_row+1}:M{last_row}")
    yesno_validation.add(f"I{header_row+1}:I{last_row}")

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()


def page_scorecards() -> None:
    hero("Scorecards", "Translator vs reviewer quality score", "Compare translator output with reviewer/final output and generate an Excel-only LQA scorecard.")
    source = st.file_uploader("Source file (optional)", type=["xlsx", "csv", "docx", "txt"], key="score_source")
    translator = st.file_uploader("Translator file", type=["xlsx", "csv", "docx", "txt"], key="score_translator")
    reviewer = st.file_uploader("Reviewer/final file", type=["xlsx", "csv", "docx", "txt"], key="score_reviewer")
    anonymize_export = st.checkbox("Anonymize translator/reviewer IDs in Excel export", value=True, key="scorecard_anonymize")

    st.info("Scorecard output is always generated as an Excel workbook with: QA Eval Sheet, Quality Evaluation, and LQA Instructions.")
    source_options: List[Dict[str, Any]] = []
    trans_options: List[Dict[str, Any]] = []
    rev_options: List[Dict[str, Any]] = []
    source_mapping: Dict[str, Optional[int]] = {"table_idx": None, "source_col": None, "target_col": None}
    trans_mapping: Dict[str, Optional[int]] = {"table_idx": None, "source_col": None, "target_col": None}
    rev_mapping: Dict[str, Optional[int]] = {"table_idx": None, "source_col": None, "target_col": None}

    if translator is not None or reviewer is not None or source is not None:
        with st.expander("Preview detected source/target mapping", expanded=True):
            st.caption("Confirm the detected table and columns before generating. Descriptions and definitions are treated as reference context and removed from source scoring.")
            map_cols = st.columns(3)
            if translator is not None:
                with map_cols[0]:
                    trans_options, trans_mapping = render_scorecard_mapping_controls("Translator file", translator, "translator", "score_trans")
            if reviewer is not None:
                with map_cols[1]:
                    rev_options, rev_mapping = render_scorecard_mapping_controls("Reviewer file", reviewer, "reviewer", "score_rev")
            if source is not None:
                with map_cols[2]:
                    source_options, source_mapping = render_scorecard_mapping_controls("Source file", source, "source", "score_src")

    if st.button("Generate Excel Scorecard", use_container_width=True, disabled=translator is None or reviewer is None):
        trans_rows = _rows_from_mapping(trans_options, int(trans_mapping["table_idx"]), "translator", trans_mapping["source_col"], trans_mapping["target_col"]) if trans_options and trans_mapping["table_idx"] is not None else extract_rows_from_upload(translator, mode="translator")
        rev_rows = _rows_from_mapping(rev_options, int(rev_mapping["table_idx"]), "reviewer", rev_mapping["source_col"], rev_mapping["target_col"]) if rev_options and rev_mapping["table_idx"] is not None else extract_rows_from_upload(reviewer, mode="reviewer")
        src_rows = []
        if source:
            src_rows = _rows_from_mapping(source_options, int(source_mapping["table_idx"]), "source", source_mapping["source_col"], None) if source_options and source_mapping["table_idx"] is not None else extract_rows_from_upload(source, mode="source")
        records, summary = build_scorecard_records(trans_rows, rev_rows, src_rows)
        if not records:
            st.error("No scorecard-ready source rows were detected. Please review the mapping columns and try again.")
            return

        metrics([
            ("LQA Score", summary["score"], summary["result"]),
            ("Segments", summary["segments"], "compared"),
            ("Changed", summary["changed_count"], "reviewer edits"),
            ("Penalty", summary["total_penalty"], "points"),
        ])
        render_lqa_visuals(records, summary)

        preview_cols = ["Item No.", "Source Text", "Original Translation", "Suggested Translation", "Error Category", "Error Severity", "Reviewer's Comment"]
        filter_cols = st.columns(2)
        severity_filter = filter_cols[0].multiselect("Filter severity", ERROR_SEVERITIES, default=ERROR_SEVERITIES)
        category_filter = filter_cols[1].multiselect("Filter category", ERROR_CATEGORIES, default=ERROR_CATEGORIES)
        filtered_records = [
            rec for rec in records
            if (not rec.get("Error Severity") or rec.get("Error Severity") in severity_filter)
            and (not rec.get("Error Category") or rec.get("Error Category") in category_filter)
        ]
        if filtered_records:
            st.dataframe(pd.DataFrame(filtered_records)[preview_cols], use_container_width=True, hide_index=True)
        else:
            st.info("No scorecard rows match the selected filters.")
        changed_records = [r for r in records if r.get("Changed") == "Yes"]
        with st.expander("Drill down changed/error segments", expanded=False):
            if changed_records:
                selected = st.selectbox("Changed segment", [f"{r['Item No.']} - {safe_text(r.get('Error Category', ''))} / {safe_text(r.get('Error Severity', ''))}" for r in changed_records])
                detail = changed_records[[f"{r['Item No.']} - {safe_text(r.get('Error Category', ''))} / {safe_text(r.get('Error Severity', ''))}" for r in changed_records].index(selected)]
                st.markdown("##### Inline reviewer diff")
                st.markdown(inline_diff_html(detail.get("Original Translation", ""), detail.get("Suggested Translation", "")), unsafe_allow_html=True)
                st.caption(safe_text(detail.get("Reviewer's Comment", "")))
            else:
                st.caption("No changed/error segments detected.")

        xlsx_bytes = create_scorecard_excel(records, summary, anonymized=anonymize_export)
        st.download_button(
            "Download Anonymized Excel Scorecard" if anonymize_export else "Download Excel Scorecard",
            xlsx_bytes,
            file_name="ErrorSweep_Translator_Scorecard_Anonymized.xlsx" if anonymize_export else "ErrorSweep_Translator_Scorecard.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

def page_memory_rules() -> None:
    hero("Memory & Rules", "Reusable language assets", "Manage translation memory, glossary, DNT terms, and client instructions.")
    tab_intel, tab_tm, tab_gloss, tab_dnt, tab_inst = st.tabs(["Rules ZIP Analyzer", "Translation Memory", "Glossary", "DNT", "Instructions"])

    with tab_intel:
        st.markdown("### Rules ZIP Analyzer")
        rules_zip = st.file_uploader("Upload client rules ZIP", type=["zip"], key="memory_rules_zip")
        render_rules_zip_warning(rules_zip)
        if rules_zip is None:
            st.info("Upload a client ZIP containing TXT, MD, CSV, TSV, DOCX, or XLSX rules. ErrorSweep will extract glossary, DNT, and instruction hints before QA or translation.")
        else:
            parsed_rules = enrich_rules_from_chunks(parse_rules_zip(rules_zip))
            metrics([
                ("Files", len(parsed_rules.get("chunks", [])), "parsed"),
                ("Glossary", len(parsed_rules.get("glossary", [])), "pairs"),
                ("DNT", len(parsed_rules.get("dnt", [])), "locked terms"),
                ("Instructions", len(parsed_rules.get("instructions", [])), "hints"),
            ])
            summary = rules_summary_for_ai(parsed_rules)
            if summary:
                st.markdown("### AI / QA rule brief")
                st.code(summary, language="text")

            c1, c2 = st.columns(2)
            with c1:
                st.markdown("### Extracted glossary")
                glossary_rows = [
                    {"source": item.get("source_term", ""), "target": item.get("target_term", ""), "notes": f"Imported from {item.get('source', 'Rules ZIP')}"}
                    for item in parsed_rules.get("glossary", [])
                ]
                if glossary_rows:
                    st.dataframe(pd.DataFrame(glossary_rows), use_container_width=True, hide_index=True)
                else:
                    st.info("No glossary pairs detected.")
            with c2:
                st.markdown("### Extracted DNT")
                dnt_rows = [{"term": item.get("term", ""), "source": item.get("source", "Rules ZIP")} for item in parsed_rules.get("dnt", [])]
                if dnt_rows:
                    st.dataframe(pd.DataFrame(dnt_rows), use_container_width=True, hide_index=True)
                else:
                    st.info("No DNT terms detected.")

            st.markdown("### Extracted instructions")
            instruction_rows = [
                {"text": item.get("text", ""), "source": item.get("source", "Rules ZIP")}
                for item in parsed_rules.get("instructions", [])
            ]
            if instruction_rows:
                st.dataframe(pd.DataFrame(instruction_rows), use_container_width=True, hide_index=True)
            else:
                st.info("No instruction hints detected.")

            import_cols = st.columns(3)
            if import_cols[0].button("Import glossary", use_container_width=True, disabled=not glossary_rows):
                existing = {(safe_text(g.get("source")).lower(), safe_text(g.get("target")).lower()) for g in st.session_state.glossary}
                added = 0
                for row in glossary_rows:
                    key = (safe_text(row.get("source")).lower(), safe_text(row.get("target")).lower())
                    if key[0] and key not in existing:
                        st.session_state.glossary.append(row)
                        existing.add(key)
                        added += 1
                trim_session_list("glossary")
                add_audit("Rules glossary imported", f"{added} terms")
                st.success(f"Imported {added} glossary term(s).")
            if import_cols[1].button("Import DNT", use_container_width=True, disabled=not dnt_rows):
                existing_terms = {safe_text(term).lower() for term in st.session_state.dnt}
                added = 0
                for row in dnt_rows:
                    term = safe_text(row.get("term"))
                    if term and term.lower() not in existing_terms:
                        st.session_state.dnt.append(term)
                        existing_terms.add(term.lower())
                        added += 1
                trim_session_list("dnt")
                add_audit("Rules DNT imported", f"{added} terms")
                st.success(f"Imported {added} DNT term(s).")
            if import_cols[2].button("Import instructions", use_container_width=True, disabled=not instruction_rows):
                existing_text = {safe_text(item.get("text")).lower() for item in st.session_state.rule_instructions}
                added = 0
                for row in instruction_rows:
                    text = safe_text(row.get("text"))
                    if text and text.lower() not in existing_text:
                        st.session_state.rule_instructions.append({"text": text, "source": row.get("source", "Rules ZIP")})
                        existing_text.add(text.lower())
                        added += 1
                trim_session_list("rule_instructions")
                add_audit("Rules instructions imported", f"{added} instructions")
                st.success(f"Imported {added} instruction(s).")

    with tab_tm:
        st.markdown("### Translation Memory")
        with st.form("add_tm"):
            c1, c2 = st.columns(2)
            src = c1.text_area("Source", height=90)
            tgt = c2.text_area("Target", height=90)
            lang = st.text_input("Target language")
            submitted = st.form_submit_button("Add TM entry", use_container_width=True)
        if submitted and src and tgt:
            st.session_state.tm.append({"source": src, "target": tgt, "language": lang, "created": now_stamp(), "approved_by": (current_user() or {}).get("email","")})
            trim_session_list("tm")
            st.success("TM entry added.")
        st.dataframe(pd.DataFrame(st.session_state.tm), use_container_width=True, hide_index=True)

    with tab_gloss:
        st.markdown("### Glossary")
        render_tag_cloud([f"{g.get('source', '')} -> {g.get('target', '')}" for g in st.session_state.glossary], "No glossary terms yet.")
        with st.expander("Bulk import glossary pairs", expanded=False):
            bulk_gloss = st.text_area("Paste one pair per line as source -> target", height=120, key="bulk_glossary_pairs")
            if st.button("Import glossary pairs", use_container_width=True, disabled=not bulk_gloss.strip()):
                existing = {(safe_text(g.get("source")).lower(), safe_text(g.get("target")).lower()) for g in st.session_state.glossary}
                added = 0
                for line in bulk_gloss.splitlines():
                    if "->" not in line and "=>" not in line:
                        continue
                    sep = "->" if "->" in line else "=>"
                    src_term, tgt_term = [part.strip() for part in line.split(sep, 1)]
                    key = (src_term.lower(), tgt_term.lower())
                    if src_term and tgt_term and key not in existing:
                        st.session_state.glossary.append({"source": src_term, "target": tgt_term, "notes": "Bulk import"})
                        existing.add(key)
                        added += 1
                trim_session_list("glossary")
                st.success(f"Imported {added} glossary pair(s).")
        with st.form("add_gloss"):
            c1, c2, c3 = st.columns(3)
            src = c1.text_input("Source term")
            tgt = c2.text_input("Target term")
            notes = c3.text_input("Notes")
            submitted = st.form_submit_button("Add glossary term", use_container_width=True)
        if submitted and src:
            st.session_state.glossary.append({"source": src, "target": tgt, "notes": notes})
            trim_session_list("glossary")
            st.success("Glossary term added.")
        st.dataframe(pd.DataFrame(st.session_state.glossary), use_container_width=True, hide_index=True)

    with tab_dnt:
        st.markdown("### Do-not-translate terms")
        render_tag_cloud(st.session_state.dnt, "No DNT terms yet.")
        with st.expander("Bulk import DNT terms", expanded=False):
            bulk_dnt = st.text_area("Paste DNT terms separated by new lines, commas, or semicolons", height=100, key="bulk_dnt_terms")
            if st.button("Import DNT terms", use_container_width=True, disabled=not bulk_dnt.strip()):
                existing = {safe_text(term).lower() for term in st.session_state.dnt}
                added = 0
                for term in _split_rule_terms(bulk_dnt.replace("\n", ";")):
                    if term and term.lower() not in existing:
                        st.session_state.dnt.append(term)
                        existing.add(term.lower())
                        added += 1
                trim_session_list("dnt")
                st.success(f"Imported {added} DNT term(s).")
        term = st.text_input("Add DNT term")
        if st.button("Add DNT", use_container_width=True) and term:
            st.session_state.dnt.append(term)
            trim_session_list("dnt")
            st.success("DNT term added.")
        st.dataframe(pd.DataFrame({"DNT term": st.session_state.dnt}), use_container_width=True, hide_index=True)

    with tab_inst:
        st.markdown("### Client instructions")
        with st.form("add_rule_instruction"):
            instruction = st.text_area("Instruction", height=100, placeholder="Example: Keep product names in English and use formal tone for all UI strings.")
            source = st.text_input("Source / note", value="Manual")
            submitted = st.form_submit_button("Add instruction", use_container_width=True)
        if submitted and instruction:
            st.session_state.rule_instructions.append({"text": instruction, "source": source or "Manual"})
            trim_session_list("rule_instructions")
            st.success("Instruction added.")
        st.dataframe(pd.DataFrame(st.session_state.rule_instructions), use_container_width=True, hide_index=True)
        merged = workspace_rules()
        st.markdown("### Current enforced rule brief")
        st.code(rules_summary_for_ai(merged) or "No saved rules yet.", language="text")


def page_team_roles() -> None:
    hero("Team & Roles", "Workspace access control", "Manage workspace users and role-level access.")
    if current_role() not in ("Platform Owner", "Workspace Owner", "Workspace Admin", "Project Manager"):
        st.error("You do not have permission to manage team roles.")
        return
    workspace = safe_text((current_user() or {}).get("workspace") or "Demo Workspace")
    seat_state = workspace_seat_state(workspace)
    metrics([
        ("Plan", seat_state["plan"]["name"], safe_text(seat_state["subscription"].get("status", "Active"))),
        ("Seats", f"{seat_state['used']:,}/{seat_state['limit']:,}", "active + invited"),
        ("Available", seat_state["available"], "seats left"),
    ])
    if safe_text(workspace).lower() != "platform":
        st.progress(min(1.0, seat_state["used"] / max(1, seat_state["limit"])), text="Seat usage")
        if seat_state["used"] >= max(1, seat_state["limit"]):
            st.warning("This workspace is at the current plan seat limit. Upgrade from Billing or suspend an inactive user before adding another active user.")
    visible_users = [{k: v for k, v in u.items() if k != "password_hash"} for u in st.session_state.users]
    cards = []
    for user in visible_users:
        email = safe_text(user.get("email", ""))
        role = safe_text(user.get("role", "User"))
        status = safe_text(user.get("status", "Active"))
        cards.append(
            f'<div class="es-avatar-card"><div class="es-avatar">{escape(monogram(email or role))}</div>'
            f'<div><b>{escape(email or "Unknown user")}</b><br><span class="es-small">{escape(safe_text(user.get("workspace", "")))} · {escape(status)}</span><br>'
            f'<span class="es-role-badge">{escape(role)}</span></div></div>'
        )
    st.html(f'<div class="es-avatar-grid">{"".join(cards)}</div>')
    st.dataframe(pd.DataFrame(display_records(visible_users)), use_container_width=True, hide_index=True)
    with st.form("add_user", enter_to_submit=False):
        c1, c2, c3 = st.columns(3)
        email = c1.text_input("User email")
        role = c2.selectbox("Role", ["Workspace Owner", "Workspace Admin", "Project Manager", "Translator", "Reviewer", "Client Viewer", "Billing Admin", "User"])
        status = c3.selectbox("Status", ["Active", "Suspended"])
        submitted = st.form_submit_button("Add user", use_container_width=True)
    if submitted:
        if not safe_text(email):
            st.error("Please enter a user email before adding a workspace user.")
            return
        allowed, seat_message, latest_seat_state = check_workspace_seat_allowance(workspace, email, status)
        if not allowed:
            st.error(seat_message)
            return
        user_record = persist_saas_record("users", {"email": email, "workspace": workspace, "role": role, "plan": latest_seat_state["plan"]["name"], "status": status})
        st.session_state.users.append(user_record)
        trim_session_list("users")
        for workspace_record in st.session_state.get("workspaces", []):
            if safe_text(workspace_record.get("workspace")) == workspace:
                workspace_record["users"] = workspace_user_count(workspace)
                break
        add_audit("User added", email)
        queue_email_notification(
            email,
            "You were invited to ErrorSweep",
            f"You have been added to the '{workspace}' workspace as {role}. Sign in to review assigned localization work.",
            "workspace.invite",
            metadata={"role": role, "status": status},
            workspace=workspace,
        )
        st.success("User added.")


def page_billing() -> None:
    hero("Billing", "Plans and usage", "Workspace plan, credits, invoices, and payment gateway setup.")
    user = current_user() or {}
    workspace = safe_text(user.get("workspace") or "Demo Workspace")
    subscription = workspace_subscription(workspace)
    allowance = workspace_usage_allowance(workspace)
    usage = workspace_usage_totals(workspace)
    seat_state = workspace_seat_state(workspace)
    billing_provider = billing_provider_label()
    metrics([
        ("Plan", subscription.get("plan", "Trial"), safe_text(subscription.get("status", "Active"))),
        ("Seats", f"{seat_state['used']:,}/{seat_state['limit']:,}", "active + invited"),
        ("Segments", f"{usage['segments']:,}/{allowance['segments']:,}", "usage allowance"),
        ("Gateway", billing_provider, "ready" if billing_provider_ready(billing_provider) else "needs setup"),
    ])
    st.progress(min(1.0, seat_state["used"] / max(1, seat_state["limit"])), text="Seat usage")
    st.progress(min(1.0, usage["segments"] / max(1, allowance["segments"])), text="Segment usage")
    st.progress(min(1.0, usage["characters"] / max(1, allowance["characters"])), text="Character usage")
    st.caption(f"Plan allowance: {seat_state['used']:,}/{seat_state['limit']:,} seats, {usage['segments']:,}/{allowance['segments']:,} segments, and {usage['characters']:,}/{allowance['characters']:,} characters.")
    if safe_text(workspace).lower() != "platform" and seat_state["used"] >= max(1, seat_state["limit"]):
        st.warning("This workspace is at the current plan seat limit. Upgrade before inviting more active users.")
    if max(usage["segments"] / max(1, allowance["segments"]), usage["characters"] / max(1, allowance["characters"])) >= 0.85:
        st.warning("This workspace is above 85% of the current plan allowance. Upgrade before large QA or Pro runs.")

    st.markdown("### Choose a plan")
    render_pricing_graphic(subscription.get("plan", "Trial"), subscription.get("billing_cycle", "monthly"))

    with st.form("billing_checkout_intent", enter_to_submit=False):
        c1, c2 = st.columns(2)
        plan_names = [plan["name"] for plan in PLAN_CATALOG]
        current_index = plan_names.index(subscription.get("plan", "Trial")) if subscription.get("plan", "Trial") in plan_names else 0
        selected_plan = c1.selectbox("Plan", plan_names, index=current_index)
        paid_plan_names = [plan["name"] for plan in PLAN_CATALOG if plan["name"] != "Trial"]
        post_trial_plan = ""
        if selected_plan == "Trial":
            c2.text_input("Billing cycle after trial", value="monthly", disabled=True)
            billing_cycle = "monthly"
            post_trial_plan = st.selectbox("Subscription after trial", paid_plan_names, index=0)
            post_plan = plan_record(post_trial_plan)
            trial_days = configured_trial_days()
            if post_plan["name"] == "Enterprise":
                st.info(f"Trial requires a card/UPI mandate or custom payment authorization. Cancel anytime before the {trial_days}-day trial ends.")
            else:
                st.info(
                    f"Trial requires card or UPI mandate. Nothing is charged today. "
                    f"After {trial_days} days, {post_plan['name']} starts at {format_money(post_plan['monthly'], post_plan['currency'])}/month unless cancelled before the trial ends."
                )
            default_link = trial_mandate_link_for_plan(post_trial_plan, billing_cycle)
        else:
            c2.text_input("Billing cycle", value="monthly", disabled=True)
            billing_cycle = "monthly"
            selected_paid_plan = plan_record(selected_plan)
            if selected_plan == "Enterprise":
                st.info("Enterprise uses a custom card/UPI monthly mandate or payment authorization. The agreed amount is deducted every month until cancelled.")
            else:
                st.info(
                    f"{selected_plan} uses a card/UPI monthly mandate. "
                    f"{format_money(selected_paid_plan['monthly'], selected_paid_plan['currency'])} will be deducted every month until cancelled."
                )
            default_link = monthly_mandate_link_for_plan(selected_plan, billing_cycle)
        payment_link = st.text_input(
            "Card/UPI monthly mandate link",
            value="",
            placeholder=default_link or "https://checkout.stripe.com/... or https://rzp.io/...",
            help="Paste a hosted monthly mandate link, or configure ERRORSWEEP_MONTHLY_MANDATE_LINK_PRO / ERRORSWEEP_CARD_UPI_MANDATE_LINK_PRO style secrets.",
        )
        accept_trial_terms = True
        if selected_plan == "Trial":
            accept_trial_terms = st.checkbox(
                f"I understand the trial can be cancelled anytime before day {configured_trial_days()}, and the selected subscription starts afterward.",
                value=False,
            )
        create_checkout = st.form_submit_button(
            "Start trial with card/UPI mandate" if selected_plan == "Trial" else "Create card/UPI monthly mandate",
            use_container_width=True,
        )
    if create_checkout:
        if safe_text(payment_link) and not sanitize_payment_link(payment_link):
            st.error("Please enter a valid http(s) card/UPI monthly mandate link.")
            return
        if selected_plan == "Trial":
            if post_trial_plan not in paid_plan_names:
                st.error("Please select the subscription that should start after the trial.")
                return
            if not accept_trial_terms:
                st.error("Please confirm the trial cancellation and post-trial subscription terms.")
                return
            if not (sanitize_payment_link(payment_link) or sanitize_payment_link(trial_mandate_link_for_plan(post_trial_plan, billing_cycle))):
                st.error("Trial requires a valid card/UPI mandate link. Paste one here or configure ERRORSWEEP_TRIAL_MANDATE_LINK.")
                return
        else:
            if not (sanitize_payment_link(payment_link) or sanitize_payment_link(monthly_mandate_link_for_plan(selected_plan, billing_cycle))):
                st.error("This subscription requires a valid card/UPI monthly mandate link. Paste one here or configure ERRORSWEEP_MONTHLY_MANDATE_LINK.")
                return
        intent = create_checkout_intent(selected_plan, billing_cycle, payment_link=payment_link, post_trial_plan=post_trial_plan)
        if safe_text(intent.get("checkout_url")):
            if selected_plan == "Trial":
                st.success("Trial mandate recorded. Open the link below to add card/UPI authorization. User can cancel anytime before the trial ends.")
                st.link_button("Open card / UPI mandate link", intent["checkout_url"], use_container_width=True)
            else:
                st.success("Monthly mandate recorded. Open it below to authorize recurring card/UPI deduction.")
                st.link_button("Open card / UPI monthly mandate link", intent["checkout_url"], use_container_width=True)
        elif safe_text(intent.get("status")) == "manual_pending":
            st.info("Checkout intent recorded. Add a hosted mandate link above, or configure Stripe/Razorpay mandate-link secrets to show a live link.")
        elif safe_text(intent.get("status")) == "trial_mandate_link_missing":
            st.warning("Trial intent recorded, but card/UPI mandate link is missing.")
        elif safe_text(intent.get("status")) == "monthly_mandate_link_missing":
            st.warning("Mandate intent recorded, but card/UPI monthly mandate link is missing.")
        else:
            st.warning("Checkout intent recorded, but provider credentials are incomplete.")

    st.markdown("### Cancel trial or subscription")
    pending_checkout = next(
        (
            item for item in st.session_state.get("checkout_sessions", [])
            if safe_text(item.get("workspace")) == workspace
            and safe_text(item.get("status")).lower() not in {"cancelled", "canceled", "expired", "paid", "completed"}
        ),
        None,
    )
    subscription_status = safe_text(subscription.get("status")).lower()
    has_persisted_subscription = bool(safe_text(subscription.get("id")))
    can_cancel_billing = (
        (has_persisted_subscription and subscription_status not in {"cancelled", "canceled", "expired"})
        or pending_checkout is not None
    )
    if can_cancel_billing:
        active_label = (
            f"{subscription.get('plan', 'subscription')} subscription"
            if has_persisted_subscription and subscription_status not in {"cancelled", "canceled", "expired"}
            else f"{pending_checkout.get('plan', 'Trial')} pending mandate"
        )
        st.caption(f"Cancel the current {active_label}. Trial users can cancel anytime before the trial end date to stop the post-trial monthly mandate.")
        with st.form("cancel_billing_path", enter_to_submit=False):
            cancel_reason = st.text_area("Cancellation reason", value="No longer needed", height=80)
            confirm_cancel = st.checkbox("I understand this cancels the current trial/payment mandate or active subscription.", value=False)
            cancel_submit = st.form_submit_button("Cancel trial / subscription", use_container_width=True)
        if cancel_submit:
            if not confirm_cancel:
                st.error("Please confirm cancellation before continuing.")
                return
            cancelled_kind, cancelled_record = cancel_current_billing(cancel_reason)
            if cancelled_record:
                st.success("Cancellation recorded. Billing status and notification outbox were updated.")
                st.rerun()
            else:
                st.info("No active subscription or pending mandate was found to cancel.")
    else:
        st.info("No active subscription or pending trial mandate is currently available to cancel.")

    if is_owner():
        with st.expander("Owner action: activate selected subscription locally", expanded=False):
            st.caption("Use this only after manual payment confirmation or while testing provider webhooks locally.")
            c1, c2 = st.columns(2)
            activate_plan = c1.selectbox("Activate plan", [plan["name"] for plan in PLAN_CATALOG], key="activate_subscription_plan")
            c2.text_input("Activate cycle", value="monthly", disabled=True, key="activate_subscription_cycle_display")
            activate_cycle = "monthly"
            if st.button("Mark subscription active", use_container_width=True):
                activate_subscription(activate_plan, activate_cycle, billing_provider)
                st.success("Subscription record activated.")
                st.rerun()

    st.markdown("### Subscription records")
    subscription_rows = [
        item for item in st.session_state.get("subscriptions", [])
        if safe_text(item.get("workspace")) == workspace or is_owner()
    ]
    if subscription_rows:
        st.dataframe(pd.DataFrame(display_records(subscription_rows)), use_container_width=True, hide_index=True)
    else:
        st.info("No persisted subscription records yet. New public trials should complete card/UPI mandate authorization before activation.")

    st.markdown("### Checkout intents")
    checkout_rows = [
        item for item in st.session_state.get("checkout_sessions", [])
        if safe_text(item.get("workspace")) == workspace or is_owner()
    ]
    if checkout_rows:
        st.dataframe(pd.DataFrame(display_records(checkout_rows)), use_container_width=True, hide_index=True)
    else:
        st.info("No checkout intents yet.")

    if is_owner():
        st.markdown("### Provider webhook reconciliation")
        st.caption("Paste a Stripe/Razorpay webhook event here for local testing. A deployed webhook endpoint can reuse the same normalization logic later.")
        with st.form("billing_webhook_reconcile", enter_to_submit=False):
            c1, c2 = st.columns(2)
            webhook_provider = c1.selectbox("Provider", ["auto", "razorpay", "stripe", "manual"], index=0)
            apply_updates = c2.checkbox("Apply subscription/payment updates", value=True)
            signature_header = st.text_input("Signature header (optional)", placeholder="Stripe-Signature or X-Razorpay-Signature value")
            webhook_secret = st.text_input("Webhook secret override (optional)", type="password", help="Leave blank to use STRIPE_WEBHOOK_SECRET, RAZORPAY_WEBHOOK_SECRET, or ERRORSWEEP_BILLING_WEBHOOK_SECRET.")
            raw_event = st.text_area("Webhook JSON", height=180, placeholder='{"event":"payment.captured","payload":{"payment":{"entity":{"amount":399900,"currency":"INR","notes":{"workspace":"Demo Workspace","plan":"Pro"}}}}}')
            webhook_submit = st.form_submit_button("Record webhook event", use_container_width=True)
        if webhook_submit:
            if not raw_event.strip():
                st.error("Paste a webhook JSON payload first.")
                return
            try:
                event_record, messages = record_billing_webhook_event(
                    webhook_provider,
                    raw_event,
                    signature_header=signature_header,
                    webhook_secret=webhook_secret,
                    apply_updates=apply_updates,
                )
                st.success(f"Billing event recorded: {event_record.get('provider')} / {event_record.get('event_type')} / {event_record.get('status')}")
                for message in messages:
                    st.caption(message)
                st.rerun()
            except Exception as exc:
                LOGGER.warning("Billing webhook reconciliation failed: %s", exc)
                st.error(f"Unable to record webhook event: {exc}")

    st.markdown("### Billing events")
    billing_event_rows = [
        item for item in st.session_state.get("billing_events", [])
        if safe_text(item.get("workspace")) == workspace or is_owner()
    ]
    if billing_event_rows:
        preview_events = []
        for item in billing_event_rows[:100]:
            row = dict(item)
            if isinstance(row.get("metadata_json"), (dict, list)):
                row["metadata_json"] = json.dumps(row["metadata_json"], ensure_ascii=False)[:260]
            preview_events.append(row)
        st.dataframe(pd.DataFrame(display_records(preview_events)), use_container_width=True, hide_index=True)
    else:
        st.info("No billing webhook events recorded yet.")


def page_account() -> None:
    hero("Account", "Profile and workspace preferences", "Manage user profile, workspace settings, and notification preferences.")
    user = current_user() or {}
    st.write("Email:", user.get("email"))
    st.write("Role:", user.get("role"))
    st.write("Workspace:", user.get("workspace"))

    st.markdown("### AI access")
    st.caption("Use included AI, or add any OpenAI-compatible API key, model, and base URL available to you. Your key is kept only in this session for the MVP.")
    current_mode = current_ai_route_label()
    st.info(f"Current route: {current_mode}")
    with st.form("byo_key_form"):
        presets = {
            "OpenAI": "",
            "OpenAI-compatible / Custom": st.session_state.get("byo_ai_base_url", ""),
            "OpenRouter": "https://openrouter.ai/api/v1",
            "Groq": "https://api.groq.com/openai/v1",
            "Together AI": "https://api.together.xyz/v1",
            "Fireworks AI": "https://api.fireworks.ai/inference/v1",
            "Google Gemini OpenAI-compatible": "https://generativelanguage.googleapis.com/v1beta/openai",
            "Local vLLM / LM Studio": "http://127.0.0.1:8000/v1",
        }
        provider = st.selectbox("API provider", list(presets.keys()), index=list(presets.keys()).index(st.session_state.get("byo_ai_provider", "OpenAI-compatible / Custom")) if st.session_state.get("byo_ai_provider") in presets else 1)
        default_base = presets.get(provider, "")
        byo_base_url = st.text_input(
            "Base URL (leave blank for OpenAI)",
            value=st.session_state.get("byo_ai_base_url", default_base) or default_base,
            help="Use any OpenAI-compatible chat/completions endpoint. Examples: Groq, OpenRouter, Together, Fireworks, Gemini OpenAI-compatible, vLLM, LM Studio.",
        )
        byo_key = st.text_input("Your API key", type="password", placeholder="Paste provider API key", help="Leave blank to use included/self-hosted routes if configured.")
        byo_model = st.text_input("Model name", value=st.session_state.get("byo_openai_model", secret("ERRORSWEEP_OPENAI_DEFAULT_MODEL", DEFAULT_MODEL)))
        col_a, col_b = st.columns(2)
        save_key = col_a.form_submit_button("Use this key", use_container_width=True)
        clear_key = col_b.form_submit_button("Clear BYO key", use_container_width=True)
    if save_key:
        if byo_key.strip():
            st.session_state["byo_openai_api_key"] = byo_key.strip()
            st.session_state["byo_openai_model"] = byo_model.strip() or DEFAULT_MODEL
            st.session_state["byo_ai_provider"] = provider
            st.session_state["byo_ai_base_url"] = byo_base_url.strip()
            st.success(f"BYO {provider} key activated for this session.")
        else:
            st.warning("No key entered. Included AI will be used if configured.")
    if clear_key:
        st.session_state.pop("byo_openai_api_key", None)
        st.session_state.pop("byo_openai_model", None)
        st.session_state.pop("byo_ai_provider", None)
        st.session_state.pop("byo_ai_base_url", None)
        st.success("BYO key cleared. Included AI will be used if configured.")

    st.checkbox("Email notifications", value=True)
    st.checkbox("Show review hints", value=True)
    my_email = safe_text(user.get("email", "")).lower()
    my_notifications = [
        item for item in st.session_state.get("notifications", [])
        if safe_text(item.get("recipient", "")).lower() == my_email
    ]
    st.markdown("### Notification history")
    if my_notifications:
        rows = []
        for item in my_notifications[:25]:
            rows.append({
                "created": item.get("created", ""),
                "event_type": item.get("event_type", ""),
                "subject": item.get("subject", ""),
                "status": item.get("status", ""),
                "provider": item.get("provider", ""),
                "sent_at": item.get("sent_at", ""),
            })
        st.dataframe(pd.DataFrame(display_records(rows)), use_container_width=True, hide_index=True)
    else:
        st.info("No notifications for this account yet.")


def page_admin() -> None:
    hero("Admin", "Workspace admin", "Workspace-level configuration and maintenance.")
    if current_role() not in ("Platform Owner", "Workspace Owner", "Workspace Admin"):
        st.error("Admin access is restricted.")
        return
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("### Workspace summary")
        metrics([
            ("Projects", len(st.session_state.projects), ""),
            ("Jobs", len(st.session_state.jobs), ""),
            ("Review rows", len(st.session_state.review_segments), ""),
            ("TM", len(st.session_state.tm), ""),
        ])
    with c2:
        st.markdown("### Maintenance")
        if st.button("Clear demo jobs/review only", use_container_width=True):
            st.session_state.jobs = []
            st.session_state.review_segments = []
            st.session_state.subtitle_segments = []
            st.success("Demo jobs and review rows cleared.")
        if st.button("Clear all demo workspace data", use_container_width=True):
            for key in ["projects", "jobs", "tm", "review_segments", "subtitle_segments", "last_pro_review_segments", "latest_human_review_segments", "pro_review_rows"]:
                st.session_state[key] = []
            st.session_state["pro_post_editing_ready"] = False
            st.success("Demo workspace data cleared.")


# Owner pages

def page_owner_console() -> None:
    hero("Owner Console", "Private platform owner view", "Only your master account can see global payments, users, workspaces, usage, and platform controls.")
    metrics([
        ("Workspaces", len(st.session_state.workspaces), "all customer/client spaces"),
        ("Users", len(st.session_state.users), "all access records"),
        ("Payments", len(st.session_state.payments), "received or demo records"),
        ("Audit Logs", len(st.session_state.audit_logs), "platform events"),
    ])

    st.markdown("### Release persistence")
    if persistence_health is not None:
        health = persistence_health()
        render_topology_map(health)
        h1, h2, h3, h4 = st.columns(4)
        h1.metric("Storage", health.get("storage_mode", "unknown"))
        h2.metric("Supabase", "Ready" if health.get("supabase_configured") else "Fallback")
        h3.metric("Jobs table", health.get("editor_jobs_table", "unknown"))
        h4.metric("Usage table", health.get("usage_events_table", "unknown"))
        with st.expander("Persistence diagnostics", expanded=False):
            diagnostic_rows: List[Dict[str, str]] = []

            def collect_diagnostics(prefix: str, value: Any) -> None:
                if isinstance(value, dict):
                    for child_key, child_value in value.items():
                        collect_diagnostics(f"{prefix}.{child_key}" if prefix else safe_text(child_key), child_value)
                else:
                    diagnostic_rows.append({"Check": prefix, "Value": safe_text(value)})

            collect_diagnostics("", health)
            st.dataframe(pd.DataFrame(diagnostic_rows), use_container_width=True, hide_index=True)
    else:
        st.warning("production_persistence.py is not available. Editor jobs are using session/local fallback only.")

    st.markdown("### Current / recent task details")
    active_job_id = st.session_state.get("active_review_session_id", "")
    active_rows = st.session_state.get("review_segments") or st.session_state.get("last_pro_review_segments") or []
    last_task = st.session_state.get("last_pro_task_details") or {}
    session_jobs = st.session_state.get("owner_recent_editor_jobs", [])

    if active_job_id or last_task or active_rows:
        t1, t2, t3, t4 = st.columns(4)
        t1.metric("Active job", str(active_job_id or last_task.get("id", "—"))[:10] if (active_job_id or last_task.get("id")) else "—")
        t2.metric("File", last_task.get("file_name", st.session_state.get("review_workspace_file_name", "—")) or "—")
        t3.metric("Rows", len(active_rows) or int(last_task.get("row_count") or 0))
        t4.metric("Target", last_task.get("target_language", st.session_state.get("review_workspace_language", "—")) or "—")

        with st.expander("Current task row preview", expanded=False):
            preview_rows = []
            for i, row in enumerate(active_rows[:25], start=1):
                preview_rows.append({
                    "No": i,
                    "Source": safe_text(row.get("source", ""))[:180],
                    "Target": safe_text(row.get("target", row.get("translation", "")))[:180],
                    "Status": safe_text(row.get("status", "")),
                    "Match": safe_text(row.get("match", "")),
                    "Location": safe_text(row.get("location", "")),
                })
            if preview_rows:
                st.dataframe(pd.DataFrame(preview_rows), use_container_width=True, hide_index=True)
            else:
                st.info("The active task exists, but no row preview is available in session.")
    else:
        st.info("No active Pro review task is currently stored in this browser session.")

    if session_jobs:
        with st.expander("Session-created editor jobs", expanded=True):
            st.dataframe(pd.DataFrame(display_records(session_jobs)), use_container_width=True, hide_index=True)

    st.markdown("### Owner actions")
    c1, c2, c3 = st.columns(3)
    c1.info("Review all workspace access from User Access Matrix.")
    c2.info("Track received payments from Payments Received.")
    c3.info("Control global feature flags from Platform Settings.")

    st.markdown("### Translation / AI usage")
    persistent_usage = []
    if fetch_persistent_usage_events is not None:
        try:
            persistent_usage = fetch_persistent_usage_events(300)
        except Exception as exc:
            LOGGER.error("Unable to fetch persistent usage events: %s", exc)
            persistent_usage = []
    usage_rows = persistent_usage or st.session_state.get("ai_usage_events", [])
    if usage_rows:
        usage_df = pd.DataFrame(display_records(usage_rows))
        usage_values = []
        for row in usage_rows[:24]:
            usage_values.append(int(row.get("characters", row.get("segments", 0)) or 0))
        st.html(
            f"""
            <div class="es-dashboard-panel">
              <div class="es-dashboard-title"><h3>Usage velocity</h3><span class="es-code-chip">AI / MT</span></div>
              {area_chart_svg(list(reversed(usage_values or [0])), "ownerUsage")}
            </div>
            """
        )
        st.dataframe(usage_df, use_container_width=True, hide_index=True)
        if "characters" in usage_df.columns:
            st.caption(f"Total characters logged: {int(pd.to_numeric(usage_df['characters'], errors='coerce').fillna(0).sum())}")
    else:
        st.info("No usage logged yet.")

    st.markdown("### Recent editor jobs")
    editor_jobs = []
    if fetch_persistent_editor_jobs is not None:
        try:
            editor_jobs = fetch_persistent_editor_jobs(100)
        except Exception as exc:
            LOGGER.error("Unable to fetch persistent editor jobs: %s", exc)
            editor_jobs = []
    combined_jobs = []
    if editor_jobs:
        combined_jobs.extend(editor_jobs)
    for job in st.session_state.get("owner_recent_editor_jobs", []):
        if not any(str(j.get("id")) == str(job.get("id")) for j in combined_jobs):
            combined_jobs.append(job)

    if combined_jobs:
        st.dataframe(pd.DataFrame(display_records(combined_jobs)), use_container_width=True, hide_index=True)
    else:
        st.info("No persisted editor jobs found yet. Run Pro, click Open Human Review Editor, then return here.")


def page_payments_received() -> None:
    hero("Payments Received", "Revenue and payment records", "Owner-only list of payments, plans, access granted, and payment status.")
    with st.form("add_payment"):
        c1, c2, c3, c4 = st.columns(4)
        workspace = c1.text_input("Workspace")
        user = c2.text_input("User email")
        plan = c3.selectbox("Plan", [p["name"] for p in PLAN_CATALOG])
        amount = c4.number_input("Amount (INR)", min_value=0.0, value=0.0)
        submitted = st.form_submit_button("Add payment record", use_container_width=True)
    if submitted:
        plan_currency = plan_record(plan).get("currency", "INR")
        payment_record = persist_saas_record("payments", {"date": now_stamp(), "workspace": workspace, "user": user, "plan": plan, "amount": amount, "currency": plan_currency, "status": "Recorded"})
        st.session_state.payments.insert(0, payment_record)
        trim_session_list("payments")
        add_audit("Payment record added", f"{workspace}: {amount}")
        queue_email_notification(
            user,
            "ErrorSweep payment recorded",
            f"Payment record added for workspace '{workspace}': {format_money(amount, plan_currency)} on plan {plan}.",
            "billing.payment_recorded",
            metadata={"workspace": workspace, "plan": plan, "amount": amount, "currency": plan_currency},
            workspace=workspace,
        )
        st.success("Payment record added.")
    st.dataframe(pd.DataFrame(display_records(st.session_state.payments)), use_container_width=True, hide_index=True)


def page_user_access_matrix() -> None:
    hero("User Access Matrix", "Who has what access", "Owner-only view of user roles, workspaces, plans, statuses, and allowed pages.")
    rows = []
    for u in st.session_state.users:
        role = u.get("role", "User")
        safe_user = {k: v for k, v in u.items() if k != "password_hash"}
        rows.append({**safe_user, "allowed_pages": ", ".join(ROLE_PAGE_ACCESS.get(role, []))})
    st.dataframe(pd.DataFrame(display_records(rows)), use_container_width=True, hide_index=True)


def page_all_workspaces() -> None:
    hero("All Workspaces", "Customer workspace overview", "Owner-only list of all organizations, plans, users, and job counts.")
    st.dataframe(pd.DataFrame(display_records(st.session_state.workspaces)), use_container_width=True, hide_index=True)


def page_platform_settings() -> None:
    hero("Platform Settings", "Global feature controls", "Owner-only controls for platform features and public availability.")
    health = persistence_health() if persistence_health is not None else {}
    st.markdown("### Public launch readiness")
    readiness_rows = launch_readiness_rows(health)
    ready_count = sum(1 for row in readiness_rows if safe_text(row.get("Launch gate")).lower() in {"ready", "ready to send", "ready to test provider"} or "configured" in safe_text(row.get("Launch gate")).lower())
    metrics([
        ("Launch items", len(readiness_rows), "tracked"),
        ("Ready/configured", ready_count, "local or external"),
        ("Needs setup", max(0, len(readiness_rows) - ready_count), "before public launch"),
        ("Mode", "Production" if is_production_mode() else "Local", "environment"),
    ])
    st.dataframe(pd.DataFrame(readiness_rows), use_container_width=True, hide_index=True)

    st.markdown("### Feature flags")
    settings = {
        "Main API translation": True,
        "Pro post-editing Human Review": True,
        "Scorecards": True,
        "Subtitle / Transcription Editor": True,
        "Public registration": False,
        "Billing collection": False,
        "Self-hosted engines": False,
    }
    for label, val in settings.items():
        st.checkbox(label, value=val)

    st.markdown("### Object storage")
    storage_health = object_storage_status() if object_storage_status is not None else {"provider": "local", "bucket": "local", "configured": False, "mode": "local_fallback"}
    metrics([
        ("Provider", storage_health.get("provider", "local"), storage_health.get("mode", "local_fallback")),
        ("Bucket", storage_health.get("bucket", "local"), "configured" if storage_health.get("configured") else "needs setup"),
        ("Local cache", "Available", "fallback/cache"),
        ("Cloud ready", "Yes" if storage_health.get("mode") == "cloud" else "No", "Supabase/S3/GCS"),
    ])
    st.caption(f"Local cache path: {storage_health.get('local_root', '')}")
    if storage_health.get("mode") == "cloud":
        st.success("Cloud object storage is configured. New job attachments and media previews will be stored through the object storage adapter.")
    else:
        st.info("Object storage is in local fallback mode. Set ERRORSWEEP_OBJECT_STORAGE_PROVIDER plus Supabase Storage, S3, or GCS bucket secrets before public multi-instance launch.")

    st.markdown("### File manifest")
    file_rows = st.session_state.get("files", [])
    if file_rows:
        metrics([
            ("Tracked files", len(file_rows), "workspace manifests"),
            ("Stored bytes", sum(int(item.get("size_bytes") or 0) for item in file_rows), "object metadata"),
            ("Job uploads", sum(1 for item in file_rows if safe_text(item.get("purpose")) == "job_assignment"), "assignment files"),
            ("Media previews", sum(1 for item in file_rows if safe_text(item.get("purpose")) == "media_preview"), "preview files"),
        ])
        display_files = []
        for item in file_rows[:100]:
            row = dict(item)
            if row.get("storage_key"):
                row["storage_key"] = safe_text(row["storage_key"])[:180]
            display_files.append(row)
        st.dataframe(pd.DataFrame(display_records(display_files)), use_container_width=True, hide_index=True)
    else:
        st.info("No file manifests yet. Manual job assignment uploads and generated downloads will appear here.")

    st.markdown("### Async task queue")
    async_health = async_backend_status() if async_backend_status is not None else {"provider": "local", "ready": False, "mode": "local_inline"}
    task_rows = st.session_state.get("task_queue", [])
    task_summary = task_status_summary(task_rows)
    metrics([
        ("Backend", async_health.get("provider", "local"), async_health.get("mode", "local_inline")),
        ("Worker ready", "Yes" if async_health.get("ready") and async_health.get("mode") == "external" else "No", "external queue"),
        ("Tracked tasks", len(task_rows), "workspace lifecycle records"),
        ("Running", task_summary.get("running", 0), "active"),
    ])
    metrics([
        ("Failed", task_summary.get("failed", 0), "needs retry"),
        ("Completed", task_summary.get("completed", 0), "finished"),
        ("Redis queue", async_health.get("redis_queue", "errorsweep:tasks"), "if Redis/Celery"),
        ("HTTP worker", "Configured" if safe_text(async_health.get("worker_url")) else "Not set", "if endpoint"),
    ])
    if async_health.get("mode") == "external" and async_health.get("ready"):
        st.success("External async handoff is configured. New QA and Pro jobs will queue to the worker instead of running inside the Streamlit request.")
    else:
        st.info("Async jobs are running inline locally. Set ERRORSWEEP_ASYNC_WORKER_URL or REDIS_URL/CELERY_BROKER_URL to hand heavy QA and Pro work to a worker.")
    if task_rows:
        display_rows = []
        for item in task_rows[:100]:
            row = dict(item)
            if isinstance(row.get("metadata_json"), (dict, list)):
                row["metadata_json"] = json.dumps(row["metadata_json"], ensure_ascii=False)
            display_rows.append(row)
        st.dataframe(pd.DataFrame(display_records(display_rows)), use_container_width=True, hide_index=True)
    else:
        st.info("No task queue records yet. QA and Pro workflows will create lifecycle entries as they run.")

    st.markdown("### Email notification outbox")
    notifications = st.session_state.get("notifications", [])
    provider = email_provider_label()
    pending_count = sum(1 for item in notifications if safe_text(item.get("status")).lower() in {"queued", "provider_pending", "failed"})
    sent_count = sum(1 for item in notifications if safe_text(item.get("status")).lower() == "sent")
    metrics([
        ("Outbox", len(notifications), "workspace notifications"),
        ("Pending", pending_count, "queued/provider pending"),
        ("Sent", sent_count, "delivered"),
        ("Provider", provider, "resend/sendgrid/smtp/manual"),
    ])
    if provider in {"", "manual", "not_configured"}:
        st.info("Email dispatch is in outbox-only mode. Set ERRORSWEEP_EMAIL_PROVIDER to resend, sendgrid, or smtp with provider credentials to send real emails.")
    else:
        if st.button("Send pending notifications", use_container_width=True):
            sent, failed = dispatch_pending_notifications()
            if sent:
                st.success(f"Sent {sent} notification(s).")
            if failed:
                st.warning(f"{failed} notification(s) could not be sent. Check the error column.")
            if not sent and not failed:
                st.info("No queued notifications found.")
    if notifications:
        notification_rows = []
        for item in notifications[:100]:
            row = dict(item)
            if isinstance(row.get("metadata_json"), (dict, list)):
                row["metadata_json"] = json.dumps(row["metadata_json"], ensure_ascii=False)
            if row.get("body"):
                row["body"] = safe_text(row["body"])[:260]
            notification_rows.append(row)
        st.dataframe(pd.DataFrame(display_records(notification_rows)), use_container_width=True, hide_index=True)
    else:
        st.info("No notification records yet. Signup, team invites, QA completion, Pro completion, and payment records will create outbox entries.")

    st.markdown("### Release readiness diagnostics")
    if persistence_health is not None:
        render_topology_map(health)
        diagnostic_rows: List[Dict[str, str]] = []

        def collect_diagnostics(prefix: str, value: Any) -> None:
            if isinstance(value, dict):
                for child_key, child_value in value.items():
                    collect_diagnostics(f"{prefix}.{child_key}" if prefix else safe_text(child_key), child_value)
            else:
                raw = safe_text(value)
                state = "Ready" if raw.lower() in {"true", "ok", "connected"} else "Needs setup" if raw.lower() in {"false", "not_checked", "not configured", "not_configured", ""} else raw
                diagnostic_rows.append({"Check": prefix, "Status": state, "Value": raw})

        collect_diagnostics("", health)
        st.dataframe(pd.DataFrame(diagnostic_rows), use_container_width=True, hide_index=True)
        if health.get("supabase_configured") and health.get("editor_jobs_table") == "ok" and health.get("usage_events_table") == "ok":
            st.success("Production persistence is connected. Editor jobs and usage events can survive app reboot.")
        else:
            st.warning("Production persistence is not fully connected. Run the v42 Supabase SQL schema and add SUPABASE_URL + SUPABASE_SERVICE_ROLE_KEY in Streamlit secrets.")
    else:
        st.warning("production_persistence.py is missing.")


def page_platform_audit_logs() -> None:
    hero("Platform Audit Logs", "Owner event trail", "Owner-only view of sign-ins, payments, access changes, and administrative actions.")
    if st.session_state.audit_logs:
        st.dataframe(pd.DataFrame(display_records(st.session_state.audit_logs)), use_container_width=True, hide_index=True)
    else:
        st.info("No audit logs yet.")


PAGE_RENDERERS = {
    "Dashboard": page_dashboard,
    "Projects": page_projects,
    "Jobs": page_jobs,
    "ErrorSweep QA": page_qa,
    "ErrorSweep Pro": page_pro,
    "Subtitle / Transcription Editor": page_subtitle_transcription_editor,
    "Human Review Workspace": page_human_review_workspace,
    "Subtitle Workspace": page_subtitle_workspace,
    "Transcription Workspace": page_transcription_workspace,
    "Scorecards": page_scorecards,
    "Memory & Rules": page_memory_rules,
    "Team & Roles": page_team_roles,
    "Billing": page_billing,
    "Account": page_account,
    "Admin": page_admin,
    "Owner Console": page_owner_console,
    "Payments Received": page_payments_received,
    "User Access Matrix": page_user_access_matrix,
    "All Workspaces": page_all_workspaces,
    "Platform Settings": page_platform_settings,
    "Platform Audit Logs": page_platform_audit_logs,
}


# ==========================================================
# Main app
# ==========================================================

def render_app() -> None:
    maybe_cleanup_media_preview_files()
    if query_get("es_logout"):
        logout()
        return
    user = current_user()
    if not user:
        render_public_app()
        return
    hydrate_saas_state_for_user()

    # v41: external editor routes open as full-window pages in a new tab.
    # They must render before normal dashboard routing/navigation.
    if render_external_editor_router():
        return

    # Restore selected page from URL query when navigation links are used.
    requested_page = query_get("es_page")
    if requested_page in allowed_pages():
        st.session_state.page = requested_page

    # Ensure selected page is allowed.
    if st.session_state.page not in allowed_pages():
        st.session_state.page = allowed_pages()[0] if allowed_pages() else "Dashboard"

    page = st.session_state.page
    renderer = PAGE_RENDERERS.get(page, page_dashboard)

    # Dedicated editor workspaces should feel like full web applications, not
    # like a normal dashboard page squeezed beside the platform navigation.
    # This is especially important for the CAT-style Human Review editor.
    if page in {"Human Review Workspace", "Subtitle Workspace", "Transcription Workspace"}:
        renderer()
        return

    render_navigation()
    renderer()


render_app()
